from __future__ import annotations

import base64
import csv
import html
import hashlib
import json
import os
import re
import shutil
import subprocess
import sys
import time
from datetime import datetime, timezone
from io import BytesIO, StringIO
from pathlib import Path
from typing import Any
from urllib.parse import urljoin, urlparse

import httpx
try:
    from bs4 import BeautifulSoup
except Exception:  # pragma: no cover - optional parser fallback
    BeautifulSoup = None  # type: ignore[assignment]
from fastapi import APIRouter, File, Form, HTTPException, Request, UploadFile
from fastapi.responses import FileResponse, HTMLResponse, JSONResponse, Response
from pydantic import BaseModel, Field
from .facebook_session_runtime import ensure_runtime_session, latest_saved_session, runtime_session_available

from .utils.env import load_dotenv
from .security import (
    DEFAULT_PERMISSIONS,
    current_user_from_auth_header,
    deactivate_user,
    list_public_users,
    upsert_user,
)

load_dotenv()

try:
    from selenium import webdriver
    from selenium.webdriver.chrome.service import Service as ChromeService
    from selenium.webdriver.common.by import By
except Exception:  # pragma: no cover - selenium is optional at import time
    webdriver = None
    ChromeService = None
    By = None

ROOT_DIR = Path(__file__).resolve().parents[1]
DATA_DIR = ROOT_DIR / "data"
RUNTIME_DIR = ROOT_DIR / "runtime"
FACEBOOK_POSTS_DIR = RUNTIME_DIR / "facebook_posts"
FML_DIR = ROOT_DIR / "automation" / "facebook-marketplace-lister"
FML_ACCOUNTS_PATH = FML_DIR / "accounts.json"
FML_IMAGES_DIR = FML_DIR / "images"
FML_DRIVERS_DIR = FML_DIR / "drivers"
ADMIN_BUNDLE_INDEX = ROOT_DIR / "app" / "static" / "admin" / "index.html"
SALES_FRONTEND_INDEX = ROOT_DIR / "sales-assistant" / "frontend" / "dist" / "index.html"
SALES_BACKEND_ENTRYPOINT = ROOT_DIR / "sales-assistant" / "backend" / "dist" / "index.js"
DEFAULT_SALES_ASSISTANT_BACKEND_URL = "http://127.0.0.1:4300"
INVENTORY_SNAPSHOT_PATH = DATA_DIR / "latest" / "inventory_porsche_wb.json"
INVENTORY_LIVE_CACHE_PATH = DATA_DIR / "latest" / "inventory_dealership_live.json"
INVENTORY_LIVE_META_PATH = DATA_DIR / "latest" / "inventory_dealership_live_meta.json"
INVENTORY_LIVE_BACKUP_PATH = DATA_DIR / "latest" / "inventory_dealership_live.backup.json"
INVENTORY_LIVE_META_BACKUP_PATH = DATA_DIR / "latest" / "inventory_dealership_live_meta.backup.json"
INVENTORY_MANUAL_PATH = DATA_DIR / "latest" / "inventory_manual.json"
JD_POWER_VALUATIONS_PATH = DATA_DIR / "latest" / "jd_power_trade_values.json"
DEFAULT_DEALERSHIP_INVENTORY_URL = "https://www.tavernachryslerdodgejeepramfiat.com/used-inventory/index.htm"
DEFAULT_DEALERSHIP_NEW_INVENTORY_URL = "https://www.tavernachryslerdodgejeepramfiat.com/new-inventory/index.htm"
DEFAULT_BANK_TAX_RATE = 0.06
DEFAULT_BANK_FEES = 2400
FACEBOOK_MARKETPLACE_DEFAULT_LOCATION = os.getenv("FACEBOOK_MARKETPLACE_DEFAULT_LOCATION", "33317").strip() or "33317"
FACEBOOK_MARKETPLACE_LOCATION_LABEL = os.getenv("FACEBOOK_MARKETPLACE_LOCATION_LABEL", "Plantation, FL 33317").strip() or "Plantation, FL 33317"
FACEBOOK_SELLER_MESSENGER_LINK = os.getenv("FACEBOOK_SELLER_MESSENGER_LINK", "https://m.me/AniTheCarGuy").strip() or "https://m.me/AniTheCarGuy"
try:
    FACEBOOK_MARKETPLACE_PRICE_BUMP = int(os.getenv("FACEBOOK_MARKETPLACE_PRICE_BUMP", "2400") or "2400")
except ValueError:
    FACEBOOK_MARKETPLACE_PRICE_BUMP = 2400
try:
    FACEBOOK_MARKETPLACE_DOWN_PAYMENT = int(os.getenv("FACEBOOK_MARKETPLACE_DOWN_PAYMENT", "999") or "999")
except ValueError:
    FACEBOOK_MARKETPLACE_DOWN_PAYMENT = 999
FACEBOOK_POST_STATUS_PATH = RUNTIME_DIR / "facebook_post_status.json"
FACEBOOK_MARKETPLACE_SYNC_PATH = RUNTIME_DIR / "facebook_marketplace_sync.json"
VEHICLE_ASSETS_CACHE_DIR = RUNTIME_DIR / "vehicle_assets"
VEHICLE_ASSET_FILE_CACHE_DIR = RUNTIME_DIR / "vehicle_asset_files"
BANK_BRAIN_HISTORY_PATH = RUNTIME_DIR / "bank_brain_history.json"
BANK_BRAIN_AUDIT_PATH = RUNTIME_DIR / "bank_brain_audit.json"
CARFAX_CHROME_PROFILE_DIR = Path(os.getenv("CARFAX_CHROME_PROFILE_DIR", str(RUNTIME_DIR / "carfax_auth_profile"))).resolve()
BANK_PROFILES_GENERATED_PATH = DATA_DIR / "bank_profiles.generated.json"
SALES_ASSISTANT_BANKS_PATH = ROOT_DIR / "sales-assistant" / "data" / "banks.json"
BANK_DOCS_ROOT = Path(os.getenv("BANK_DOCS_ROOT", str(ROOT_DIR / "Bank"))).resolve()
BANK_DOCS_DECODED_DIR = RUNTIME_DIR / "routeone_docs" / "decoded"
BANK_DOCS_INDEX_PATH = RUNTIME_DIR / "routeone_docs" / "decoded_index.json"
BANK_DOCS_GENERATED_INDEX_PATH = DATA_DIR / "routeone_docs.decoded_index.generated.json"
BANK_DOCS_LINK_CACHE_DIR = RUNTIME_DIR / "routeone_docs" / "linked_cache"
XCONSOLE_STATE_DIR = Path(
    os.getenv("XCONSOLE_STATE_DIR", str(BANK_DOCS_ROOT / "_xconsole"))
).resolve()
DEALERSHIPS_CONFIG_PATH = XCONSOLE_STATE_DIR / "dealerships.json"
LEADS_PATH = DATA_DIR / "post_lead" / "leads.json"
LEAD_RESPONSES_PATH = XCONSOLE_STATE_DIR / "lead_responses.json"
LEAD_SYNC_STATE_PATH = XCONSOLE_STATE_DIR / "lead_sync_state.json"


def _facebook_token(raw_value: Any) -> str:
    token = str(raw_value or "")
    return re.sub(r"\s+", "", token).strip()


def _facebook_token_candidates(*names: str) -> list[str]:
    tokens: list[str] = []
    for name in names:
        token = _facebook_token(os.getenv(name))
        if token and token not in tokens:
            tokens.append(token)
    return tokens


def _facebook_named_token_candidates(*names: str) -> list[tuple[str, str]]:
    candidates: list[tuple[str, str]] = []
    seen: set[str] = set()
    for name in names:
        token = _facebook_token(os.getenv(name))
        if token and token not in seen:
            candidates.append((name, token))
            seen.add(token)
    return candidates


def _token_preview(token: str) -> str:
    clean = str(token or "")
    if len(clean) <= 14:
        return clean
    return f"{clean[:10]}...{clean[-4:]}"


def _facebook_graph_diagnose_token(
    token: str,
    *,
    scope: str,
    entity_id: str = "me",
    fields: str = "id,name",
    timeout: float = 10.0,
) -> dict[str, Any]:
    if not token:
        return {
            "ok": False,
            "mode": "missing_token",
            "status": "missing",
            "scope": scope,
            "fields": fields,
            "entity_id": entity_id,
            "status_code": 0,
        }

    try:
        with httpx.Client(timeout=timeout) as client:
            response = client.get(
                f"https://graph.facebook.com/v19.0/{entity_id}",
                params={"fields": fields, "access_token": token},
            )
    except Exception as exc:
        return {
            "ok": False,
            "mode": "connection_error",
            "scope": scope,
            "entity_id": entity_id,
            "status_code": 0,
            "error": str(exc),
        }

    try:
        payload = response.json()
    except Exception:
        payload = {}
    status_code = response.status_code
    if status_code >= 400:
        graph_error = payload.get("error", {}) if isinstance(payload, dict) else {}
        raw_error_code = graph_error.get("code") if isinstance(graph_error, dict) else None
        try:
            error_code = int(raw_error_code) if raw_error_code is not None else status_code
        except Exception:
            error_code = status_code
        if error_code == 190:
            mode = "token_invalid"
        elif error_code in {298, 100}:
            mode = "conversation_permission_missing"
        else:
            mode = "http_error"
        return {
            "ok": False,
            "mode": mode,
            "scope": scope,
            "entity_id": entity_id,
            "status_code": status_code,
            "graph_status": error_code,
            "payload": payload.get("error") if isinstance(payload, dict) else payload,
        }

    return {
        "ok": True,
        "mode": "ok",
        "scope": scope,
        "entity_id": entity_id,
        "status_code": status_code,
        "payload": {
            "id": payload.get("id"),
            "name": payload.get("name"),
        },
    }
OFFERUP_STATUS_PATH = XCONSOLE_STATE_DIR / "offerup_status.json"
OFFERUP_POSTS_DIR = XCONSOLE_STATE_DIR / "offerup_posts"
VIN_DECODE_CACHE_DIR = XCONSOLE_STATE_DIR / "vin_decode"
CARFAX_SUMMARY_DIR = DATA_DIR / "carfax_summaries"
XCONSOLE_RELEASE_TAG = "marketplace-state-carfax-leads-2026-04-30"

router = APIRouter()


def _running_on_windows() -> bool:
    return os.name == "nt" or sys.platform.startswith("win")


def _env_flag(name: str, default: bool = False) -> bool:
    raw = os.getenv(name)
    if raw is None:
        return default
    return raw.strip().lower() not in {"0", "false", "no", "off"}


class FacebookPostRequest(BaseModel):
    vin: str = Field(..., min_length=3)
    title: str = Field(..., min_length=3)
    price: str | int | float = Field(...)
    model: str | None = None
    mileage: str | int | None = None
    body_style: str | None = None
    fuel_type: str | None = None
    condition: str | None = None
    drivetrain: str | None = None
    engine: str | None = None
    transmission: str | None = None
    location: str | None = None
    exterior: str | None = None
    interior: str | None = None
    detail_url: str | None = None
    description: str | None = None
    highlights: list[str] = Field(default_factory=list)
    images: list[str] = Field(default_factory=list)
    account_id: str | None = None
    suggested_down_payment: int | None = None
    mode: str = Field(default="draft", pattern="^(draft|live)$")


class FacebookPreflightRequest(BaseModel):
    account_id: str | None = None
    images: list[str] = Field(default_factory=list)
    vin: str | None = None


class CarfaxReportTextRequest(BaseModel):
    report_text: str = Field(..., min_length=20)
    source_url: str | None = None


class FacebookBootstrapRequest(BaseModel):
    create_template_account_if_missing: bool = True


class FacebookVehicleImageImportRequest(BaseModel):
    vin: str = Field(..., min_length=3)
    limit: int = Field(default=20, ge=1, le=60)
    overwrite: bool = False


class FacebookPrepareLivePostRequest(BaseModel):
    vin: str = Field(..., min_length=3)
    account_id: str | None = None
    import_missing_images: bool = True
    image_limit: int = Field(default=20, ge=1, le=60)
    overwrite_images: bool = False


class FacebookRelinkImagesRequest(BaseModel):
    vin: str = Field(..., min_length=3)
    images: list[str] = Field(default_factory=list)
    include_vin_matches: bool = False
    overwrite: bool = False
    delete_source: bool = False


class FacebookFullRepairRequest(BaseModel):
    vin: str | None = None
    ensure_placeholder_images: bool = True
    placeholder_count: int = Field(default=6, ge=1, le=30)


class WireEverythingRequest(BaseModel):
    vin: str | None = None
    ensure_placeholder_images: bool = True
    placeholder_count: int = Field(default=6, ge=1, le=30)
    reload_sales_data: bool = True


class InventoryLiveSyncRequest(BaseModel):
    source_url: str | None = None
    timeout_seconds: int = Field(default=180, ge=5, le=300)
    persist: bool = True


class DealershipRequest(BaseModel):
    id: str | None = None
    name: str = Field(..., min_length=2)
    preowned_url: str | None = None
    used_url: str | None = None
    new_url: str | None = None
    active: bool = True


class ManualVehicleAddRequest(BaseModel):
    vin: str = Field(..., min_length=3)
    title: str = Field(..., min_length=3)
    price: str | int | float | None = None
    mileage: str | int | float | None = None
    drivetrain: str | None = None
    engine: str | None = None
    transmission: str | None = None
    location: str | None = None
    detail_url: str | None = None
    exterior: str | None = None
    interior: str | None = None
    photos: list[str] = Field(default_factory=list)


class FacebookOneClickPostRequest(BaseModel):
    vin: str = Field(..., min_length=3)
    account_id: str | None = None
    selected_photo_indexes: list[int] = Field(default_factory=list)
    skip_photo_indexes: list[int] = Field(default_factory=lambda: [0, 2])
    caption_override: str | None = None
    mode: str = Field(default="live", pattern="^(draft|live)$")
    auto_import_photos: bool = True
    photo_limit: int = Field(default=24, ge=1, le=60)


class FacebookBatchPostRequest(BaseModel):
    vins: list[str] = Field(..., min_length=1, max_length=25)
    account_id: str | None = None
    skip_photo_indexes: list[int] = Field(default_factory=lambda: [0, 2])
    auto_import_photos: bool = True
    photo_limit: int = Field(default=24, ge=1, le=60)


class FacebookMarketplaceSyncRequest(BaseModel):
    verify_live_urls: bool = True
    processing_review_minutes: int = Field(default=45, ge=1, le=1440)


class BankBrainAnalyzeRequest(BaseModel):
    report_text: str | None = None
    structured_data: dict[str, Any] = Field(default_factory=dict)
    requested_amount: float | None = None


class CreditStructureRequest(BaseModel):
    vin: str | None = None
    vehicle_price: float = Field(..., ge=0)
    book_value: float | None = Field(default=None, ge=0)
    taxes: float = Field(default=0, ge=0)
    tax_rate: float = Field(default=DEFAULT_BANK_TAX_RATE, ge=0, le=1)
    fees: float = Field(default=DEFAULT_BANK_FEES, ge=0)
    backend_products: float = Field(default=0, ge=0)
    down_payment: float = Field(default=0, ge=0)
    term_months: int = Field(default=72, ge=12, le=96)
    apr: float = Field(default=9.99, ge=0, le=35)
    monthly_income: float | None = Field(default=None, ge=0)
    current_dti: float | None = Field(default=None, ge=0, le=100)
    credit_score: int | None = Field(default=None, ge=300, le=850)
    tradelines: int | None = Field(default=None, ge=0, le=200)
    derogatories: int | None = Field(default=None, ge=0, le=50)
    utilization: float | None = Field(default=None, ge=0, le=100)


class BankBrainVehicleRecommendationRequest(BaseModel):
    score: int | None = Field(default=None, ge=300, le=850)
    monthly_income: float | None = Field(default=None, ge=0)
    current_dti: float | None = Field(default=None, ge=0, le=100)
    down_payment: float = Field(default=0, ge=0)
    desired_payment: float | None = Field(default=None, ge=0)
    max_results: int = Field(default=12, ge=1, le=50)


class BankBrainDecisionRequest(BaseModel):
    vin: str | None = None
    bank_code: str = Field(..., min_length=2)
    outcome: str = Field(..., pattern="^(approved|declined|countered)$")
    notes: str | None = None
    metrics: dict[str, Any] = Field(default_factory=dict)


class BankBrainDocsRebuildRequest(BaseModel):
    reload_sales_data: bool = True
    max_link_depth: int = Field(default=1, ge=0, le=3)
    max_links_per_resource: int = Field(default=12, ge=0, le=80)


class XconsoleUserRequest(BaseModel):
    username: str = Field(..., min_length=2)
    password: str | None = None
    display_name: str | None = None
    role: str = Field(default="operator", pattern="^(admin|manager|operator)$")
    permissions: list[str] = Field(default_factory=list)
    active: bool = True


class LeadManualAddRequest(BaseModel):
    customer_name: str = Field(default="Unknown Lead")
    channel: str = Field(default="facebook")
    message: str = Field(..., min_length=1)
    vehicle_vin: str | None = None
    source: str = Field(default="manual")


class LeadRespondRequest(BaseModel):
    lead_id: str = Field(..., min_length=1)
    response_text: str = Field(default="")
    channel: str = Field(default="facebook")
    mark_status: str = Field(default="responded")
    attachment_url: str | None = None
    attachment_type: str | None = None


class OfferUpPostRequest(BaseModel):
    vin: str = Field(..., min_length=3)
    caption_override: str | None = None
    mode: str = Field(default="draft", pattern="^(draft|live)$")


def _safe_read_json(path: Path, fallback: Any) -> Any:
    try:
        return json.loads(path.read_text(encoding="utf-8"))
    except Exception:
        return fallback


def _safe_write_json(path: Path, payload: Any) -> None:
    path.parent.mkdir(parents=True, exist_ok=True)
    serialized = json.dumps(payload, indent=2)
    tmp_path = path.with_name(f"{path.name}.tmp-{time.time_ns()}")
    tmp_path.write_text(serialized, encoding="utf-8")
    tmp_path.replace(path)


def _utc_now() -> str:
    return datetime.now(timezone.utc).isoformat()


def _deployment_fingerprint() -> dict[str, Any]:
    return {
        "release": XCONSOLE_RELEASE_TAG,
        "project_name": os.getenv("RAILWAY_PROJECT_NAME") or "zealous-caring",
        "service_name": os.getenv("RAILWAY_SERVICE_NAME") or "Xconsole-Dealership-Tool",
        "service_id": os.getenv("RAILWAY_SERVICE_ID"),
        "deployment_id": os.getenv("RAILWAY_DEPLOYMENT_ID") or os.getenv("RAILWAY_DEPLOY_ID"),
        "public_domain": os.getenv("RAILWAY_PUBLIC_DOMAIN") or "xconsole.up.railway.app",
        "environment": os.getenv("RAILWAY_ENVIRONMENT_NAME") or "production",
    }


def _display_path(path: Path) -> str:
    try:
        return str(path.resolve().relative_to(ROOT_DIR)).replace("\\", "/")
    except Exception:
        return str(path)


def _sanitize_doc_segment(value: str, fallback: str) -> str:
    cleaned = re.sub(r"[^A-Za-z0-9._ -]+", "_", value.strip())
    cleaned = re.sub(r"\s+", " ", cleaned).strip(" ._-")
    return cleaned[:140] or fallback


def _stable_id(value: str, fallback: str = "item") -> str:
    cleaned = re.sub(r"[^a-z0-9]+", "-", str(value or "").strip().lower()).strip("-")
    if cleaned:
        return cleaned[:80]
    return f"{fallback}-{hashlib.sha1(str(time.time()).encode()).hexdigest()[:8]}"


def _normalize_dealership(raw: Any) -> dict[str, Any] | None:
    if not isinstance(raw, dict):
        return None
    name = str(raw.get("name") or raw.get("label") or "").strip()
    if not name:
        return None
    dealer_id = str(raw.get("id") or "").strip() or _stable_id(name, "dealer")
    urls = {
        "preowned_url": str(raw.get("preowned_url") or raw.get("pre_owned_url") or "").strip(),
        "used_url": str(raw.get("used_url") or raw.get("inventory_url") or "").strip(),
        "new_url": str(raw.get("new_url") or "").strip(),
    }
    active = bool(raw.get("active", True))
    return {
        "id": dealer_id,
        "name": name,
        **urls,
        "active": active,
        "source_urls": _split_inventory_source_urls(
            "\n".join(value for value in urls.values() if value)
        ),
    }


def _default_dealerships() -> list[dict[str, Any]]:
    return [
        {
            "id": "taverna-cdjr",
            "name": "Taverna CDJR",
            "preowned_url": DEFAULT_DEALERSHIP_INVENTORY_URL,
            "used_url": DEFAULT_DEALERSHIP_INVENTORY_URL,
            "new_url": DEFAULT_DEALERSHIP_NEW_INVENTORY_URL,
            "active": True,
            "source_urls": _split_inventory_source_urls(
                f"{DEFAULT_DEALERSHIP_INVENTORY_URL}\n{DEFAULT_DEALERSHIP_NEW_INVENTORY_URL}"
            ),
        }
    ]


def _load_dealerships() -> list[dict[str, Any]]:
    payload = _safe_read_json(DEALERSHIPS_CONFIG_PATH, {"items": []})
    raw_items = payload.get("items", []) if isinstance(payload, dict) else payload
    if not isinstance(raw_items, list):
        raw_items = []
    normalized = [item for item in (_normalize_dealership(raw) for raw in raw_items) if item]
    return normalized or _default_dealerships()


def _save_dealership(request: DealershipRequest) -> dict[str, Any]:
    dealer_id = str(request.id or "").strip() or _stable_id(request.name, "dealer")
    row = {
        "id": dealer_id,
        "name": request.name.strip(),
        "preowned_url": str(request.preowned_url or "").strip(),
        "used_url": str(request.used_url or "").strip(),
        "new_url": str(request.new_url or "").strip(),
        "active": bool(request.active),
        "updated_at": _utc_now(),
    }
    normalized = _normalize_dealership(row)
    if not normalized or not normalized.get("source_urls"):
        raise HTTPException(
            status_code=400,
            detail={"message": "Add at least one valid http(s) inventory URL."},
        )

    existing = [item for item in _load_dealerships() if item.get("id") != dealer_id]
    next_items = existing + [normalized]
    _safe_write_json(
        DEALERSHIPS_CONFIG_PATH,
        {
            "items": next_items,
            "updated_at": _utc_now(),
        },
    )
    return normalized


def _configured_dealership_source_urls() -> list[str]:
    urls: list[str] = []
    seen: set[str] = set()
    for dealership in _load_dealerships():
        if not dealership.get("active", True):
            continue
        for url in dealership.get("source_urls") or []:
            lowered = str(url).lower()
            if lowered in seen:
                continue
            seen.add(lowered)
            urls.append(str(url))
    return urls


def _default_inventory_source_url() -> str:
    return ", ".join(_default_inventory_source_urls())


def _split_inventory_source_urls(raw: str | None) -> list[str]:
    text = str(raw or "").strip()
    if not text:
        return []
    candidates = re.split(r"[\r\n,]+", text)
    urls = []
    seen = set()
    for candidate in candidates:
        url = candidate.strip()
        if not url or not re.match(r"^https?://", url, flags=re.IGNORECASE):
            continue
        lowered = url.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        urls.append(url)
    return urls


def _default_inventory_source_urls() -> list[str]:
    urls: list[str] = []
    seen: set[str] = set()

    def add_many(values: list[str]) -> None:
        for url in values:
            lowered = url.lower()
            if lowered in seen:
                continue
            seen.add(lowered)
            urls.append(url)

    add_many(_split_inventory_source_urls(os.getenv("DEALERSHIP_INVENTORY_URLS")))
    used = str(os.getenv("DEALERSHIP_INVENTORY_URL", DEFAULT_DEALERSHIP_INVENTORY_URL)).strip()
    new = str(os.getenv("DEALERSHIP_NEW_INVENTORY_URL", DEFAULT_DEALERSHIP_NEW_INVENTORY_URL)).strip()
    add_many(_split_inventory_source_urls(", ".join([used, new])))
    add_many(_configured_dealership_source_urls())
    return urls or [DEFAULT_DEALERSHIP_INVENTORY_URL]


def _to_price_text(value: str | int | float) -> str:
    if isinstance(value, (int, float)):
        return f"${value:,.0f}"
    stripped = str(value).strip()
    if stripped.startswith("$"):
        return stripped
    digits = re.sub(r"[^\d.]", "", stripped)
    if not digits:
        return stripped
    try:
        numeric = float(digits)
        return f"${numeric:,.0f}"
    except ValueError:
        return stripped


def _facebook_marketplace_price(value: Any) -> str | int | float:
    numeric = _to_float(value)
    if numeric is None:
        return value if value is not None else ""
    return int(round(numeric + FACEBOOK_MARKETPLACE_PRICE_BUMP))


def _facebook_marketplace_price_text(value: Any) -> str:
    return _to_price_text(_facebook_marketplace_price(value))


def _facebook_suggested_down_payment(vehicle_or_price: Any = None) -> int:
    ltv: float | None = None
    price = vehicle_or_price
    jd_trade = None
    if isinstance(vehicle_or_price, dict):
        price = vehicle_or_price.get("price")
        jd_trade = vehicle_or_price.get("jd_power_trade_in")
        ltv = _to_float(vehicle_or_price.get("jd_power_ltv"))
    if ltv is None:
        price_number = _to_float(price)
        trade_number = _to_float(jd_trade)
        if price_number and trade_number and trade_number > 0:
            financed_basis = price_number + DEFAULT_BANK_FEES
            ltv = ((financed_basis + (financed_basis * DEFAULT_BANK_TAX_RATE)) / trade_number) * 100
    if ltv is None:
        suggested = FACEBOOK_MARKETPLACE_DOWN_PAYMENT
    elif ltv <= 90:
        suggested = 750
    elif ltv <= 98:
        suggested = 999
    elif ltv <= 108:
        suggested = 1499
    elif ltv <= 118:
        suggested = 1999
    else:
        suggested = 2999
    return int(max(750, min(2999, round(suggested))))


def _facebook_post_price(value: Any = None) -> int:
    return _facebook_suggested_down_payment(value)


def _facebook_post_price_text(_value: Any = None) -> str:
    return _to_price_text(_facebook_post_price(_value))


def _facebook_listing_location(value: Any = None) -> str:
    lowered = str(value or "").strip().lower()
    if "fort lauderdale" in lowered:
        return "Fort Lauderdale, FL 33317"
    if "33317" in lowered or "plantation" in lowered:
        return "Plantation, FL 33317"
    return FACEBOOK_MARKETPLACE_LOCATION_LABEL


def _facebook_caption_mileage(value: Any) -> str:
    digits = re.sub(r"[^\d]", "", str(value or ""))
    if digits:
        return f"{int(digits):,} mi"
    return str(value or "").strip()


def _render_listing_text(payload: FacebookPostRequest) -> str:
    description = str(payload.description or "").strip()
    messenger_cta = (
        "Message my seller page for the full walkaround, CARFAX, and the quickest reply: "
        f"{FACEBOOK_SELLER_MESSENGER_LINK}"
    )
    if description:
        description = re.sub(r"^https?://\S+\s*$", "", description, flags=re.IGNORECASE | re.MULTILINE).strip()
        down_text = _to_price_text(payload.suggested_down_payment or _facebook_post_price(payload.price))
        description = re.sub(r"\$[1-9]\d{1,2},\d{3}(?:\.\d+)?(?:\s*(?:plus|\\+)?\s*tax(?:es)?\.?)?", down_text, description, flags=re.IGNORECASE)
        if "down payment" not in description.lower():
            description = f"{payload.title.strip()}\n{down_text} down payment options for qualified buyers.\n{description}"
        if FACEBOOK_SELLER_MESSENGER_LINK.lower() not in description.lower():
            description = f"{description.strip()}\n\n{messenger_cta}"
        return description.strip() + "\n"

    lines: list[str] = []
    lines.append(payload.title.strip())
    lines.append(f"{_to_price_text(payload.suggested_down_payment or _facebook_post_price(payload.price))} down payment options for qualified buyers.")

    if payload.mileage:
        lines.append(f"Mileage: {payload.mileage} miles")
    if payload.drivetrain:
        lines.append(f"Drivetrain: {payload.drivetrain}")
    if payload.engine:
        lines.append(f"Engine: {payload.engine}")
    if payload.transmission:
        lines.append(f"Transmission: {payload.transmission}")
    if payload.location:
        lines.append(f"Location: {payload.location}")
    if payload.exterior:
        lines.append(f"Exterior: {payload.exterior}")
    if payload.interior:
        lines.append(f"Interior: {payload.interior}")

    for item in payload.highlights[:3]:
        cleaned = item.strip()
        if cleaned:
            lines.append(f"- {cleaned}")

    lines.extend(["", messenger_cta])

    return "\n".join(lines) + "\n"


def _write_listing_text(vin: str, text: str) -> Path:
    timestamp = int(time.time())
    target_dir = FACEBOOK_POSTS_DIR / f"{vin}_{timestamp}"
    target_dir.mkdir(parents=True, exist_ok=True)
    target_file = target_dir / "facebook_listing.txt"
    target_file.write_text(text, encoding="utf-8")
    return target_file


def _extract_line_value(lines: list[str], label: str) -> str | None:
    needle = f"{label.lower()}:"
    for line in lines:
        if line.lower().startswith(needle):
            return line.split(":", 1)[1].strip()
    return None


def _parse_listing_file(path: Path) -> dict[str, Any]:
    text = path.read_text(encoding="utf-8", errors="ignore")
    lines = [line.strip() for line in text.splitlines() if line.strip()]

    title = lines[0] if lines else path.parent.name
    vin = path.parent.name.split("_", 1)[0]
    ts_raw = path.parent.name.split("_", 1)[1] if "_" in path.parent.name else "0"

    return {
        "vin": vin,
        "timestamp": int(re.sub(r"[^\d]", "", ts_raw) or "0"),
        "title": title,
        "price": _extract_line_value(lines, "Price"),
        "mileage": _extract_line_value(lines, "Mileage"),
        "drivetrain": _extract_line_value(lines, "Drivetrain"),
        "engine": _extract_line_value(lines, "Engine"),
        "transmission": _extract_line_value(lines, "Transmission"),
        "location": _extract_line_value(lines, "Location"),
        "detail_url": next((line for line in reversed(lines) if line.startswith("http")), None),
        "file": str(path.relative_to(ROOT_DIR)).replace("\\", "/"),
        "text": text,
    }


def _to_float(value: Any) -> float | None:
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        return float(value)
    if value is None:
        return None
    cleaned = re.sub(r"[^0-9.\-]+", "", str(value))
    if cleaned in {"", ".", "-", "-."}:
        return None
    try:
        return float(cleaned)
    except ValueError:
        return None


def _public_permission_catalog() -> list[dict[str, str]]:
    labels = {
        "inventory.view": "View inventory",
        "inventory.edit": "Edit inventory",
        "dealerships.manage": "Manage dealerships",
        "facebook.post": "Post to Facebook",
        "facebook.leads": "Read/respond to Facebook leads",
        "assets.view": "View vehicle assets",
        "stickers.view": "View stickers",
        "carfax.view": "View Carfax links",
        "offerup.post": "Create OfferUp drafts",
        "bankbrain.view": "Use Bank Brain",
        "bankbrain.train": "Upload RouteOne forms",
        "users.manage": "Manage users",
        "admin.full": "Full admin",
    }
    return [{"id": item, "label": labels.get(item, item)} for item in DEFAULT_PERMISSIONS]


def _lead_id(seed: str) -> str:
    return hashlib.sha1(seed.encode("utf-8", errors="ignore")).hexdigest()[:16]


def _normalize_lead(raw: Any, index: int = 0) -> dict[str, Any] | None:
    if not isinstance(raw, dict):
        return None
    message = str(
        raw.get("message")
        or raw.get("last_message")
        or raw.get("body")
        or raw.get("text")
        or ""
    ).strip()
    customer_name = str(
        raw.get("customer_name")
        or raw.get("name")
        or raw.get("sender_name")
        or raw.get("from")
        or "Unknown Lead"
    ).strip()
    vehicle_vin = str(raw.get("vehicle_vin") or raw.get("vin") or "").strip().upper()
    channel = str(raw.get("channel") or raw.get("source") or "facebook").strip().lower()
    created_at = str(raw.get("created_at") or raw.get("timestamp") or raw.get("created_time") or _utc_now())
    seed = "|".join([customer_name, vehicle_vin, message, created_at, str(index)])
    messages = raw.get("messages") if isinstance(raw.get("messages"), list) else []
    responses = raw.get("responses") if isinstance(raw.get("responses"), list) else []
    thread = raw.get("thread") if isinstance(raw.get("thread"), list) else []
    return {
        "id": str(raw.get("id") or raw.get("lead_id") or _lead_id(seed)),
        "customer_name": customer_name,
        "channel": channel,
        "message": message or "No message captured yet.",
        "vehicle_vin": vehicle_vin,
        "source": str(raw.get("source") or channel).strip(),
        "status": str(raw.get("status") or "new").strip().lower(),
        "created_at": created_at,
        "last_message_at": raw.get("last_message_at") or created_at,
        "conversation_id": raw.get("conversation_id"),
        "profile_id": raw.get("profile_id"),
        "messages": messages,
        "responses": responses,
        "thread": thread,
    }


def _load_lead_responses() -> list[dict[str, Any]]:
    payload = _safe_read_json(LEAD_RESPONSES_PATH, {"responses": []})
    responses = payload.get("responses", []) if isinstance(payload, dict) else []
    if not isinstance(responses, list):
        return []
    return [item for item in responses if isinstance(item, dict)]


def _lead_thread_from_history(lead: dict[str, Any], responses: list[dict[str, Any]]) -> list[dict[str, Any]]:
    incoming = lead.get("messages") if isinstance(lead.get("messages"), list) else []
    thread: list[dict[str, Any]] = []
    for item in incoming:
        if not isinstance(item, dict):
            continue
        text = str(item.get("text") or item.get("message") or "").strip()
        attachments = _normalize_thread_attachments(item.get("attachments"))
        if not text and not attachments:
            continue
        thread.append(
            {
                "direction": str(item.get("direction") or "inbound"),
                "text": text,
                "created_at": str(item.get("created_at") or lead.get("created_at") or _utc_now()),
                "author": str(item.get("author") or item.get("from_name") or lead.get("customer_name") or "Buyer"),
                "delivery_status": str(item.get("delivery_status") or "received"),
                "attachments": attachments,
            }
        )
    for item in responses:
        text = str(item.get("response_text") or "").strip()
        attachments = _normalize_thread_attachments(item.get("attachments"))
        if not text and not attachments:
            continue
        if not _is_visible_thread_response(item):
            continue
        thread.append(
            {
                "direction": "outbound",
                "text": text,
                "created_at": str(item.get("created_at") or _utc_now()),
                "author": str(item.get("author") or "xconsole"),
                "delivery_status": str(item.get("delivery_status") or "logged"),
                "attachments": attachments,
            }
        )
    thread.sort(key=lambda row: str(row.get("created_at") or ""))
    return thread


def _load_leads() -> list[dict[str, Any]]:
    payload = _safe_read_json(LEADS_PATH, [])
    raw_items = payload.get("items", []) if isinstance(payload, dict) else payload
    if not isinstance(raw_items, list):
        raw_items = []
    leads = [lead for lead in (_normalize_lead(item, index) for index, item in enumerate(raw_items)) if lead]
    responses = _load_lead_responses()
    responses_by_lead: dict[str, list[dict[str, Any]]] = {}
    responded_ids = {str(item.get("lead_id")) for item in responses if isinstance(item, dict)}
    for item in responses:
        lead_id = str(item.get("lead_id") or "").strip()
        if not lead_id:
            continue
        responses_by_lead.setdefault(lead_id, []).append(item)
    for lead in leads:
        if lead["id"] in responded_ids and lead.get("status") == "new":
            lead["status"] = "responded"
        lead_responses = responses_by_lead.get(str(lead.get("id") or ""), [])
        lead["responses"] = lead_responses
        lead["thread"] = _lead_thread_from_history(lead, lead_responses)
    return sorted(leads, key=lambda item: str(item.get("last_message_at") or ""), reverse=True)


def _save_leads(leads: list[dict[str, Any]]) -> None:
    _safe_write_json(
        LEADS_PATH,
        {
            "items": leads,
            "updated_at": _utc_now(),
        },
    )


def _append_lead_response(
    *,
    lead_id: str,
    channel: str,
    response_text: str,
    author: str | None,
    delivery_status: str = "logged",
    provider_message_id: str | None = None,
    error_detail: str | None = None,
    attachments: list[dict[str, Any]] | None = None,
) -> dict[str, Any]:
    payload = _safe_read_json(LEAD_RESPONSES_PATH, {"responses": []})
    responses = payload.get("responses", []) if isinstance(payload, dict) else []
    if not isinstance(responses, list):
        responses = []
    entry = {
        "id": _lead_id(f"{lead_id}|{response_text}|{time.time()}"),
        "lead_id": lead_id,
        "channel": channel,
        "response_text": response_text,
        "author": author or "xconsole",
        "created_at": _utc_now(),
        "delivery_status": delivery_status,
    }
    if provider_message_id:
        entry["provider_message_id"] = provider_message_id
    if error_detail:
        entry["error_detail"] = error_detail
    if attachments:
        entry["attachments"] = attachments
    responses.append(entry)
    _safe_write_json(LEAD_RESPONSES_PATH, {"responses": responses[-1000:], "updated_at": _utc_now()})
    return entry


def _normalize_thread_attachments(raw_value: Any) -> list[dict[str, Any]]:
    values = raw_value if isinstance(raw_value, list) else []
    normalized: list[dict[str, Any]] = []
    for item in values:
        if not isinstance(item, dict):
            continue
        url = str(item.get("url") or item.get("attachment_url") or item.get("file_url") or "").strip()
        attachment_type = str(item.get("type") or item.get("attachment_type") or "file").strip().lower() or "file"
        title = str(item.get("title") or item.get("name") or "").strip()
        if not url and not title:
            continue
        normalized.append(
            {
                "type": attachment_type,
                "url": url or None,
                "title": title or None,
            }
        )
    return normalized


def _is_visible_thread_response(item: dict[str, Any]) -> bool:
    status = str(item.get("delivery_status") or "").strip().lower()
    provider_message_id = str(item.get("provider_message_id") or "").strip()
    attachments = _normalize_thread_attachments(item.get("attachments"))
    if status in {"sent", "delivered"}:
        return True
    if provider_message_id:
        return True
    if attachments and status in {"sent", "delivered"}:
        return True
    return False


def _attachments_from_facebook_message(message: dict[str, Any]) -> list[dict[str, Any]]:
    attachments = message.get("attachments") if isinstance(message.get("attachments"), dict) else {}
    rows = attachments.get("data") if isinstance(attachments, dict) else []
    if not isinstance(rows, list):
        return []
    normalized: list[dict[str, Any]] = []
    for item in rows:
        if not isinstance(item, dict):
            continue
        attachment_type = str(item.get("type") or "").strip().lower() or "file"
        image_data = item.get("image_data") if isinstance(item.get("image_data"), dict) else {}
        video_data = item.get("video_data") if isinstance(item.get("video_data"), dict) else {}
        file_url = str(item.get("file_url") or image_data.get("url") or video_data.get("url") or "").strip()
        title = str(item.get("title") or item.get("name") or attachment_type).strip()
        normalized.append(
            {
                "type": attachment_type,
                "url": file_url or None,
                "title": title or None,
            }
        )
    return normalized


def _facebook_lead_connection_status() -> dict[str, Any]:
    page_tokens = _facebook_token_candidates(
        "FACEBOOK_PAGE_ACCESS_TOKEN",
        "FACEBOOK_MESSENGER_ACCESS_TOKEN",
        "MESSENGER_WEBHOOK_TOKEN",
        "MESSERNGER_WEBHOOK_TOKEN",
    )
    token = page_tokens[0] if page_tokens else ""
    user_tokens = _facebook_token_candidates(
        "FACEBOOK_USER_ACCESS_TOKEN",
        "FACEBOOK_PERSONAL_ACCESS_TOKEN",
        "FACEBOOK_PERSONAL_TOKEN",
        "FACEBOOK_LONG_LIVED_USER_TOKEN",
        "MESSENGER_USER_ACCESS_TOKEN",
    )
    user_token = user_tokens[0] if user_tokens else ""
    page_id = str(os.getenv("FACEBOOK_PAGE_ID", "")).strip()
    app_id = str(os.getenv("FACEBOOK_APP_ID") or os.getenv("FB_APP_ID") or "").strip()
    client_token = str(os.getenv("FACEBOOK_CLIENT_TOKEN") or os.getenv("FB_CLIENT_TOKEN") or "").strip()
    app_secret = str(os.getenv("FACEBOOK_APP_SECRET") or os.getenv("FB_APP_SECRET") or "").strip()
    instagram_token = str(
        os.getenv("INSTAGRAM_ACCESS_TOKEN") or os.getenv("INSTAGRAM_MESSAGING_TOKEN") or ""
    ).strip()
    missing = []
    if not token:
        missing.append("FACEBOOK_PAGE_ACCESS_TOKEN")
    return {
        "configured": bool(token),
        "connected": bool(token),
        "missing": missing,
        "page_id": page_id,
        "page_name": os.getenv("FACEBOOK_PAGE_NAME", "").strip(),
        "page_id_configured": bool(page_id),
        "token_configured": bool(token),
        "user_token_configured": bool(user_token),
        "app_id_configured": bool(app_id),
        "client_token_configured": bool(client_token),
        "app_secret_configured": bool(app_secret),
        "instagram_token_configured": bool(instagram_token),
    }


def _facebook_send_page_message(*, recipient_id: str, message_text: str) -> dict[str, Any]:
    page_tokens = _facebook_token_candidates(
        "FACEBOOK_PAGE_ACCESS_TOKEN",
        "FACEBOOK_MESSENGER_ACCESS_TOKEN",
        "MESSENGER_WEBHOOK_TOKEN",
        "MESSERNGER_WEBHOOK_TOKEN",
    )
    page_token = page_tokens[0] if page_tokens else ""
    page_id = str(os.getenv("FACEBOOK_PAGE_ID", "")).strip()
    clean_recipient = str(recipient_id or "").strip()
    clean_message = str(message_text or "").strip()
    if not page_token:
        return {"ok": False, "mode": "missing_page_token", "message": "FACEBOOK_PAGE_ACCESS_TOKEN is not configured."}
    if not page_id:
        return {"ok": False, "mode": "missing_page_id", "message": "FACEBOOK_PAGE_ID is not configured."}
    if not clean_recipient:
        return {"ok": False, "mode": "missing_recipient", "message": "Lead has no Messenger recipient id."}
    if not clean_message:
        return {"ok": False, "mode": "missing_message", "message": "Response text is empty."}

    endpoint = f"https://graph.facebook.com/v19.0/{page_id}/messages"
    payload = {
        "messaging_type": "RESPONSE",
        "recipient": {"id": clean_recipient},
        "message": {"text": clean_message[:1900]},
    }
    try:
        with httpx.Client(timeout=18.0) as client:
            response = client.post(endpoint, params={"access_token": page_token}, json=payload)
            data = response.json() if response.content else {}
    except Exception as exc:
        return {"ok": False, "mode": "connection_error", "message": str(exc)}

    if response.status_code >= 400:
        error = data.get("error") if isinstance(data, dict) else {}
        message = str(error.get("message") or f"Graph rejected send with status {response.status_code}")
        mode = "graph_rejected"
        if "outside of allowed window" in message.lower():
            mode = "outside_reply_window"
        return {
            "ok": False,
            "mode": mode,
            "http_status": response.status_code,
            "message": message,
            "error": error,
        }
    return {
        "ok": True,
        "mode": "sent",
        "recipient_id": clean_recipient,
        "message_id": str((data or {}).get("message_id") or ""),
        "raw": data,
    }


def _guess_messenger_attachment_type(content_type: str | None, filename: str | None = None) -> str:
    lowered_type = str(content_type or "").strip().lower()
    lowered_name = str(filename or "").strip().lower()
    if lowered_type.startswith("image/") or re.search(r"\.(png|jpe?g|webp|gif|bmp)$", lowered_name):
        return "image"
    if lowered_type.startswith("video/") or re.search(r"\.(mp4|mov|avi|mkv|webm)$", lowered_name):
        return "video"
    if lowered_type.startswith("audio/") or re.search(r"\.(mp3|wav|m4a|aac)$", lowered_name):
        return "audio"
    return "file"


def _facebook_send_page_attachment_url(
    *,
    recipient_id: str,
    attachment_url: str,
    attachment_type: str = "image",
) -> dict[str, Any]:
    page_tokens = _facebook_token_candidates(
        "FACEBOOK_PAGE_ACCESS_TOKEN",
        "FACEBOOK_MESSENGER_ACCESS_TOKEN",
        "MESSENGER_WEBHOOK_TOKEN",
        "MESSERNGER_WEBHOOK_TOKEN",
    )
    page_token = page_tokens[0] if page_tokens else ""
    page_id = str(os.getenv("FACEBOOK_PAGE_ID", "")).strip()
    clean_recipient = str(recipient_id or "").strip()
    clean_url = str(attachment_url or "").strip()
    clean_type = str(attachment_type or "image").strip().lower() or "image"
    if not page_token:
        return {"ok": False, "mode": "missing_page_token", "message": "FACEBOOK_PAGE_ACCESS_TOKEN is not configured."}
    if not page_id:
        return {"ok": False, "mode": "missing_page_id", "message": "FACEBOOK_PAGE_ID is not configured."}
    if not clean_recipient:
        return {"ok": False, "mode": "missing_recipient", "message": "Lead has no Messenger recipient id."}
    if not clean_url:
        return {"ok": False, "mode": "missing_attachment_url", "message": "Attachment URL is empty."}

    endpoint = f"https://graph.facebook.com/v19.0/{page_id}/messages"
    payload = {
        "messaging_type": "RESPONSE",
        "recipient": {"id": clean_recipient},
        "message": {
            "attachment": {
                "type": clean_type,
                "payload": {
                    "url": clean_url,
                    "is_reusable": False,
                },
            }
        },
    }
    try:
        with httpx.Client(timeout=24.0) as client:
            response = client.post(endpoint, params={"access_token": page_token}, json=payload)
            data = response.json() if response.content else {}
    except Exception as exc:
        return {"ok": False, "mode": "connection_error", "message": str(exc)}

    if response.status_code >= 400:
        error = data.get("error") if isinstance(data, dict) else {}
        message = str(error.get("message") or f"Graph rejected attachment send with status {response.status_code}")
        mode = "graph_rejected"
        if "outside of allowed window" in message.lower():
            mode = "outside_reply_window"
        return {
            "ok": False,
            "mode": mode,
            "http_status": response.status_code,
            "message": message,
            "error": error,
        }
    return {
        "ok": True,
        "mode": "sent",
        "recipient_id": clean_recipient,
        "message_id": str((data or {}).get("message_id") or ""),
        "attachment_url": clean_url,
        "attachment_type": clean_type,
        "raw": data,
    }


def _facebook_send_page_attachment_upload(
    *,
    recipient_id: str,
    filename: str,
    content: bytes,
    content_type: str | None,
) -> dict[str, Any]:
    page_tokens = _facebook_token_candidates(
        "FACEBOOK_PAGE_ACCESS_TOKEN",
        "FACEBOOK_MESSENGER_ACCESS_TOKEN",
        "MESSENGER_WEBHOOK_TOKEN",
        "MESSERNGER_WEBHOOK_TOKEN",
    )
    page_token = page_tokens[0] if page_tokens else ""
    page_id = str(os.getenv("FACEBOOK_PAGE_ID", "")).strip()
    clean_recipient = str(recipient_id or "").strip()
    clean_name = Path(str(filename or "attachment")).name
    mime = str(content_type or "application/octet-stream").strip() or "application/octet-stream"
    attachment_type = _guess_messenger_attachment_type(mime, clean_name)
    if not page_token:
        return {"ok": False, "mode": "missing_page_token", "message": "FACEBOOK_PAGE_ACCESS_TOKEN is not configured."}
    if not page_id:
        return {"ok": False, "mode": "missing_page_id", "message": "FACEBOOK_PAGE_ID is not configured."}
    if not clean_recipient:
        return {"ok": False, "mode": "missing_recipient", "message": "Lead has no Messenger recipient id."}
    if not content:
        return {"ok": False, "mode": "missing_attachment_file", "message": "Attachment file is empty."}

    endpoint = f"https://graph.facebook.com/v19.0/{page_id}/messages"
    form_data = {
        "recipient": json.dumps({"id": clean_recipient}),
        "messaging_type": "RESPONSE",
        "message": json.dumps({"attachment": {"type": attachment_type, "payload": {"is_reusable": False}}}),
    }
    files = {"filedata": (clean_name, content, mime)}
    try:
        with httpx.Client(timeout=40.0) as client:
            response = client.post(endpoint, params={"access_token": page_token}, data=form_data, files=files)
            data = response.json() if response.content else {}
    except Exception as exc:
        return {"ok": False, "mode": "connection_error", "message": str(exc)}

    if response.status_code >= 400:
        error = data.get("error") if isinstance(data, dict) else {}
        message = str(error.get("message") or f"Graph rejected attachment upload with status {response.status_code}")
        mode = "graph_rejected"
        if "outside of allowed window" in message.lower():
            mode = "outside_reply_window"
        return {
            "ok": False,
            "mode": mode,
            "http_status": response.status_code,
            "message": message,
            "error": error,
        }
    return {
        "ok": True,
        "mode": "sent",
        "recipient_id": clean_recipient,
        "message_id": str((data or {}).get("message_id") or ""),
        "attachment_type": attachment_type,
        "title": clean_name,
        "raw": data,
    }


def _resolve_facebook_page_id_from_token(token: str) -> tuple[str, dict[str, Any]]:
    token_value = str(token or "").strip()
    if not token_value:
        return "", {}
    preferred_page_name = os.getenv("FACEBOOK_PAGE_NAME", "").strip().lower()
    try:
        with httpx.Client(timeout=10.0) as client:
            response = client.get(
                "https://graph.facebook.com/v19.0/me/accounts",
                params={"fields": "id,name", "access_token": token_value},
            )
        payload = response.json()
        if response.status_code < 400:
            data = payload.get("data") if isinstance(payload, dict) else None
            if isinstance(data, list):
                for entry in data:
                    if not isinstance(entry, dict):
                        continue
                    page_id = str(entry.get("id") or "").strip()
                    page_name = str(entry.get("name") or "").strip()
                    if preferred_page_name and page_name.lower() == preferred_page_name:
                        return page_id, {"source": "graph_accounts_preferred", "name": page_name}
                    if page_id:
                        return page_id, {"source": "graph_accounts_first", "name": page_name}
    except Exception:
        return "", {}
    return "", {}


def _infer_vin_from_lead_text(text: str) -> str:
    vin_match = re.search(r"\b([A-HJ-NPR-Z0-9]{17})\b", str(text or "").upper())
    if vin_match:
        return vin_match.group(1)
    lowered = str(text or "").lower()
    if not lowered:
        return ""
    candidates = _enrich_inventory_items(_load_inventory_candidates())
    for vehicle in candidates[:1500]:
        clean_vin = str(vehicle.get("vin") or "").strip().upper()
        title = str(vehicle.get("title") or "").strip().lower()
        model = str(vehicle.get("model") or "").strip().lower()
        year = str(vehicle.get("year") or "").strip().lower()
        tokens = [token for token in re.findall(r"[a-z0-9]+", title) if len(token) > 2]
        if clean_vin and clean_vin.lower() in lowered:
            return clean_vin
        if title and all(token in lowered for token in tokens[:3]):
            return clean_vin
        if model and year and model in lowered and year in lowered:
            return clean_vin
    return ""


def _conversation_to_lead(
    conversation: dict[str, Any],
    *,
    skip_participant_id: str = "",
    source: str = "facebook_messenger",
) -> dict[str, Any] | None:
    messages_payload = conversation.get("messages") if isinstance(conversation.get("messages"), dict) else {}
    messages = messages_payload.get("data") if isinstance(messages_payload, dict) else []
    latest: dict[str, Any] = {}
    if isinstance(messages, list) and messages:
        message_rows: list[tuple[datetime, dict[str, Any]]] = []
        for msg in messages:
            if not isinstance(msg, dict):
                continue
            created_at = _parse_iso_datetime(msg.get("created_time")) or datetime.min.replace(tzinfo=timezone.utc)
            message_rows.append((created_at, msg))
        if message_rows:
            latest = sorted(message_rows, key=lambda row: row[0])[-1][1]
    message_text = str(latest.get("message") or "").strip()
    if not message_text:
        attachments = latest.get("attachments") if isinstance(latest.get("attachments"), dict) else {}
        attachment_data = attachments.get("data") if isinstance(attachments, dict) else []
        if isinstance(attachment_data, list) and attachment_data:
            first_attachment = attachment_data[0]
            if isinstance(first_attachment, dict):
                message_text = str(
                    first_attachment.get("title")
                    or first_attachment.get("type")
                    or first_attachment.get("subattachments", {}).get("data", [{}])[0].get("media_type")
                    or ""
                ).strip()
    created_at = str(
        latest.get("created_time")
        or conversation.get("updated_time")
        or conversation.get("created_time")
        or _utc_now()
    )
    participants_payload = conversation.get("participants") if isinstance(conversation.get("participants"), dict) else {}
    participants = participants_payload.get("data") if isinstance(participants_payload, dict) else []
    if not isinstance(participants, list) and isinstance(conversation.get("to"), dict):
        participants = conversation.get("to", {}).get("data") if isinstance(conversation.get("to", {}).get("data"), list) else participants
    customer_name = "Facebook Lead"
    profile_id = ""
    skip_id = skip_participant_id.strip()
    if isinstance(participants, list):
        for participant in participants:
            if not isinstance(participant, dict):
                continue
            pid = str(participant.get("id") or "").strip()
            if skip_id and pid == skip_id:
                continue
            customer_name = str(participant.get("name") or customer_name).strip() or customer_name
            profile_id = pid
            break
    if customer_name == "Facebook Lead" and isinstance(latest.get("from"), dict):
        latest_sender = latest.get("from", {})
        customer_name = str(latest_sender.get("name") or customer_name).strip() or customer_name
        if not profile_id:
            profile_id = str(latest_sender.get("id") or profile_id).strip()
    blob = " ".join(
        [
            message_text,
            str(conversation.get("link") or ""),
            json.dumps(latest.get("attachments") or {}, ensure_ascii=False),
            json.dumps(latest.get("shares") or {}, ensure_ascii=False),
        ]
    )
    normalized_source = (
        "facebook_messenger_personal" if source == "facebook_messenger_personal" else source
    )
    if source not in {"facebook_messenger_personal", "facebook_marketplace"}:
        normalized_source = "facebook_marketplace" if "marketplace" in blob.lower() else source
    conversation_id = str(conversation.get("id") or "").strip()
    thread_messages: list[dict[str, Any]] = []
    if isinstance(messages, list):
        sorted_messages = sorted(
            [msg for msg in messages if isinstance(msg, dict)],
            key=lambda row: str(row.get("created_time") or ""),
        )
        for msg in sorted_messages:
            text = str(msg.get("message") or "").strip()
            attachments = _attachments_from_facebook_message(msg)
            if not text and not attachments:
                continue
            if not text:
                text = str((attachments[0] or {}).get("title") or "Attachment").strip()
            sender = msg.get("from") if isinstance(msg.get("from"), dict) else {}
            sender_id = str(sender.get("id") or "").strip()
            sender_name = str(sender.get("name") or customer_name).strip() or customer_name
            direction = "outbound" if skip_id and sender_id == skip_id else "inbound"
            thread_messages.append(
                {
                    "direction": direction,
                    "text": text,
                    "created_at": str(msg.get("created_time") or created_at),
                    "author": sender_name,
                    "delivery_status": "received" if direction == "inbound" else "sent",
                    "attachments": attachments,
                }
            )
    return {
        "id": _lead_id(conversation_id or f"{customer_name}|{created_at}|{message_text}"),
        "customer_name": customer_name,
        "channel": normalized_source,
        "message": message_text or "Messenger conversation synced with no readable latest text.",
        "vehicle_vin": _infer_vin_from_lead_text(blob),
        "source": normalized_source,
        "status": "new",
        "created_at": created_at,
        "last_message_at": created_at,
        "conversation_id": conversation_id,
        "profile_id": profile_id,
        "messages": thread_messages,
    }


def _facebook_graph_get_all_conversations(
    entity_id: str,
    token: str,
    *,
    fields: str,
    page_size: int = 50,
    include_platform: bool = True,
    endpoint_suffix: str = "conversations",
) -> tuple[list[dict[str, Any]], dict[str, Any] | None]:
    conversations: list[dict[str, Any]] = []
    diagnostics: dict[str, Any] = {
        "ok": True,
        "requests": 0,
        "pages": 0,
        "stopped_at": None,
        "next": None,
        "scope": "platform" if include_platform else "personal",
    }
    if not entity_id or not token:
        diagnostics["ok"] = False
        diagnostics["mode"] = "missing_token"
        return conversations, diagnostics

    endpoint = f"https://graph.facebook.com/v19.0/{entity_id}/{endpoint_suffix}"
    params: dict[str, Any] = {
        "fields": fields,
        "limit": max(1, min(100, int(page_size or 50))),
        "access_token": token,
    }
    if include_platform:
        params["platform"] = "messenger"
    next_url: str | None = endpoint

    try:
        with httpx.Client(timeout=18.0) as client:
            visited: set[str] = {endpoint}
            while next_url:
                response = client.get(next_url, params=params if next_url == endpoint else None)
                diagnostics["requests"] += 1
                try:
                    payload = response.json()
                except Exception:
                    payload = {}
                if response.status_code >= 400:
                    diagnostics["stopped_at"] = f"http_{response.status_code}"
                    diagnostics["next"] = next_url
                    diagnostics["http_status"] = response.status_code
                    diagnostics["ok"] = False
                    diagnostics["mode"] = "graph_rejected"
                    graph_error = payload.get("error") if isinstance(payload, dict) else {}
                    if isinstance(graph_error, dict):
                        graph_code = graph_error.get("code")
                        diagnostics["graph_status"] = graph_code if isinstance(graph_code, int) else response.status_code
                        if isinstance(graph_code, int) and graph_code == 190:
                            diagnostics["mode"] = "token_invalid"
                        elif isinstance(graph_code, int) and graph_code == 298:
                            diagnostics["mode"] = "conversation_permission_missing"
                            error_message = str(graph_error.get("message") or "")
                            if "read_mailbox" in error_message.lower():
                                diagnostics["permission_hint"] = "Token needs extended read_mailbox permission."
                                diagnostics["required_permissions"] = ["read_mailbox"]
                            diagnostics["message"] = error_message
                        elif isinstance(graph_code, int) and graph_code == 100:
                            error_message = str(graph_error.get("message") or "")
                            if "inbox" in error_message.lower() and endpoint_suffix == "inbox":
                                diagnostics["mode"] = "conversation_endpoint_missing"
                                diagnostics["permission_hint"] = "Try /conversations endpoint first."
                                diagnostics["message"] = error_message
                            else:
                                diagnostics["mode"] = "conversation_pull_error"
                                diagnostics["message"] = error_message
                    else:
                        diagnostics["graph_status"] = response.status_code
                    diagnostics["graph_response"] = payload
                    diagnostics["error"] = graph_error
                    return conversations, diagnostics

                try:
                    page_items = payload.get("data") if isinstance(payload, dict) else []
                except Exception:
                    page_items = []
                diagnostics["ok"] = True
                diagnostics["mode"] = "conversation_pull_ok"

                if isinstance(page_items, list):
                    diagnostics["pages"] += 1
                    conversations.extend(item for item in page_items if isinstance(item, dict))

                paging = payload.get("paging") if isinstance(payload, dict) else {}
                next_url = paging.get("next") if isinstance(paging, dict) else None
                if not next_url:
                    break
                if next_url in visited:
                    break
                visited.add(next_url)
                params = None

        diagnostics["stopped_at"] = "complete"
        diagnostics["mode"] = diagnostics.get("mode") or "conversation_pull_ok"
        diagnostics.setdefault("graph_status", 200)
        return conversations, diagnostics
    except Exception as exc:  # pragma: no cover - transport failures handled by caller
        diagnostics["stopped_at"] = f"exception:{type(exc).__name__}"
        diagnostics["ok"] = False
        diagnostics["mode"] = "conversation_pull_error"
        diagnostics["error"] = str(exc)
        diagnostics["next"] = next_url
        return conversations, diagnostics


def _carfax_report_payload_richness(result: dict[str, Any]) -> int:
    if not isinstance(result, dict):
        return 0
    if not result.get("ok"):
        return 1
    metrics = [
        "accident_events",
        "service_events",
        "title_brand",
        "accident_damage",
        "owner_count",
        "value_badge",
        "carfax_value",
        "market_position",
        "market_delta",
        "usage",
    ]
    score = 0
    for key in metrics:
        value = result.get(key)
        if isinstance(value, list):
            score += 1 if value else 0
        elif value:
            text = str(value).strip().lower()
            if text and text not in {"none", "not parsed", "n/a", "no service records parsed", "no title brand/issues parsed from report text."}:
                score += 1
    return score


def _needs_carfax_browser_fallback(parsed: dict[str, Any] | None) -> bool:
    if not isinstance(parsed, dict):
        return True
    if parsed.get("source") == "carfax_report_blocked" or parsed.get("blocked"):
        return True
    if not parsed.get("ok"):
        return True
    if _carfax_report_payload_richness(parsed) < 2:
        return True
    return False


def _carfax_extract_vin(value: Any) -> str:
    match = re.search(r"\b([A-HJ-NPR-Z0-9]{17})\b", str(value or "").upper())
    return match.group(1) if match else ""


def _carfax_normalize_vehicle_text(value: Any) -> str:
    text = re.sub(r"[^a-z0-9]+", " ", str(value or "").lower())
    return " ".join(part for part in text.split() if part)


def _carfax_target_vehicle_title(vehicle: dict[str, Any] | None) -> str:
    if not isinstance(vehicle, dict):
        return ""
    explicit_title = str(vehicle.get("title") or "").strip()
    if explicit_title:
        return explicit_title
    return " ".join(
        str(vehicle.get(part) or "").strip()
        for part in ("year", "make", "model", "trim")
        if str(vehicle.get(part) or "").strip()
    ).strip()


def _validate_carfax_report_identity(
    report: dict[str, Any] | None,
    *,
    vehicle: dict[str, Any] | None = None,
    expected_vin: str = "",
) -> dict[str, Any]:
    if not isinstance(report, dict):
        return {"ok": False, "reason": "missing_report"}

    target_vin = str(expected_vin or (vehicle or {}).get("vin") or "").strip().upper()
    target_title = _carfax_target_vehicle_title(vehicle)
    target_year = str((vehicle or {}).get("year") or "").strip()
    target_make = str((vehicle or {}).get("make") or "").strip().lower()
    target_model = str((vehicle or {}).get("model") or "").strip().lower()

    source_text_candidates = [
        report.get("browser_title"),
        report.get("report_vehicle_title"),
        report.get("summary"),
        " ".join(str(item) for item in (report.get("highlights") or [])[:5]),
    ]
    source_text = " ".join(str(item or "") for item in source_text_candidates if item)
    normalized_source = _carfax_normalize_vehicle_text(source_text)
    report_vin = _carfax_extract_vin(report.get("report_vin")) or _carfax_extract_vin(source_text)

    if target_vin and report_vin and report_vin != target_vin:
        return {
            "ok": False,
            "reason": "vin_mismatch",
            "target_vin": target_vin,
            "report_vin": report_vin,
            "target_title": target_title,
            "report_title": str(report.get("browser_title") or report.get("report_vehicle_title") or "").strip(),
        }

    if normalized_source and any(token in normalized_source for token in ("vehicle history report", "carfax report")):
        title_tokens = [
            token
            for token in [target_year, target_make, target_model]
            if token and len(str(token).strip()) >= 2
        ]
        missing_title_tokens = [token for token in title_tokens if str(token).lower() not in normalized_source]
        if len(missing_title_tokens) >= 2:
            return {
                "ok": False,
                "reason": "title_mismatch",
                "target_vin": target_vin,
                "report_vin": report_vin,
                "target_title": target_title,
                "report_title": str(report.get("browser_title") or report.get("report_vehicle_title") or "").strip(),
                "missing_tokens": missing_title_tokens,
            }

    return {"ok": True, "target_vin": target_vin, "report_vin": report_vin, "target_title": target_title}


def _reject_carfax_identity_mismatch(
    report: dict[str, Any] | None,
    *,
    vehicle: dict[str, Any] | None = None,
    expected_vin: str = "",
) -> dict[str, Any]:
    validation = _validate_carfax_report_identity(report, vehicle=vehicle, expected_vin=expected_vin)
    if validation.get("ok"):
        return dict(report or {})

    target_vin = str(validation.get("target_vin") or expected_vin or (vehicle or {}).get("vin") or "").strip().upper()
    target_title = str(validation.get("target_title") or _carfax_target_vehicle_title(vehicle) or target_vin).strip()
    report_title = str(validation.get("report_title") or "").strip()
    report_vin = str(validation.get("report_vin") or "").strip().upper()
    reason = str(validation.get("reason") or "identity_mismatch")
    mismatch_parts = [f"Target VIN {target_vin}"]
    if target_title:
        mismatch_parts.append(target_title)
    if report_vin:
        mismatch_parts.append(f"report VIN {report_vin}")
    if report_title:
        mismatch_parts.append(f"report title {report_title}")
    explanation = "CARFAX report identity mismatch. " + " | ".join(mismatch_parts)

    existing = dict(report or {})
    return {
        **existing,
        "ok": False,
        "blocked": False,
        "source": "carfax_report_identity_mismatch",
        "report_access": explanation,
        "identity_mismatch": True,
        "identity_validation": validation,
        "highlights": [explanation],
    }


def _carfax_report_matches_vehicle(
    report: dict[str, Any] | None,
    *,
    vehicle: dict[str, Any] | None = None,
    expected_vin: str = "",
) -> bool:
    return bool(_validate_carfax_report_identity(report, vehicle=vehicle, expected_vin=expected_vin).get("ok"))


def _strip_carfax_structured_report_fields(facts: dict[str, Any] | None) -> dict[str, Any]:
    cleaned = dict(facts or {})
    for key in (
        "ok",
        "highlights",
        "summary",
        "source_url",
        "carfax_value",
        "market_position",
        "market_delta",
        "accident_damage",
        "accident_events",
        "accident_counts",
        "title_brand",
        "service_history",
        "service_records_count",
        "last_service_date",
        "service_events",
        "usage",
        "report_vin",
        "report_vehicle_title",
        "browser_title",
        "browser_text_chars",
        "updated_at",
        "http_status",
        "extract_meta",
    ):
        cleaned.pop(key, None)
    return cleaned


def _sync_facebook_leads(source: str = "page") -> dict[str, Any]:
    mode = (source or "page").strip().lower()
    if mode not in {"page", "personal", "all"}:
        mode = "page"
    include_page = mode in {"page", "all"}
    include_personal = mode in {"personal", "all"}

    connection = _facebook_lead_connection_status()
    page_id = str(os.getenv("FACEBOOK_PAGE_ID", "")).strip()
    page_tokens = _facebook_named_token_candidates(
        "FACEBOOK_PAGE_ACCESS_TOKEN",
        "FACEBOOK_MESSENGER_ACCESS_TOKEN",
        "MESSENGER_WEBHOOK_TOKEN",
        "MESSERNGER_WEBHOOK_TOKEN",
    )
    if not page_tokens:
        raw_token = str(os.getenv("FACEBOOK_PAGE_ACCESS_TOKEN", "")).strip()
        if raw_token:
            page_tokens = [("FACEBOOK_PAGE_ACCESS_TOKEN", _facebook_token(raw_token))]
    user_tokens = _facebook_token_candidates(
        "FACEBOOK_USER_ACCESS_TOKEN",
        "FACEBOOK_PERSONAL_ACCESS_TOKEN",
        "FACEBOOK_PERSONAL_TOKEN",
        "FACEBOOK_LONG_LIVED_USER_TOKEN",
        "MESSENGER_USER_ACCESS_TOKEN",
    )
    user_token = user_tokens[0] if user_tokens else ""

    if not include_page and not include_personal:
        return {
            "ok": False,
            "mode": "not_connected",
            "connection": connection,
            "imported": 0,
            "guidance": ["Set 'source' to page, personal, or all."],
        }

    def _fetch_entity_conversations(
        entity_id: str,
        fetch_token: str,
        scope: str,
    ) -> tuple[list[dict[str, Any]], dict[str, Any], dict[str, Any], int, int]:
        if not entity_id or not fetch_token:
            return [], {"mode": "missing_token", "ok": False, "scope": scope}, {"ok": False, "mode": "missing_token"}, 0, 0

        conversation_attempts: list[dict[str, Any]] = []
        conversation_endpoints = ["conversations"] if scope == "page" else ["conversations", "inbox"]
        fields_variants = [
            "id,updated_time,created_time,link,participants,messages.limit(1).order(reverse_chronological){message,created_time,from,to,attachments,shares}",
            "id,updated_time,created_time,link,participants,messages.limit(20){message,created_time,from,to,attachments,shares}",
        ]

        conversations: list[dict[str, Any]] = []
        seen_conversation_ids: set[str] = set()
        total_pages = 0
        total_requests = 0
        best_diagnostics: dict[str, Any] = {"ok": False, "scope": scope, "mode": "conversation_pull_error"}

        for endpoint_suffix in conversation_endpoints:
            endpoint_found = False
            for fields in fields_variants:
                endpoint_convos, endpoint_diag = _facebook_graph_get_all_conversations(
                    entity_id,
                    fetch_token,
                    fields=fields,
                    page_size=50,
                    include_platform=(scope == "page"),
                    endpoint_suffix=endpoint_suffix,
                )
                if not isinstance(endpoint_diag, dict):
                    endpoint_diag = {"ok": False, "scope": scope, "mode": "conversation_pull_error"}
                total_pages += int(endpoint_diag.get("pages", 0))
                total_requests += int(endpoint_diag.get("requests", 0))
                conversation_attempts.append(
                    {
                        "endpoint": endpoint_suffix,
                        "fields": fields,
                        "ok": bool(endpoint_diag.get("ok", False)),
                        "mode": endpoint_diag.get("mode"),
                        "pages": int(endpoint_diag.get("pages", 0)),
                        "requests": int(endpoint_diag.get("requests", 0)),
                        "graph_status": endpoint_diag.get("graph_status"),
                    }
                )

                if endpoint_diag.get("ok", True):
                    endpoint_found = True
                    if isinstance(endpoint_convos, list):
                        for convo in endpoint_convos:
                            if not isinstance(convo, dict):
                                continue
                            conv_id = str(convo.get("id") or "").strip()
                            if conv_id:
                                if conv_id in seen_conversation_ids:
                                    continue
                                seen_conversation_ids.add(conv_id)
                            conversations.append(convo)
                    if endpoint_suffix == conversation_endpoints[-1] or scope == "page":
                        best_diagnostics = dict(endpoint_diag)
                        best_diagnostics["endpoint_suffix"] = endpoint_suffix
                        best_diagnostics["fields_variant_used"] = fields
                        break
                    if conversations:
                        continue
                    best_diagnostics = dict(endpoint_diag)
                    best_diagnostics["endpoint_suffix"] = endpoint_suffix
                    best_diagnostics["fields_variant_used"] = fields
                    continue

                if endpoint_diag.get("graph_status") == 100:
                    # Not query supported for this field shape/entity; still try another variant.
                    continue

                # Keep the last meaningful error so we can diagnose why each token failed.
                best_diagnostics = {
                    "ok": False,
                    "scope": scope,
                    "mode": endpoint_diag.get("mode"),
                    "endpoint_suffix": endpoint_suffix,
                    "fields_variant_used": fields,
                    "graph_status": endpoint_diag.get("graph_status"),
                    "graph_response": endpoint_diag.get("graph_response"),
                    "error": endpoint_diag.get("error"),
                    "requests": total_requests,
                    "pages": total_pages,
                }

            if endpoint_found:
                if scope == "page":
                    # Page endpoints often mirror each other; avoid expensive extras if we already have results.
                    if conversations:
                        break
                else:
                    # Personal token can benefit from both endpoint shapes; keep scanning.
                    continue

        if not conversations:
            best_diagnostics.setdefault("mode", "empty_result")
            best_diagnostics.setdefault("graph_status", 200)

        return conversations, best_diagnostics, {
            "endpoint_attempts": conversation_attempts,
        }, total_pages, total_requests

    def _add_source_metadata(
        conversations: list[dict[str, Any]],
        source_name: str,
        participant_id: str,
    ) -> tuple[list[dict[str, Any]], int, int]:
        if not conversations:
            return [], 0, 0
        prepared: list[dict[str, Any]] = []
        for conversation in conversations:
            if not isinstance(conversation, dict):
                continue
            conversation["_facebook_source_override"] = source_name
            conversation["_facebook_participant_skip_id"] = participant_id
            prepared.append(conversation)
        return prepared, len(prepared), 0

    if include_page:
        if not connection.get("connected") and not include_personal:
            return {
                "ok": False,
                "mode": "not_connected",
                "connection": connection,
                "imported": 0,
                "guidance": [
                    "Set FACEBOOK_PAGE_ID and a valid page token before Messenger page sync.",
                    "Until connected, use Manual Lead to keep Messenger conversations visible in Xconsole.",
                ],
            }

    if include_page and not page_id and page_tokens:
        for _, candidate_token in page_tokens:
            candidate_page_id, resolved = _resolve_facebook_page_id_from_token(candidate_token)
            if candidate_page_id:
                page_id = candidate_page_id
                connection["page_id"] = candidate_page_id
                connection["page_id_resolved"] = True
                connection["resolved_page"] = resolved
                break

    all_conversations: list[dict[str, Any]] = []
    pagination_pages_total = 0
    pagination_requests_total = 0
    diagnostics: list[dict[str, Any]] = []
    personal_warnings: list[str] = []
    token_diagnostics: list[dict[str, Any]] = []
    used_personal_tokens = user_tokens if include_personal else []

    if include_page:
        page_identity_checked = False
        for slot, (env_name, candidate_token) in enumerate(page_tokens, start=1):
            if not candidate_token:
                continue

            resolved_entity_id = page_id
            identity = _facebook_graph_diagnose_token(
                candidate_token,
                scope="page",
                entity_id=resolved_entity_id or "me",
                fields="id,name",
            )
            if not resolved_entity_id and identity.get("ok") and isinstance(identity.get("payload"), dict):
                # me/accounts payload is authoritative only when called with a page token; identity with this call may be a profile.
                page_id_candidate = str(identity.get("payload", {}).get("id") or "")
                if page_id_candidate:
                    resolved_entity_id = page_id_candidate
            page_identity_checked = True
            token_record: dict[str, Any] = {
                "slot": slot,
                "scope": "page",
                "env_name": env_name,
                "token_preview": _token_preview(candidate_token),
                "identity": identity,
                "entity_id": resolved_entity_id,
            }

            if not identity.get("ok"):
                token_record["conversations"] = {"ok": False, "count": 0, "mode": identity.get("mode")}
                token_diagnostics.append(token_record)
                continue
            if not resolved_entity_id and not page_id:
                token_record["conversations"] = {"ok": False, "count": 0, "mode": "missing_page_identity"}
                token_diagnostics.append(token_record)
                continue

            convs, conv_diag, conv_meta, pages_count, request_count = _fetch_entity_conversations(
                resolved_entity_id,
                candidate_token,
                "page",
            )
            token_record["conversations"] = {
                "ok": bool(conv_diag.get("ok", False)),
                "count": len(convs or []),
                "mode": conv_diag.get("mode"),
                "pages": conv_diag.get("pages"),
                "requests": conv_diag.get("requests"),
                "endpoint_suffix": conv_diag.get("endpoint_suffix"),
                "fields_variant_used": conv_diag.get("fields_variant_used"),
            }
            token_record["attempts"] = conv_meta.get("endpoint_attempts")
            token_diagnostics.append(token_record)
            diagnostics.append(conv_diag)
            if isinstance(pages_count, int):
                pagination_pages_total += pages_count
            if isinstance(request_count, int):
                pagination_requests_total += request_count

            if convs:
                prepared, _, _ = _add_source_metadata(convs, "facebook_messenger", resolved_entity_id or "")
                all_conversations.extend(prepared)
                # Keep collecting in case additional page tokens expose older/other inboxes.
                continue

        if not page_identity_checked and not include_personal:
            return {
                "ok": False,
                "mode": "not_connected",
                "connection": connection,
                "imported": 0,
                "guidance": [
                    "No valid page token identity found. Add FACEBOOK_PAGE_ID and set a valid PAGE access token.",
                    "Until connected, use Manual Lead to keep Messenger conversations visible in Xconsole.",
                ],
                "token_diagnostics": token_diagnostics,
            }
        if not all_conversations and not include_personal and page_tokens:
            personal_warnings.append("Page conversations did not return any conversations with the active token set.")

    if include_personal:
        personal_identity_failures = []
        for slot, personal_token in enumerate(used_personal_tokens, start=1):
            if not personal_token:
                continue
            identity = _facebook_graph_diagnose_token(
                personal_token,
                scope="personal",
                entity_id="me",
                fields="id,name",
            )
            if not identity.get("ok"):
                personal_identity_failures.append({"slot": slot, "identity": identity})
                if identity.get("mode") == "token_invalid":
                    personal_warnings.append(
                        "Configured personal token was invalid. Trying alternate personal token aliases."
                    )
                else:
                    personal_warnings.append(
                        f"Personal token {slot} not usable yet: {identity.get('mode', 'token_issue')}."
                    )
                continue

            token_convs, candidate_diag, conv_meta, token_pages_count, token_request_count = _fetch_entity_conversations(
                "me",
                personal_token,
                "personal",
            )
            diagnostics.append(candidate_diag)
            personal_meta = {
                "scope": "personal",
                "slot": slot,
                "identity": identity,
                "conversations": {
                    "ok": bool(candidate_diag.get("ok", False)),
                    "count": len(token_convs or []),
                    "mode": candidate_diag.get("mode"),
                    "pages": candidate_diag.get("pages"),
                    "requests": candidate_diag.get("requests"),
                    "endpoint_suffix": candidate_diag.get("endpoint_suffix"),
                    "fields_variant_used": candidate_diag.get("fields_variant_used"),
                },
                "attempts": conv_meta.get("endpoint_attempts"),
            }
            token_diagnostics.append(personal_meta)
            if isinstance(token_pages_count, int):
                pagination_pages_total += token_pages_count
            if isinstance(token_request_count, int):
                pagination_requests_total += token_request_count

            if token_convs:
                personal_id = str(identity.get("payload", {}).get("id", "")) if isinstance(identity.get("payload"), dict) else ""
                prepared, _, _ = _add_source_metadata(token_convs, "facebook_messenger_personal", personal_id)
                all_conversations.extend(prepared)

            if not candidate_diag.get("ok", True) and include_personal:
                if candidate_diag.get("mode") in {"conversation_permission_missing", "conversation_pull_error", "graph_rejected", "conversation_endpoint_missing"}:
                    permission_hint = candidate_diag.get("permission_hint") or candidate_diag.get("message") or ""
                    if candidate_diag.get("mode") == "conversation_permission_missing" and not permission_hint:
                        permission_hint = "Token is not authorized for mailbox access."
                    personal_warnings.append(
                        f"Personal token slot {slot}: {candidate_diag.get('mode')} ({candidate_diag.get('graph_status')}){f' - {permission_hint}' if permission_hint else ''}."
                    )
        if include_personal and not all_conversations and not used_personal_tokens:
            personal_warnings.append("No personal token supplied; only page conversations were synced if configured.")

        if include_personal and all(c.get("mode") in {"token_invalid", "missing_token"} for c in personal_identity_failures):
            if not include_page:
                return {
                    "ok": False,
                    "mode": "not_connected",
                    "connection": connection,
                    "imported": 0,
                    "guidance": [
                        "Set FACEBOOK_USER_ACCESS_TOKEN (or FACEBOOK_PERSONAL_ACCESS_TOKEN) for personal Messenger chats.",
                        "Until connected, use Manual Lead to keep Messenger conversations visible in Xconsole.",
                    ],
                    "token_diagnostics": token_diagnostics,
                }

    # Normalize duplicates by conversation hash/id.
    unique_conversations: list[dict[str, Any]] = []
    seen_lead_ids: set[str] = set()
    for conversation in all_conversations:
        if not isinstance(conversation, dict):
            continue
        if "id" in conversation and conversation["id"]:
            key = str(conversation["id"]).strip()
        else:
            message_text = str(conversation.get("messages", {}).get("data", [{}])[0].get("message", "")).strip()
            key = _lead_id(f"{message_text}|{conversation.get('updated_time')}|{conversation.get('link')}")
        if key in seen_lead_ids:
            continue
        seen_lead_ids.add(key)
        unique_conversations.append(conversation)

    existing = _load_leads()
    by_id = {str(item.get("id")): item for item in existing if isinstance(item, dict)}
    imported = 0
    updated = 0
    if not isinstance(unique_conversations, list):
        unique_conversations = []

    for conversation in unique_conversations:
        if not isinstance(conversation, dict):
            continue
        lead = _conversation_to_lead(
            conversation,
            skip_participant_id=str(conversation.get("_facebook_participant_skip_id") or ""),
            source=str(conversation.get("_facebook_source_override") or "facebook_messenger"),
        )
        if not lead:
            continue
        if lead["id"] in by_id:
            by_id[lead["id"]].update({key: value for key, value in lead.items() if value not in (None, "")})
            updated += 1
        else:
            existing.append(lead)
            by_id[lead["id"]] = lead
            imported += 1

    _save_leads(list(by_id.values()))

    return {
        "ok": True,
        "mode": "conversations_synced",
        "connection": connection,
        "imported": imported,
        "updated": updated,
        "warnings": personal_warnings,
        "pagination": {
            "pages": pagination_pages_total,
            "requests": pagination_requests_total,
            "sources": [item.get("scope") for item in diagnostics if isinstance(item, dict) and item.get("scope")],
            "diagnostics": diagnostics,
            "token_diagnostics": token_diagnostics,
        },
        "guidance": [
            f"Facebook Messenger sync complete: {imported} new, {updated} updated across {pagination_pages_total} page(s), {pagination_requests_total} request(s).",
            *personal_warnings,
        ],
    }

def _sync_facebook_leads_if_stale(
    *, min_seconds: int = 45, source: str | None = None, force: bool = False
) -> dict[str, Any] | None:
    connection = _facebook_lead_connection_status()
    requested_source = (source or os.getenv("FACEBOOK_LEAD_SYNC_SOURCE") or "all").strip().lower()
    if not requested_source:
        requested_source = "all"
    requested_mode = requested_source if requested_source in {"page", "personal", "all"} else "page"

    if requested_mode == "page":
        if not connection.get("connected"):
            return None
    elif requested_mode == "personal":
        if not bool(connection.get("user_token_configured")):
            return None
    else:
        if not connection.get("connected") and not connection.get("user_token_configured"):
            return None

    state = _safe_read_json(LEAD_SYNC_STATE_PATH, {})
    last_sync = _parse_iso_datetime(state.get("last_sync_at") if isinstance(state, dict) else None)
    now = datetime.now(timezone.utc)
    if last_sync and last_sync.tzinfo is None:
        last_sync = last_sync.replace(tzinfo=timezone.utc)
    if not force and last_sync and (now - last_sync).total_seconds() < min_seconds:
        return {
            "ok": True,
            "mode": "lead_sync_throttled",
            "last_sync_at": last_sync.isoformat(),
        }
    result = _sync_facebook_leads(source=requested_source)
    _safe_write_json(
        LEAD_SYNC_STATE_PATH,
        {
            "last_sync_at": now.isoformat(),
            "mode": result.get("mode"),
            "imported": result.get("imported"),
            "updated": result.get("updated"),
        },
    )
    return result


def _diagnose_facebook_lead_sync_candidates(*, requested_source: str = "all") -> dict[str, Any]:
    include_page = requested_source in {"page", "all"}
    include_personal = requested_source in {"personal", "all"}

    page_token_candidates = _facebook_named_token_candidates(
        "FACEBOOK_PAGE_ACCESS_TOKEN",
        "FACEBOOK_MESSENGER_ACCESS_TOKEN",
        "MESSENGER_WEBHOOK_TOKEN",
        "MESSERNGER_WEBHOOK_TOKEN",
    )
    personal_token_candidates = _facebook_named_token_candidates(
        "FACEBOOK_USER_ACCESS_TOKEN",
        "FACEBOOK_PERSONAL_ACCESS_TOKEN",
        "FACEBOOK_PERSONAL_TOKEN",
        "FACEBOOK_LONG_LIVED_USER_TOKEN",
        "MESSENGER_USER_ACCESS_TOKEN",
    )
    page_id = str(os.getenv("FACEBOOK_PAGE_ID", "")).strip()

    def _personalized_conversation_check(entity_id: str, token: str, *, scope: str) -> dict[str, Any]:
        endpoint_candidates = ["conversations"]
        if scope == "personal":
            endpoint_candidates.append("inbox")
        for endpoint in endpoint_candidates:
            convs, conv_diag = _facebook_graph_get_all_conversations(
                entity_id,
                token,
                fields="id,updated_time,created_time,link,participants,messages.limit(1).order(reverse_chronological){message,created_time,from,to,attachments,shares}",
                page_size=10,
                include_platform=(scope == "page"),
                endpoint_suffix=endpoint,
            )
            if conv_diag.get("ok"):
                return {
                    "endpoint": endpoint,
                    "ok": True,
                    "count": len(convs or []),
                    "mode": conv_diag.get("mode"),
                    "pages": conv_diag.get("pages"),
                    "requests": conv_diag.get("requests"),
                    "stopped_at": conv_diag.get("stopped_at"),
                    "graph_response": conv_diag.get("graph_response"),
                    "permission_hint": conv_diag.get("permission_hint"),
                    "message": conv_diag.get("message"),
                }
            if conv_diag.get("graph_status") in {298, 100, 10, 190}:
                return {
                    "endpoint": endpoint,
                    "ok": False,
                    "count": 0,
                    "mode": conv_diag.get("mode"),
                    "graph_status": conv_diag.get("graph_status"),
                    "error": conv_diag.get("graph_response"),
                    "permission_hint": conv_diag.get("permission_hint"),
                    "message": conv_diag.get("message"),
                    "pages": conv_diag.get("pages"),
                    "requests": conv_diag.get("requests"),
                }
        return {
            "endpoint": endpoint_candidates[-1],
            "ok": False,
            "count": 0,
            "mode": "conversation_pull_error",
        }

    page_results: list[dict[str, Any]] = []
    personal_results: list[dict[str, Any]] = []

    if include_page:
        for slot, (env_name, token) in enumerate(page_token_candidates, start=1):
            identity = _facebook_graph_diagnose_token(token, scope="page", entity_id=page_id or "me")
            row: dict[str, Any] = {
                "slot": slot,
                "scope": "page",
                "env_name": env_name,
                "token_preview": _token_preview(token),
                "identity": identity,
            }
            if identity.get("ok") and page_id:
                row["conversations"] = _personalized_conversation_check(page_id, token, scope="page")
            else:
                row["conversations"] = {
                    "ok": False,
                    "count": 0,
                    "mode": identity.get("mode"),
                    "graph_status": identity.get("graph_status"),
                }
            page_results.append(row)

    if include_personal:
        for slot, (env_name, token) in enumerate(personal_token_candidates, start=1):
            identity = _facebook_graph_diagnose_token(token, scope="personal", entity_id="me")
            row = {
                "slot": slot,
                "scope": "personal",
                "env_name": env_name,
                "token_preview": _token_preview(token),
                "identity": identity,
            }
            if identity.get("ok"):
                row["conversations"] = _personalized_conversation_check("me", token, scope="personal")
            else:
                row["conversations"] = {
                    "ok": False,
                    "count": 0,
                    "mode": identity.get("mode"),
                    "graph_status": identity.get("graph_status"),
                }
            personal_results.append(row)

    diagnostics: dict[str, Any] = {
        "page": {
            "configured_page_id": page_id,
            "token_slots": page_results,
            "ok_count": len([row for row in page_results if row.get("identity", {}).get("ok")]),
        },
        "personal": {
            "token_slots": personal_results,
            "ok_count": len([row for row in personal_results if row.get("identity", {}).get("ok")]),
        },
        "summary": {
            "requested_source": requested_source,
            "page_tokens_checked": len(page_results),
            "personal_tokens_checked": len(personal_results),
            "page_conversations_ok": len([row for row in page_results if row.get("conversations", {}).get("ok")]),
            "personal_conversations_ok": len([row for row in personal_results if row.get("conversations", {}).get("ok")]),
        },
    }

    if not include_page:
        diagnostics["page"].pop("configured_page_id", None)
    return diagnostics


def _load_offerup_status() -> dict[str, Any]:
    payload = _safe_read_json(OFFERUP_STATUS_PATH, {"posts": {}})
    posts = payload.get("posts", {}) if isinstance(payload, dict) else {}
    if not isinstance(posts, dict):
        posts = {}
    return {
        "ok": True,
        "ready_for_live": False,
        "mode": "draft",
        "reason": "OfferUp has no connected live API/session in this build; Xconsole creates one-click drafts.",
        "posts": posts,
        "drafts_count": len(posts),
        "updated_at": payload.get("updated_at") if isinstance(payload, dict) else None,
    }


def _save_offerup_post(vin: str, post: dict[str, Any]) -> None:
    status = _load_offerup_status()
    posts = status.get("posts", {})
    if not isinstance(posts, dict):
        posts = {}
    posts[vin] = post
    _safe_write_json(OFFERUP_STATUS_PATH, {"posts": posts, "updated_at": _utc_now()})


def _offerup_post_from_inventory(request: OfferUpPostRequest) -> dict[str, Any]:
    clean_vin = str(request.vin or "").strip().upper()
    vehicle = _find_vehicle_by_vin(clean_vin)
    if not vehicle:
        raise HTTPException(status_code=404, detail={"message": f"Vehicle not found for VIN {clean_vin}"})
    caption = _build_caption_from_vehicle(vehicle, caption_override=request.caption_override)
    OFFERUP_POSTS_DIR.mkdir(parents=True, exist_ok=True)
    target = OFFERUP_POSTS_DIR / f"{clean_vin}_{int(time.time())}.txt"
    target.write_text(caption, encoding="utf-8")
    post = {
        "vin": clean_vin,
        "mode": request.mode,
        "status": "drafted" if request.mode == "draft" else "live_not_connected",
        "draft_file": _display_path(target),
        "caption": caption,
        "created_at": _utc_now(),
    }
    _save_offerup_post(clean_vin, post)
    _append_audit_event("offerup_post", {"vin": clean_vin, "mode": request.mode, "status": post["status"]})
    return {
        "ok": request.mode == "draft",
        "vin": clean_vin,
        "post": post,
        "live_success": False,
        "live_detail": "OfferUp live posting is not connected. Draft was created for operator review.",
        "status": _load_offerup_status(),
    }


def _vin_decode_cache_path(vin: str) -> Path:
    clean_vin = re.sub(r"[^A-Z0-9]", "", str(vin or "").upper()) or "UNKNOWN"
    VIN_DECODE_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    return VIN_DECODE_CACHE_DIR / f"{clean_vin}.json"


def _decode_vin_values(vin: str) -> dict[str, Any]:
    clean_vin = re.sub(r"[^A-Z0-9]", "", str(vin or "").upper())
    if len(clean_vin) < 11:
        raise HTTPException(status_code=400, detail={"message": "VIN must be at least 11 characters"})
    cache_path = _vin_decode_cache_path(clean_vin)
    cached = _safe_read_json(cache_path, None)
    if isinstance(cached, dict) and cached.get("vin") == clean_vin:
        return cached

    decoded: dict[str, Any] = {
        "vin": clean_vin,
        "ok": False,
        "source": "fallback",
        "decoded_at": _utc_now(),
        "fields": {},
        "raw": None,
    }
    try:
        with httpx.Client(timeout=10.0, follow_redirects=True) as client:
            response = client.get(
                f"https://vpic.nhtsa.dot.gov/api/vehicles/DecodeVinValuesExtended/{clean_vin}",
                params={"format": "json"},
            )
        payload = response.json()
        result = (payload.get("Results") or [{}])[0] if isinstance(payload, dict) else {}
        if isinstance(result, dict):
            fields = {
                "year": result.get("ModelYear"),
                "make": result.get("Make"),
                "model": result.get("Model"),
                "trim": result.get("Trim") or result.get("Series"),
                "body_class": result.get("BodyClass"),
                "vehicle_type": result.get("VehicleType"),
                "drive_type": result.get("DriveType"),
                "engine": " ".join(
                    str(part)
                    for part in [
                        result.get("EngineCylinders") and f"{result.get('EngineCylinders')} cyl",
                        result.get("DisplacementL") and f"{result.get('DisplacementL')}L",
                        result.get("FuelTypePrimary"),
                    ]
                    if part
                ),
                "manufacturer": result.get("Manufacturer"),
                "plant_country": result.get("PlantCountry"),
                "gvwr": result.get("GVWR"),
            }
            decoded.update({"ok": True, "source": "nhtsa_vpic", "fields": fields, "raw": result})
    except Exception as exc:
        decoded["error"] = str(exc)

    if not decoded.get("ok"):
        vehicle = _find_vehicle_by_vin(clean_vin)
        title = str(vehicle.get("title") or "") if vehicle else ""
        parts = title.split()
        decoded["fields"] = {
            "year": next((part for part in parts if re.fullmatch(r"20[0-9]{2}|19[0-9]{2}", part)), None),
            "make": parts[1] if len(parts) > 1 else None,
            "model": parts[2] if len(parts) > 2 else None,
            "trim": " ".join(parts[3:]) if len(parts) > 3 else None,
        }
    _safe_write_json(cache_path, decoded)
    return decoded


def _carfax_summary_for_vin(vin: str) -> dict[str, Any] | None:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        return None
    candidates = [CARFAX_SUMMARY_DIR / f"{clean_vin}.json", CARFAX_SUMMARY_DIR / f"{clean_vin.lower()}.json"]
    for path in candidates:
        payload = _safe_read_json(path, None)
        if isinstance(payload, dict):
            normalized = dict(payload)
            facts = normalized.get("facts")
            if isinstance(facts, dict):
                facts = _normalize_carfax_parsed_payload(facts)
                normalized["facts"] = facts
                normalized["summary"] = str(facts.get("summary") or normalized.get("summary") or "").strip()
                normalized["highlights"] = list(facts.get("highlights") or normalized.get("highlights") or [])
            if normalized != payload:
                _safe_write_json(path, normalized)
            return normalized
    return None


def _load_jd_power_valuations() -> dict[str, dict[str, Any]]:
    payload = _safe_read_json(JD_POWER_VALUATIONS_PATH, {"items": []})
    items = payload.get("items", []) if isinstance(payload, dict) else payload
    if not isinstance(items, list):
        return {}
    output: dict[str, dict[str, Any]] = {}
    for item in items:
        if not isinstance(item, dict):
            continue
        vin = str(item.get("vin") or "").strip().upper()
        if not vin:
            continue
        output[vin] = item
    return output


def _jd_power_valuation_for_vin(vin: str | None) -> dict[str, Any] | None:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        return None
    return _load_jd_power_valuations().get(clean_vin)


def _bank_sale_price_from_inventory_price(inventory_price: float | int | str | None) -> float | None:
    price = _to_float(inventory_price)
    if price is None:
        return None
    return float(price + DEFAULT_BANK_FEES)


def _jd_power_ltv_from_pricing(*, inventory_price: float | int | str | None, jd_trade_value: float | int | str | None) -> dict[str, float | None]:
    bank_sale_price = _bank_sale_price_from_inventory_price(inventory_price)
    jd_trade = _to_float(jd_trade_value)
    if bank_sale_price is None:
        return {"bank_sale_price": None, "taxes": None, "ltv_basis": None, "ltv": None}
    taxes = round(bank_sale_price * DEFAULT_BANK_TAX_RATE, 2)
    ltv_basis = round(bank_sale_price + taxes, 2)
    ltv = round((ltv_basis / jd_trade) * 100.0, 2) if jd_trade and jd_trade > 0 else None
    return {
        "bank_sale_price": round(bank_sale_price, 2),
        "taxes": taxes,
        "ltv_basis": ltv_basis,
        "ltv": ltv,
    }


VIN_PATTERN = re.compile(r"\b[A-HJ-NPR-Z0-9]{17}\b", re.IGNORECASE)


def _normalize_table_cell(value: Any) -> str:
    if value is None:
        return ""
    return re.sub(r"\s+", " ", str(value).replace("\n", " ").strip())


def _normalize_header(value: Any) -> str:
    return re.sub(r"[^a-z0-9]+", " ", _normalize_table_cell(value).lower()).strip()


def _rows_from_upload_table(raw: bytes, filename: str, content_type: str | None = None) -> list[list[Any]]:
    name = str(filename or "").lower()
    def decode_text() -> str:
        for encoding in ("utf-8-sig", "utf-8", "latin-1"):
            try:
                return raw.decode(encoding, errors="ignore")
            except Exception:
                continue
        return raw.decode("utf-8", errors="ignore")

    if name.endswith(".xls"):
        try:
            import xlrd  # type: ignore
        except Exception as exc:
            raise HTTPException(status_code=500, detail={"message": "xlrd is required to import JD Power .xls files", "error": str(exc)})
        workbook = xlrd.open_workbook(file_contents=raw)
        sheet = workbook.sheet_by_index(0)
        return [[sheet.cell_value(r, c) for c in range(sheet.ncols)] for r in range(sheet.nrows)]

    if name.endswith(".xlsx") or name.endswith(".xlsm"):
        try:
            import openpyxl
        except Exception as exc:
            raise HTTPException(status_code=500, detail={"message": "openpyxl is required to import JD Power .xlsx files", "error": str(exc)})

        workbook = openpyxl.load_workbook(BytesIO(raw), read_only=True, data_only=True)
        sheet = workbook[workbook.sheetnames[0]]
        return [list(row) for row in sheet.iter_rows(values_only=True)]

    lowered_type = str(content_type or "").lower()
    if name.endswith(".csv") or "csv" in lowered_type:
        text = decode_text()
        return [row for row in csv.reader(StringIO(text)) if any(_normalize_table_cell(cell) for cell in row)]

    if name.endswith(".tsv") or "tab-separated" in lowered_type:
        text = decode_text()
        return [row for row in csv.reader(StringIO(text), delimiter="\t") if any(_normalize_table_cell(cell) for cell in row)]

    text = _extract_text_from_upload(raw, filename, content_type)
    rows: list[list[Any]] = []
    for line in text.splitlines():
        clean = line.strip()
        if not clean:
            continue
        if "\t" in clean:
            row = [cell.strip() for cell in clean.split("\t")]
        elif "," in clean and clean.count(",") >= 2:
            row = next(csv.reader([clean]))
        else:
            row = [cell.strip() for cell in re.split(r"\s{2,}", clean) if cell.strip()]
        rows.append(row)
    return rows


def _header_col_score(header: str, kind: str) -> int:
    compact = header.replace(" ", "")
    if kind == "vin":
        return 100 if header in {"vin", "vehicle identification number"} or "vin" in header.split() else (80 if "vehicle identification" in header else 0)
    if kind == "vehicle":
        return 90 if any(token in header for token in ("vehicle", "description", "year make model")) else 0
    if kind == "stock":
        return 90 if "stock" in header else 0
    if kind == "price":
        if header == "price":
            return 100
        if header == "sale price":
            return 65
        if "price" in header and not any(token in header for token in ("trade", "book", "jd", "j d", "nada", "value")):
            return 80
        return 20 if header in {"sale", "retail", "internet"} else 0
    if kind == "jd_trade":
        score = 0
        if any(token in header for token in ("jd power", "j d power", "jdpower", "nada", "book", "black book")) or any(token in compact for token in ("jdpower", "jdptrade")):
            score += 60
        if any(token in header for token in ("trade in", "trade", "wholesale", "loan", "clean", "value", "advance")):
            score += 45
        if "price" in header and score < 80:
            score -= 30
        return max(0, score)
    if kind == "class":
        return 80 if header == "class" or "class" in header else 0
    if kind == "new_used":
        return 80 if "new used" in header or ("new" in header and "used" in header) else 0
    return 0


def _best_header_col(headers: list[str], kind: str, minimum: int = 50) -> int | None:
    scored = [(_header_col_score(header, kind), index) for index, header in enumerate(headers)]
    scored.sort(reverse=True)
    best_score, best_index = scored[0] if scored else (0, -1)
    return best_index if best_score >= minimum else None


def _find_jd_header_row(rows: list[list[Any]]) -> tuple[int, list[str], dict[str, int | None]]:
    best: tuple[int, int, list[str], dict[str, int | None]] | None = None
    for row_index, row in enumerate(rows[:80]):
        headers = [_normalize_header(value) for value in row]
        vin_col = _best_header_col(headers, "vin", 60)
        jd_col = _best_header_col(headers, "jd_trade", 70)
        row_text = " ".join(headers)
        score = (100 if vin_col is not None else 0) + (100 if jd_col is not None else 0)
        if "trade" in row_text and ("jd" in row_text or "book" in row_text or "nada" in row_text):
            score += 30
        if score <= 0:
            continue
        columns = {
            "vin": vin_col,
            "jd_trade": jd_col,
            "vehicle": _best_header_col(headers, "vehicle", 40),
            "stock": _best_header_col(headers, "stock", 40),
            "price": _best_header_col(headers, "price", 40),
            "class": _best_header_col(headers, "class", 40),
            "new_used": _best_header_col(headers, "new_used", 40),
        }
        if best is None or score > best[0]:
            best = (score, row_index, headers, columns)
    if not best or best[3]["vin"] is None:
        raise HTTPException(status_code=400, detail={"message": "Could not find a VIN column in the valuation file."})
    if best[3]["jd_trade"] is None:
        raise HTTPException(status_code=400, detail={"message": "Could not find a JD Power / book trade-in value column in the valuation file."})
    _, row_index, headers, columns = best
    return row_index, headers, columns


def _row_value(row: list[Any], index: int | None) -> Any:
    if index is None or index < 0 or index >= len(row):
        return None
    return row[index]


def _extract_vin_from_cell(value: Any) -> str | None:
    match = VIN_PATTERN.search(str(value or "").upper())
    return match.group(0).upper() if match else None


def _fallback_trade_value(row: list[Any], skip_indexes: set[int]) -> float | None:
    values: list[float] = []
    for index, value in enumerate(row):
        if index in skip_indexes:
            continue
        parsed = _to_float(value)
        if parsed is not None and 1000 <= parsed <= 200000:
            values.append(parsed)
    return max(values) if values else None


def _parse_jd_power_file(raw: bytes, filename: str, content_type: str | None = None) -> dict[str, Any]:
    rows = _rows_from_upload_table(raw, filename, content_type)

    if not rows:
        return {"items": [], "diagnostics": {"rows_seen": 0, "message": "No readable rows found."}}

    header_row, headers, columns = _find_jd_header_row(rows)
    vin_col = columns["vin"]
    jd_col = columns["jd_trade"]

    items: list[dict[str, Any]] = []
    skipped = 0
    for row in rows[header_row + 1 :]:
        vin = _extract_vin_from_cell(_row_value(row, vin_col))
        if not vin:
            row_text = " ".join(_normalize_table_cell(cell) for cell in row)
            vin = _extract_vin_from_cell(row_text)
        if not vin:
            skipped += 1
            continue
        trade_value = _to_float(_row_value(row, jd_col))
        if trade_value is None:
            skipped += 1
            continue
        item = {
            "vin": vin,
            "vehicle": _normalize_table_cell(_row_value(row, columns.get("vehicle"))),
            "stock_number": _normalize_table_cell(_row_value(row, columns.get("stock"))),
            "class": _normalize_table_cell(_row_value(row, columns.get("class"))),
            "new_used": _normalize_table_cell(_row_value(row, columns.get("new_used"))),
            "dealer_price": _to_float(_row_value(row, columns.get("price"))),
            "jd_power_trade_in": trade_value,
            "source_file": filename,
        }
        items.append(item)
    diagnostics = {
        "rows_seen": len(rows),
        "header_row": header_row + 1,
        "skipped_rows": skipped,
        "detected_columns": {key: (headers[value] if value is not None and value < len(headers) else None) for key, value in columns.items()},
    }
    return {"items": items, "diagnostics": diagnostics}


def _save_jd_power_valuations(items: list[dict[str, Any]], source_file: str) -> dict[str, Any]:
    deduped: dict[str, dict[str, Any]] = {}
    for item in items:
        vin = str(item.get("vin") or "").strip().upper()
        if vin:
            deduped[vin] = item
    payload = {
        "items": list(deduped.values()),
        "count": len(deduped),
        "source_file": source_file,
        "updated_at": _utc_now(),
    }
    _safe_write_json(JD_POWER_VALUATIONS_PATH, payload)
    return payload


def _vehicle_bank_brain(vin: str) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    vehicle = _find_vehicle_by_vin(clean_vin)
    if not vehicle:
        raise HTTPException(status_code=404, detail={"message": f"Vehicle not found for VIN {clean_vin}"})
    decoded = _decode_vin_values(clean_vin)
    fields = decoded.get("fields") if isinstance(decoded, dict) else {}
    if not isinstance(fields, dict):
        fields = {}
    inventory_price = _to_float(vehicle.get("price")) or 0.0
    sale_price = float(_facebook_marketplace_price(inventory_price) or inventory_price or 0.0)
    valuation = _jd_power_valuation_for_vin(clean_vin)
    jd_trade_value = _to_float(valuation.get("jd_power_trade_in")) if valuation else None
    mileage = _to_float(vehicle.get("mileage")) or 0.0
    year = _to_float(fields.get("year"))
    current_year = datetime.now(timezone.utc).year
    age = max(0, current_year - int(year)) if year else None
    down = 0.0

    structure_request = CreditStructureRequest(
        vin=clean_vin,
        vehicle_price=sale_price,
        book_value=jd_trade_value,
        taxes=round(sale_price * DEFAULT_BANK_TAX_RATE, 2),
        tax_rate=DEFAULT_BANK_TAX_RATE,
        fees=DEFAULT_BANK_FEES,
        backend_products=0,
        down_payment=down,
        term_months=72,
        apr=11.99,
        monthly_income=None,
        current_dti=None,
        credit_score=None,
        tradelines=None,
        derogatories=None,
        utilization=None,
    )
    structure_result = _simulate_credit_structure(structure_request)
    recommendation = dict(structure_result.get("recommendation") or {})
    ranked = list(recommendation.get("ranked_banks") or [])

    collateral_flags: list[str] = []
    adjustment = 0.0
    if age is not None and age >= 8:
        collateral_flags.append("Older collateral: some prime banks may cap term/LTV.")
        adjustment -= 7.0
    if mileage >= 100_000:
        collateral_flags.append("High-mileage collateral: expect stricter advance and term caps.")
        adjustment -= 8.0
    if sale_price >= 75_000:
        collateral_flags.append("High-ticket unit: down payment and bureau strength matter more.")
        adjustment -= 3.0

    adjusted_ranked: list[dict[str, Any]] = []
    for item in ranked:
        if not isinstance(item, dict):
            continue
        row = dict(item)
        base_confidence = _to_float(row.get("confidence")) or 0.0
        row["confidence"] = max(1.0, min(99.0, round(base_confidence + adjustment, 1)))
        row.setdefault("reasons", [])
        if collateral_flags:
            row["reasons"] = list(row.get("reasons") or []) + collateral_flags[:1]
        adjusted_ranked.append(row)
    adjusted_ranked.sort(key=lambda item: item.get("confidence", 0), reverse=True)
    if adjusted_ranked:
        recommendation["ranked_banks"] = adjusted_ranked
        recommendation["best_bank"] = adjusted_ranked[0]
        recommendation["backup_bank"] = adjusted_ranked[1] if len(adjusted_ranked) > 1 else None
    recommendation["collateral_flags"] = collateral_flags

    carfax_summary = _carfax_summary_for_vin(clean_vin)
    packet = [
        "Pull/attach credit app before final lender selection.",
        "Attach bookout, invoice/sticker if available, proof of income, proof of residence, insurance, and trade payoff if applicable.",
    ]
    if collateral_flags:
        packet.append("Pre-empt collateral concerns with stronger down payment or shorter term.")

    return {
        "ok": True,
        "vin": clean_vin,
        "vehicle": vehicle,
        "decode": decoded,
        "carfax_summary": carfax_summary,
        "valuation": valuation,
        "default_structure": structure_result.get("structure"),
        "recommendation": recommendation,
        "packet_guidance": packet,
        "assumptions": [
            "No customer bureau attached yet, so confidence is collateral/deal-structure based.",
            "Final approval probability should be recalculated after credit report upload.",
        ],
    }


def _load_runtime_posts() -> list[dict[str, Any]]:
    if not FACEBOOK_POSTS_DIR.exists():
        return []
    posts: list[dict[str, Any]] = []
    for file_path in FACEBOOK_POSTS_DIR.glob("*/*facebook_listing.txt"):
        try:
            posts.append(_parse_listing_file(file_path))
        except Exception:
            continue
    posts.sort(key=lambda item: item.get("timestamp", 0), reverse=True)
    return posts


def _vin_candidate(value: Any) -> str:
    raw = str(value or "").strip().upper()
    if not raw:
        return ""
    if re.fullmatch(r"[A-HJ-NPR-Z0-9]{11,17}", raw):
        return raw
    return ""


def _extract_vin_from_text(value: Any) -> str:
    text = str(value or "").upper()
    match = re.search(r"\b[A-HJ-NPR-Z0-9]{17}\b", text)
    return match.group(0) if match else ""


def _value_for_keys(item: dict[str, Any], keys: tuple[str, ...]) -> Any:
    lower_map = {str(key).lower(): key for key in item.keys()}
    for key in keys:
        found = lower_map.get(key.lower())
        if found is not None:
            return item.get(found)
    return None


def _schema_types(item: dict[str, Any]) -> set[str]:
    value = item.get("@type")
    if isinstance(value, str):
        return {value.lower()}
    if isinstance(value, list):
        return {str(entry).lower() for entry in value}
    return set()


def _extract_json_block_after(text: str, start_index: int) -> str | None:
    length = len(text)
    index = start_index
    while index < length and text[index] in " \t\r\n=:":
        index += 1
    if index >= length or text[index] not in "{[":
        return None

    opener = text[index]
    closer = "}" if opener == "{" else "]"
    depth = 0
    in_string = False
    escape_next = False

    for cursor in range(index, length):
        char = text[cursor]
        if in_string:
            if escape_next:
                escape_next = False
            elif char == "\\":
                escape_next = True
            elif char == '"':
                in_string = False
            continue

        if char == '"':
            in_string = True
            continue
        if char == opener:
            depth += 1
            continue
        if char == closer:
            depth -= 1
            if depth == 0:
                return text[index : cursor + 1]

    return None


def _walk_json_dicts(payload: Any) -> list[dict[str, Any]]:
    found: list[dict[str, Any]] = []

    def _walk(node: Any) -> None:
        if isinstance(node, dict):
            found.append(node)
            for value in node.values():
                _walk(value)
            return
        if isinstance(node, list):
            for value in node:
                _walk(value)

    _walk(payload)
    return found


def _looks_like_vehicle_record(item: dict[str, Any]) -> bool:
    keys = {str(key).lower() for key in item.keys()}
    schema_types = _schema_types(item)

    # JSON-LD vehicle blocks (for example Product/Car with offers)
    if {"car", "product"}.intersection(schema_types):
        has_vehicle_payload = bool(
            keys.intersection(
                {
                    "name",
                    "vehicleidentificationnumber",
                    "offers",
                    "image",
                    "vehicleconfiguration",
                    "vehiclemodeldate",
                    "vehicleengine",
                }
            )
        )
        if has_vehicle_payload:
            return True

    if {"year", "make", "model"}.issubset(keys):
        return True
    if "vin" in keys or "vehicle_vin" in keys or "vehiclevin" in keys or "vehicleidentificationnumber" in keys:
        return True
    score = 0
    for bucket in [
        {"price", "internet_price", "msrp", "sale_price"},
        {"mileage", "odometer", "mileagefromodometer"},
        {"photos", "images", "media", "gallery", "image"},
        {"detail_url", "vehicle_url", "vdp_url", "url", "href"},
    ]:
        if keys.intersection(bucket):
            score += 1
    return score >= 2


def _extract_inventory_dicts_from_payload(payload: Any) -> list[dict[str, Any]]:
    candidates: list[dict[str, Any]] = []
    for node in _walk_json_dicts(payload):
        if _looks_like_vehicle_record(node):
            candidates.append(node)
    return candidates



def _extract_inventory_dicts_from_html(html_text: str, source_url: str | None = None) -> tuple[list[dict[str, Any]], list[str]]:
    """Extract inventory records from Dealer.com / Dealer Inspire style pages.

    The old version only looked for JSON in script tags. Taverna/Dealer.com often
    renders the visible inventory with client-side widgets, so the raw HTML can
    contain either JSON bootstraps, vehicle-card markup, or only VDP links.
    This parser tries all three before giving up.
    """
    payloads: list[Any] = []
    notes: list[str] = []

    script_matches = re.findall(
        r"<script(?P<attrs>[^>]*)>(?P<body>.*?)</script>",
        html_text,
        flags=re.IGNORECASE | re.DOTALL,
    )

    assignment_tokens = [
        "__NEXT_DATA__",
        "__PRELOADED_STATE__",
        "__INITIAL_STATE__",
        "__APOLLO_STATE__",
        "inventoryData",
        "inventoryState",
        "vehicleData",
        "window.DDC",
        "DDC.data",
        "digitalData",
    ]

    for attrs, body in script_matches:
        attrs_lower = attrs.lower()
        body_text = html.unescape(body.strip())
        if not body_text:
            continue

        is_json_script = "application/ld+json" in attrs_lower or "application/json" in attrs_lower
        if is_json_script or "__next_data__" in attrs_lower:
            try:
                payloads.append(json.loads(body_text))
                continue
            except Exception:
                pass

        # Common JS assignment payloads.
        for token in assignment_tokens:
            if token not in body_text:
                continue
            token_index = body_text.find(token)
            equals_index = body_text.find("=", token_index)
            if equals_index < 0:
                continue
            json_block = _extract_json_block_after(body_text, equals_index + 1)
            if not json_block:
                continue
            try:
                payloads.append(json.loads(json_block))
                break
            except Exception:
                continue

        # Dealer widgets sometimes inline arrays/objects without a stable global.
        if any(needle in body_text.lower() for needle in ("vin", "stocknumber", "internetprice", "vehiclecard", "inventory")):
            for json_block in re.findall(r"(\{[^{}]{0,4000}(?:vin|VIN|stockNumber|internetPrice|salePrice|vehicleName)[\s\S]{0,4000}?\})", body_text):
                try:
                    payloads.append(json.loads(json_block))
                except Exception:
                    continue

    records: list[dict[str, Any]] = []
    for payload in payloads:
        records.extend(_extract_inventory_dicts_from_payload(payload))

    # HTML/card fallback for pages where JSON bootstraps are stripped or empty.
    html_records: list[dict[str, Any]] = []
    if BeautifulSoup is not None:
        try:
            soup = BeautifulSoup(html_text, "html.parser")

            # Find likely vehicle cards first.
            card_selectors = [
                "[data-vin]",
                "[data-vehicle-vin]",
                "[data-stock]",
                "[class*='vehicle-card']",
                "[class*='inventory-listing']",
                "[class*='vehicle-listing']",
                "[class*='inventory-item']",
                "[class*='hproduct']",
            ]
            candidate_cards = []
            for selector in card_selectors:
                candidate_cards.extend(soup.select(selector))

            # If cards are not available, build records from VDP links.
            vdp_links = soup.find_all(
                "a",
                href=re.compile(r"/(?:used|new|certified)/.*\.htm|VehicleDetails|vin=|stock=", re.I),
            )

            seen_cards: set[int] = set()
            for link in vdp_links:
                card = None
                for parent in [link] + list(link.parents)[:6]:
                    text_value = parent.get_text(" ", strip=True) if hasattr(parent, "get_text") else ""
                    if re.search(r"\b(?:19|20)\d{2}\b", text_value) and (
                        re.search(r"\$\s?[0-9][0-9,]{2,}", text_value) or re.search(r"\b[A-HJ-NPR-Z0-9]{17}\b", text_value.upper())
                    ):
                        card = parent
                        break
                if card is not None and id(card) not in seen_cards:
                    candidate_cards.append(card)
                    seen_cards.add(id(card))

            seen_detail_urls: set[str] = set()
            for card in candidate_cards:
                try:
                    text_value = html.unescape(card.get_text(" ", strip=True))
                    if not text_value:
                        continue
                    href = ""
                    link = card.find("a", href=re.compile(r"/(?:used|new|certified)/.*\.htm|VehicleDetails|vin=|stock=", re.I))
                    if link is not None:
                        href = str(link.get("href") or "").strip()
                    detail_url = urljoin(source_url or "", href) if href else None
                    if detail_url:
                        if detail_url.lower() in seen_detail_urls:
                            continue
                        seen_detail_urls.add(detail_url.lower())

                    vin = (
                        _vin_candidate(card.get("data-vin"))
                        or _vin_candidate(card.get("data-vehicle-vin"))
                        or _extract_vin_from_text(text_value)
                    )
                    stock = (
                        str(card.get("data-stock") or card.get("data-stock-number") or "").strip()
                        or (re.search(r"\b(?:stock|stk)\s*#?\s*:?\s*([A-Z0-9-]{3,24})\b", text_value, re.I).group(1)
                            if re.search(r"\b(?:stock|stk)\s*#?\s*:?\s*([A-Z0-9-]{3,24})\b", text_value, re.I)
                            else "")
                    )
                    title = ""
                    title_node = card.find(["h1", "h2", "h3", "h4"])
                    if title_node:
                        title = title_node.get_text(" ", strip=True)
                    if not title and link is not None:
                        title = link.get_text(" ", strip=True) or str(link.get("title") or "")
                    if not title:
                        title_match = re.search(
                            r"\b((?:19|20)\d{2}\s+(?:Chrysler|Dodge|Jeep|Ram|FIAT|Fiat|Kia|Toyota|Honda|Ford|Chevrolet|GMC|Nissan|Hyundai|Mazda|Volkswagen|Subaru|BMW|Mercedes-Benz|Audi|Lexus|Acura|Cadillac|Land Rover|Porsche|Volvo|Genesis|Tesla)[A-Za-z0-9 .'\-]{2,100})",
                            text_value,
                            flags=re.I,
                        )
                        title = title_match.group(1).strip() if title_match else ""

                    price_match = re.search(r"\$\s?([0-9][0-9,]{2,})", text_value)
                    mileage_match = re.search(r"\b([0-9][0-9,]{1,8})\s*(?:miles|mi)\b", text_value, re.I)
                    images = []
                    for image in card.find_all("img"):
                        for attr in ("src", "data-src", "data-original", "data-lazy", "data-lazy-src"):
                            src = str(image.get(attr) or "").strip()
                            if src:
                                images.append(urljoin(source_url or "", src))
                                break

                    if not (vin or title or detail_url):
                        continue
                    html_records.append(
                        {
                            "vin": vin or (_synthetic_vin_from_detail_url(detail_url) if detail_url else "UNKNOWN"),
                            "title": title or vin or "Vehicle",
                            "price": int(price_match.group(1).replace(",", "")) if price_match else None,
                            "mileage": int(mileage_match.group(1).replace(",", "")) if mileage_match else None,
                            "stock_number": stock or None,
                            "detail_url": detail_url,
                            "photos": _dedupe_urls(images),
                            "status_label": "In Stock",
                        }
                    )
                except Exception:
                    continue
        except Exception as exc:
            notes.append(f"html_card_parse_error={exc}")

    records.extend(html_records)

    notes.append(f"script_payloads={len(payloads)}")
    notes.append(f"script_candidate_records={len(records) - len(html_records)}")
    notes.append(f"html_candidate_records={len(html_records)}")
    notes.append(f"candidate_records={len(records)}")
    return records, notes


def _proxy_markdown_url_for_source(source_url: str) -> str:
    raw = str(source_url or "").strip()
    if not raw:
        return "https://r.jina.ai/http://"
    without_scheme = re.sub(r"^https?://", "", raw, flags=re.IGNORECASE)
    return f"https://r.jina.ai/http://{without_scheme}"


def _synthetic_vin_from_detail_url(detail_url: str) -> str:
    parsed = urlparse(detail_url)
    tail = parsed.path.rsplit("/", 1)[-1]
    tail = tail.rsplit(".", 1)[0]
    token = re.sub(r"[^A-Z0-9]", "", str(tail).upper())
    digest = hashlib.sha1(str(detail_url).encode("utf-8")).hexdigest().upper()
    synthetic = (token + digest)[:17]
    return synthetic if synthetic else digest[:17]


def _extract_inventory_dicts_from_markdown_proxy(
    markdown_text: str,
    *,
    source_url: str | None = None,
) -> tuple[list[dict[str, Any]], list[str]]:
    text = str(markdown_text or "")
    records: list[dict[str, Any]] = []
    notes: list[str] = []

    image_by_detail: dict[str, list[str]] = {}
    image_pattern = re.compile(
        r"\[!\[[^\]]*\]\((https?://[^)\s]+)\)\]\((https?://[^)\s]+/(?:used|new)/[^)\s]+\.htm)\)",
        flags=re.IGNORECASE,
    )
    for image_url, detail_url in image_pattern.findall(text):
        clean_detail = detail_url.strip()
        clean_image = image_url.strip()
        if clean_detail.startswith("http://"):
            clean_detail = "https://" + clean_detail[len("http://") :]
        if clean_image.startswith("http://"):
            clean_image = "https://" + clean_image[len("http://") :]
        bucket = image_by_detail.setdefault(clean_detail, [])
        lowered = {entry.lower() for entry in bucket}
        if clean_image.lower() not in lowered:
            bucket.append(clean_image)

    listing_pattern = re.compile(
        r"\[(\d{4}\s+[^\]]+?)\]\((https?://[^)\s]+/(?:used|new)/[^)\s]+\.htm)\)\s*\$([0-9,]+)",
        flags=re.IGNORECASE,
    )

    seen: set[str] = set()
    for match in listing_pattern.finditer(text):
        title = str(match.group(1) or "").strip()
        detail_url = str(match.group(2) or "").strip()
        price_raw = str(match.group(3) or "").replace(",", "").strip()

        if detail_url.startswith("http://"):
            detail_url = "https://" + detail_url[len("http://") :]
        if source_url and detail_url.startswith("/"):
            detail_url = urljoin(source_url, detail_url)

        unique_key = detail_url.lower() or title.lower()
        if not unique_key or unique_key in seen:
            continue
        seen.add(unique_key)

        mileage_raw: int | None = None
        window = text[match.end() : match.end() + 320]
        mileage_match = re.search(r"([0-9][0-9,]{0,9})\s+miles", window, flags=re.IGNORECASE)
        if mileage_match:
            try:
                mileage_raw = int(mileage_match.group(1).replace(",", ""))
            except Exception:
                mileage_raw = None

        price_value: int | str | None
        try:
            price_value = int(price_raw) if price_raw else None
        except Exception:
            price_value = price_raw or None

        photos = image_by_detail.get(detail_url, [])
        records.append(
            {
                "vin": _synthetic_vin_from_detail_url(detail_url),
                "title": title,
                "price": price_value,
                "mileage": mileage_raw,
                "detail_url": detail_url,
                "photos": photos,
                "status_label": "In Stock",
            }
        )

    notes.append(f"proxy_image_links={sum(len(value) for value in image_by_detail.values())}")
    notes.append(f"proxy_candidate_records={len(records)}")
    return records, notes


def _dedupe_urls(values: list[str]) -> list[str]:
    normalized: list[str] = []
    seen: set[str] = set()
    for raw in values:
        candidate = html.unescape(str(raw or "").strip())
        if not candidate:
            continue
        if candidate.startswith("//"):
            candidate = f"https:{candidate}"
        if candidate.startswith("http://"):
            candidate = "https://" + candidate[len("http://") :]
        lowered = candidate.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        normalized.append(candidate)
    return normalized


def _dedupe_strings(values: list[str]) -> list[str]:
    normalized: list[str] = []
    seen: set[str] = set()
    for raw in values:
        candidate = re.sub(r"\s+", " ", str(raw or "").strip())
        if not candidate:
            continue
        lowered = candidate.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        normalized.append(candidate)
    return normalized


def _extract_vehicle_photo_urls_from_html(html_text: str, *, base_url: str | None = None) -> list[str]:
    candidates = re.findall(
        r"""(?:src|data-src|data-zoom-image|content)\s*=\s*["']([^"']+)["']""",
        html_text,
        flags=re.IGNORECASE,
    )
    candidates.extend(
        match.replace("\\/", "/").replace("\\u0026", "&")
        for match in re.findall(
            r"""https?:\\?/\\?/pictures\.dealer\.com[^"'\s<>)\\]+""",
            html_text,
            flags=re.IGNORECASE,
        )
    )
    candidates.extend(
        re.findall(
            r"""https?://pictures\.dealer\.com[^"'\s<>)\\]+""",
            html_text,
            flags=re.IGNORECASE,
        )
    )

    parsed_base = urlparse(base_url or "")
    account_hint = ""
    account_match = re.search(r"accountId=([A-Za-z0-9_-]+)", html_text)
    if account_match:
        account_hint = account_match.group(1).strip().lower()

    absolute: list[str] = []
    for raw in candidates:
        candidate = str(raw or "").strip()
        if not candidate:
            continue
        candidate = candidate.replace("\\/", "/").replace("\\u0026", "&")
        candidate = html.unescape(candidate).rstrip("\\")
        if candidate.startswith("//"):
            candidate = f"https:{candidate}"
        elif base_url and candidate.startswith("/"):
            candidate = urljoin(base_url, candidate)
        lowered = candidate.lower()
        if "pictures.dealer.com" not in lowered:
            continue
        if account_hint and f"/{account_hint}/" not in lowered:
            continue
        if not account_hint and parsed_base.netloc and "dealer.com" not in parsed_base.netloc and "tavernacdjrfllccllc" not in lowered:
            continue
        if "/thumb_" in lowered or "sprite" in lowered or "logo" in lowered:
            continue
        if not re.search(r"\.(?:jpe?g|png|webp)(?:[?#].*)?$", lowered):
            continue
        candidate = re.sub(r"([.](?:jpe?g|png|webp))(?:[?#].*)?$", r"\1", candidate, flags=re.IGNORECASE)
        absolute.append(candidate)

    deduped = _dedupe_urls(absolute)
    primary = [url for url in deduped if "/tavernacdjrfllccllc/" in url.lower()]
    return primary or deduped


def _extract_detail_facts_from_markdown_proxy(markdown_text: str) -> dict[str, Any]:
    text = str(markdown_text or "")
    compact = re.sub(r"\s+", " ", text)
    facts: dict[str, Any] = {}

    title_match = re.search(r"##\s+(?:Used\s+)?(\d{4}\s+[^\n]+)", text)
    if title_match:
        facts["title"] = title_match.group(1).strip()

    vin_match = re.search(r"\bVIN\s+([A-HJ-NPR-Z0-9]{17})\b", compact, flags=re.IGNORECASE)
    if vin_match:
        facts["vin"] = vin_match.group(1).strip().upper()

    stock_match = re.search(r"\bStock Number\s+([A-Z0-9-]+)\b", compact, flags=re.IGNORECASE)
    if stock_match:
        facts["stock_number"] = stock_match.group(1).strip().upper()

    mileage_match = re.search(r"\bOdometer\s+([0-9,]+)\s+miles\b", compact, flags=re.IGNORECASE)
    if mileage_match:
        try:
            facts["mileage"] = int(mileage_match.group(1).replace(",", ""))
        except Exception:
            pass

    field_patterns = {
        "exterior": r"\bExterior Color\s+(.+?)\s+Interior Color\b",
        "interior": r"\bInterior Color\s+(.+?)\s+Odometer\b",
        "transmission": r"\bTransmission\s+(.+?)\s+Drivetrain\b",
        "drivetrain": r"\bDrivetrain\s+(.+?)\s+Engine\b",
        "engine": r"\bEngine\s+(.+?)\s+VIN\b",
    }
    for key, pattern in field_patterns.items():
        match = re.search(pattern, compact, flags=re.IGNORECASE)
        if match:
            facts[key] = match.group(1).strip()

    highlights_match = re.search(
        r"###\s+Highlighted Features\s+(.*?)(?:###\s+Included Packages|###\s+Detailed Specifications|###\s+KBB\.com Consumer Reviews|\Z)",
        text,
        flags=re.IGNORECASE | re.DOTALL,
    )
    if highlights_match:
        highlights = [
            re.sub(r"\s+", " ", line).strip(" -*")
            for line in highlights_match.group(1).splitlines()
            if line.strip().startswith("*")
        ]
        highlights = [item for item in highlights if item]
        if highlights:
            facts["highlights"] = highlights[:24]

    return facts


def _fetch_detail_markdown_via_proxy(
    *,
    detail_url: str,
    client: httpx.Client,
) -> tuple[str | None, list[str]]:
    proxy_url = _proxy_markdown_url_for_source(detail_url)
    response = client.get(proxy_url)
    notes = [
        f"detail_proxy_url={proxy_url}",
        f"detail_proxy_status={response.status_code}",
    ]
    if response.status_code >= 400:
        return None, notes
    return response.text, notes


def _merge_detail_payload_into_record(
    record: dict[str, Any],
    detail_payload: dict[str, Any],
) -> dict[str, Any]:
    merged = dict(record)
    if detail_payload.get("vin"):
        merged["vin"] = detail_payload["vin"]

    for key in (
        "title",
        "mileage",
        "drivetrain",
        "engine",
        "transmission",
        "exterior",
        "interior",
        "stock_number",
        "highlights",
    ):
        value = detail_payload.get(key)
        if value in (None, "", [], {}):
            continue
        if key == "mileage":
            merged[key] = value
            continue
        if not merged.get(key):
            merged[key] = value
            continue
        if key == "title" and str(merged.get(key) or "").strip().lower() == str(record.get("vin") or "").strip().lower():
            merged[key] = value

    return merged


def _enrich_inventory_records_from_proxy_detail_pages(
    records: list[dict[str, Any]],
    *,
    timeout_seconds: int,
) -> tuple[list[dict[str, Any]], list[str]]:
    enriched: list[dict[str, Any]] = []
    diagnostics: list[str] = []
    started_at = time.time()
    budget_seconds = max(12.0, float(timeout_seconds) * 1.25)

    if not records:
        return enriched, diagnostics

    with httpx.Client(
        timeout=max(4.0, min(6.0, float(timeout_seconds))),
        follow_redirects=True,
        headers={
            "user-agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/126.0.0.0 Safari/537.36"
            ),
            "accept": "text/plain,text/markdown,text/html;q=0.9,*/*;q=0.8",
        },
    ) as client:
        for index, record in enumerate(records):
            if time.time() - started_at > budget_seconds:
                diagnostics.append(f"detail_proxy_budget_exhausted_at={index}")
                enriched.extend(dict(item) for item in records[index:])
                break
            detail_url = str(record.get("detail_url") or "").strip()
            if not detail_url:
                enriched.append(dict(record))
                continue

            merged = dict(record)
            try:
                markdown_text, notes = _fetch_detail_markdown_via_proxy(detail_url=detail_url, client=client)
                diagnostics.extend(notes)
                if markdown_text:
                    merged = _merge_detail_payload_into_record(
                        merged,
                        _extract_detail_facts_from_markdown_proxy(markdown_text),
                    )
            except Exception as exc:
                diagnostics.append(f"detail_proxy_error={detail_url}|{exc}")
            enriched.append(merged)

    diagnostics.append(f"detail_proxy_enriched={len(enriched)}")
    return enriched, diagnostics


def _fetch_live_inventory_records_via_proxy_markdown(*, source_url: str, timeout_seconds: int) -> dict[str, Any]:
    proxy_url = _proxy_markdown_url_for_source(source_url)
    with httpx.Client(
        timeout=float(timeout_seconds),
        follow_redirects=True,
        headers={
            "user-agent": (
                "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                "AppleWebKit/537.36 (KHTML, like Gecko) "
                "Chrome/126.0.0.0 Safari/537.36"
            ),
            "accept": "text/plain,text/markdown,text/html,application/json;q=0.9,*/*;q=0.8",
        },
    ) as client:
        response = client.get(proxy_url)

    content_type = str(response.headers.get("content-type", "")).lower()
    notes: list[str] = [
        "source_mode=proxy_markdown",
        f"proxy_url={proxy_url}",
        f"proxy_status={response.status_code}",
        f"proxy_content_type={content_type}",
    ]
    if response.status_code >= 400:
        return {
            "source_url": source_url,
            "fetched_at": datetime.now(timezone.utc).isoformat(),
            "items": [],
            "items_count": 0,
            "diagnostics": notes,
        }

    records, parse_notes = _extract_inventory_dicts_from_markdown_proxy(
        response.text,
        source_url=source_url,
    )
    notes.extend(parse_notes)
    records, detail_notes = _enrich_inventory_records_from_proxy_detail_pages(
        records,
        timeout_seconds=timeout_seconds,
    )
    notes.extend(detail_notes)
    return {
        "source_url": source_url,
        "fetched_at": datetime.now(timezone.utc).isoformat(),
        "items": records,
        "items_count": len(records),
        "diagnostics": notes,
    }


def _normalize_photo_entries(value: Any) -> list[str]:
    raw_items = value if isinstance(value, list) else [value]
    normalized: list[str] = []
    seen: set[str] = set()

    for entry in raw_items:
        candidate: str | None = None
        if isinstance(entry, str):
            stripped = entry.strip()
            if stripped.startswith(("http://", "https://")):
                candidate = stripped
        elif isinstance(entry, dict):
            for key in ("url", "src", "image", "photo"):
                value_at_key = entry.get(key)
                if isinstance(value_at_key, str):
                    stripped = value_at_key.strip()
                    if stripped.startswith(("http://", "https://")):
                        candidate = stripped
                        break
        if not candidate:
            continue
        lowered = candidate.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        normalized.append(candidate)
    return normalized


def _normalize_mileage_value(raw: Any) -> Any:
    if isinstance(raw, dict):
        for key in ("value", "amount"):
            candidate = raw.get(key)
            if candidate is not None:
                return candidate
        return None
    return raw


def _normalize_status_label(item: dict[str, Any], offers: dict[str, Any] | None) -> str:
    raw = _value_for_keys(item, ("status_label", "status", "availability", "inventory_status"))
    if raw is None and offers:
        raw = offers.get("availability")
    text = str(raw or "").strip()
    lowered = text.lower()
    if "instock" in lowered or "in stock" in lowered:
        return "In Stock"
    if "outofstock" in lowered or "out of stock" in lowered:
        return "Out of Stock"
    if "preorder" in lowered or "pre-order" in lowered:
        return "Preorder"
    if text:
        return text
    return "Ready"


def _normalize_text_value(raw: Any) -> Any:
    if isinstance(raw, dict):
        for key in ("name", "label", "text", "description", "value"):
            candidate = raw.get(key)
            if isinstance(candidate, (str, int, float)) and str(candidate).strip():
                return candidate
        return None
    return raw


def _normalize_inventory_records(records: list[dict[str, Any]], *, source_url: str | None = None) -> list[dict[str, Any]]:
    normalized: list[dict[str, Any]] = []
    seen: set[str] = set()

    for item in records:
        if not isinstance(item, dict):
            continue

        offers_raw = _value_for_keys(item, ("offers", "offer"))
        offers: dict[str, Any] | None = None
        if isinstance(offers_raw, dict):
            offers = offers_raw
        elif isinstance(offers_raw, list):
            offers = next((entry for entry in offers_raw if isinstance(entry, dict)), None)

        offer_url = offers.get("url") if isinstance(offers, dict) else None
        offer_price = offers.get("price") if isinstance(offers, dict) else None
        seller_name = None
        if isinstance(offers, dict) and isinstance(offers.get("seller"), dict):
            seller_name = offers["seller"].get("name")

        vin = _vin_candidate(
            _value_for_keys(
                item,
                (
                    "vin",
                    "vehicle_vin",
                    "vehicleVin",
                    "VIN",
                    "vehicleIdentificationNumber",
                    "vehicle_identification_number",
                ),
            )
        ) or _extract_vin_from_text(_value_for_keys(item, ("title", "name", "vehicleName", "description")))

        year_from_model_date = _value_for_keys(item, ("vehicleModelDate",))
        year_candidate = ""
        if year_from_model_date:
            matched_year = re.search(r"\b(19|20)\d{2}\b", str(year_from_model_date))
            year_candidate = matched_year.group(0) if matched_year else ""

        title = str(
            _value_for_keys(item, ("title", "name", "vehicleName", "vehicle_title"))
            or " ".join(
                str(_value_for_keys(item, (part,)) or "").strip()
                for part in ("year", "make", "model", "trim", "vehicleConfiguration")
                if str(_value_for_keys(item, (part,)) or "").strip()
            )
        ).strip()
        if not title and year_candidate:
            make = str(_value_for_keys(item, ("make", "brand")) or "").strip()
            model = str(_value_for_keys(item, ("model", "vehicleConfiguration")) or "").strip()
            title = " ".join(part for part in [year_candidate, make, model] if part).strip()

        if not vin and not title:
            continue

        detail_url_raw = _value_for_keys(
            item,
            ("detail_url", "vehicle_url", "vdp_url", "url", "href", "permalink"),
        )
        detail_url = str(detail_url_raw).strip() if detail_url_raw else (str(offer_url).strip() if offer_url else None)
        if detail_url and source_url and detail_url.startswith("/"):
            detail_url = urljoin(source_url, detail_url)

        photos = _normalize_photo_entries(
            _value_for_keys(item, ("photos", "images", "media", "gallery", "image"))
        )
        carfax_url, carfax_facts = _extract_carfax_link_from_inventory_item(item)
        direct_carfax_url = _value_for_keys(item, ("carfax_url", "carfaxUrl", "carfax", "vehicleHistoryUrl"))
        if direct_carfax_url and not carfax_url:
            carfax_url = str(direct_carfax_url).strip()
        inventory_category = str(_value_for_keys(item, ("inventory_category", "inventoryCategory", "type", "condition")) or "").strip()

        output = {
            "vin": vin or "UNKNOWN",
            "title": title or (vin or "UNKNOWN"),
            "price": _value_for_keys(
                item,
                ("price", "internet_price", "internetPrice", "sale_price", "msrp"),
            )
            or offer_price,
            "mileage": _normalize_mileage_value(
                _value_for_keys(item, ("mileage", "odometer", "mileageFromOdometer"))
            ),
            "drivetrain": _normalize_text_value(
                _value_for_keys(item, ("drivetrain", "drive_train", "driveWheelConfiguration"))
            ),
            "engine": _normalize_text_value(
                _value_for_keys(item, ("engine", "engine_description", "vehicleEngine"))
            ),
            "transmission": _normalize_text_value(
                _value_for_keys(item, ("transmission", "transmission_description", "vehicleTransmission"))
            ),
            "location": _normalize_text_value(
                _value_for_keys(item, ("location", "dealer_name", "store_name")) or seller_name
            ),
            "detail_url": detail_url,
            "exterior": _value_for_keys(item, ("exterior", "ext_color", "exterior_color", "color")),
            "interior": _value_for_keys(
                item,
                ("interior", "int_color", "interior_color", "vehicleInteriorColor"),
            ),
            "photos": photos,
            "status_label": _normalize_status_label(item, offers),
            "inventory_category": inventory_category,
            "stock_number": _value_for_keys(item, ("stock_number", "stockNumber", "stock")),
            "carfax_url": carfax_url,
            "carfax_facts": carfax_facts,
        }

        unique_key = output["vin"] if output["vin"] != "UNKNOWN" else (
            output["detail_url"] or f"{output['title']}|{output.get('price')}"
        )
        unique_key = str(unique_key).strip().lower()
        if not unique_key or unique_key in seen:
            continue
        seen.add(unique_key)
        normalized.append(output)

    return normalized


def _normalize_inventory_blob(blob: Any, *, source_url: str | None = None) -> list[dict[str, Any]]:
    if isinstance(blob, list):
        return _normalize_inventory_records(
            [entry for entry in blob if isinstance(entry, dict)],
            source_url=source_url,
        )

    if isinstance(blob, dict):
        for key in ("items", "vehicles", "inventory", "results", "data"):
            value = blob.get(key)
            if isinstance(value, list):
                return _normalize_inventory_records(
                    [entry for entry in value if isinstance(entry, dict)],
                    source_url=source_url,
                )
        return _normalize_inventory_records([blob], source_url=source_url)

    return []


def _read_inventory_blob_with_backup(primary_path: Path, backup_path: Path) -> Any:
    primary_payload = _safe_read_json(primary_path, None)
    primary_items = _normalize_inventory_blob(primary_payload)
    if primary_items:
        return primary_payload
    backup_payload = _safe_read_json(backup_path, None)
    backup_items = _normalize_inventory_blob(backup_payload)
    if backup_items:
        return backup_payload
    return primary_payload if primary_payload is not None else backup_payload


def _read_vehicle_assets_cache(vin: str) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        return {}
    cache_path = _vehicle_assets_cache_path(clean_vin)
    cached = _safe_read_json(cache_path, {})
    if not isinstance(cached, dict):
        return {}
    changed = False
    for key in ("carfax_report", "carfax_facts"):
        value = cached.get(key)
        if isinstance(value, dict):
            normalized = _normalize_carfax_parsed_payload(value)
            if normalized != value:
                cached[key] = normalized
                changed = True
    summary = cached.get("carfax_summary")
    if isinstance(summary, dict):
        facts = summary.get("facts")
        normalized_facts = _normalize_carfax_parsed_payload(facts) if isinstance(facts, dict) else facts
        normalized_summary = dict(summary)
        if isinstance(normalized_facts, dict):
            normalized_summary["facts"] = normalized_facts
            normalized_summary["summary"] = str(normalized_facts.get("summary") or normalized_summary.get("summary") or "").strip()
            normalized_summary["highlights"] = list(normalized_facts.get("highlights") or normalized_summary.get("highlights") or [])
        if normalized_summary != summary:
            cached["carfax_summary"] = normalized_summary
            changed = True
    if isinstance(cached.get("carfax_facts"), dict):
        refreshed_summary = _carfax_summary_from_assets(clean_vin, cached)
        if refreshed_summary != cached.get("carfax_summary"):
            cached["carfax_summary"] = refreshed_summary
            changed = True
    if changed:
        _safe_write_json(cache_path, cached)
    return cached


def _merge_cached_vehicle_assets(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    merged: list[dict[str, Any]] = []
    for item in items:
        row = dict(item)
        cached = _read_vehicle_assets_cache(str(row.get("vin") or ""))
        photos = cached.get("photos")
        if isinstance(photos, list) and photos:
            row["photos"] = photos
        for key in ("sticker_url", "carfax_url", "main_photo", "photos_count", "carfax_facts"):
            value = cached.get(key)
            if value not in (None, "", [], {}):
                row[key] = value

        clean_vin = str(row.get("vin") or "").upper()
        cached_summary = cached.get("carfax_summary")
        if isinstance(cached_summary, dict):
            if _is_generic_carfax_summary(cached_summary):
                cached_summary = _carfax_summary_from_assets(
                    clean_vin,
                    {
                        "vin": clean_vin,
                        "carfax_url": row.get("carfax_url"),
                        "carfax_facts": row.get("carfax_facts") if isinstance(row.get("carfax_facts"), dict) else {},
                    },
                )
            row["carfax_summary"] = cached_summary
        elif cached.get("carfax_facts"):
            row["carfax_summary"] = _carfax_summary_from_assets(
                clean_vin,
                {
                    "vin": clean_vin,
                    "carfax_url": row.get("carfax_url"),
                    "carfax_facts": row.get("carfax_facts") if isinstance(row.get("carfax_facts"), dict) else {},
                },
            )
        merged.append(row)
    return merged


def _load_inventory_candidates() -> list[dict[str, Any]]:
    manual_items = _load_manual_inventory()
    live_normalized = _normalize_inventory_blob(
        _read_inventory_blob_with_backup(INVENTORY_LIVE_CACHE_PATH, INVENTORY_LIVE_BACKUP_PATH)
    )
    if live_normalized:
        return _merge_cached_vehicle_assets(_merge_inventory_sources(live_normalized, manual_items))

    snapshot_normalized = _normalize_inventory_blob(_safe_read_json(INVENTORY_SNAPSHOT_PATH, []))
    if snapshot_normalized:
        return _merge_cached_vehicle_assets(_merge_inventory_sources(snapshot_normalized, manual_items))

    latest_by_vin: dict[str, dict[str, Any]] = {}
    for post in _load_runtime_posts():
        vin = str(post.get("vin", "")).strip().upper()
        if not vin:
            continue
        if vin in latest_by_vin:
            continue
        latest_by_vin[vin] = {
            "vin": vin,
            "title": post.get("title"),
            "price": post.get("price"),
            "mileage": post.get("mileage"),
            "drivetrain": post.get("drivetrain"),
            "engine": post.get("engine"),
            "transmission": post.get("transmission"),
            "location": post.get("location"),
            "detail_url": post.get("detail_url"),
            "photos": [],
            "status_label": "Ready",
        }
    return _merge_cached_vehicle_assets(_merge_inventory_sources(list(latest_by_vin.values()), manual_items))


def _is_active_inventory_item(item: dict[str, Any]) -> bool:
    status = str(item.get("status_label") or item.get("status") or "").strip().lower()
    if "transit" in status or "factory" in status:
        return False
    return True


def _inventory_count_summary(items: list[dict[str, Any]]) -> dict[str, int]:
    return {
        "total": len(items),
        "active": sum(1 for item in items if _is_active_inventory_item(item)),
        "in_transit": sum(
            1
            for item in items
            if "transit" in str(item.get("status_label") or item.get("status") or "").lower()
        ),
    }



def _fetch_live_inventory_records(*, source_url: str, timeout_seconds: int) -> dict[str, Any]:
    """Fetch and normalize live dealership inventory.

    Fixes:
    - Taverna/Dealer.com pages are often JS-rendered, so direct HTML may return
      a shell or a few placeholder records. We now fall back when the direct
      parse is weak, not only when it is zero.
    - normalized_proxy was previously referenced before assignment if the proxy
      fetch failed.
    - We choose the best result among direct HTML, proxy markdown, and browser
      rendering instead of accepting a bad first parse.
    """
    diagnostics: list[str] = []
    raw_records: list[dict[str, Any]] = []
    normalized: list[dict[str, Any]] = []
    normalized_proxy: list[dict[str, Any]] = []
    normalized_browser: list[dict[str, Any]] = []
    fetched_at = datetime.now(timezone.utc).isoformat()

    try:
        with httpx.Client(
            timeout=float(timeout_seconds),
            follow_redirects=True,
            headers={
                "user-agent": (
                    "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                    "AppleWebKit/537.36 (KHTML, like Gecko) "
                    "Chrome/126.0.0.0 Safari/537.36"
                ),
                "accept": "text/html,application/json;q=0.9,*/*;q=0.8",
                "accept-language": "en-US,en;q=0.9",
                "cache-control": "no-cache",
            },
        ) as client:
            response = client.get(source_url)

        content_type = str(response.headers.get("content-type", "")).lower()
        diagnostics.extend([f"http_status={response.status_code}", f"content_type={content_type}"])

        if response.status_code < 400:
            if "json" in content_type:
                payload = response.json()
                raw_records = _extract_inventory_dicts_from_payload(payload)
                diagnostics.append(f"payload_records={len(raw_records)}")
            else:
                html_text = response.text
                raw_records, html_notes = _extract_inventory_dicts_from_html(html_text, source_url=source_url)
                diagnostics.extend(html_notes)

            normalized = _normalize_inventory_records(raw_records, source_url=source_url)
            diagnostics.append(f"direct_normalized_records={len(normalized)}")
        else:
            diagnostics.append(f"direct_http_error={response.status_code}")
    except Exception as exc:
        diagnostics.append(f"direct_fetch_error={exc}")

    # Direct raw HTML from Dealer.com can show only a rendered shell. Treat a
    # tiny direct result as weak and continue to fallbacks.
    try:
        min_direct_records = int(os.getenv("DEALERSHIP_INVENTORY_MIN_DIRECT_RECORDS", "20") or "20")
    except ValueError:
        min_direct_records = 20

    should_fallback = len(normalized) < min_direct_records

    if should_fallback:
        try:
            proxy_payload = _fetch_live_inventory_records_via_proxy_markdown(
                source_url=source_url,
                timeout_seconds=timeout_seconds,
            )
            diagnostics.extend(list(proxy_payload.get("diagnostics") or []))
            proxy_records = list(proxy_payload.get("items") or [])
            normalized_proxy = _normalize_inventory_records(proxy_records, source_url=source_url)
            diagnostics.append(f"proxy_normalized_records={len(normalized_proxy)}")
        except Exception as exc:
            diagnostics.append(f"proxy_error={exc}")

        try:
            browser_payload = _fetch_live_inventory_records_via_browser_html(
                source_url=source_url,
                timeout_seconds=timeout_seconds,
            )
            diagnostics.extend(list(browser_payload.get("diagnostics") or []))
            browser_records = list(browser_payload.get("items") or [])
            normalized_browser = _normalize_inventory_records(browser_records, source_url=source_url)
            diagnostics.append(f"browser_normalized_records={len(normalized_browser)}")
        except Exception as exc:
            diagnostics.append(f"browser_error={exc}")

    candidates = [
        ("direct", normalized),
        ("proxy", normalized_proxy),
        ("browser", normalized_browser),
    ]
    best_source, best_items = max(candidates, key=lambda row: len(row[1]))
    diagnostics.append(f"selected_source={best_source}")
    diagnostics.append(f"selected_records={len(best_items)}")

    if not best_items and diagnostics and any("http_status=4" in item or "direct_http_error=4" in item for item in diagnostics):
        # Preserve the useful diagnostics but avoid throwing before fallbacks are tried.
        diagnostics.append("no_inventory_records_after_all_fallbacks")

    return {
        "source_url": source_url,
        "fetched_at": fetched_at,
        "items": best_items,
        "items_count": len(best_items),
        "diagnostics": diagnostics,
    }


def _persist_inventory_live_cache(
    *,
    source_url: str,
    fetched_at: str,
    items: list[dict[str, Any]],
    diagnostics: list[str],
) -> None:
    _safe_write_json(INVENTORY_LIVE_CACHE_PATH, items)
    _safe_write_json(INVENTORY_LIVE_BACKUP_PATH, items)
    _safe_write_json(
        INVENTORY_LIVE_META_PATH,
        {
            "source_url": source_url,
            "fetched_at": fetched_at,
            "items_count": len(items),
            "diagnostics": diagnostics,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        },
    )
    _safe_write_json(
        INVENTORY_LIVE_META_BACKUP_PATH,
        {
            "source_url": source_url,
            "fetched_at": fetched_at,
            "items_count": len(items),
            "diagnostics": diagnostics,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        },
    )


def _prime_inventory_asset_summaries(items: list[dict[str, Any]]) -> dict[str, Any]:
    """Seed reusable per-VIN asset/CARFAX summaries during inventory ingestion.

    This keeps the UI from depending on a manual asset refresh for every car.
    Full CARFAX report detail still upgrades this cache when the report is
    reachable through the normal asset pipeline.
    """
    stats = {
        "primed": 0,
        "skipped": 0,
        "with_carfax": 0,
        "with_report_facts": 0,
        "with_report_fetch": 0,
        "with_discovered_link": 0,
        "errors": 0,
    }
    try:
        report_fetch_limit = int(os.getenv("CARFAX_PRIME_REPORT_FETCH_LIMIT", "0") or "0")
    except ValueError:
        report_fetch_limit = 0
    report_fetch_all = report_fetch_limit <= 0
    if report_fetch_limit < 0:
        report_fetch_all = True
        report_fetch_limit = 10**9
    try:
        detail_discovery_limit = int(os.getenv("CARFAX_PRIME_DETAIL_DISCOVERY_LIMIT", "120") or "120")
    except ValueError:
        detail_discovery_limit = 120
    if detail_discovery_limit < 0:
        detail_discovery_limit = 10**9
    if detail_discovery_limit > 5000:
        detail_discovery_limit = 5000
    now = datetime.now(timezone.utc).isoformat()
    for item in items:
        if not isinstance(item, dict):
            continue
        clean_vin = str(item.get("vin") or "").strip().upper()
        if not clean_vin or clean_vin == "UNKNOWN":
            stats["skipped"] += 1
            continue
        try:
            cache_path = _vehicle_assets_cache_path(clean_vin)
            existing = _safe_read_json(cache_path, {})
            if not isinstance(existing, dict):
                existing = {}

            photos = _dedupe_urls(
                [
                    url
                    for url in (_extract_photo_url(entry) for entry in (item.get("photos") or []))
                    if url
                ]
            )
            existing_photos = existing.get("photos") if isinstance(existing.get("photos"), list) else []
            final_photos = _dedupe_urls(
                [
                    url
                    for url in (
                        _extract_photo_url(entry)
                        for entry in list(existing_photos or []) + list(photos or [])
                    )
                    if url
                ]
            )

            carfax_facts = existing.get("carfax_facts") if isinstance(existing.get("carfax_facts"), dict) else {}
            item_carfax_facts = item.get("carfax_facts") if isinstance(item.get("carfax_facts"), dict) else {}
            if item_carfax_facts:
                carfax_facts = {**item_carfax_facts, **carfax_facts}

            payload: dict[str, Any] = {
                **existing,
                "vin": clean_vin,
                "detail_url": existing.get("detail_url") or item.get("detail_url"),
                "photos": final_photos,
                "photos_count": len(final_photos),
                "main_photo": existing.get("main_photo") or (final_photos[0] if final_photos else None),
                "sticker_url": existing.get("sticker_url") or item.get("sticker_url"),
                "carfax_url": existing.get("carfax_url") or item.get("carfax_url"),
                "carfax_facts": carfax_facts,
                "loaded_at": existing.get("loaded_at") or now,
                "last_ingested_at": now,
                "source_mode": existing.get("source_mode") or "inventory_ingestion",
            }

            if isinstance(existing.get("carfax_report"), dict) and existing["carfax_report"].get("ok"):
                stats["with_report_facts"] += 1
            if payload.get("carfax_url") or payload.get("carfax_facts"):
                stats["with_carfax"] += 1

            if (
                not payload.get("carfax_url")
                and payload.get("detail_url")
                and stats["with_discovered_link"] < detail_discovery_limit
            ):
                detail_url = str(payload.get("detail_url") or "")
                discovered_url, discovered_facts, discovery_meta = _discover_carfax_from_detail_page(
                    detail_url,
                    timeout_seconds=10.0,
                )
                if discovered_url:
                    payload["carfax_url"] = discovered_url
                    stats["with_discovered_link"] += 1
                if discovered_facts:
                    merged_facts = dict(payload.get("carfax_facts") if isinstance(payload.get("carfax_facts"), dict) else {})
                    merged_facts.update(discovered_facts)
                    payload["carfax_facts"] = {
                        **merged_facts,
                    }
                if discovered_url or discovered_facts or discovery_meta:
                    payload["carfax_discovery"] = {
                        **discovery_meta,
                        "detail_url": detail_url,
                    }
            elif (
                not payload.get("carfax_url")
                and payload.get("detail_url")
                and stats["with_discovered_link"] >= detail_discovery_limit
            ):
                payload["carfax_discovery"] = {
                    "method": "detail_discovery_skipped",
                    "reason": "detail_discovery_limit_reached",
                }

            if (
                payload.get("carfax_url")
                and (
                    report_fetch_all
                    or not (isinstance(payload.get("carfax_report"), dict) and payload["carfax_report"].get("ok"))
                    or _needs_carfax_browser_fallback(payload.get("carfax_report"))
                )
                and (report_fetch_all or stats["with_report_fetch"] < report_fetch_limit)
            ):
                report_details = _fetch_carfax_report_details(
                    str(payload.get("carfax_url") or ""),
                    vehicle=item,
                    expected_vin=clean_vin,
                )
                payload["carfax_report"] = report_details
                payload["carfax_facts"] = _merge_carfax_facts(
                    payload.get("carfax_facts") if isinstance(payload.get("carfax_facts"), dict) else {},
                    report_details,
                )
                if isinstance(payload.get("carfax_report"), dict):
                    stats["with_report_fetch"] += 1
                    if payload["carfax_report"].get("ok"):
                        stats["with_report_facts"] += 1

            payload.update(_vehicle_asset_summary_payload(vehicle=item, assets=payload))
            _safe_write_json(cache_path, payload)
            if isinstance(payload.get("carfax_summary"), dict):
                CARFAX_SUMMARY_DIR.mkdir(parents=True, exist_ok=True)
                _safe_write_json(CARFAX_SUMMARY_DIR / f"{clean_vin}.json", payload["carfax_summary"])
            stats["primed"] += 1
        except Exception:
            stats["errors"] += 1
    return stats


def _inventory_source_status() -> dict[str, Any]:
    live_payload = _read_inventory_blob_with_backup(INVENTORY_LIVE_CACHE_PATH, INVENTORY_LIVE_BACKUP_PATH)
    snapshot_payload = _safe_read_json(INVENTORY_SNAPSHOT_PATH, [])
    meta = _safe_read_json(INVENTORY_LIVE_META_PATH, {})
    if not isinstance(meta, dict) or not meta.get("items_count"):
        backup_meta = _safe_read_json(INVENTORY_LIVE_META_BACKUP_PATH, {})
        if isinstance(backup_meta, dict):
            meta = backup_meta

    live_items = _normalize_inventory_blob(live_payload)
    snapshot_items = _normalize_inventory_blob(snapshot_payload)
    live_counts = _inventory_count_summary(live_items)
    snapshot_counts = _inventory_count_summary(snapshot_items)

    source = "runtime_posts"
    if live_items:
        source = "live_cache"
    elif snapshot_items:
        source = "snapshot"

    return {
        "configured_url": _default_inventory_source_url(),
        "active_source": source,
        "live_cache_path": str(INVENTORY_LIVE_CACHE_PATH),
        "snapshot_path": str(INVENTORY_SNAPSHOT_PATH),
        "live_cache_exists": INVENTORY_LIVE_CACHE_PATH.exists(),
        "snapshot_exists": INVENTORY_SNAPSHOT_PATH.exists(),
        "live_cache_count": len(live_items),
        "live_cache_active_count": live_counts["active"],
        "live_cache_in_transit_count": live_counts["in_transit"],
        "snapshot_count": len(snapshot_items),
        "snapshot_active_count": snapshot_counts["active"],
        "snapshot_in_transit_count": snapshot_counts["in_transit"],
        "last_synced_at": meta.get("fetched_at") if isinstance(meta, dict) else None,
        "last_source_url": meta.get("source_url") if isinstance(meta, dict) else None,
        "last_sync_diagnostics": meta.get("diagnostics") if isinstance(meta, dict) else None,
    }



def _sync_live_inventory(
    *,
    source_url: str | None,
    timeout_seconds: int,
    persist: bool,
) -> dict[str, Any]:
    target_sources = _split_inventory_source_urls(source_url) or _default_inventory_source_urls()
    fetched_payloads: list[dict[str, Any]] = []
    diagnostics: list[str] = []
    errors: list[dict[str, Any]] = []

    for target_source in target_sources:
        try:
            fetched = _fetch_live_inventory_records(
                source_url=target_source,
                timeout_seconds=timeout_seconds,
            )
            # Keep payloads even if zero so diagnostics tell the truth.
            fetched_payloads.append(fetched)
            diagnostics.extend([f"{target_source}: {item}" for item in list(fetched.get("diagnostics") or [])])
        except HTTPException as exc:
            detail = exc.detail if isinstance(exc.detail, dict) else {"message": str(exc.detail)}
            errors.append({"source_url": target_source, "detail": detail})
            diagnostics.append(f"{target_source}: error={detail}")
        except Exception as exc:
            errors.append({"source_url": target_source, "error": str(exc)})
            diagnostics.append(f"{target_source}: error={exc}")

    if not fetched_payloads:
        raise HTTPException(
            status_code=502,
            detail={
                "message": "Failed to fetch every dealership inventory source",
                "source_urls": target_sources,
                "errors": errors,
            },
        )

    merged_items = _merge_inventory_sources(
        [
            item
            for fetched in fetched_payloads
            for item in list(fetched.get("items") or [])
            if isinstance(item, dict)
        ],
        [],
    )
    fetched_at = datetime.now(timezone.utc).isoformat()
    joined_source = ", ".join(target_sources)
    asset_prime = _prime_inventory_asset_summaries(merged_items) if merged_items else {
        "primed": 0,
        "skipped": 0,
        "with_carfax": 0,
        "with_report_facts": 0,
        "with_discovered_link": 0,
        "errors": 0,
    }
    diagnostics.append(
        "asset_summary_prime="
        f"primed:{asset_prime.get('primed', 0)},"
        f"carfax:{asset_prime.get('with_carfax', 0)},"
        f"discovered_links:{asset_prime.get('with_discovered_link', 0)},"
        f"report:{asset_prime.get('with_report_facts', 0)},"
        f"errors:{asset_prime.get('errors', 0)}"
    )

    if persist and merged_items:
        _persist_inventory_live_cache(
            source_url=joined_source,
            fetched_at=fetched_at,
            items=merged_items,
            diagnostics=diagnostics,
        )
    elif persist and not merged_items:
        # Do not overwrite a good existing cache with zero records. Save the
        # failed diagnostics to meta so the admin screen can show why sync failed.
        _safe_write_json(
            INVENTORY_LIVE_META_PATH,
            {
                "source_url": joined_source,
                "fetched_at": fetched_at,
                "items_count": 0,
                "diagnostics": diagnostics,
                "errors": errors,
                "updated_at": datetime.now(timezone.utc).isoformat(),
                "persisted": False,
                "reason": "No website inventory records parsed; existing live cache was preserved.",
            },
        )

    return {
        "ok": bool(merged_items),
        "source_url": joined_source,
        "source_urls": target_sources,
        "fetched_at": fetched_at,
        "items_count": len(merged_items),
        "items": merged_items,
        "persisted": bool(persist and merged_items),
        "asset_summary_prime": asset_prime,
        "diagnostics": diagnostics,
        "errors": errors,
        "source_status": _inventory_source_status(),
    }


def _load_manual_inventory() -> list[dict[str, Any]]:
    payload = _safe_read_json(INVENTORY_MANUAL_PATH, {"items": []})
    if isinstance(payload, dict):
        source_items = payload.get("items", [])
    else:
        source_items = payload
    if not isinstance(source_items, list):
        source_items = []
    return _normalize_inventory_blob(source_items)


def _merge_inventory_sources(
    base_items: list[dict[str, Any]],
    manual_items: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    merged_by_key: dict[str, dict[str, Any]] = {}
    order: list[str] = []

    def _put(item: dict[str, Any]) -> None:
        vin = str(item.get("vin", "")).strip().upper()
        key = vin if vin and vin != "UNKNOWN" else str(item.get("detail_url") or item.get("title") or "")
        key = key.strip().lower()
        if not key:
            return
        if key not in merged_by_key:
            order.append(key)
        merged_by_key[key] = item

    for entry in base_items:
        _put(entry)
    for entry in manual_items:
        _put(entry)

    return [merged_by_key[key] for key in order]


def _load_facebook_post_status() -> dict[str, dict[str, Any]]:
    payload = _safe_read_json(FACEBOOK_POST_STATUS_PATH, {"posts": {}})
    if not isinstance(payload, dict):
        return {}
    posts = payload.get("posts", {})
    if not isinstance(posts, dict):
        return {}
    normalized: dict[str, dict[str, Any]] = {}
    for key, value in posts.items():
        if not isinstance(value, dict):
            continue
        vin = str(key or "").strip().upper()
        if not vin:
            continue
        normalized[vin] = value
    return normalized


def _save_facebook_post_status(posts: dict[str, dict[str, Any]]) -> None:
    _safe_write_json(
        FACEBOOK_POST_STATUS_PATH,
        {
            "posts": posts,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        },
    )


def _facebook_post_status_label(marketplace_status: str) -> str:
    status = str(marketplace_status or "").strip().lower()
    if status == "live":
        return "Live"
    if status == "processing":
        return "Processing"
    if status == "needs_review":
        return "Needs Review"
    if status == "draft":
        return "Draft"
    if status == "failed":
        return "Failed"
    return "Not Posted"


def _set_facebook_vehicle_status(
    *,
    vin: str,
    mode: str,
    marketplace_status: str,
    detail: str | None = None,
    listing_url: str | None = None,
    confirmation: dict[str, Any] | None = None,
) -> dict[str, Any] | None:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        return None
    status = str(marketplace_status or "needs_review").strip().lower()
    if status not in {"live", "processing", "needs_review", "draft", "failed"}:
        status = "needs_review"
    now = datetime.now(timezone.utc).isoformat()
    posts = _load_facebook_post_status()
    existing = dict(posts.get(clean_vin) or {})
    row = {
        **existing,
        "posted": status == "live",
        "posted_status": _facebook_post_status_label(status),
        "marketplace_status": status,
        "mode": mode,
        "detail": detail,
        "listing_url": listing_url or existing.get("listing_url"),
        "last_attempt_at": now,
        "updated_at": now,
    }
    if status == "live":
        row["posted_at"] = existing.get("posted_at") or now
        row["live_verified_at"] = now
    else:
        row["posted_at"] = existing.get("posted_at") if existing.get("posted") else None
    if confirmation:
        row["confirmation"] = confirmation
    posts[clean_vin] = row
    _save_facebook_post_status(posts)
    return row


def _parse_iso_datetime(value: Any) -> datetime | None:
    raw = str(value or "").strip()
    if not raw:
        return None
    try:
        if raw.endswith("Z"):
            raw = raw[:-1] + "+00:00"
        return datetime.fromisoformat(raw)
    except Exception:
        return None


def _sync_marketplace_post_statuses(*, verify_live_urls: bool = True, processing_review_minutes: int = 45) -> dict[str, Any]:
    posts = _load_facebook_post_status()
    now = datetime.now(timezone.utc)
    updates: list[dict[str, Any]] = []
    changed = False
    for clean_vin, row in list(posts.items()):
        if not isinstance(row, dict):
            continue
        status = str(row.get("marketplace_status") or ("live" if row.get("posted") else "not_posted")).lower()
        listing_url = str(row.get("listing_url") or "").strip()
        next_status = status
        reason = ""

        if row.get("posted") and status != "live":
            next_status = "needs_review"
            reason = "posted flag existed without live Marketplace status"
        elif status == "live" and (not listing_url or "/marketplace/item/" not in listing_url.lower()):
            next_status = "needs_review"
            reason = "live status had no verified marketplace item URL"
        elif status == "processing":
            attempted_at = _parse_iso_datetime(row.get("last_attempt_at") or row.get("updated_at"))
            if attempted_at and attempted_at.tzinfo is None:
                attempted_at = attempted_at.replace(tzinfo=timezone.utc)
            age_minutes = ((now - attempted_at).total_seconds() / 60) if attempted_at else processing_review_minutes + 1
            if age_minutes >= processing_review_minutes:
                next_status = "needs_review"
                reason = f"processing for {age_minutes:.0f} minutes without visible listing confirmation"

        if verify_live_urls and status == "live" and listing_url:
            # Facebook may require auth to read the public item page. A blocked
            # response does not prove the listing is gone, so keep Live only when
            # we already have the item URL captured by Marketplace automation.
            row["live_verified_at"] = row.get("live_verified_at") or now.isoformat()

        if next_status != status:
            row["posted"] = next_status == "live"
            row["marketplace_status"] = next_status
            row["posted_status"] = _facebook_post_status_label(next_status)
            row["detail"] = reason or row.get("detail")
            row["updated_at"] = now.isoformat()
            posts[clean_vin] = row
            changed = True
            updates.append({"vin": clean_vin, "status": next_status, "reason": reason})

    if changed:
        _save_facebook_post_status(posts)
    _safe_write_json(
        FACEBOOK_MARKETPLACE_SYNC_PATH,
        {
            "updated_at": now.isoformat(),
            "updates": updates,
            "count": len(posts),
        },
    )
    counts: dict[str, int] = {}
    for row in posts.values():
        if not isinstance(row, dict):
            continue
        key = str(row.get("marketplace_status") or ("live" if row.get("posted") else "not_posted"))
        counts[key] = counts.get(key, 0) + 1
    return {"ok": True, "counts": counts, "updates": updates, "synced_at": now.isoformat()}


def _mark_vehicle_posted(
    *,
    vin: str,
    mode: str,
    status_label: str = "Posted",
    detail: str | None = None,
) -> None:
    _set_facebook_vehicle_status(
        vin=vin,
        mode=mode,
        marketplace_status="live",
        detail=detail or status_label,
    )


def _enrich_inventory_items(items: list[dict[str, Any]]) -> list[dict[str, Any]]:
    status_map = _load_facebook_post_status()
    valuations = _load_jd_power_valuations()
    enriched: list[dict[str, Any]] = []
    for item in items:
        row = dict(item)
        vin = str(row.get("vin", "")).strip().upper()
        state = status_map.get(vin, {})
        posted = bool(state.get("posted"))
        row["posted"] = posted
        row["posted_status"] = state.get("posted_status") or ("Live" if posted else "Not Posted")
        row["marketplace_status"] = state.get("marketplace_status") or ("live" if posted else "not_posted")
        row["listing_url"] = state.get("listing_url")
        row["post_detail"] = state.get("detail")
        row["posted_at"] = state.get("posted_at")
        row["last_post_attempt_at"] = state.get("last_attempt_at")
        valuation = valuations.get(vin)
        jd_trade = _to_float(valuation.get("jd_power_trade_in")) if valuation else None
        row["has_jd_power_trade_in"] = bool(jd_trade and jd_trade > 0)
        row["jd_power_trade_in"] = round(jd_trade, 2) if jd_trade else None
        sale_price = _to_float(row.get("price"))
        if jd_trade and jd_trade > 0 and sale_price is not None:
            pricing = _jd_power_ltv_from_pricing(inventory_price=sale_price, jd_trade_value=jd_trade)
            row["jd_power_ltv"] = pricing["ltv"]
            row["bank_sale_price"] = pricing["bank_sale_price"]
            row["bank_ltv_taxes"] = pricing["taxes"]
            row["bank_ltv_basis"] = pricing["ltv_basis"]
            row["default_bank_fees"] = DEFAULT_BANK_FEES
        enriched.append(row)
    return sorted(
        enriched,
        key=lambda row: (
            0 if row.get("has_jd_power_trade_in") else 1,
            _to_float(row.get("jd_power_ltv")) if row.get("jd_power_ltv") is not None else 9999,
            str(row.get("title") or ""),
        ),
    )


def _find_vehicle_by_vin(vin: str) -> dict[str, Any] | None:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        return None
    return next(
        (
            item
            for item in _load_inventory_candidates()
            if str(item.get("vin", "")).strip().upper() == clean_vin
        ),
        None,
    )


def _vehicle_assets_cache_path(vin: str) -> Path:
    clean_vin = re.sub(r"[^A-Z0-9]", "", str(vin or "").upper()) or "UNKNOWN"
    VEHICLE_ASSETS_CACHE_DIR.mkdir(parents=True, exist_ok=True)
    return VEHICLE_ASSETS_CACHE_DIR / f"{clean_vin}.json"


def _decode_embedded_asset_url(value: Any, *, base_url: str | None = None) -> str:
    candidate = str(value or "").strip()
    if not candidate:
        return ""
    candidate = html.unescape(candidate)
    candidate = candidate.replace("\\/", "/").replace("\\u002F", "/").replace("\\u002f", "/")
    candidate = candidate.replace("\\u0026", "&").replace("\\u0026", "&")
    candidate = candidate.replace("\\x3A", ":").replace("\\x2F", "/")
    candidate = candidate.strip(" '\"")
    if candidate.startswith("//"):
        candidate = f"https:{candidate}"
    elif base_url and candidate.startswith("/"):
        candidate = urljoin(base_url, candidate)
    return candidate


def _extract_asset_links_from_text_candidates(raw_text: str, *, base_url: str | None = None) -> dict[str, str | None]:
    text = str(raw_text or "")
    carfax_patterns = [
        r"https?:\\?/\\?/[^\"'\\s<>]*carfax[^\"'\\s<>]*",
        r"(?:carfaxUrl|carfaxURL|vehicleHistoryUrl|vehicle_history_url|historyUrl)\s*[\"'=:\s]+\s*[\"']([^\"']+)[\"']",
    ]
    sticker_patterns = [
        r"https?:\\?/\\?/[^\"'\\s<>]*(?:sticker|monroney)[^\"'\\s<>]*",
        r"(?:windowStickerUrl|window_sticker_url|stickerUrl|monroneyUrl)\s*[\"'=:\s]+\s*[\"']([^\"']+)[\"']",
    ]

    def _first_match(patterns: list[str]) -> str | None:
        for pattern in patterns:
            for match in re.finditer(pattern, text, flags=re.IGNORECASE):
                groups = match.groups()
                candidate = groups[0] if groups else match.group(0)
                decoded = _decode_embedded_asset_url(candidate, base_url=base_url)
                if decoded.lower().startswith(("http://", "https://")):
                    return decoded
        return None

    return {
        "sticker_url": _first_match(sticker_patterns),
        "carfax_url": _first_match(carfax_patterns),
    }


def _extract_asset_links_from_html(html_text: str, *, base_url: str | None = None) -> dict[str, str | None]:
    hrefs = re.findall(
        r"""(?:href|src)\s*=\s*["']([^"']+)["']""",
        html_text,
        flags=re.IGNORECASE,
    )
    absolute_candidates: list[str] = []
    seen: set[str] = set()
    for href in hrefs:
        candidate = html.unescape(href.strip())
        if not candidate:
            continue
        if candidate.startswith("//"):
            candidate = f"https:{candidate}"
        elif base_url and candidate.startswith("/"):
            candidate = urljoin(base_url, candidate)
        lowered = candidate.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        absolute_candidates.append(candidate)

    sticker_url = next(
        (
            url
            for url in absolute_candidates
            if ("sticker" in url.lower() or "monroney" in url.lower())
            and (url.lower().startswith("http://") or url.lower().startswith("https://"))
        ),
        None,
    )
    carfax_url = next(
        (
            url
            for url in absolute_candidates
            if "carfax" in url.lower()
            and (url.lower().startswith("http://") or url.lower().startswith("https://"))
        ),
        None,
    )
    if not sticker_url or not carfax_url:
        supplemental = _extract_asset_links_from_text_candidates(html_text, base_url=base_url)
        sticker_url = sticker_url or supplemental.get("sticker_url")
        carfax_url = carfax_url or supplemental.get("carfax_url")
    return {"sticker_url": sticker_url, "carfax_url": carfax_url}


def _normalize_marketplace_body_style(value: Any) -> str:
    lowered = str(value or "").strip().lower()
    if not lowered:
        return ""
    if any(token in lowered for token in ("minivan", "van")) and "truck" not in lowered:
        return "Minivan" if "minivan" in lowered else "Van"
    if any(token in lowered for token in ("convertible", "roadster", "cabriolet", "spyder", "soft top")):
        return "Convertible"
    if any(token in lowered for token in ("coupe",)):
        return "Coupe"
    if any(token in lowered for token in ("hatchback", "liftback")):
        return "Hatchback"
    if any(token in lowered for token in ("wagon",)):
        return "Wagon"
    if any(token in lowered for token in ("truck", "pickup", "supercrew", "crew cab", "super cab", "quad cab", "king cab", "double cab")):
        return "Truck"
    if any(token in lowered for token in ("suv", "sport utility", "utility", "crossover")):
        return "SUV"
    if any(token in lowered for token in ("sedan",)):
        return "Sedan"
    return str(value or "").strip()


def _normalize_marketplace_fuel_type(value: Any) -> str:
    lowered = str(value or "").strip().lower()
    if not lowered:
        return ""
    if any(token in lowered for token in ("electric", "bev", "ev")):
        return "Electric"
    if any(token in lowered for token in ("hybrid", "plug-in", "plug in", "phev", "4xe")):
        return "Hybrid"
    if any(token in lowered for token in ("diesel", "tdi", "duramax", "cummins", "power stroke", "ecodiesel")):
        return "Diesel"
    if any(token in lowered for token in ("gasoline", "regular unleaded", "premium unleaded", "unleaded", "flex fuel")):
        return "Gasoline"
    return str(value or "").strip()


def _extract_ddc_state_payloads(html_text: str, bucket: str) -> list[dict[str, Any]]:
    pattern = rf"DDC\.WS\.state\['{re.escape(bucket)}'\]\['[^']+'\]\s*=\s*(\{{.*?\}}\s*);"
    payloads: list[dict[str, Any]] = []
    for match in re.finditer(pattern, html_text, flags=re.IGNORECASE | re.DOTALL):
        raw = str(match.group(1) or "").strip()
        if raw.endswith(";"):
            raw = raw[:-1].strip()
        try:
            parsed = json.loads(raw)
        except Exception:
            continue
        if isinstance(parsed, dict):
            payloads.append(parsed)
    return payloads


def _extract_quick_specs_from_html(html_text: str) -> dict[str, Any]:
    candidates = _extract_ddc_state_payloads(html_text, "ws-quick-specs")
    candidates.extend(_extract_ddc_state_payloads(html_text, "ws-vehicle-ctas"))
    merged: dict[str, Any] = {}
    for payload in candidates:
        quick_specs = payload.get("quickSpecs") if isinstance(payload.get("quickSpecs"), dict) else {}
        vehicle = payload.get("vehicle") if isinstance(payload.get("vehicle"), dict) else {}

        def _pick(*values: Any) -> str:
            for value in values:
                if isinstance(value, dict):
                    inner = value.get("en_US") or value.get("value")
                    if inner not in (None, ""):
                        return str(inner).strip()
                elif value not in (None, ""):
                    return str(value).strip()
            return ""

        spec = {
            "year": _pick(vehicle.get("year")),
            "make": _pick(vehicle.get("make")),
            "model": _pick(vehicle.get("model"), quick_specs.get("model")),
            "body_style": _pick(quick_specs.get("bodyStyle"), vehicle.get("bodyStyle"), vehicle.get("normalBodyStyle")),
            "marketplace_body_style": _normalize_marketplace_body_style(
                _pick(quick_specs.get("bodyStyle"), vehicle.get("bodyStyle"), vehicle.get("normalBodyStyle"))
            ),
            "fuel_type": _pick(quick_specs.get("fuelType"), vehicle.get("fuelType"), vehicle.get("normalFuelType")),
            "marketplace_fuel_type": _normalize_marketplace_fuel_type(
                _pick(quick_specs.get("fuelType"), vehicle.get("fuelType"), vehicle.get("normalFuelType"))
            ),
            "drivetrain": _pick(quick_specs.get("driveLine"), vehicle.get("driveLine"), vehicle.get("normalDriveLine")),
            "transmission": _pick(quick_specs.get("transmission"), vehicle.get("transmission"), vehicle.get("normalTransmission")),
            "engine": _pick(quick_specs.get("engine"), vehicle.get("engine")),
            "engine_size": _pick(vehicle.get("engineSize")),
            "exterior": _pick(quick_specs.get("exteriorColor"), vehicle.get("exteriorColor"), quick_specs.get("normalExteriorColor")),
            "interior": _pick(quick_specs.get("interiorColor"), vehicle.get("interiorColor"), quick_specs.get("normalInteriorColor")),
            "stock_number": _pick(quick_specs.get("stockNumber"), vehicle.get("stockNumber")),
        }
        city = quick_specs.get("cityFuelEconomy") or vehicle.get("cityFuelEconomy")
        highway = quick_specs.get("highwayFuelEconomy") or vehicle.get("highwayFuelEconomy")
        combined = quick_specs.get("combinedFuelEfficiency") or vehicle.get("combinedFuelEfficiency")
        if isinstance(city, dict):
            spec["mpg_city"] = city.get("imperial") or city.get("value")
        elif city not in (None, ""):
            spec["mpg_city"] = city
        if isinstance(highway, dict):
            spec["mpg_hwy"] = highway.get("imperial") or highway.get("value")
        elif highway not in (None, ""):
            spec["mpg_hwy"] = highway
        if isinstance(combined, dict):
            spec["mpg_combined"] = combined.get("imperial") or combined.get("value")
        elif combined not in (None, ""):
            spec["mpg_combined"] = combined
        for key, value in spec.items():
            if value not in (None, "", [], {}):
                merged[key] = value
    return merged


def _quick_spec_highlights(quick_specs: dict[str, Any]) -> list[str]:
    if not isinstance(quick_specs, dict):
        return []
    highlights: list[str] = []
    body_style = str(quick_specs.get("body_style") or "").strip()
    if body_style:
        highlights.append(f"Body: {body_style}")
    drivetrain = str(quick_specs.get("drivetrain") or "").strip()
    if drivetrain:
        highlights.append(f"Drive: {drivetrain}")
    fuel_type = str(quick_specs.get("fuel_type") or "").strip()
    if fuel_type:
        highlights.append(f"Fuel: {fuel_type}")
    mpg_city = quick_specs.get("mpg_city")
    mpg_hwy = quick_specs.get("mpg_hwy")
    mpg_combined = quick_specs.get("mpg_combined")
    if mpg_city not in (None, "") and mpg_hwy not in (None, ""):
        highlights.append(f"EPA: {mpg_city} city / {mpg_hwy} hwy")
    elif mpg_combined not in (None, ""):
        highlights.append(f"EPA combined: {mpg_combined}")
    return _dedupe_strings(highlights)


def _extract_carfax_facts_from_html(html_text: str) -> dict[str, Any]:
    badge_chunks: list[str] = []
    for match in re.finditer(r"<img\b[^>]*(?:carfax|valuebadge)[^>]*>", html_text, flags=re.IGNORECASE):
        tag = match.group(0)
        for attr in ("alt", "title", "src"):
            attr_match = re.search(rf"""{attr}\s*=\s*["']([^"']+)["']""", tag, flags=re.IGNORECASE)
            if attr_match:
                badge_chunks.append(html.unescape(attr_match.group(1)))

    compact = re.sub(r"\s+", " ", " ".join(badge_chunks)).strip()
    lowered = compact.lower()
    if "carfax" not in lowered and "1own" not in lowered and "valuebadge" not in lowered:
        return {}

    facts: dict[str, Any] = {
        "source": "dealer_page_carfax_badge",
        "badge_text": compact,
        "owner_count": None,
        "accident_damage": None,
        "title_brand": None,
        "service_history": None,
        "usage": None,
        "value_badge": None,
        "report_access": "CARFAX link detected from the dealer listing; Xconsole will fetch and cache full report text when CARFAX exposes it to the automated browser.",
    }
    highlights: list[str] = []

    if re.search(r"\b1[- ]?owner\b|1own", lowered):
        facts["owner_count"] = "1 owner"
        highlights.append("CARFAX badge shows 1 owner.")
    elif owner_match := re.search(r"\b([0-9]+)[- ]?owner\b", lowered):
        facts["owner_count"] = f"{owner_match.group(1)} owners"
        highlights.append(f"CARFAX badge shows {facts['owner_count']}.")

    if "great" in lowered and "value" in lowered:
        facts["value_badge"] = "Great Value"
        highlights.append("CARFAX badge shows Great Value.")
    elif "good" in lowered and "value" in lowered:
        facts["value_badge"] = "Good Value"
        highlights.append("CARFAX badge shows Good Value.")

    if re.search(r"no[-_ ]?accidents?|accident[-_ ]?free", lowered):
        facts["accident_damage"] = "No accidents reported by CARFAX badge."
        highlights.append("CARFAX badge indicates no accidents.")

    facts["highlights"] = highlights or [compact]
    return facts


def _extract_carfax_link_from_inventory_item(item: dict[str, Any]) -> tuple[str | None, dict[str, Any]]:
    callouts = item.get("callout") if isinstance(item.get("callout"), list) else []
    chunks: list[str] = []
    carfax_url: str | None = None
    for callout in callouts:
        if not isinstance(callout, dict):
            continue
        href = str(callout.get("href") or "").strip()
        image_src = str(callout.get("imageSrc") or callout.get("image") or "").strip()
        image_alt = str(callout.get("imageAlt") or callout.get("alt") or "").strip()
        image_title = str(callout.get("imageTitle") or callout.get("title") or "").strip()
        blob = " ".join([href, image_src, image_alt, image_title]).lower()
        if "carfax" not in blob and "valuebadge" not in blob and "1own" not in blob:
            continue
        if href and not carfax_url:
            carfax_url = href
        chunks.extend(part for part in [image_alt, image_title, image_src, href] if part)

    if not chunks:
        return carfax_url, {}

    html_bits = []
    if carfax_url:
        html_bits.append(f'<a href="{html.escape(carfax_url)}">CARFAX</a>')
    for chunk in chunks:
        if re.search(r"\.(?:svg|png|jpe?g|webp)(?:[?#].*)?$", chunk, flags=re.IGNORECASE):
            html_bits.append(f'<img src="{html.escape(chunk)}" alt="{html.escape(" ".join(chunks))}">')
        else:
            html_bits.append(f'<span>{html.escape(chunk)}</span>')
    facts = _extract_carfax_facts_from_html(" ".join(html_bits))
    if facts and carfax_url:
        facts["source_url"] = carfax_url
    return carfax_url, facts


def _html_to_visible_text(raw: str) -> str:
    text = html.unescape(str(raw or ""))
    if BeautifulSoup is not None:
        try:
            soup = BeautifulSoup(text, "html.parser")
            for tag in soup(["script", "style", "noscript", "svg"]):
                tag.decompose()
            text = soup.get_text("\n")
        except Exception:
            text = re.sub(r"<[^>]+>", "\n", text)
    else:
        text = re.sub(r"<[^>]+>", "\n", text)
    text = text.replace("\xa0", " ")
    text = re.sub(r"[ \t\r\f\v]+", " ", text)
    text = re.sub(r"\n{2,}", "\n", text)
    return text.strip()


def _carfax_visible_lines(raw: str) -> list[str]:
    text = _html_to_visible_text(raw)
    lines = [re.sub(r"\s+", " ", line).strip(" -\t") for line in text.splitlines()]
    return [line for line in lines if line]


def _looks_like_carfax_block(raw: str) -> bool:
    lowered = str(raw or "").lower()
    blocked_markers = [
        "please enable js",
        "disable any ad blocker",
        "captcha-delivery",
        "datadome",
        "verify you are human",
        "access denied",
    ]
    return any(marker in lowered for marker in blocked_markers)


def _extract_text_from_carfax_response(response: Any, *, fallback_url: str = "") -> tuple[str, dict[str, Any]]:
    text_chunks: list[str] = []
    diagnostics: dict[str, Any] = {
        "source_content_type": "",
        "extract_method": "http_text",
        "extract_meta": {},
    }
    try:
        content_type = str(response.headers.get("content-type", "") or "").lower()
    except Exception:
        content_type = ""
    diagnostics["source_content_type"] = content_type

    if not response:
        return "", diagnostics

    body = response.text or ""
    raw_body = response.content
    fallback_source = str(response.url if hasattr(response, "url") else fallback_url) if (response and hasattr(response, "url")) else fallback_url

    if "application/pdf" in content_type or (str(fallback_source).lower().endswith(".pdf")) or str(raw_body[:4] or b"") == "b'%PDF'":
        pdf_text, pdf_meta = _extract_pdf_upload_text(raw_body)
        if pdf_text:
            diagnostics["extract_method"] = "pdf_ocr_text"
            diagnostics["extract_meta"] = {"pdf_chars": len(pdf_text), **pdf_meta}
            return pdf_text, diagnostics
        diagnostics["extract_meta"] = {"pdf_chars": 0, **pdf_meta}

    if content_type.startswith("image/") or re.search(r"\.(?:png|jpe?g|webp)$", str(fallback_source).lower()):
        image_text, image_meta = _extract_image_upload_text(raw_body)
        if image_text:
            diagnostics["extract_method"] = "image_ocr_text"
            diagnostics["extract_meta"] = {**image_meta, "image_chars": len(image_text)}
            return image_text, diagnostics
        diagnostics["extract_meta"] = {**image_meta}

    if body:
        visible = _html_to_visible_text(body)
        if visible:
            diagnostics["extract_method"] = "visible_text"
            diagnostics["extract_meta"] = {"visible_chars": len(visible)}
            return visible, diagnostics

    text_chunks.append(str(body or "").strip())
    return "\n".join(part for part in text_chunks if part), diagnostics


_CARFAX_DATE_PATTERN = re.compile(
    r"\b(?:Jan(?:uary)?|Feb(?:ruary)?|Mar(?:ch)?|Apr(?:il)?|May|Jun(?:e)?|Jul(?:y)?|"
    r"Aug(?:ust)?|Sep(?:t(?:ember)?)?|Oct(?:ober)?|Nov(?:ember)?|Dec(?:ember)?)\s+"
    r"\d{1,2},\s+\d{4}\b|\b\d{1,2}/\d{1,2}/\d{2,4}\b",
    flags=re.IGNORECASE,
)


def _sentence_window(lines: list[str], index: int, *, radius: int = 3) -> str:
    start = max(0, index - 1)
    end = min(len(lines), index + radius + 1)
    return re.sub(r"\s+", " ", " ".join(lines[start:end])).strip()


def _classify_accident_severity(text: str) -> str:
    lowered = text.lower()
    airbag_major = ("airbag deployed" in lowered or "airbag deployment" in lowered) and not any(
        marker in lowered for marker in ["airbags did not deploy", "airbag did not deploy", "not deployed"]
    )
    if any(marker in lowered for marker in ["severe", "major", "structural", "frame damage", "total loss"]) or airbag_major:
        return "major"
    if any(marker in lowered for marker in ["moderate", "functional damage"]):
        return "moderate"
    if any(marker in lowered for marker in ["minor", "cosmetic"]):
        return "minor"
    if "damage" in lowered:
        return "damage"
    return "reported"


def _short_fact(text: str, *, limit: int = 180) -> str:
    clean = re.sub(r"\s+", " ", str(text or "")).strip(" -")
    return clean if len(clean) <= limit else f"{clean[: limit - 1].rstrip()}..."


def _carfax_is_generic_heading(line: str) -> bool:
    compact = re.sub(r"\s+", " ", str(line or "")).strip(" -").lower()
    if not compact:
        return True
    generic = {
        "accident",
        "accident / damage",
        "accident / damage history",
        "event 1",
        "event 2",
        "event 3",
        "event 4",
        "total loss",
        "structural damage",
        "airbag deployment",
        "service history",
        "additional history",
        "ownership history",
        "title history",
        "detailed history",
        "vehicle details",
        "report",
        "check availability",
    }
    return compact in generic


def _is_carfax_explanatory_accident_line(text: str) -> bool:
    compact = " ".join(str(text or "").split()).strip()
    lowered = compact.lower()
    if not lowered:
        return False
    blocked_phrases = (
        "according to the national safety council",
        "injury facts, 2021 edition",
        "carfax recommends getting a pre-purchase inspection",
        "carfax recommends that you have this vehicle inspected",
        "learn more about this",
    )
    return any(phrase in lowered for phrase in blocked_phrases)


def _clean_carfax_accident_description(text: str) -> str:
    compact = " ".join(str(text or "").split()).strip()
    if not compact:
        return ""
    compact = re.sub(r"More\s*information\b", "", compact, flags=re.IGNORECASE).strip()
    compact = re.sub(r"\s{2,}", " ", compact).strip()
    return compact


def _carfax_positive_accident_signal_count(*parts: Any) -> int:
    combined = " ".join(str(part or "") for part in parts if str(part or "").strip())
    lowered = combined.lower()
    if not lowered:
        return 0
    lowered = lowered.replace("no accidents reported", "")
    lowered = lowered.replace("no accidents or damage reported", "")
    patterns = (
        r"\baccident reported\b",
        r"\bminor damage\b",
        r"\bmoderate damage\b",
        r"\bmajor damage\b",
        r"\bdamage reported\b",
        r"\bcollision damage\b",
        r"\bstructural damage\b",
        r"\bairbag deployed\b",
    )
    return sum(len(re.findall(pattern, lowered, flags=re.IGNORECASE)) for pattern in patterns)


def _normalize_carfax_parsed_payload(report: dict[str, Any]) -> dict[str, Any]:
    if not isinstance(report, dict):
        return report

    normalized = dict(report)
    accident_events_raw = normalized.get("accident_events")
    accident_events = accident_events_raw if isinstance(accident_events_raw, list) else []
    cleaned_accident_events: list[dict[str, str]] = []

    for item in accident_events:
        if not isinstance(item, dict):
            continue
        date_text = " ".join(str(item.get("date") or "").split()).strip() or "date not parsed"
        severity = " ".join(str(item.get("severity") or "").split()).strip().lower() or "reported"
        description = _clean_carfax_accident_description(str(item.get("description") or ""))
        lowered = description.lower()
        if not description or _is_carfax_explanatory_accident_line(description):
            continue
        if description.lower() == "accident reported":
            continue
        if "accident reported:" in lowered:
            extracted_date = re.sub(r"^accident reported:\s*", "", description, flags=re.IGNORECASE).strip().rstrip(".")
            if extracted_date and re.fullmatch(r"\d{1,2}/\d{1,2}/\d{2,4}", extracted_date):
                continue
        cleaned_accident_events.append(
            {
                "date": date_text,
                "severity": severity,
                "description": description,
            }
        )

    deduped_accident_events: list[dict[str, str]] = []
    seen_signatures: set[tuple[str, str]] = set()
    dated_descriptions = {
        re.sub(r"[^a-z0-9]+", " ", str(item.get("description") or "").lower()).strip()
        for item in cleaned_accident_events
        if str(item.get("date") or "").strip().lower() != "date not parsed"
    }
    for item in cleaned_accident_events:
        normalized_desc = re.sub(r"[^a-z0-9]+", " ", str(item.get("description") or "").lower()).strip()
        date_text = str(item.get("date") or "").strip()
        if date_text.lower() == "date not parsed" and normalized_desc in dated_descriptions:
            continue
        signature = (date_text.lower(), normalized_desc)
        if signature in seen_signatures:
            continue
        seen_signatures.add(signature)
        deduped_accident_events.append(item)

    if not deduped_accident_events:
        highlights_raw = normalized.get("highlights")
        highlight_lines = highlights_raw if isinstance(highlights_raw, list) else []
        recovered_events: list[dict[str, str]] = []
        for line in highlight_lines:
            text = _clean_carfax_accident_description(str(line or ""))
            compact = " ".join(text.split()).strip()
            match = re.match(
                r"^(?P<date>\d{1,2}/\d{1,2}/\d{2,4}|date not parsed):\s*(?P<severity>minor|moderate|major|damage|reported)\s*-\s*(?P<desc>.+)$",
                compact,
                flags=re.IGNORECASE,
            )
            if not match:
                continue
            description = _clean_carfax_accident_description(match.group("desc"))
            lowered_desc = description.lower()
            if not description or _is_carfax_explanatory_accident_line(description):
                continue
            if "accident reported:" in lowered_desc:
                extracted_date = re.sub(r"^accident reported:\s*", "", description, flags=re.IGNORECASE).strip().rstrip(".")
                if extracted_date and re.fullmatch(r"\d{1,2}/\d{1,2}/\d{2,4}", extracted_date):
                    continue
            recovered_events.append(
                {
                    "date": match.group("date"),
                    "severity": match.group("severity").lower(),
                    "description": description,
                }
            )
        if recovered_events:
            deduped_accident_events = recovered_events

    normalized["accident_events"] = deduped_accident_events
    existing_confidence = str(normalized.get("accident_confidence") or "").strip().lower()
    highlights_raw = normalized.get("highlights")
    highlight_lines = highlights_raw if isinstance(highlights_raw, list) else []
    positive_signals = _carfax_positive_accident_signal_count(
        normalized.get("summary"),
        normalized.get("accident_damage"),
        *highlight_lines,
    )
    explicit_no_accident = any(
        phrase in " ".join(str(item or "") for item in highlight_lines).lower()
        for phrase in ("no accidents reported", "no accidents or damage reported")
    ) or any(
        phrase in str(normalized.get(key) or "").lower()
        for key in ("summary", "accident_damage", "badge_text")
        for phrase in ("no accidents reported", "no accidents or damage reported")
    )
    if deduped_accident_events:
        severity_counts: dict[str, int] = {}
        for event in deduped_accident_events:
            severity_name = str(event.get("severity") or "reported")
            severity_counts[severity_name] = severity_counts.get(severity_name, 0) + 1
        order = ("minor", "moderate", "major", "damage", "reported")
        ordered_parts = [f"{severity_counts[key]} {key}" for key in order if severity_counts.get(key)]
        normalized["accident_damage"] = f"{len(deduped_accident_events)} accident/damage event(s) parsed: {', '.join(ordered_parts)}."
        normalized["accident_counts"] = {
            "total": len(deduped_accident_events),
            "minor": severity_counts.get("minor", 0),
            "moderate": severity_counts.get("moderate", 0),
            "major": severity_counts.get("major", 0),
            "damage": severity_counts.get("damage", 0),
        }
        normalized["accident_confidence"] = "incidents_found"
    else:
        normalized["accident_counts"] = {"total": 0, "minor": 0, "moderate": 0, "major": 0, "damage": 0}
        if explicit_no_accident and positive_signals:
            accident_confidence = "conflict"
        elif explicit_no_accident:
            accident_confidence = "explicit_clear"
        elif positive_signals:
            accident_confidence = "signal_without_timeline"
        elif existing_confidence in {"conflict", "signal_without_timeline", "unknown"}:
            accident_confidence = existing_confidence
        else:
            accident_confidence = "unknown"
        normalized["accident_confidence"] = accident_confidence
        if accident_confidence == "explicit_clear":
            normalized["accident_damage"] = "No accidents reported in parsed CARFAX text."
        else:
            normalized["accident_damage"] = "Accident history needs manual confirmation from CARFAX."

    highlights = highlight_lines
    rebuilt_highlights: list[str] = []
    owner_seen = False
    service_seen = False
    market_seen = False
    value_seen = False
    title_seen = False
    use_seen = False
    last_owned_seen = False

    for line in highlights:
        text = _clean_carfax_accident_description(str(line or ""))
        lowered = text.lower()
        if not text or _is_carfax_explanatory_accident_line(text):
            continue
        if re.match(r"^(?:\d{1,2}/\d{1,2}/\d{2,4}|date not parsed):", lowered):
            if " - " in text:
                continue
            accident_desc = re.sub(r"^(?:\d{1,2}/\d{1,2}/\d{2,4}|date not parsed):\s*[a-z]+\s*-\s*", "", lowered).strip()
            if "accident reported:" in accident_desc:
                extracted_date = re.sub(r"^accident reported:\s*", "", accident_desc, flags=re.IGNORECASE).strip().rstrip(".")
                if extracted_date and re.fullmatch(r"\d{1,2}/\d{1,2}/\d{2,4}", extracted_date):
                    continue
        if lowered.startswith("owners:"):
            if owner_seen:
                continue
            owner_seen = True
            rebuilt_highlights.append(text)
            continue
        if (
            "accident/damage event(s) parsed" in lowered
            or lowered.startswith("accident/damage:")
            or lowered.startswith("no accidents reported in parsed carfax text")
            or lowered.startswith("accident history needs manual confirmation from carfax")
        ):
            continue
        if lowered.startswith("service history:"):
            if service_seen:
                continue
            service_seen = True
            rebuilt_highlights.append(text)
            continue
        if lowered.startswith("market/value:"):
            if market_seen:
                continue
            market_seen = True
            rebuilt_highlights.append(text)
            continue
        if lowered.startswith("carfax value:"):
            if value_seen:
                continue
            value_seen = True
            rebuilt_highlights.append(text)
            continue
        if "title brand/issues" in lowered or lowered.startswith("title/brand:"):
            if title_seen:
                continue
            title_seen = True
            rebuilt_highlights.append(text)
            continue
        if lowered.startswith("use type:"):
            if use_seen:
                continue
            use_seen = True
            rebuilt_highlights.append(text)
            continue
        if lowered.startswith("last owned in:"):
            if last_owned_seen:
                continue
            last_owned_seen = True
            rebuilt_highlights.append(text)
            continue
        rebuilt_highlights.append(text)

    accident_timeline = [
        f"{item['date']}: {item['severity']} - {item['description']}"
        for item in deduped_accident_events[:5]
    ]
    final_highlights: list[str] = []
    accident_summary_inserted = False
    accident_timeline_inserted = False
    for line in rebuilt_highlights:
        lowered = line.lower()
        final_highlights.append(line)
        if lowered.startswith("owners:") and not accident_summary_inserted:
            final_highlights.append(str(normalized["accident_damage"]))
            accident_summary_inserted = True
            if accident_timeline and not accident_timeline_inserted:
                final_highlights.extend(accident_timeline)
                accident_timeline_inserted = True
    if not accident_summary_inserted:
        final_highlights.insert(0, str(normalized["accident_damage"]))
        if accident_timeline:
            final_highlights[1:1] = accident_timeline
    normalized["highlights"] = _dedupe_strings(final_highlights)

    summary_parts = []
    owner_count = str(normalized.get("owner_count") or "").strip()
    accident_damage = str(normalized.get("accident_damage") or "").strip()
    service_history = str(normalized.get("service_history") or "").strip()
    joined_value = " / ".join(
        part
        for part in [
            str(normalized.get("value_badge") or "").strip(),
            str(normalized.get("market_position") or "").strip(),
            str(normalized.get("market_delta") or "").strip(),
        ]
        if part
    )
    if owner_count:
        summary_parts.append(owner_count)
    if accident_damage:
        summary_parts.append(accident_damage)
    if service_history:
        summary_parts.append(service_history)
    if joined_value:
        summary_parts.append(joined_value)
    if summary_parts:
        normalized["summary"] = "CARFAX report parsed: " + "; ".join(summary_parts) + "."

    return normalized


def _parse_carfax_event_date(value: str) -> datetime | None:
    raw = str(value or "").strip()
    if not raw or raw == "date not parsed":
        return None
    for fmt in ("%m/%d/%Y", "%m/%d/%y", "%B %d, %Y", "%b %d, %Y"):
        try:
            return datetime.strptime(raw, fmt)
        except ValueError:
            continue
    return None


def _extract_carfax_report_facts(raw_text_or_html: str, *, source_url: str | None = None, source: str = "carfax_report") -> dict[str, Any]:
    raw = str(raw_text_or_html or "")
    if _looks_like_carfax_block(raw):
        return {
            "ok": False,
            "source": "carfax_report_blocked",
            "blocked": True,
            "report_access": "CARFAX returned browser verification. Xconsole cached the linked report and badge facts, but full accident, service, title, and market-value detail was not exposed to the automated fetch.",
            "source_url": source_url,
            "highlights": [
                "Linked CARFAX detail was blocked by browser verification.",
                "Badge facts from the dealer listing remain available.",
            ],
        }

    lines = _carfax_visible_lines(raw)
    full_text = re.sub(r"\s+", " ", " ".join(lines)).strip()
    lowered = full_text.lower()
    report_vin = _carfax_extract_vin(full_text)
    report_vehicle_title = ""
    for line in lines[:10]:
        compact = " ".join(str(line or "").split()).strip()
        if re.search(r"\b(19|20)\d{2}\b", compact) and len(compact) <= 120:
            report_vehicle_title = compact
            break
    if len(full_text) < 120 or "carfax" not in lowered:
        return {
            "ok": False,
            "source": source,
            "blocked": False,
            "report_access": "Xconsole did not find enough CARFAX report text to parse.",
            "source_url": source_url,
            "highlights": ["Linked CARFAX report did not expose parseable report text."],
        }

    owner_count = None
    if owner_match := re.search(r"\b([1-9][0-9]?)\s*(?:previous\s+)?[- ]?\s*owners?\b", full_text, flags=re.IGNORECASE):
        count = int(owner_match.group(1))
        owner_count = "1 owner" if count == 1 else f"{count} owners"
    elif re.search(r"\bone owner\b", lowered):
        owner_count = "1 owner"

    value_badge = None
    carfax_value = None
    market_position = None
    market_delta = None
    listing_price = None
    detail_records_available = None
    last_owned_location = None
    regular_oil_changes = bool(re.search(r"\bregular oil changes\b", lowered))

    value_card_match = re.search(
        r"\b(good value|great value|fair value|high price)\b\s+(\$[\d,]+)\s+(\$[\d,]+)\s+"
        r"(below|under|less than|above|over|higher than)\s+(\$[\d,]+)\s+carfax\s+value\b",
        full_text,
        flags=re.IGNORECASE,
    )
    if value_card_match:
        value_badge = value_card_match.group(1).title()
        listing_price = value_card_match.group(2)
        market_delta = value_card_match.group(3)
        direction = value_card_match.group(4).lower()
        carfax_value = value_card_match.group(5)
        market_position = "below market value" if direction in {"below", "under", "less than"} else "above market value"
    else:
        if "good value" in lowered:
            value_badge = "Good Value"
        elif "great value" in lowered:
            value_badge = "Great Value"
        elif "fair value" in lowered:
            value_badge = "Fair Value"
        elif "high price" in lowered or "overpriced" in lowered:
            value_badge = "High Price"

        if value_match := re.search(r"carfax\s+value[^$]{0,40}(\$[\d,]+)", full_text, flags=re.IGNORECASE):
            carfax_value = value_match.group(1)

        if delta_match := re.search(
            r"(\$[\d,]+)\s+(below|under|less than|above|over|higher than)\s+(?:the\s+)?(?:carfax\s+)?(?:market\s+)?value",
            full_text,
            flags=re.IGNORECASE,
        ):
            market_delta = delta_match.group(1)
            direction = delta_match.group(2).lower()
            market_position = "below market value" if direction in {"below", "under", "less than"} else "above market value"
        elif delta_match := re.search(
            r"(below|under|less than|above|over|higher than)[^$]{0,18}(\$[\d,]+)\s+carfax\s+value",
            full_text,
            flags=re.IGNORECASE,
        ):
            direction = delta_match.group(1).lower()
            market_position = "below market value" if direction in {"below", "under", "less than"} else "above market value"
        elif delta_match := re.search(
            r"(below|under|less than|above|over|higher than)[^$]{0,45}(\$[\d,]+)",
            full_text,
            flags=re.IGNORECASE,
        ):
            direction = delta_match.group(1).lower()
            market_delta = delta_match.group(2)
            market_position = "below market value" if direction in {"below", "under", "less than"} else "above market value"

    if records_match := re.search(r"\b([0-9]+)\s+detailed\s+records\s+available\b", full_text, flags=re.IGNORECASE):
        detail_records_available = int(records_match.group(1))
    for line in lines:
        compact_line = " ".join(str(line or "").split()).strip()
        if not compact_line:
            continue
        if detail_records_available is None:
            records_match = re.match(r"^([0-9]+)\s+detailed\s+records\s+available$", compact_line, flags=re.IGNORECASE)
            if records_match:
                detail_records_available = int(records_match.group(1))
        if last_owned_location is None:
            location_match = re.match(r"^last\s+owned\s+in\s+([A-Za-z][A-Za-z .'-]+)$", compact_line, flags=re.IGNORECASE)
            if location_match:
                last_owned_location = " ".join(location_match.group(1).split())

    accident_events: list[dict[str, str]] = []
    service_events: list[dict[str, str]] = []
    accident_keywords = re.compile(r"\b(accident|collision|damage reported|damage event|structural|airbag|total loss|frame damage)\b", re.IGNORECASE)
    service_keywords = re.compile(
        r"\b(service|maintenance|oil|filter|tire|brake|battery|inspection|recommended maintenance|vehicle serviced)\b",
        re.IGNORECASE,
    )
    negative_accident = re.compile(
        r"\b(no|not)\s+[^.]{0,80}(accident|damage|damage reported|accidents reported|total loss|structural damage|airbag deployment|frame damage)[^.]{0,80}(reported|found|to carfax)?\b",
        re.IGNORECASE,
    )
    positive_accident_signal_count = 0
    for index, line in enumerate(lines):
        window = _sentence_window(lines, index, radius=4)
        line_date_match = _CARFAX_DATE_PATTERN.search(line)
        date_match = line_date_match or _CARFAX_DATE_PATTERN.search(window)
        if (
            accident_keywords.search(line)
            and not negative_accident.search(line)
            and not _carfax_is_generic_heading(line)
        ):
            line_clean = _short_fact(line)
            if (
                not line_date_match
                and len(line_clean) < 45
                and re.fullmatch(r"[A-Za-z /&-]+", line_clean)
            ):
                continue
            if (
                not line_date_match
                and any(token in line_clean.lower() for token in ["total loss", "structural damage", "airbag deployment"])
            ):
                continue
            if (
                "no issues reported" in window.lower()
                or "guaranteed no problem" in window.lower()
            ):
                continue
            positive_accident_signal_count += 1
            date_text = date_match.group(0) if date_match else ""
            event = {
                "date": date_text or "date not parsed",
                "severity": _classify_accident_severity(line),
                "description": line_clean,
            }
            key = (event["date"], event["severity"], event["description"][:80])
            if key not in {(item["date"], item["severity"], item["description"][:80]) for item in accident_events}:
                accident_events.append(event)
        if service_keywords.search(line):
            if not line_date_match and re.search(r"\bservice\s+(?:history\s+)?records?\b", line, flags=re.IGNORECASE):
                continue
            if re.fullmatch(r"(recent service highlights|service & comments|service|oil|regular oil changes)", line.strip(), flags=re.IGNORECASE):
                continue
            if "recall requires service" in line.lower():
                continue
            if _carfax_is_generic_heading(line):
                continue
            date_text = date_match.group(0) if date_match else ""
            event = {
                "date": date_text or "date not parsed",
                "description": _short_fact(line),
            }
            key = (event["date"], event["description"][:80])
            if key not in {(item["date"], item["description"][:80]) for item in service_events}:
                service_events.append(event)

    cleaned_accident_events: list[dict[str, str]] = []
    seen_accident_descriptions: set[tuple[str, str]] = set()
    for item in accident_events:
        desc = str(item.get("description") or "").strip()
        if not desc or _carfax_is_generic_heading(desc):
            continue
        signature = (
            str(item.get("date") or "").strip(),
            re.sub(r"[^a-z0-9]+", " ", desc.lower()).strip(),
        )
        if signature in seen_accident_descriptions:
            continue
        seen_accident_descriptions.add(signature)
        cleaned_accident_events.append(item)
    accident_events = cleaned_accident_events[:8]
    service_events = [item for item in service_events if "carfax" not in item["description"].lower()[:35]][:12]

    service_records_count = None
    for pattern in [
        r"\b([0-9]+)\s+service\s+history\s+records?\b",
        r"\b([0-9]+)\s+service\s+records?\b",
        r"\bservice\s+history\s+records?\s*[:\-]?\s*([0-9]+)\b",
    ]:
        if match := re.search(pattern, full_text, flags=re.IGNORECASE):
            service_records_count = int(match.group(1))
            break
    if service_records_count is None and service_events:
        service_records_count = len(service_events)

    dated_service_events = [
        (parsed, item["date"])
        for item in service_events
        if (parsed := _parse_carfax_event_date(str(item.get("date") or ""))) is not None
    ]
    last_service_date = max(dated_service_events, key=lambda item: item[0])[1] if dated_service_events else None
    service_descriptors: list[str] = []
    if regular_oil_changes:
        service_descriptors.append("regular oil changes")
    if detail_records_available:
        service_descriptors.append(f"{detail_records_available} detailed records available")

    if detail_records_available:
        if detail_records_available >= 10:
            service_quality = f"good service history ({detail_records_available} detailed records available)"
        elif detail_records_available >= 4:
            service_quality = f"service history available ({detail_records_available} detailed records available)"
        else:
            service_quality = f"limited history available ({detail_records_available} detailed records available)"
    elif service_records_count is None:
        service_quality = "service history not parsed"
    elif service_records_count >= 10:
        service_quality = f"strong service history ({service_records_count} records parsed)"
    elif service_records_count >= 4:
        service_quality = f"good service history ({service_records_count} records parsed)"
    elif service_records_count >= 1:
        service_quality = f"limited service history ({service_records_count} record(s) parsed)"
    else:
        service_quality = "no service records parsed"
    if regular_oil_changes and "regular oil changes" not in service_quality.lower():
        if _is_placeholder_carfax_text(service_quality) or service_quality == "no service records parsed":
            service_quality = "regular oil changes"
        else:
            service_quality = f"{service_quality}; regular oil changes"

    explicit_no_accidents = bool(re.search(r"\bno\s+accidents?\s+(reported|or damage reported)?\b|\bno\s+accidents?\s+or\s+damage\s+reported\b|\baccident\s+free\b", lowered))
    accident_confidence = "unknown"
    if explicit_no_accidents and not accident_events and positive_accident_signal_count == 0:
        accident_confidence = "explicit_clear"
        accident_events = []
    elif accident_events:
        accident_confidence = "incidents_found"
    elif explicit_no_accidents and positive_accident_signal_count > 0:
        accident_confidence = "conflict"
    elif positive_accident_signal_count > 0:
        accident_confidence = "signal_without_timeline"

    if accident_confidence == "explicit_clear":
        accident_damage = "No accidents reported in parsed CARFAX text."
    elif accident_events:
        severity_counts: dict[str, int] = {}
        for event in accident_events:
            severity_counts[event["severity"]] = severity_counts.get(event["severity"], 0) + 1
        accident_damage = ", ".join(f"{count} {severity}" for severity, count in sorted(severity_counts.items()))
        accident_damage = f"{len(accident_events)} accident/damage event(s) parsed: {accident_damage}."
    else:
        accident_damage = "Accident history could not be confirmed from cached CARFAX text."

    title_markers = ["salvage", "rebuilt", "lemon", "manufacturer buyback", "flood", "fire", "hail", "junk", "not actual mileage", "odometer rollback", "total loss"]
    title_flags = [
        marker
        for marker in title_markers
        if re.search(rf"(branded title|title brand|problem found|alert)[^.{{}}]{{0,220}}\b{re.escape(marker)}\b|\b{re.escape(marker)}\b[^.{{}}]{{0,220}}(problem found|alert)", lowered)
    ]
    if "none of these title problems were reported" in lowered or "guaranteed - none of these title problems" in lowered:
        title_brand = "No title brand/issues parsed from CARFAX text."
    elif title_flags:
        title_brand = "Potential title/brand flags parsed: " + ", ".join(_dedupe_strings(title_flags)) + "."
    elif re.search(r"\b(no|not)\s+(branded title|title issues|title problem|salvage|lemon)\b|no problems reported|no title issues|no title problems", lowered):
        title_brand = "No title brand/issues parsed from CARFAX text."
    else:
        title_brand = "Title/brand detail not clearly parsed from report text."

    usage_matches = [
        label
        for label, pattern in [
            ("personal", r"\bpersonal vehicle\b|\bpersonal use\b|\bpersonal lease\b|\btypes? of owners?:[^.]{0,120}\bpersonal\b"),
            ("lease", r"\blease vehicle\b|\bleased\b|\bpersonal lease\b|\btypes? of owners?:[^.]{0,120}\blease\b"),
            ("rental", r"\brental\b"),
            ("fleet", r"\bfleet\b"),
            ("commercial", r"\bcommercial\b"),
            ("corporate", r"\bcorporate vehicle\b|\btype of owner\s+corporate\b"),
        ]
        if re.search(pattern, lowered)
    ]
    usage = ", ".join(_dedupe_strings(usage_matches)) if usage_matches else "Use type not clearly parsed."

    major_count = sum(1 for item in accident_events if item["severity"] == "major")
    minor_count = sum(1 for item in accident_events if item["severity"] == "minor")
    moderate_count = sum(1 for item in accident_events if item["severity"] == "moderate")
    damage_count = sum(1 for item in accident_events if item["severity"] == "damage")

    highlights: list[str] = []
    if owner_count:
        highlights.append(f"Owners: {owner_count}.")
    highlights.append(accident_damage)
    if accident_events:
        highlights.extend(
            f"{item['date']}: {item['severity']} - {item['description']}"
            for item in accident_events[:5]
        )
    highlights.append(f"Service history: {service_quality}.")
    if last_service_date:
        highlights.append(f"Last parsed service date: {last_service_date}.")
    if regular_oil_changes:
        highlights.append("Maintenance badge: Regular Oil Changes.")
    if detail_records_available:
        highlights.append(f"Detailed records available: {detail_records_available}.")
    if value_badge or market_position or market_delta:
        value_line = "Market/value: "
        value_parts = [part for part in [value_badge, market_position, market_delta] if part]
        highlights.append(value_line + " / ".join(value_parts) + ".")
    if carfax_value:
        highlights.append(f"CARFAX value: {carfax_value}.")
    highlights.append(title_brand)
    highlights.append(f"Use type: {usage}")
    if last_owned_location:
        highlights.append(f"Last owned in: {last_owned_location}.")

    summary_bits = [
        owner_count,
        accident_damage,
        service_quality,
        " / ".join(part for part in [value_badge, market_position, market_delta] if part) or None,
    ]
    summary = "CARFAX report parsed: " + "; ".join(str(part) for part in summary_bits if part) + "."

    return _normalize_carfax_parsed_payload({
        "ok": True,
        "source": source,
        "source_url": source_url,
        "report_vin": report_vin,
        "report_vehicle_title": report_vehicle_title,
        "summary": summary,
        "owner_count": owner_count,
        "value_badge": value_badge,
        "carfax_value": carfax_value,
        "market_position": market_position,
        "market_delta": market_delta,
        "listing_price": listing_price,
        "accident_damage": accident_damage,
        "accident_events": accident_events,
        "accident_confidence": accident_confidence,
        "accident_counts": {
            "total": len(accident_events),
            "minor": minor_count,
            "moderate": moderate_count,
            "major": major_count,
            "damage": damage_count,
        },
        "title_brand": title_brand,
        "service_history": service_quality,
        "service_records_count": service_records_count,
        "detail_records_available": detail_records_available,
        "last_service_date": last_service_date,
        "service_events": service_events[:8],
        "usage": usage,
        "last_owned_location": last_owned_location,
        "regular_oil_changes": regular_oil_changes,
        "report_access": "CARFAX report text parsed by Xconsole.",
        "highlights": _dedupe_strings(highlights)[:14],
        "updated_at": datetime.now(timezone.utc).isoformat(),
    })


def _merge_carfax_facts(existing: dict[str, Any] | None, report: dict[str, Any] | None) -> dict[str, Any]:
    merged: dict[str, Any] = dict(existing or {})
    if not isinstance(report, dict):
        return merged
    if not report.get("ok"):
        if report.get("identity_mismatch"):
            merged = _strip_carfax_structured_report_fields(merged)
            merged["source"] = "dealer_page_carfax_badge"
        merged["report_access"] = report.get("report_access") or merged.get("report_access")
        merged["report_blocked"] = bool(report.get("blocked"))
        merged["report_source"] = report.get("source")
        if report.get("identity_mismatch"):
            merged["identity_mismatch"] = True
            merged["identity_validation"] = report.get("identity_validation")
        return merged

    preserved_badge = {
        key: merged.get(key)
        for key in ("badge_text", "value_badge", "owner_count")
        if merged.get(key)
    }
    merged.update(report)
    for key, value in preserved_badge.items():
        if not merged.get(key):
            merged[key] = value
    merged["source"] = report.get("source") or "carfax_report"
    return _normalize_carfax_parsed_payload(merged)


def _fetch_carfax_report_details(
    carfax_url: str | None,
    *,
    visible_browser: bool = False,
    vehicle: dict[str, Any] | None = None,
    expected_vin: str = "",
) -> dict[str, Any]:
    source_url = str(carfax_url or "").strip()
    if not source_url:
        return {"ok": False, "source": "missing_carfax_url", "report_access": "No CARFAX URL available."}
    headers = {
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/122.0.0.0 Safari/537.36"
        ),
        "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
    }
    try:
        with httpx.Client(timeout=16.0, follow_redirects=True, headers=headers) as client:
            response = client.get(source_url)
        parsed_text, text_meta = _extract_text_from_carfax_response(
            response, fallback_url=source_url
        )
        parsed = _extract_carfax_report_facts(parsed_text, source_url=str(response.url), source="carfax_report")
        parsed["http_status"] = response.status_code
        parsed["extract_meta"] = text_meta
        parsed = _reject_carfax_identity_mismatch(parsed, vehicle=vehicle, expected_vin=expected_vin)
        if parsed.get("ok") or _env_flag("CARFAX_DISABLE_BROWSER_FALLBACK", False):
            return parsed
    except Exception:
        if visible_browser:
            return {
                "ok": False,
                "source": "carfax_report_fetch_error",
                "source_url": source_url,
                "report_access": "CARFAX report fetch failed. Retry with visible browser from the CARFAX viewer.",
                "highlights": ["CARFAX linked report could not be fetched automatically."],
            }
        parsed = {"ok": False, "source": "carfax_report_fetch_error", "source_url": source_url, "report_access": "CARFAX report fetch failed."}

    if not _needs_carfax_browser_fallback(parsed):
        return parsed

    browser_result = _fetch_carfax_report_details_with_browser(
        source_url,
        visible=visible_browser if not _env_flag("CARFAX_BROWSER_FORCE_VISIBLE", False) else True,
        vehicle=vehicle,
        expected_vin=expected_vin,
    )
    if browser_result.get("ok"):
        return browser_result
    if not visible_browser:
        if browser_result.get("source") in {"carfax_report_blocked", "carfax_report_fetch_error"} or not browser_result.get("ok"):
            visible_retry = _fetch_carfax_report_details_with_browser(
                source_url,
                visible=True,
                vehicle=vehicle,
                expected_vin=expected_vin,
            )
            if visible_retry.get("ok") or visible_retry.get("source") == "carfax_report_plain_chrome":
                return visible_retry
    if visible_browser:
        return browser_result
    merged = browser_result or parsed
    if isinstance(merged, dict):
        return merged
    return {
        "ok": False,
        "source": "carfax_report_fetch_error",
        "source_url": source_url,
        "report_access": "CARFAX report fetch failed.",
        "highlights": ["CARFAX linked report could not be fetched automatically."],
    }


def _fetch_carfax_report_details_with_browser(
    carfax_url: str,
    *,
    visible: bool = False,
    vehicle: dict[str, Any] | None = None,
    expected_vin: str = "",
) -> dict[str, Any]:
    if webdriver is None or ChromeService is None:
        return {"ok": False, "source": "carfax_browser_unavailable", "source_url": carfax_url}
    if visible:
        plain_result = _fetch_carfax_report_details_with_plain_chrome(
            carfax_url,
            vehicle=vehicle,
            expected_vin=expected_vin,
        )
        if plain_result.get("ok"):
            return plain_result
    driver = None
    try:
        driver = _open_browser(
            timeout_seconds=90.0 if visible else 28.0,
            headless=not visible,
            profile_dir=CARFAX_CHROME_PROFILE_DIR if visible else None,
            images_enabled=True,
        )
        driver.get(carfax_url)
        body_text = ""
        page_source = ""
        parsed: dict[str, Any] | None = None
        deadline = time.time() + (90 if visible else 8)
        while time.time() < deadline:
            time.sleep(3 if visible else 2)
            page_source = str(driver.page_source or "")
            try:
                body_text = str(driver.find_element(By.TAG_NAME, "body").text or "")
            except Exception:
                body_text = ""
            raw_candidate = body_text if len(body_text) > len(page_source) else page_source
            parsed = _extract_carfax_report_facts(
                raw_candidate,
                source_url=str(driver.current_url or carfax_url),
                source="carfax_report_browser_visible" if visible else "carfax_report_browser",
            )
            if parsed.get("ok") or not visible:
                break
        page_source = str(driver.page_source or "")
        raw = body_text if len(body_text) > len(page_source) else page_source
        parsed = parsed or _extract_carfax_report_facts(
            raw,
            source_url=str(driver.current_url or carfax_url),
            source="carfax_report_browser_visible" if visible else "carfax_report_browser",
        )
        parsed["browser_title"] = str(driver.title or "")
        parsed["browser_text_chars"] = len(body_text)
        parsed = _reject_carfax_identity_mismatch(parsed, vehicle=vehicle, expected_vin=expected_vin)
        if visible and not parsed.get("ok"):
            parsed["report_access"] = (
                "CARFAX still did not expose the readable report after opening a visible Chrome session. "
                "Leave that Chrome window on the completed report page and run Retry Full Parse again."
            )
        return parsed
    except Exception as exc:
        return {
            "ok": False,
            "source": "carfax_browser_fetch_error",
            "source_url": carfax_url,
            "report_access": f"CARFAX browser fetch failed: {exc}",
        }
    finally:
        if driver is not None:
            try:
                driver.quit()
            except Exception:
                pass


def _fetch_carfax_report_details_with_plain_chrome(
    carfax_url: str,
    *,
    vehicle: dict[str, Any] | None = None,
    expected_vin: str = "",
) -> dict[str, Any]:
    chrome_binary = _find_chrome_binary()
    chromedriver = _find_chromedriver()
    if webdriver is None or ChromeService is None or not chrome_binary or not chromedriver:
        return {"ok": False, "source": "carfax_plain_chrome_unavailable", "source_url": carfax_url}
    port = int(os.getenv("CARFAX_REMOTE_DEBUGGING_PORT", "9234") or "9234")
    CARFAX_CHROME_PROFILE_DIR.mkdir(parents=True, exist_ok=True)
    command = [
        str(chrome_binary),
        f"--remote-debugging-port={port}",
        f"--user-data-dir={CARFAX_CHROME_PROFILE_DIR}",
        "--new-window",
        carfax_url,
    ]
    process = None
    driver = None
    try:
        process = subprocess.Popen(command, cwd=ROOT_DIR, stdout=subprocess.DEVNULL, stderr=subprocess.DEVNULL)
        time.sleep(5)
        chrome_options = webdriver.ChromeOptions()
        chrome_options.add_experimental_option("debuggerAddress", f"127.0.0.1:{port}")
        driver = webdriver.Chrome(
            service=ChromeService(executable_path=str(chromedriver)),
            options=chrome_options,
        )
        deadline = time.time() + 120
        parsed: dict[str, Any] | None = None
        body_text = ""
        page_source = ""
        while time.time() < deadline:
            time.sleep(4)
            handles = list(driver.window_handles)
            if handles:
                driver.switch_to.window(handles[-1])
            if str(driver.current_url or "").strip() != carfax_url and "carfax.com" not in str(driver.current_url or "").lower():
                try:
                    driver.get(carfax_url)
                except Exception:
                    pass
            page_source = str(driver.page_source or "")
            try:
                body_text = str(driver.find_element(By.TAG_NAME, "body").text or "")
            except Exception:
                body_text = ""
            raw = body_text if len(body_text) > len(page_source) else page_source
            parsed = _extract_carfax_report_facts(
                raw,
                source_url=str(driver.current_url or carfax_url),
                source="carfax_report_plain_chrome",
            )
            if parsed.get("ok"):
                break
        parsed = parsed or {
            "ok": False,
            "source": "carfax_report_plain_chrome",
            "source_url": carfax_url,
            "report_access": "Plain Chrome did not expose readable CARFAX report text before timeout.",
        }
        parsed["browser_title"] = str(driver.title or "")
        parsed["browser_text_chars"] = len(body_text)
        parsed = _reject_carfax_identity_mismatch(parsed, vehicle=vehicle, expected_vin=expected_vin)
        if not parsed.get("ok"):
            parsed["report_access"] = (
                "CARFAX still did not expose the readable report in plain Chrome. "
                "Complete any verification in the opened Chrome window, leave the report open, then click Retry Full Parse again."
            )
        return parsed
    except Exception as exc:
        return {
            "ok": False,
            "source": "carfax_plain_chrome_fetch_error",
            "source_url": carfax_url,
            "report_access": f"Plain Chrome CARFAX fetch failed: {exc}",
        }
    finally:
        if driver is not None:
            try:
                driver.quit()
            except Exception:
                pass
        if process is not None and process.poll() is None and not _env_flag("CARFAX_KEEP_BROWSER_OPEN", True):
            try:
                process.terminate()
            except Exception:
                pass


def _build_browser_options(*, headless: bool = True, profile_dir: Path | None = None, images_enabled: bool = False) -> Any:
    chrome_binary = _find_chrome_binary()
    if webdriver is None or ChromeService is None or not chrome_binary:
        return None

    chrome_options = webdriver.ChromeOptions()
    chrome_options.binary_location = str(chrome_binary)
    if headless:
        chrome_options.add_argument("--headless=new")
    if profile_dir is not None:
        profile_dir.mkdir(parents=True, exist_ok=True)
        chrome_options.add_argument(f"--user-data-dir={profile_dir}")
    chrome_options.add_argument("--no-sandbox")
    chrome_options.add_argument("--disable-dev-shm-usage")
    chrome_options.add_argument("--disable-gpu")
    chrome_options.add_argument("--window-size=1600,1200")
    chrome_options.add_argument("--disable-background-networking")
    chrome_options.add_argument("--disable-notifications")
    if not images_enabled:
        chrome_options.add_argument("--blink-settings=imagesEnabled=false")
    chrome_options.page_load_strategy = "eager"
    return chrome_options


def _open_headless_browser(timeout_seconds: float = 18.0) -> Any:
    return _open_browser(timeout_seconds=timeout_seconds, headless=True)


def _open_browser(*, timeout_seconds: float = 18.0, headless: bool = True, profile_dir: Path | None = None, images_enabled: bool = False) -> Any:
    if webdriver is None or ChromeService is None:
        raise RuntimeError("selenium_unavailable")
    chromedriver = _find_chromedriver()
    if not chromedriver:
        raise RuntimeError("chromedriver_missing")
    chrome_options = _build_browser_options(headless=headless, profile_dir=profile_dir, images_enabled=images_enabled)
    if chrome_options is None:
        raise RuntimeError("chrome_binary_missing")

    driver = webdriver.Chrome(
        service=ChromeService(executable_path=str(chromedriver)),
        options=chrome_options,
    )
    driver.set_page_load_timeout(max(10, int(timeout_seconds)))
    return driver


def _extract_ws_inventory_bootstrap(page_source: str) -> tuple[str, dict[str, Any]] | None:
    match = re.search(
        r'fetch\("(/api/widget/ws-inv-data/getInventory)".*?body:decodeURI\("([^"]+)"\)',
        page_source,
        flags=re.DOTALL,
    )
    if not match:
        return None
    try:
        from urllib.parse import unquote

        payload = json.loads(unquote(match.group(2)))
    except Exception:
        return None
    if not isinstance(payload, dict):
        return None
    return match.group(1), payload


def _tracking_attributes_map(item: dict[str, Any]) -> dict[str, Any]:
    mapped: dict[str, Any] = {}
    raw = item.get("trackingAttributes")
    if not isinstance(raw, list):
        return mapped
    for entry in raw:
        if not isinstance(entry, dict):
            continue
        name = str(entry.get("name") or "").strip()
        if not name:
            continue
        mapped[name] = entry.get("value")
        normalized = str(entry.get("normalizedValue") or "").strip()
        if normalized:
            mapped[f"normal{name[:1].upper()}{name[1:]}"] = normalized
    return mapped


def _ws_inventory_item_to_record(item: dict[str, Any], *, source_url: str) -> dict[str, Any] | None:
    if not isinstance(item, dict):
        return None
    attrs = _tracking_attributes_map(item)
    tracking_pricing = item.get("trackingPricing") if isinstance(item.get("trackingPricing"), dict) else {}
    pricing = item.get("pricing") if isinstance(item.get("pricing"), dict) else {}

    title_raw = item.get("title")
    if isinstance(title_raw, list):
        title = " ".join(str(part).strip() for part in title_raw if str(part).strip())
    else:
        title = str(title_raw or "").strip()
    if not title:
        title = " ".join(
            str(item.get(part) or "").strip()
            for part in ("year", "make", "model", "trim")
            if str(item.get(part) or "").strip()
        )

    link = str(item.get("link") or item.get("url") or "").strip()
    detail_url = urljoin(source_url, link) if link else None
    images_raw = item.get("images") if isinstance(item.get("images"), list) else []
    photos = _dedupe_urls(
        [
            str(image.get("uri") or image.get("src") or image.get("url") or "").strip()
            for image in images_raw
            if isinstance(image, dict)
        ]
    )

    price = (
        tracking_pricing.get("internetPrice")
        or tracking_pricing.get("salePrice")
        or pricing.get("internetPrice")
        or pricing.get("salePrice")
        or item.get("internetPrice")
        or item.get("salePrice")
        or item.get("askingPrice")
    )
    carfax_url, carfax_facts = _extract_carfax_link_from_inventory_item(item)
    direct_carfax_url = item.get("carfax_url") or item.get("carfaxUrl") or item.get("vehicleHistoryUrl")
    if direct_carfax_url and not carfax_url:
        carfax_url = str(direct_carfax_url).strip()
    raw_inventory_category = (
        item.get("inventory_category")
        or item.get("inventoryCategory")
        or item.get("type")
        or item.get("condition")
        or item.get("vehicleType")
    )
    inventory_category = str(raw_inventory_category or "").strip()
    if not inventory_category:
        source_path = urlparse(source_url).path.lower()
        if "/new-" in source_path or "/new/" in source_path or "new-inventory" in source_path:
            inventory_category = "new"
        elif "/used-" in source_path or "/used/" in source_path or "used-inventory" in source_path or "pre-owned" in source_path:
            inventory_category = "used"

    engine_parts = [
        attrs.get("engineSize") or item.get("engineSize"),
        attrs.get("engine") or item.get("engine"),
    ]
    engine = " ".join(str(part).strip() for part in engine_parts if str(part).strip())

    vin_value = _vin_candidate(item.get("vin")) or _extract_vin_from_text(json.dumps(item)[:4000])
    if not vin_value and detail_url:
        vin_value = _synthetic_vin_from_detail_url(detail_url)
    if not vin_value and not title:
        return None

    record = {
        "vin": vin_value or "UNKNOWN",
        "title": title or vin_value or "Vehicle",
        "price": price,
        "mileage": _normalize_mileage_value(attrs.get("odometer") or item.get("odometer") or item.get("mileage")),
        "drivetrain": _normalize_text_value(attrs.get("driveLine") or item.get("driveLine")),
        "engine": _normalize_text_value(engine),
        "transmission": _normalize_text_value(attrs.get("transmission") or item.get("transmission")),
        "location": _normalize_text_value(item.get("accountName")),
        "detail_url": detail_url,
        "exterior": attrs.get("exteriorColor") or item.get("exteriorColor") or item.get("extColor"),
        "interior": attrs.get("interiorColor") or item.get("interiorColor") or item.get("intColor"),
        "photos": photos,
        "status_label": str(item.get("status") or "In Stock"),
        "inventory_category": inventory_category,
        "stock_number": item.get("stockNumber"),
        "days_on_lot": item.get("daysOnLot") or attrs.get("daysOnLot"),
    }
    if carfax_url:
        record["carfax_url"] = carfax_url
    if carfax_facts:
        record["carfax_facts"] = carfax_facts
    return record


def _fetch_ws_inventory_pages_from_browser(
    driver: Any,
    *,
    api_path: str,
    bootstrap_payload: dict[str, Any],
    source_url: str,
    budget_seconds: float,
    started_at: float,
) -> tuple[list[dict[str, Any]], list[str]]:
    notes: list[str] = ["source_mode=browser_ws_inventory_api"]
    driver.set_script_timeout(max(30, int(min(120, budget_seconds))))
    script = """
const done = arguments[arguments.length - 1];
fetch(arguments[0], {
  method: "POST",
  headers: {"Content-Type": "application/json"},
  body: JSON.stringify(arguments[1])
}).then(async (response) => {
  done({status: response.status, text: await response.text()});
}).catch((error) => done({error: String(error)}));
"""

    def fetch_page(start: int, payload: dict[str, Any]) -> dict[str, Any]:
        page_payload = json.loads(json.dumps(payload))
        if start:
            page_payload["inventoryParameters"] = dict(page_payload.get("inventoryParameters") or {})
            page_payload["inventoryParameters"]["start"] = [str(start)]
        result = driver.execute_async_script(script, api_path, page_payload)
        if not isinstance(result, dict):
            raise RuntimeError("inventory API returned invalid browser result")
        if result.get("error"):
            raise RuntimeError(str(result.get("error")))
        if int(result.get("status") or 0) >= 400:
            raise RuntimeError(f"inventory API HTTP {result.get('status')}")
        return json.loads(str(result.get("text") or "{}"))

    records: list[dict[str, Any]] = []
    seen: set[str] = set()
    first_payload = fetch_page(0, bootstrap_payload)
    page_info = first_payload.get("pageInfo") if isinstance(first_payload.get("pageInfo"), dict) else {}
    page_size = int(page_info.get("pageSize") or bootstrap_payload.get("preferences", {}).get("pageSize") or 18)
    total_count = int(page_info.get("totalCount") or 0)
    starts = list(range(0, total_count, page_size)) if total_count and page_size else [0]
    notes.append(f"ws_inventory_total_count={total_count or 'unknown'}")
    notes.append(f"ws_inventory_page_size={page_size}")
    notes.append(f"ws_inventory_pages_planned={len(starts)}")

    def add_inventory(payload: dict[str, Any], start: int) -> int:
        added = 0
        inventory = payload.get("inventory") if isinstance(payload.get("inventory"), list) else []
        for raw_item in inventory:
            record = _ws_inventory_item_to_record(raw_item, source_url=source_url) if isinstance(raw_item, dict) else None
            if not record:
                continue
            key = str(record.get("vin") or record.get("detail_url") or record.get("title") or "").lower()
            if not key or key in seen:
                continue
            seen.add(key)
            records.append(record)
            added += 1
        notes.append(f"ws_inventory_start={start}_new={added}")
        return added

    add_inventory(first_payload, 0)
    for start in starts[1:]:
        if time.time() - started_at > budget_seconds - 8:
            notes.append(f"ws_inventory_budget_exhausted_at_start={start}")
            break
        try:
            add_inventory(fetch_page(start, bootstrap_payload), start)
        except Exception as exc:
            notes.append(f"ws_inventory_page_error_start={start}|{exc}")
            continue
    notes.append(f"ws_inventory_records={len(records)}")
    return records, notes


def _fetch_live_inventory_records_via_browser_html(
    *,
    source_url: str,
    timeout_seconds: int,
) -> dict[str, Any]:
    notes: list[str] = ["source_mode=browser_html"]
    driver = None
    records: list[dict[str, Any]] = []
    seen: set[str] = set()
    started_at = time.time()
    budget_seconds = max(18.0, float(timeout_seconds))

    def _records_from_current_page() -> list[dict[str, Any]]:
        page_records: list[dict[str, Any]] = []
        cards = driver.find_elements(
            By.XPATH,
            "//*[contains(@class,'vehicle-card') or contains(@class,'vehicle-card-details-container') or @data-vin]",
        ) if driver else []
        for card in cards:
            try:
                links = card.find_elements(
                    By.XPATH,
                    ".//a[contains(@href,'/used/') or contains(@href,'/new/')]",
                )
                card_vin = _vin_candidate(card.get_attribute("data-vin")) or _extract_vin_from_text(card.text)
                detail_url = ""
                title = ""
                for link in links:
                    candidate = str(link.get_attribute("href") or "").strip()
                    if not candidate:
                        continue
                    if candidate.startswith("/"):
                        candidate = urljoin(source_url, candidate)
                    key = candidate.lower()
                    if key in seen:
                        continue
                    detail_url = candidate
                    title = str(link.text or link.get_attribute("title") or "").strip()
                    break
                if not detail_url:
                    continue

                card_text = str(card.text or "").strip()
                if not title:
                    title = next(
                        (
                            line.strip()
                            for line in card_text.splitlines()
                            if re.search(r"\b(?:19|20)\d{2}\b", line)
                        ),
                        "",
                    )
                if not title:
                    title = _synthetic_vin_from_detail_url(detail_url)

                price_match = re.search(r"\$([0-9,]+)", card_text)
                mileage_match = re.search(
                    r"([0-9][0-9,]{0,9})\s+miles",
                    card_text,
                    flags=re.IGNORECASE,
                )
                photo_urls = _dedupe_urls(
                    [
                        src
                        for image in card.find_elements(By.TAG_NAME, "img")
                        if (src := image.get_attribute("src"))
                        and "pictures.dealer.com" in str(src).lower()
                    ]
                )

                seen.add(detail_url.lower())
                page_records.append(
                    {
                        "vin": card_vin or _synthetic_vin_from_detail_url(detail_url),
                        "title": title,
                        "price": int(price_match.group(1).replace(",", "")) if price_match else None,
                        "mileage": int(mileage_match.group(1).replace(",", "")) if mileage_match else None,
                        "detail_url": detail_url,
                        "photos": photo_urls,
                        "status_label": "In Stock",
                    }
                )
            except Exception:
                continue
        return page_records

    def _total_vehicle_count_from_page() -> int | None:
        if driver is None:
            return None
        text = str(driver.find_element(By.TAG_NAME, "body").text or "")
        matches = re.findall(r"\b([0-9]{1,5})\s+Vehicles\b", text, flags=re.IGNORECASE)
        if not matches:
            return None
        try:
            return max(int(item.replace(",", "")) for item in matches)
        except Exception:
            return None

    def _page_starts_from_pagination(page_size: int, total_count: int | None) -> list[int]:
        starts: set[int] = {0}
        if driver is not None:
            for link in driver.find_elements(By.TAG_NAME, "a"):
                href = str(link.get_attribute("href") or "")
                match = re.search(r"[?&]start=([0-9]+)", href)
                if match:
                    starts.add(int(match.group(1)))
        if total_count and page_size > 0:
            starts.update(range(0, total_count, page_size))
        return sorted(starts)

    try:
        driver = _open_headless_browser(timeout_seconds=max(18.0, float(timeout_seconds)))
        first_url = source_url
        if "start=" not in first_url:
            separator = "&" if "?" in first_url else "?"
            first_url = f"{first_url}{separator}start=0"
        driver.get(first_url)
        time.sleep(4)

        bootstrap = _extract_ws_inventory_bootstrap(driver.page_source)
        if bootstrap:
            api_path, api_payload = bootstrap
            try:
                api_records, api_notes = _fetch_ws_inventory_pages_from_browser(
                    driver,
                    api_path=api_path,
                    bootstrap_payload=api_payload,
                    source_url=source_url,
                    budget_seconds=budget_seconds,
                    started_at=started_at,
                )
                notes.extend(api_notes)
                if api_records:
                    return {
                        "source_url": source_url,
                        "fetched_at": datetime.now(timezone.utc).isoformat(),
                        "items": api_records,
                        "items_count": len(api_records),
                        "diagnostics": notes,
                    }
            except Exception as exc:
                notes.append(f"ws_inventory_api_error={exc}")

        first_page_records = _records_from_current_page()
        records.extend(first_page_records)
        page_size = max(1, len(first_page_records))
        total_count = _total_vehicle_count_from_page()
        starts = _page_starts_from_pagination(page_size=page_size, total_count=total_count)
        notes.append(f"browser_first_page_records={len(first_page_records)}")
        notes.append(f"browser_total_vehicle_count={total_count or 'unknown'}")
        notes.append(f"browser_page_starts_planned={len(starts)}")

        for page_index, start in enumerate(starts):
            if start == 0:
                continue
            if time.time() - started_at > budget_seconds - 8:
                notes.append(f"browser_budget_exhausted_at_start={start}")
                break
            separator = "&" if "?" in source_url else "?"
            page_url = re.sub(r"([?&])start=[0-9]+", rf"\1start={start}", source_url)
            if page_url == source_url:
                page_url = f"{source_url}{separator}start={start}"
            driver.get(page_url)
            time.sleep(1.2)
            before = len(records)
            records.extend(_records_from_current_page())
            notes.append(f"browser_page_{page_index + 1}_start={start}_new={len(records) - before}")
        notes.append(f"browser_listing_records={len(records)}")
    except Exception as exc:
        notes.append(f"browser_html_error={exc}")
        return {
            "source_url": source_url,
            "fetched_at": datetime.now(timezone.utc).isoformat(),
            "items": [],
            "items_count": 0,
            "diagnostics": notes,
        }
    finally:
        if driver is not None:
            try:
                driver.quit()
            except Exception:
                pass

    if len(records) <= 80:
        records, detail_notes = _enrich_inventory_records_from_proxy_detail_pages(
            records,
            timeout_seconds=timeout_seconds,
        )
        notes.extend(detail_notes)
    else:
        notes.append("detail_proxy_skipped_for_large_inventory")
    return {
        "source_url": source_url,
        "fetched_at": datetime.now(timezone.utc).isoformat(),
        "items": records,
        "items_count": len(records),
        "diagnostics": notes,
    }


def _fetch_vehicle_asset_bundle_from_browser(
    *,
    detail_url: str,
    timeout_seconds: float = 18.0,
) -> dict[str, Any]:
    if webdriver is None or ChromeService is None:
        return {"ok": False, "error": "selenium_unavailable"}

    driver = None
    try:
        driver = _open_headless_browser(timeout_seconds=timeout_seconds)
        driver.get(detail_url)
        deadline = time.time() + max(6.0, float(timeout_seconds))
        html_text = ""
        links: dict[str, str | None] = {"sticker_url": None, "carfax_url": None}
        photos: list[str] = []
        carfax_facts: dict[str, Any] = {}
        resource_text = ""
        while time.time() < deadline:
            time.sleep(1.5)
            html_text = str(driver.page_source or "")
            try:
                resource_urls = driver.execute_script(
                    "return (window.performance && performance.getEntriesByType) ? performance.getEntriesByType('resource').map((entry) => entry.name) : [];"
                ) or []
            except Exception:
                resource_urls = []
            if isinstance(resource_urls, list) and resource_urls:
                resource_text = "\n".join(str(item or "") for item in resource_urls)
            scan_blob = html_text if not resource_text else f"{html_text}\n{resource_text}"
            links = _extract_asset_links_from_html(scan_blob, base_url=detail_url)
            photos = _extract_vehicle_photo_urls_from_html(scan_blob, base_url=detail_url)
            carfax_facts = _extract_carfax_facts_from_html(scan_blob)
            if links.get("carfax_url") or links.get("sticker_url") or len(photos) >= 3:
                break
        return {
            "ok": True,
            "source_mode": "browser_html",
            "current_url": driver.current_url,
            "page_title": driver.title,
            "photos": photos,
            "photos_count": len(photos),
            "main_photo": photos[0] if photos else None,
            "sticker_url": links.get("sticker_url"),
            "carfax_url": links.get("carfax_url"),
            "carfax_facts": carfax_facts,
        }
    except Exception as exc:
        return {
            "ok": False,
            "error": str(exc),
            "source_mode": "browser_html",
        }
    finally:
        if driver is not None:
            try:
                driver.quit()
            except Exception:
                pass


def _discover_carfax_from_detail_page(detail_url: str, *, timeout_seconds: float = 12.0) -> tuple[str | None, dict[str, Any], dict[str, Any]]:
    """Best-effort extraction of CARFAX URL/facts from listing detail pages."""
    details: dict[str, Any] = {"method": "none", "ok": False, "error": None}
    carfax_url: str | None = None
    carfax_facts: dict[str, Any] = {}

    raw_detail_url = str(detail_url or "").strip()
    if not raw_detail_url:
        details["error"] = "missing_detail_url"
        return carfax_url, carfax_facts, details

    try:
        headers = {
            "User-Agent": "Mozilla/5.0 XconsoleDealershipAssetCache/1.0",
            "Accept": "text/html,application/json,*/*;q=0.8",
        }
        with httpx.Client(timeout=timeout_seconds, follow_redirects=True, headers=headers) as client:
            response = client.get(raw_detail_url)
        details["method"] = "httpx"
        details["http_status"] = response.status_code
        details["ok"] = response.status_code < 400
        if response.status_code < 400:
            links = _extract_asset_links_from_html(response.text, base_url=response.url.__str__())
            discovered = links.get("carfax_url")
            if discovered:
                carfax_url = str(discovered).strip()
            discovered_facts = _extract_carfax_facts_from_html(response.text)
            if discovered_facts:
                carfax_facts = dict(discovered_facts)
            if carfax_url or carfax_facts:
                return carfax_url, carfax_facts, details
        else:
            details["error"] = f"detail_http_status_{response.status_code}"
    except Exception as exc:
        details["error"] = str(exc)

    bundle = _fetch_vehicle_asset_bundle_from_browser(detail_url=raw_detail_url, timeout_seconds=max(12.0, timeout_seconds))
    if bundle.get("ok"):
        details["method"] = "browser"
        details["ok"] = True
        details["source_mode"] = bundle.get("source_mode")
        discovered = bundle.get("carfax_url")
        if discovered:
            carfax_url = str(discovered).strip()
        discovered_facts = bundle.get("carfax_facts")
        if isinstance(discovered_facts, dict):
            merged_facts = dict(carfax_facts)
            merged_facts.update(discovered_facts)
            carfax_facts = merged_facts
        return carfax_url, carfax_facts, details

    details["error"] = bundle.get("error") or details.get("error") or "carfax_discovery_failed"
    return carfax_url, carfax_facts, details


def _load_vehicle_assets(vin: str, *, refresh: bool = False) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        raise HTTPException(status_code=400, detail={"message": "vin is required"})

    cache_path = _vehicle_assets_cache_path(clean_vin)
    if cache_path.exists() and not refresh:
        cached = _safe_read_json(cache_path, {})
        if isinstance(cached, dict) and cached.get("vin") == clean_vin:
            vehicle = _find_vehicle_by_vin(clean_vin) or {
                "vin": clean_vin,
                "detail_url": cached.get("detail_url"),
                "photos": cached.get("photos") if isinstance(cached.get("photos"), list) else [],
                "carfax_url": cached.get("carfax_url"),
                "carfax_facts": cached.get("carfax_facts") if isinstance(cached.get("carfax_facts"), dict) else {},
            }
            # Never return a stale/cached report if we still have a linked CARFAX URL
            # but weak parsed report details. This keeps the UI honest for vehicles
            # that were cached before CARFAX parsing matured.
            cached_report = cached.get("carfax_report")
            cached_carfax_url = str(cached.get("carfax_url") or "").strip()
            if isinstance(cached_report, dict):
                validated_cached_report = _reject_carfax_identity_mismatch(
                    cached_report,
                    vehicle=vehicle,
                    expected_vin=clean_vin,
                )
                if validated_cached_report != cached_report:
                    cached["carfax_report"] = validated_cached_report
                    cached["carfax_facts"] = _merge_carfax_facts(
                        cached.get("carfax_facts") if isinstance(cached.get("carfax_facts"), dict) else {},
                        validated_cached_report,
                    )
                    cached.update(_vehicle_asset_summary_payload(vehicle=vehicle, assets=cached))
                    _safe_write_json(cache_path, cached)
                    if isinstance(cached.get("carfax_summary"), dict):
                        CARFAX_SUMMARY_DIR.mkdir(parents=True, exist_ok=True)
                        _safe_write_json(CARFAX_SUMMARY_DIR / f"{clean_vin}.json", cached["carfax_summary"])
                    cached_report = validated_cached_report
            return cached

    vehicle = _find_vehicle_by_vin(clean_vin)
    if not vehicle:
        if cache_path.exists():
            cached = _safe_read_json(cache_path, {})
            if isinstance(cached, dict) and cached.get("vin") == clean_vin:
                return cached
        raise HTTPException(status_code=404, detail={"message": f"Vehicle not found for VIN {clean_vin}"})

    photos = vehicle.get("photos") if isinstance(vehicle.get("photos"), list) else []
    detail_url = vehicle.get("detail_url")
    existing = _safe_read_json(cache_path, {})
    if not isinstance(existing, dict) or existing.get("vin") != clean_vin:
        existing = {}
    payload: dict[str, Any] = {
        **existing,
        "vin": clean_vin,
        "detail_url": detail_url or existing.get("detail_url"),
        "photos": photos or (existing.get("photos") if isinstance(existing.get("photos"), list) else []),
        "photos_count": 0,
        "main_photo": None,
        "sticker_url": existing.get("sticker_url"),
        "carfax_url": existing.get("carfax_url") or vehicle.get("carfax_url"),
        "carfax_facts": {
            **(vehicle.get("carfax_facts") if isinstance(vehicle.get("carfax_facts"), dict) else {}),
            **(existing.get("carfax_facts") if isinstance(existing.get("carfax_facts"), dict) else {}),
        },
        "loaded_at": existing.get("loaded_at") or datetime.now(timezone.utc).isoformat(),
    }

    if detail_url and isinstance(detail_url, str):
        browser_bundle = _fetch_vehicle_asset_bundle_from_browser(detail_url=detail_url)
        if browser_bundle.get("ok"):
            payload["detail_source_mode"] = browser_bundle.get("source_mode")
            payload["photos"] = list(browser_bundle.get("photos") or payload["photos"])
            payload["photos_count"] = len(payload["photos"])
            payload["main_photo"] = payload["photos"][0] if payload["photos"] else None
            payload["sticker_url"] = browser_bundle.get("sticker_url") or payload.get("sticker_url")
            payload["carfax_url"] = browser_bundle.get("carfax_url") or payload.get("carfax_url")
            if browser_bundle.get("carfax_facts"):
                merged_badge_facts = dict(payload.get("carfax_facts") if isinstance(payload.get("carfax_facts"), dict) else {})
                merged_badge_facts.update(browser_bundle.get("carfax_facts") or {})
                payload["carfax_facts"] = merged_badge_facts
            payload["detail_fetch_url"] = browser_bundle.get("current_url")
            payload["detail_page_title"] = browser_bundle.get("page_title")
        else:
            payload["detail_fetch_error"] = browser_bundle.get("error")
        try:
            with httpx.Client(timeout=15.0, follow_redirects=True) as client:
                response = client.get(detail_url)
            if response.status_code < 400:
                links = _extract_asset_links_from_html(response.text, base_url=detail_url)
                if not payload.get("sticker_url"):
                    payload["sticker_url"] = links.get("sticker_url")
                if not payload.get("carfax_url"):
                    payload["carfax_url"] = links.get("carfax_url")
                carfax_facts = _extract_carfax_facts_from_html(response.text)
                if carfax_facts and not payload.get("carfax_facts"):
                    payload["carfax_facts"] = carfax_facts
                quick_specs = _extract_quick_specs_from_html(response.text)
                if quick_specs:
                    merged_specs = dict(payload.get("quick_specs") if isinstance(payload.get("quick_specs"), dict) else {})
                    merged_specs.update(quick_specs)
                    payload["quick_specs"] = merged_specs
                payload["detail_fetch_status"] = response.status_code
            else:
                payload["detail_fetch_status"] = response.status_code
        except Exception as exc:
            if not payload.get("detail_fetch_error"):
                payload["detail_fetch_error"] = str(exc)

    if payload.get("carfax_url"):
        should_refresh_carfax_report = True
        existing_report = existing.get("carfax_report") if isinstance(existing.get("carfax_report"), dict) else payload.get("carfax_report")
        if isinstance(existing_report, dict) and existing_report.get("ok"):
            if not _needs_carfax_browser_fallback(existing_report):
                should_refresh_carfax_report = False
        if should_refresh_carfax_report:
            report_details = _fetch_carfax_report_details(
                str(payload.get("carfax_url") or ""),
                vehicle=vehicle,
                expected_vin=clean_vin,
            )
            if (
                isinstance(existing_report, dict)
                and existing_report.get("ok")
                and _carfax_report_matches_vehicle(existing_report, vehicle=vehicle, expected_vin=clean_vin)
                and not _carfax_report_matches_vehicle(report_details, vehicle=vehicle, expected_vin=clean_vin)
            ):
                report_details = existing_report
        else:
            report_details = existing_report
        payload["carfax_report"] = report_details
        payload["carfax_facts"] = _merge_carfax_facts(
            payload.get("carfax_facts") if isinstance(payload.get("carfax_facts"), dict) else {},
            report_details,
        )

    payload["photos"] = _dedupe_urls(
        [url for url in (_extract_photo_url(entry) for entry in (payload.get("photos") or [])) if url]
    )
    payload["photos_count"] = len(payload["photos"])
    payload["main_photo"] = payload["photos"][0] if payload["photos"] else None
    payload.update(_vehicle_asset_summary_payload(vehicle=vehicle, assets=payload))

    _safe_write_json(cache_path, payload)
    if isinstance(payload.get("carfax_summary"), dict):
        CARFAX_SUMMARY_DIR.mkdir(parents=True, exist_ok=True)
        _safe_write_json(CARFAX_SUMMARY_DIR / f"{clean_vin}.json", payload["carfax_summary"])
    return payload


def _refresh_carfax_report_for_vin(vin: str, *, visible_browser: bool = False) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    vehicle = _find_vehicle_by_vin(clean_vin)
    cache_path = _vehicle_assets_cache_path(clean_vin)
    assets = _safe_read_json(cache_path, {})
    if not vehicle and isinstance(assets, dict) and assets.get("vin") == clean_vin:
        vehicle = {
            "vin": clean_vin,
            "detail_url": assets.get("detail_url"),
            "photos": assets.get("photos") if isinstance(assets.get("photos"), list) else [],
            "carfax_url": assets.get("carfax_url"),
            "carfax_facts": assets.get("carfax_facts") if isinstance(assets.get("carfax_facts"), dict) else {},
        }
    if not vehicle:
        raise HTTPException(status_code=404, detail={"message": f"Vehicle not found for VIN {clean_vin}"})
    if not isinstance(assets, dict) or assets.get("vin") != clean_vin:
        assets = _load_vehicle_assets(clean_vin, refresh=True)
    carfax_url = str(assets.get("carfax_url") or vehicle.get("carfax_url") or "").strip()
    if not carfax_url:
        fresh = _load_vehicle_assets(clean_vin, refresh=True)
        carfax_url = str(fresh.get("carfax_url") or "").strip()
        assets = fresh
    if not carfax_url:
        raise HTTPException(status_code=404, detail={"message": "No CARFAX URL found for this VIN."})

    report_details = _fetch_carfax_report_details(
        carfax_url,
        visible_browser=visible_browser,
        vehicle=vehicle,
        expected_vin=clean_vin,
    )
    assets["carfax_url"] = carfax_url
    assets["carfax_report"] = report_details
    assets["carfax_facts"] = _merge_carfax_facts(
        assets.get("carfax_facts") if isinstance(assets.get("carfax_facts"), dict) else {},
        report_details,
    )
    assets.update(_vehicle_asset_summary_payload(vehicle=vehicle, assets=assets))
    _safe_write_json(cache_path, assets)
    if isinstance(assets.get("carfax_summary"), dict):
        CARFAX_SUMMARY_DIR.mkdir(parents=True, exist_ok=True)
        _safe_write_json(CARFAX_SUMMARY_DIR / f"{clean_vin}.json", assets["carfax_summary"])
    return assets


def _save_carfax_report_text(vin: str, report_text: str, source_url: str | None = None) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        raise HTTPException(status_code=400, detail={"message": "vin is required"})
    vehicle = _find_vehicle_by_vin(clean_vin)
    cache_path = _vehicle_assets_cache_path(clean_vin)
    assets = _safe_read_json(cache_path, {})
    if not vehicle and isinstance(assets, dict) and assets.get("vin") == clean_vin:
        vehicle = {
            "vin": clean_vin,
            "detail_url": assets.get("detail_url"),
            "photos": assets.get("photos") if isinstance(assets.get("photos"), list) else [],
            "carfax_url": assets.get("carfax_url"),
            "carfax_facts": assets.get("carfax_facts") if isinstance(assets.get("carfax_facts"), dict) else {},
        }
    if not vehicle:
        raise HTTPException(status_code=404, detail={"message": f"Vehicle not found for VIN {clean_vin}"})
    vehicle = dict(vehicle)
    vehicle["vin"] = clean_vin

    if not isinstance(assets, dict) or assets.get("vin") != clean_vin:
        assets = _load_vehicle_assets(clean_vin, refresh=False)
    assets["vin"] = clean_vin
    if source_url:
        assets["carfax_url"] = source_url

    parsed = _extract_carfax_report_facts(
        report_text,
        source_url=source_url or str(assets.get("carfax_url") or ""),
        source="manual_carfax_report",
    )
    parsed = _reject_carfax_identity_mismatch(parsed, vehicle=vehicle, expected_vin=clean_vin)
    if not parsed.get("ok"):
        raise HTTPException(
            status_code=400,
            detail={
                "message": parsed.get("report_access") or "CARFAX report text could not be parsed.",
                "parsed": parsed,
            },
        )

    assets["carfax_report"] = parsed
    assets["carfax_facts"] = _merge_carfax_facts(
        assets.get("carfax_facts") if isinstance(assets.get("carfax_facts"), dict) else {},
        parsed,
    )
    assets.update(_vehicle_asset_summary_payload(vehicle=vehicle, assets=assets))
    _safe_write_json(cache_path, assets)
    if isinstance(assets.get("carfax_summary"), dict):
        CARFAX_SUMMARY_DIR.mkdir(parents=True, exist_ok=True)
        _safe_write_json(CARFAX_SUMMARY_DIR / f"{clean_vin}.json", assets["carfax_summary"])
    return {
        "ok": True,
        "vin": clean_vin,
        "parsed": parsed,
        "assets": assets,
    }


def _asset_cache_paths(vin: str, kind: str) -> tuple[Path, Path]:
    clean_vin = re.sub(r"[^A-Z0-9]", "", str(vin or "").upper()) or "UNKNOWN"
    clean_kind = re.sub(r"[^a-z0-9_-]", "", str(kind or "").lower()) or "asset"
    directory = VEHICLE_ASSET_FILE_CACHE_DIR / clean_vin
    directory.mkdir(parents=True, exist_ok=True)
    return directory / f"{clean_kind}.bin", directory / f"{clean_kind}.json"


def _extension_for_content_type(content_type: str, source_url: str) -> str:
    lowered = str(content_type or "").lower()
    if "pdf" in lowered:
        return ".pdf"
    if "html" in lowered:
        return ".html"
    if "png" in lowered:
        return ".png"
    if "webp" in lowered:
        return ".webp"
    if "jpeg" in lowered or "jpg" in lowered:
        return ".jpg"
    suffix = Path(urlparse(source_url).path).suffix.lower()
    return suffix if suffix in {".pdf", ".html", ".htm", ".png", ".jpg", ".jpeg", ".webp"} else ".bin"


def _write_cached_asset_response(
    *,
    vin: str,
    kind: str,
    source_url: str,
    response: httpx.Response,
    content_path: Path,
) -> dict[str, Any]:
    content_type = response.headers.get("content-type", "application/octet-stream").split(";")[0].strip()
    if response.status_code >= 400:
        raise RuntimeError(f"asset HTTP {response.status_code}")
    if len(response.content) > 30 * 1024 * 1024:
        raise RuntimeError("asset larger than 30MB cache limit")
    extension = _extension_for_content_type(content_type, str(response.url))
    target = content_path.with_suffix(extension)
    target.write_bytes(response.content)
    return {
        "ok": True,
        "vin": str(vin or "").upper(),
        "kind": kind,
        "source_url": source_url,
        "final_url": str(response.url),
        "content_type": content_type,
        "path": str(target),
        "bytes": len(response.content),
        "cached_at": datetime.now(timezone.utc).isoformat(),
    }


def _cache_remote_vehicle_asset_with_browser_cookie(
    *,
    vin: str,
    kind: str,
    source_url: str,
    content_path: Path,
) -> dict[str, Any] | None:
    if webdriver is None or ChromeService is None:
        return None
    vehicle = _find_vehicle_by_vin(vin)
    detail_url = str((vehicle or {}).get("detail_url") or "").strip()
    if not detail_url:
        return None

    driver = None
    try:
        driver = _open_headless_browser(timeout_seconds=24.0)
        driver.get(detail_url)
        time.sleep(2)
        cookies = {cookie["name"]: cookie["value"] for cookie in driver.get_cookies() if cookie.get("name") and cookie.get("value")}
        user_agent = driver.execute_script("return navigator.userAgent") or "Mozilla/5.0 XconsoleDealershipAssetCache/1.0"
        headers = {
            "User-Agent": str(user_agent),
            "Accept": "application/pdf,text/html,image/avif,image/webp,image/apng,*/*;q=0.8",
            "Referer": detail_url,
        }
        with httpx.Client(timeout=30.0, follow_redirects=True, headers=headers, cookies=cookies) as client:
            response = client.get(source_url)
        meta = _write_cached_asset_response(
            vin=vin,
            kind=kind,
            source_url=source_url,
            response=response,
            content_path=content_path,
        )
        meta["source_mode"] = "browser_cookie"
        return meta
    except Exception:
        return None
    finally:
        if driver is not None:
            try:
                driver.quit()
            except Exception:
                pass


def _cache_remote_vehicle_asset(vin: str, kind: str, source_url: str, *, force: bool = False) -> dict[str, Any]:
    content_path, meta_path = _asset_cache_paths(vin, kind)
    meta = _safe_read_json(meta_path, {})
    if (
        isinstance(meta, dict)
        and meta.get("ok")
        and meta.get("source_url") == source_url
        and meta.get("path")
        and Path(str(meta.get("path"))).exists()
        and not force
    ):
        return meta

    headers = {
        "User-Agent": "Mozilla/5.0 XconsoleDealershipAssetCache/1.0",
        "Accept": "text/html,application/pdf,image/avif,image/webp,image/apng,*/*;q=0.8",
    }
    try:
        with httpx.Client(timeout=30.0, follow_redirects=True, headers=headers) as client:
            response = client.get(source_url)
        meta = _write_cached_asset_response(
            vin=vin,
            kind=kind,
            source_url=source_url,
            response=response,
            content_path=content_path,
        )
    except Exception as exc:
        browser_meta = _cache_remote_vehicle_asset_with_browser_cookie(
            vin=vin,
            kind=kind,
            source_url=source_url,
            content_path=content_path,
        )
        meta = browser_meta or {
            "ok": False,
            "vin": str(vin or "").upper(),
            "kind": kind,
            "source_url": source_url,
            "error": str(exc),
            "cached_at": datetime.now(timezone.utc).isoformat(),
        }
    _safe_write_json(meta_path, meta)
    return meta


def _is_generic_carfax_summary(summary: dict[str, Any] | None) -> bool:
    if not isinstance(summary, dict):
        return True
    text = str(summary.get("summary") or "").lower()
    highlights = " ".join(str(item) for item in summary.get("highlights") or []).lower()
    combined = f"{text} {highlights}"
    generic_markers = [
        "official carfax link",
        "link cached",
        "carfax provides the most accident",
        "confirm accident history",
        "no carfax link",
        "badge facts from the dealer page",
        "linked carfax badge does not expose",
    ]
    return not text or any(marker in combined for marker in generic_markers)


_CARFAX_PLACEHOLDER_MARKERS = (
    "linked carfax badge does not expose",
    "not clearly parsed",
    "not parsed",
    "not available",
    "needs official report",
    "badge facts from the dealer listing",
    "no data",
    "not parsed from report text",
    "data not parsed",
)


def _is_placeholder_carfax_text(value: Any) -> bool:
    text = str(value or "").strip().lower()
    if not text:
        return True
    if text in {"n/a", "na", "unknown", "none", "null", ""}:
        return True
    return any(marker in text for marker in _CARFAX_PLACEHOLDER_MARKERS)


def _carfax_fact_has_value(value: Any) -> bool:
    return bool(str(value or "").strip()) and not _is_placeholder_carfax_text(value)


def _has_structured_carfax_facts(facts: dict[str, Any] | None) -> bool:
    if not isinstance(facts, dict):
        return False
    keys = (
        "accident_damage",
        "title_brand",
        "service_history",
        "usage",
        "market_position",
        "market_delta",
        "carfax_value",
    )
    for key in keys:
        value = facts.get(key)
        if isinstance(value, list):
            if value:
                return True
        elif _carfax_fact_has_value(value):
            return True
    # Owner/value alone from badge metadata are helpful, but do not count as
    # full report detail when we still need richer parsed facts.
    for key in ("accident_events", "service_events"):
        value = facts.get(key)
        if isinstance(value, list) and value:
            return True
    return False


def _carfax_summary_from_assets(clean_vin: str, assets: dict[str, Any]) -> dict[str, Any]:
    stored = _carfax_summary_for_vin(clean_vin)
    facts = assets.get("carfax_facts") if isinstance(assets, dict) else {}
    if isinstance(facts, dict) and facts.get("identity_mismatch"):
        facts = _strip_carfax_structured_report_fields(facts)
    facts_source = str(facts.get("source") or "") if isinstance(facts, dict) else ""
    report_sources = {"carfax_report", "manual_carfax_report", "carfax_report_browser", "carfax_report_browser_visible", "carfax_report_plain_chrome"}
    facts_has_report_summary = isinstance(facts, dict) and (facts_source in report_sources or bool(facts.get("summary")))
    has_structured_facts = _has_structured_carfax_facts(facts)
    if stored and not _is_generic_carfax_summary(stored) and not facts_has_report_summary and has_structured_facts:
        return stored
    if isinstance(facts, dict) and facts:
        owner = facts.get("owner_count") or None
        value = facts.get("value_badge")
        source = str(facts.get("source") or "").strip()
        if facts.get("summary"):
            summary = str(facts.get("summary"))
        elif source in report_sources:
            parts = []
            if owner:
                parts.append(f"Owners: {owner}")
            accident = str(facts.get("accident_damage") or "").strip()
            if accident and not _is_placeholder_carfax_text(accident):
                parts.append(accident)
            service = str(facts.get("service_history") or "").strip()
            if service and not _is_placeholder_carfax_text(service):
                parts.append(service)
            market = " / ".join(str(part) for part in [value, facts.get("market_position"), facts.get("market_delta")] if part)
            if market:
                parts.append(f"Market: {market}")
            summary = "CARFAX report parsed: " + "; ".join(part for part in parts if part) + "."
            if not parts:
                summary = "CARFAX report parsed, but no detailed fields were extracted yet."
        else:
            parts = []
            if owner:
                parts.append(f"Owners: {owner}")
            if value:
                parts.append(f"Value badge: {value}")
            summary = "CARFAX badge facts from the dealer page: " + "; ".join(parts) + "."
            if not parts:
                summary = "CARFAX badge linked, but detailed fields were not exposed to Xconsole yet."
        highlights = [str(item) for item in facts.get("highlights") or [] if str(item).strip()]
        for label, key in [
            ("Accident/damage", "accident_damage"),
            ("Title/brand", "title_brand"),
            ("Service history", "service_history"),
            ("Use type", "usage"),
            ("Market/value", "market_position"),
            ("Market delta", "market_delta"),
            ("CARFAX value", "carfax_value"),
            ("Report access", "report_access"),
        ]:
            value_text = str(facts.get(key) or "").strip()
            if value_text and not _is_placeholder_carfax_text(value_text):
                highlights.append(f"{label}: {value_text}")
        return {
            "vin": clean_vin,
            "summary": summary,
            "highlights": _dedupe_strings(highlights)[:14],
            "facts": facts,
            "source": facts.get("source") or "dealer_page_carfax_badge",
            "updated_at": datetime.now(timezone.utc).isoformat(),
        }

    return {
        "vin": clean_vin,
        "summary": (
            "Official CARFAX report link is present, but Xconsole has not parsed report facts yet. "
            "Refresh assets or open the official report to verify owners, accidents, title, and service history."
        )
        if assets.get("carfax_url")
        else "No CARFAX link has been found yet.",
        "highlights": ["CARFAX report link present; facts not parsed yet."] if assets.get("carfax_url") else [],
        "facts": {},
        "source": "unparsed_report_link",
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }


def _vehicle_asset_summary_payload(*, vehicle: dict[str, Any], assets: dict[str, Any]) -> dict[str, Any]:
    clean_vin = str(vehicle.get("vin") or assets.get("vin") or "").strip().upper()
    vehicle_kind, buyer, default_features = _vehicle_kind_and_buyer(vehicle)
    highlights = _vehicle_marketing_highlights(vehicle)
    quick_specs = assets.get("quick_specs") if isinstance(assets.get("quick_specs"), dict) else {}
    quick_spec_highlights = _quick_spec_highlights(quick_specs)
    sticker_highlights = [
        item
        for item in [
            vehicle.get("engine") and f"Engine: {vehicle.get('engine')}",
            vehicle.get("transmission") and f"Transmission: {vehicle.get('transmission')}",
            vehicle.get("drivetrain") and f"Drivetrain: {vehicle.get('drivetrain')}",
            vehicle.get("exterior") and f"Exterior: {vehicle.get('exterior')}",
            vehicle.get("interior") and f"Interior: {vehicle.get('interior')}",
            assets.get("sticker_url") and "Window sticker available in Xconsole",
        ]
        if item
    ]
    sticker_highlights = _dedupe_strings([*quick_spec_highlights, *sticker_highlights])
    carfax_summary = _carfax_summary_from_assets(clean_vin, assets)
    buyer_profile = {
        "kind": vehicle_kind,
        "buyer": buyer,
        "features": highlights or default_features,
    }
    carfax_lines = _carfax_buyer_facing_lines(carfax_summary if isinstance(carfax_summary, dict) else None)
    marketing_summary = [
        f"Ideal buyer: {buyer}.",
        *(sticker_highlights[:4]),
        *(f"CARFAX: {line}" for line in carfax_lines[:3]),
    ]
    return {
        "sticker_highlights": sticker_highlights[:8],
        "carfax_summary": carfax_summary,
        "buyer_profile": buyer_profile,
        "marketing_summary": [line for line in marketing_summary if line],
        "sticker_view_url": f"/api/vehicles/{clean_vin}/asset-view/sticker" if assets.get("sticker_url") else None,
        "carfax_view_url": f"/api/vehicles/{clean_vin}/asset-view/carfax" if assets.get("carfax_url") else None,
    }


def _asset_message_html(title: str, body: str, *, source_url: str | None = None) -> str:
    escaped_title = html.escape(title)
    escaped_body = html.escape(body)
    source = html.escape(source_url or "")
    link = f'<p class="source">{source}</p>' if source else ""
    return f"""<!doctype html>
<html><head><meta charset="utf-8" />
<style>
body{{margin:0;font-family:Inter,Segoe UI,Arial,sans-serif;background:#f7fafc;color:#102033;padding:28px;line-height:1.45}}
.box{{max-width:860px;margin:auto;background:#fff;border:1px solid #d9e3ec;border-radius:14px;padding:24px;box-shadow:0 12px 30px rgba(15,35,55,.08)}}
h1{{margin:0 0 10px;font-size:22px}} p{{font-size:15px}} .source{{color:#64748b;word-break:break-all;font-size:12px}}
</style></head><body><main class="box"><h1>{escaped_title}</h1><p>{escaped_body}</p>{link}</main></body></html>"""


def _carfax_summary_html(vehicle: dict[str, Any], assets: dict[str, Any]) -> str:
    summary_payload = _vehicle_asset_summary_payload(vehicle=vehicle, assets=assets)
    carfax = summary_payload.get("carfax_summary") if isinstance(summary_payload, dict) else {}
    buyer = summary_payload.get("buyer_profile") if isinstance(summary_payload, dict) else {}
    highlights = carfax.get("highlights") if isinstance(carfax, dict) else []
    features = buyer.get("features") if isinstance(buyer, dict) else []
    facts = carfax.get("facts") if isinstance(carfax, dict) else {}
    def lis(values: Any) -> str:
        if not isinstance(values, list) or not values:
            return "<li>No cached detail yet.</li>"
        return "".join(f"<li>{html.escape(str(value))}</li>" for value in values[:12])

    def fact_text(value: Any) -> str:
        if isinstance(value, dict):
            parts = [f"{key}: {val}" for key, val in value.items() if val not in (None, "", [], {})]
            return ", ".join(parts)
        if isinstance(value, list):
            return "; ".join(str(item) for item in value[:4])
        return str(value or "").strip()

    def fact_card(label: str, key: str, fallback: str = "Not parsed yet") -> str:
        value: Any = ""
        if isinstance(facts, dict):
            value = facts.get(key)
        rendered = fact_text(value)
        return f"<div><span>{html.escape(label)}</span><strong>{html.escape(rendered or fallback)}</strong></div>"

    def event_list(values: Any, *, kind: str) -> str:
        if not isinstance(values, list) or not values:
            return "<li>No parsed events yet.</li>"
        items: list[str] = []
        for event in values[:10]:
            if isinstance(event, dict):
                date = html.escape(str(event.get("date") or "date not parsed"))
                severity = html.escape(str(event.get("severity") or kind))
                description = html.escape(str(event.get("description") or ""))
                items.append(f"<li><strong>{date}</strong> <span>{severity}</span><p>{description}</p></li>")
            else:
                items.append(f"<li>{html.escape(str(event))}</li>")
        return "".join(items)

    title = html.escape(str(vehicle.get("title") or vehicle.get("vin") or "Vehicle"))
    summary = html.escape(str(carfax.get("summary") if isinstance(carfax, dict) else "No cached CARFAX summary."))
    buyer_text = html.escape(str(buyer.get("buyer") if isinstance(buyer, dict) else "buyer"))
    source = html.escape(str(assets.get("carfax_url") or ""))
    source_json = json.dumps(str(assets.get("carfax_url") or ""))
    vin_json = json.dumps(str(vehicle.get("vin") or assets.get("vin") or ""))
    clean_vin = html.escape(str(vehicle.get("vin") or assets.get("vin") or ""))
    report_access = html.escape(str(facts.get("report_access") if isinstance(facts, dict) else ""))
    accident_events = facts.get("accident_events") if isinstance(facts, dict) else []
    service_events = facts.get("service_events") if isinstance(facts, dict) else []
    return f"""<!doctype html>
<html><head><meta charset="utf-8" />
<style>
body{{margin:0;background:#f5f8fb;color:#102033;font-family:Inter,Segoe UI,Arial,sans-serif;padding:24px;line-height:1.42}}
.shell{{max-width:980px;margin:auto;display:grid;gap:16px}}
.card{{background:white;border:1px solid #dbe5ef;border-radius:14px;padding:18px;box-shadow:0 12px 28px rgba(15,35,55,.07)}}
.facts{{display:grid;grid-template-columns:repeat(2,minmax(0,1fr));gap:12px}}
.facts div{{border:1px solid #dbe5ef;border-radius:12px;background:#f8fbfe;padding:12px}}
.facts span{{display:block;color:#64748b;font-size:12px;text-transform:uppercase;letter-spacing:.08em}}
.facts strong{{display:block;margin-top:5px;font-size:18px}}
.timeline{{display:grid;gap:10px;margin:0;padding:0;list-style:none}}
.timeline li{{border:1px solid #dbe5ef;border-radius:12px;background:#f8fbfe;padding:12px;margin:0}}
.timeline li strong{{display:inline-block;margin-right:8px}}
.timeline li span{{color:#42617c;text-transform:uppercase;font-size:12px;letter-spacing:.06em}}
.timeline li p{{margin:6px 0 0}}
.source-link{{width:max-content;border:1px solid #2f6fa8;border-radius:10px;background:#1f5f98;color:white;padding:9px 12px;text-decoration:none;cursor:pointer;display:inline-block}}
.actions{{display:flex;gap:10px;flex-wrap:wrap;margin-top:12px}}
.notice{{border-color:#f1d08a;background:#fff9e8}}
h1{{font-size:24px;margin:0}} h2{{font-size:14px;text-transform:uppercase;letter-spacing:.08em;color:#42617c;margin:0 0 10px}}
p{{font-size:15px;margin:8px 0}} ul{{margin:0;padding-left:20px}} li{{margin:7px 0}} .source{{font-size:12px;color:#64748b;word-break:break-all}}
</style></head><body><main class="shell">
<section class="card"><h1>{title}</h1><p>{summary}</p><p class="source">{source}</p></section>
<section class="card"><h2>Report Facts</h2><div class="facts">
{fact_card("Owners", "owner_count")}
{fact_card("Value Badge", "value_badge")}
{fact_card("Market Position", "market_position")}
{fact_card("Market Delta", "market_delta")}
{fact_card("CARFAX Value", "carfax_value")}
{fact_card("Accident / Damage", "accident_damage")}
{fact_card("Accident Counts", "accident_counts")}
{fact_card("Title / Brand", "title_brand")}
{fact_card("Service History", "service_history")}
{fact_card("Service Records", "service_records_count")}
{fact_card("Last Service", "last_service_date")}
{fact_card("Use Type", "usage")}
</div></section>
<section class="card"><h2>Accident / Damage Timeline</h2><ul class="timeline">{event_list(accident_events, kind="accident")}</ul></section>
<section class="card"><h2>Service History Timeline</h2><ul class="timeline">{event_list(service_events, kind="service")}</ul></section>
<section class="card"><h2>CARFAX Highlights</h2><ul>{lis(highlights)}</ul></section>
<section class="card"><h2>Ideal Buyer</h2><p>{buyer_text}</p><ul>{lis(features)}</ul></section>
<section class="card notice"><h2>Linked Report Access</h2><p>{report_access or "Xconsole found the linked CARFAX report and will use any report facts exposed to the automated fetch."}</p><div class="actions">{f'<a class="source-link" href="/api/vehicles/{clean_vin}/asset-view/carfax?refresh=true&visible=true">Retry Full Parse</a>' if clean_vin else ''}{f'<a class="source-link" href="{source}" target="_blank" rel="noreferrer">View Linked CARFAX</a>' if source else ''}</div></section>
<section class="card"><h2>Use In Listing</h2><p>Use the cached facts above in Xconsole. If CARFAX exposes the full linked report to the automated browser, accident, service, title, use, and market-value detail will populate here automatically.</p></section>
</main>
</body></html>"""


def _load_accounts() -> list[dict[str, Any]]:
    accounts = _load_accounts_full()
    redacted: list[dict[str, Any]] = []
    for account in accounts:
        if not isinstance(account, dict):
            continue
        redacted.append(
            {
                "id": account.get("id"),
                "name": account.get("name"),
                "email": account.get("email"),
                "has_password": bool(account.get("password")),
            }
        )
    return redacted


def _env_facebook_account() -> dict[str, Any] | None:
    env_id = str(os.getenv("FACEBOOK_LOGIN_ACCOUNT_ID") or os.getenv("FACEBOOK_ACCOUNT_ID") or "").strip()
    env_email = str(os.getenv("FACEBOOK_LOGIN_EMAIL") or os.getenv("FACEBOOK_ACCOUNT_EMAIL") or "").strip()
    env_password = str(os.getenv("FACEBOOK_LOGIN_PASSWORD") or os.getenv("FACEBOOK_ACCOUNT_PASSWORD") or "").strip()
    if not env_id or not env_email:
        return None
    return {
        "id": env_id,
        "name": os.getenv("FACEBOOK_LOGIN_NAME") or "Facebook Login",
        "email": env_email,
        "password": env_password,
    }


def _ensure_facebook_accounts_file() -> None:
    if FML_ACCOUNTS_PATH.exists():
        return
    account = _env_facebook_account()
    if not account:
        return
    FML_ACCOUNTS_PATH.parent.mkdir(parents=True, exist_ok=True)
    FML_ACCOUNTS_PATH.write_text(
        json.dumps({"accounts": [account]}, indent=2),
        encoding="utf-8",
    )


def _load_accounts_full() -> list[dict[str, Any]]:
    _ensure_facebook_accounts_file()
    payload = _safe_read_json(FML_ACCOUNTS_PATH, {"accounts": []})
    accounts = payload.get("accounts", []) if isinstance(payload, dict) else []
    normalized = [entry for entry in accounts if isinstance(entry, dict)]

    env_account = _env_facebook_account()
    if env_account:
        env_id = str(env_account.get("id") or "").strip()
        if not any(str(entry.get("id", "")).strip() == env_id for entry in normalized):
            normalized.append(env_account)
    return normalized


def _resolve_default_account_id() -> str | None:
    accounts = _load_accounts_full()
    preferred = next(
        (
            str(entry.get("id", "")).strip()
            for entry in accounts
            if str(entry.get("id", "")).strip() and entry.get("password")
        ),
        "",
    )
    if preferred:
        return preferred

    fallback = next(
        (str(entry.get("id", "")).strip() for entry in accounts if str(entry.get("id", "")).strip()),
        "",
    )
    return fallback or None


def _facebook_session_cookie_path() -> Path:
    return RUNTIME_DIR / "facebook_session_cookies.json"


def _has_saved_facebook_session() -> bool:
    if runtime_session_available(ROOT_DIR):
        return True
    latest = latest_saved_session(ROOT_DIR)
    return bool(latest)


def _normalize_image_names(images: list[str]) -> list[str]:
    normalized: list[str] = []
    seen: set[str] = set()
    for raw in images:
        item = str(raw or "").strip()
        if not item:
            continue
        lowered = item.lower()
        if lowered in seen:
            continue
        seen.add(lowered)
        normalized.append(item)
    return normalized


def _list_facebook_images(limit: int = 200) -> tuple[list[str], int]:
    if not FML_IMAGES_DIR.exists():
        return [], 0
    allowed_ext = {".jpg", ".jpeg", ".png", ".webp", ".gif", ".bmp"}
    candidates = [
        path.name
        for path in FML_IMAGES_DIR.iterdir()
        if path.is_file()
        and not path.name.startswith(".")
        and path.suffix.lower() in allowed_ext
    ]
    candidates.sort(key=str.casefold)
    return candidates[:limit], len(candidates)


def _suggest_images_for_vin(vin: str, limit: int = 20) -> list[str]:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin or not FML_IMAGES_DIR.exists():
        return []

    first8 = clean_vin[:8]
    last8 = clean_vin[-8:] if len(clean_vin) >= 8 else clean_vin

    scored: list[tuple[int, str]] = []
    for image_path in FML_IMAGES_DIR.iterdir():
        if not image_path.is_file():
            continue
        name = image_path.name
        lowered = name.lower()
        score = 0
        if clean_vin.lower() in lowered:
            score += 100
        if first8 and first8.lower() in lowered:
            score += 20
        if last8 and last8.lower() in lowered:
            score += 30
        if score > 0:
            scored.append((score, name))

    scored.sort(key=lambda item: (-item[0], item[1].casefold()))
    return [name for _, name in scored[:limit]]


def _extract_photo_url(item: Any) -> str | None:
    if isinstance(item, str):
        candidate = item.strip()
        return candidate if candidate.startswith(("http://", "https://")) else None
    if isinstance(item, dict):
        for key in ("url", "src", "image", "photo"):
            value = item.get(key)
            if isinstance(value, str):
                candidate = value.strip()
                if candidate.startswith(("http://", "https://")):
                    return candidate
    return None


def _collect_vehicle_photo_urls(vin: str) -> list[str]:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        return []

    vehicle = next(
        (item for item in _load_inventory_candidates() if str(item.get("vin", "")).strip().upper() == clean_vin),
        None,
    )
    if not vehicle:
        return []

    def _collect_from_payload(payload: Any) -> list[str]:
        photos_raw = payload if isinstance(payload, list) else [payload]
        collected: list[str] = []
        seen: set[str] = set()
        for entry in photos_raw:
            url = _extract_photo_url(entry)
            if not url:
                continue
            lowered = url.lower()
            if lowered in seen:
                continue
            seen.add(lowered)
            collected.append(url)
        return collected

    cached = _read_vehicle_assets_cache(clean_vin)
    cached_photos = _collect_from_payload(cached.get("photos"))
    if cached_photos:
        return cached_photos

    inventory_photos = _collect_from_payload(vehicle.get("photos") or vehicle.get("images") or [])
    if inventory_photos:
        return inventory_photos

    try:
        assets = _load_vehicle_assets(clean_vin, refresh=False)
    except Exception:
        return []
    return _collect_from_payload(assets.get("photos"))


def _image_extension_from_url_and_content_type(url: str, content_type: str | None) -> str:
    path_ext = Path(urlparse(url).path).suffix.lower()
    if path_ext in {".jpg", ".jpeg", ".png", ".webp"}:
        return ".jpg" if path_ext == ".jpeg" else path_ext

    value = (content_type or "").lower()
    if "png" in value:
        return ".png"
    if "webp" in value:
        return ".webp"
    return ".jpg"


def _import_vehicle_images(
    *,
    vin: str,
    limit: int,
    overwrite: bool = False,
) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    urls = _collect_vehicle_photo_urls(clean_vin)
    if not urls:
        raise HTTPException(
            status_code=404,
            detail={
                "message": f"No source photo URLs found for VIN {clean_vin}",
                "vin": clean_vin,
            },
        )

    FML_IMAGES_DIR.mkdir(parents=True, exist_ok=True)

    imported: list[str] = []
    skipped_existing: list[str] = []
    errors: list[str] = []

    with httpx.Client(timeout=20.0, follow_redirects=True) as client:
        for index, url in enumerate(urls[:limit], start=1):
            try:
                response = client.get(url)
                if response.status_code >= 400:
                    raise RuntimeError(f"HTTP {response.status_code}")

                extension = _image_extension_from_url_and_content_type(
                    url,
                    response.headers.get("content-type"),
                )
                filename = f"{clean_vin}_{index:02d}{extension}"
                target = FML_IMAGES_DIR / filename

                if target.exists() and not overwrite:
                    skipped_existing.append(filename)
                    continue

                target.write_bytes(response.content)
                imported.append(filename)
            except Exception as exc:
                errors.append(f"{url} -> {exc}")

    return {
        "ok": len(errors) == 0,
        "vin": clean_vin,
        "source_urls_found": len(urls),
        "attempted": min(limit, len(urls)),
        "imported_count": len(imported),
        "imported": imported,
        "skipped_existing_count": len(skipped_existing),
        "skipped_existing": skipped_existing,
        "errors_count": len(errors),
        "errors": errors,
    }


def _seed_placeholder_images_for_vin(
    *,
    vin: str,
    count: int,
    overwrite: bool = False,
) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        raise HTTPException(status_code=400, detail={"message": "vin is required"})

    FML_IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    png_base64 = "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAQAAAC1HAwCAAAAC0lEQVR42mP8/x8AAwMCAO7Zk1YAAAAASUVORK5CYII="
    image_bytes = base64.b64decode(png_base64)

    created: list[str] = []
    skipped_existing: list[str] = []
    for index in range(1, max(1, int(count)) + 1):
        name = f"{clean_vin}_{index:02d}.png"
        target = FML_IMAGES_DIR / name
        if target.exists() and not overwrite:
            skipped_existing.append(name)
            continue
        target.write_bytes(image_bytes)
        created.append(name)

    return {
        "ok": True,
        "vin": clean_vin,
        "created_count": len(created),
        "created": created,
        "skipped_existing_count": len(skipped_existing),
        "skipped_existing": skipped_existing,
    }


def _relink_images_to_vin(
    *,
    vin: str,
    images: list[str],
    include_vin_matches: bool = False,
    overwrite: bool = False,
    delete_source: bool = False,
) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    if not clean_vin:
        raise HTTPException(status_code=400, detail={"message": "vin is required"})

    FML_IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    source_names = _normalize_image_names(images)
    if include_vin_matches:
        source_names = _normalize_image_names(
            source_names + _suggest_images_for_vin(vin=clean_vin, limit=200)
        )

    if not source_names:
        all_images, _ = _list_facebook_images(limit=2000)
        source_names = all_images

    if not source_names:
        return {
            "ok": False,
            "vin": clean_vin,
            "message": "No source images found to relink",
            "linked_count": 0,
            "linked": [],
            "missing_count": 0,
            "missing": [],
            "skipped_existing_count": 0,
            "skipped_existing": [],
            "already_linked_count": 0,
            "already_linked": [],
            "errors_count": 0,
            "errors": [],
        }

    linked: list[str] = []
    missing: list[str] = []
    skipped_existing: list[str] = []
    already_linked: list[str] = []
    errors: list[str] = []

    target_index = 1
    for source_name in source_names:
        source_path = FML_IMAGES_DIR / source_name
        if not source_path.exists() or not source_path.is_file():
            missing.append(source_name)
            continue

        extension = source_path.suffix.lower() or ".jpg"
        if extension == ".jpeg":
            extension = ".jpg"
        target_name = f"{clean_vin}_{target_index:02d}{extension}"
        target_index += 1
        target_path = FML_IMAGES_DIR / target_name

        if source_path.name.lower() == target_name.lower():
            already_linked.append(source_path.name)
            linked.append(target_name)
            continue

        if target_path.exists() and not overwrite:
            skipped_existing.append(target_name)
            continue

        try:
            if delete_source:
                if target_path.exists():
                    target_path.unlink()
                source_path.rename(target_path)
            else:
                shutil.copy2(source_path, target_path)
            linked.append(target_name)
        except Exception as exc:
            errors.append(f"{source_name} -> {exc}")

    return {
        "ok": len(errors) == 0,
        "vin": clean_vin,
        "linked_count": len(linked),
        "linked": linked,
        "missing_count": len(missing),
        "missing": missing,
        "skipped_existing_count": len(skipped_existing),
        "skipped_existing": skipped_existing,
        "already_linked_count": len(already_linked),
        "already_linked": already_linked,
        "errors_count": len(errors),
        "errors": errors,
    }


def _find_chromedriver() -> Path | None:
    env_driver = str(os.getenv("CHROMEDRIVER_PATH", "")).strip()
    if env_driver:
        candidate = Path(env_driver)
        if candidate.exists():
            return candidate
    if not _running_on_windows():
        for command in ("chromedriver", "chromium-driver"):
            resolved = shutil.which(command)
            if resolved:
                return Path(resolved)
    if not FML_DRIVERS_DIR.exists():
        for command in ("chromedriver", "chromedriver.exe"):
            resolved = shutil.which(command)
            if resolved:
                return Path(resolved)
        return None
    direct_names = ("chromedriver.exe", "chromedriver") if _running_on_windows() else ("chromedriver",)
    for direct_name in direct_names:
        direct = FML_DRIVERS_DIR / direct_name
        if direct.exists():
            return direct
    for candidate in FML_DRIVERS_DIR.glob("chromedriver*"):
        if candidate.is_file():
            if not _running_on_windows() and candidate.suffix.lower() == ".exe":
                continue
            return candidate
    for command in ("chromedriver", "chromedriver.exe"):
        resolved = shutil.which(command)
        if resolved:
            return Path(resolved)
    return None


def _find_chrome_binary() -> Path | None:
    env_binary = str(os.getenv("CHROME_BINARY", "")).strip()
    if env_binary:
        candidate = Path(env_binary)
        if candidate.exists():
            return candidate

    if not _running_on_windows():
        for command in ("chromium", "chromium-browser", "google-chrome", "chrome"):
            resolved = shutil.which(command)
            if resolved:
                return Path(resolved)

    local_candidates = [
        FML_DIR / "chrome-for-testing" / "chrome-win64" / "chrome.exe",
        FML_DIR / "chrome-114" / "chrome-win64" / "chrome.exe",
        FML_DIR / "chrome-win64" / "chrome.exe",
        Path(r"C:\Program Files\Google\Chrome\Application\chrome.exe"),
        Path(r"C:\Program Files (x86)\Google\Chrome\Application\chrome.exe"),
    ] if _running_on_windows() else []
    for candidate in local_candidates:
        if candidate.exists():
            return candidate

    for command in ("chrome.exe", "chrome", "chromium", "chromium-browser", "google-chrome"):
        resolved = shutil.which(command)
        if resolved:
            return Path(resolved)
    return None


def _chromedriver_details(chromedriver: Path | None) -> dict[str, Any]:
    if not chromedriver:
        return {
            "found": False,
            "path": None,
            "version": None,
            "error": None,
        }

    details: dict[str, Any] = {
        "found": True,
        "path": str(chromedriver),
        "version": None,
        "error": None,
    }
    try:
        completed = subprocess.run(
            [str(chromedriver), "--version"],
            capture_output=True,
            text=True,
            check=False,
            timeout=3,
        )
        output = (completed.stdout or completed.stderr or "").strip()
        if completed.returncode == 0 and output:
            details["version"] = output.splitlines()[0].strip()
        elif output:
            details["error"] = output.splitlines()[0].strip()
        else:
            details["error"] = f"chromedriver exited with code {completed.returncode}"
    except Exception as exc:
        details["error"] = str(exc)

    return details


def _chrome_binary_details(binary: Path | None) -> dict[str, Any]:
    if not binary:
        return {
            "found": False,
            "path": None,
            "version": None,
            "error": None,
        }

    details: dict[str, Any] = {
        "found": True,
        "path": str(binary),
        "version": None,
        "error": None,
    }
    try:
        completed = subprocess.run(
            [str(binary), "--version"],
            capture_output=True,
            text=True,
            check=False,
            timeout=3,
        )
        output = (completed.stdout or completed.stderr or "").strip()
        if completed.returncode == 0 and output:
            details["version"] = output.splitlines()[0].strip()
        elif output:
            details["error"] = output.splitlines()[0].strip()
        else:
            details["error"] = f"chrome exited with code {completed.returncode}"
    except Exception as exc:
        details["error"] = str(exc)

    return details


def _bootstrap_facebook_lister(
    *,
    create_template_account_if_missing: bool = True,
) -> dict[str, Any]:
    created_dirs: list[str] = []
    created_files: list[str] = []

    for path in [FML_DIR, FML_IMAGES_DIR, FML_DRIVERS_DIR, FML_DIR / "jobs"]:
        if not path.exists():
            path.mkdir(parents=True, exist_ok=True)
            created_dirs.append(str(path))

    if create_template_account_if_missing and not FML_ACCOUNTS_PATH.exists():
        env_account = _env_facebook_account()
        template = {"accounts": [env_account]} if env_account else {
            "accounts": [
                {
                    "id": "REPLACE_WITH_FACEBOOK_ACCOUNT_ID",
                    "name": "Replace Me",
                    "email": "you@example.com",
                    "password": "",
                }
            ]
        }
        FML_ACCOUNTS_PATH.write_text(
            json.dumps(template, indent=2),
            encoding="utf-8",
        )
        created_files.append(str(FML_ACCOUNTS_PATH))

    for marker_dir in [FML_IMAGES_DIR, FML_DRIVERS_DIR]:
        marker = marker_dir / ".keep"
        if not marker.exists():
            marker.write_text("", encoding="utf-8")
            created_files.append(str(marker))

    return {
        "ok": True,
        "created_dirs": created_dirs,
        "created_files": created_files,
        "live_requirements": _live_requirements_status(),
    }


def _live_requirements_status() -> dict[str, Any]:
    FML_IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    accounts = _load_accounts_full()
    require_saved_session = os.getenv("FACEBOOK_REQUIRE_SAVED_SESSION", "").strip().lower() in {"1", "true", "yes"}
    session_state = ensure_runtime_session(ROOT_DIR, force_restore=False)
    accounts_with_password = [
        entry for entry in accounts if str(entry.get("id", "")).strip() and entry.get("password")
    ]
    accounts_with_session_access = [
        entry
        for entry in accounts
        if str(entry.get("id", "")).strip()
        and (entry.get("password") or bool(session_state.get("ok")))
    ]
    chromedriver = _find_chromedriver()
    chromedriver_details = _chromedriver_details(chromedriver)
    chrome_binary = _find_chrome_binary()
    chrome_binary_details = _chrome_binary_details(chrome_binary)
    lister_import = _facebook_lister_import_status()
    return {
        "accounts_file_exists": FML_ACCOUNTS_PATH.exists(),
        "images_dir_exists": FML_IMAGES_DIR.exists(),
        "drivers_dir_exists": FML_DRIVERS_DIR.exists(),
        "accounts_ready": len(accounts_with_session_access),
        "accounts_with_password": len(accounts_with_password),
        "saved_session_ready": bool(session_state.get("ok")),
        "saved_session_restored": bool(session_state.get("restored")),
        "saved_session_source": session_state.get("source"),
        "lister_import_ok": bool(lister_import.get("ok")),
        "lister_import_error": lister_import.get("error"),
        "chromedriver_found": bool(chromedriver_details.get("found")),
        "chromedriver_path": chromedriver_details.get("path"),
        "chromedriver_version": chromedriver_details.get("version"),
        "chromedriver_error": chromedriver_details.get("error"),
        "chrome_binary_found": bool(chrome_binary_details.get("found")),
        "chrome_binary_path": chrome_binary_details.get("path"),
        "chrome_binary_version": chrome_binary_details.get("version"),
        "chrome_binary_error": chrome_binary_details.get("error"),
        "accounts_with_password": len(accounts_with_password),
    }


def _facebook_lister_import_status() -> dict[str, Any]:
    if not FML_DIR.exists():
        return {"ok": False, "error": f"Missing lister directory: {FML_DIR}"}
    inserted = False
    try:
        sys.path.insert(0, str(FML_DIR))
        inserted = True
        import importlib

        module = importlib.import_module("Lister")
        return {"ok": True, "module": str(getattr(module, "__file__", ""))}
    except Exception as exc:
        return {"ok": False, "error": str(exc)}
    finally:
        if inserted:
            try:
                sys.path.remove(str(FML_DIR))
            except ValueError:
                pass


def _sales_backend_base_url() -> str:
    raw = str(
        os.getenv("SALES_ASSISTANT_BACKEND_URL", DEFAULT_SALES_ASSISTANT_BACKEND_URL)
    ).strip()
    return raw.rstrip("/") or DEFAULT_SALES_ASSISTANT_BACKEND_URL


def _sales_backend_request(
    method: str,
    path: str,
    *,
    json_payload: dict[str, Any] | None = None,
) -> tuple[int, Any]:
    base_url = _sales_backend_base_url()
    target_url = f"{base_url}{path}"

    try:
        with httpx.Client(timeout=8.0) as client:
            response = client.request(method=method, url=target_url, json=json_payload)
    except httpx.RequestError as exc:
        raise HTTPException(
            status_code=502,
            detail={
                "message": "Sales-assistant backend is unreachable",
                "backend_url": base_url,
                "error": str(exc),
            },
        ) from exc

    try:
        payload: Any = response.json()
    except ValueError:
        payload = {"raw": (response.text or "").strip()}

    return response.status_code, payload


def _sales_backend_proxy(
    method: str,
    path: str,
    *,
    json_payload: dict[str, Any] | None = None,
) -> dict[str, Any]:
    status_code, payload = _sales_backend_request(method, path, json_payload=json_payload)
    if status_code >= 400:
        raise HTTPException(
            status_code=502,
            detail={
                "message": f"Sales-assistant backend returned HTTP {status_code}",
                "backend_url": _sales_backend_base_url(),
                "backend_path": path,
                "backend_response": payload,
            },
        )
    if isinstance(payload, dict):
        return payload
    return {"payload": payload}


def _sales_assistant_health_status() -> dict[str, Any]:
    health: dict[str, Any] = {
        "ok": False,
        "checked_at": datetime.now(timezone.utc).isoformat(),
        "backend_url": _sales_backend_base_url(),
        "reachable": False,
        "backend_status": None,
        "banks_count": 0,
        "factors_status": None,
        "factors_total_banks": None,
    }

    try:
        banks_status, banks_payload = _sales_backend_request("GET", "/api/banks")
    except HTTPException as exc:
        health["error"] = exc.detail
        return health

    health["reachable"] = True
    health["backend_status"] = banks_status

    if banks_status >= 400:
        health["error"] = {
            "message": "Sales-assistant banks endpoint returned an error",
            "backend_status": banks_status,
            "backend_response": banks_payload,
        }
        return health

    policies = banks_payload.get("policies") if isinstance(banks_payload, dict) else None
    if isinstance(policies, list):
        health["banks_count"] = len(policies)

    try:
        factors_status, factors_payload = _sales_backend_request("GET", "/api/banks/factors")
        health["factors_status"] = factors_status
        if factors_status >= 400:
            health["factors_error"] = {
                "message": "Sales-assistant bank factors endpoint returned an error",
                "backend_status": factors_status,
                "backend_response": factors_payload,
            }
        elif isinstance(factors_payload, dict):
            total_banks = factors_payload.get("totalBanks")
            if isinstance(total_banks, int):
                health["factors_total_banks"] = total_banks
    except HTTPException as exc:
        health["factors_error"] = exc.detail

    health["ok"] = (
        health["reachable"]
        and (health.get("backend_status") or 0) < 400
        and (health.get("factors_status") or 0) < 400
        and not health.get("factors_error")
    )
    return health


def _iter_bank_doc_files() -> list[Path]:
    allowed = {
        ".pdf",
        ".docx",
        ".doc",
        ".xlsx",
        ".xlsm",
        ".xls",
        ".pptx",
        ".ppt",
        ".csv",
        ".htm",
        ".html",
        ".json",
        ".txt",
    }
    if not BANK_DOCS_ROOT.exists():
        return []
    rows: list[Path] = []
    for path in BANK_DOCS_ROOT.rglob("*"):
        if not path.is_file() or path.suffix.lower() not in allowed or path.name.startswith("."):
            continue
        try:
            relative_parts = path.relative_to(BANK_DOCS_ROOT).parts
            if len(relative_parts) < 2:
                continue
            bank_folder = relative_parts[0]
        except Exception:
            bank_folder = ""
        if bank_folder.startswith("_") or bank_folder == "SmokeRailway":
            continue
        rows.append(path)
    return sorted(rows)


def _routeone_docs_status() -> dict[str, Any]:
    doc_files = _iter_bank_doc_files()
    index_payload = _safe_read_json(BANK_DOCS_INDEX_PATH, {})
    index_path = BANK_DOCS_INDEX_PATH
    if not index_payload and BANK_DOCS_GENERATED_INDEX_PATH.exists():
        index_payload = _safe_read_json(BANK_DOCS_GENERATED_INDEX_PATH, {})
        index_path = BANK_DOCS_GENERATED_INDEX_PATH
    generated_payload = _safe_read_json(BANK_PROFILES_GENERATED_PATH, {})
    sales_payload = _safe_read_json(SALES_ASSISTANT_BANKS_PATH, {})

    generated_profiles = generated_payload.get("profiles") if isinstance(generated_payload, dict) else None
    sales_policies = sales_payload.get("policies") if isinstance(sales_payload, dict) else None
    decoded_documents = index_payload.get("documents") if isinstance(index_payload, dict) else None

    by_bank: dict[str, int] = {}
    for path in doc_files:
        try:
            bank = path.relative_to(BANK_DOCS_ROOT).parts[0]
        except Exception:
            bank = "_Inbox"
        by_bank[bank] = by_bank.get(bank, 0) + 1
    if not by_bank and isinstance(decoded_documents, list):
        for item in decoded_documents:
            if not isinstance(item, dict):
                continue
            bank = str(item.get("bank") or "_Generated")
            by_bank[bank] = by_bank.get(bank, 0) + 1
    doc_count = len(doc_files) if doc_files else (len(decoded_documents) if isinstance(decoded_documents, list) else 0)

    return {
        "ok": bool(doc_count) and bool(generated_profiles or sales_policies),
        "bank_docs_root": _display_path(BANK_DOCS_ROOT),
        "doc_count": doc_count,
        "docs_by_bank": dict(sorted(by_bank.items())),
        "decoded_index_exists": index_path.exists(),
        "decoded_doc_count": len(decoded_documents) if isinstance(decoded_documents, list) else 0,
        "generated_profiles_count": len(generated_profiles) if isinstance(generated_profiles, list) else 0,
        "sales_assistant_policies_count": len(sales_policies) if isinstance(sales_policies, list) else 0,
        "last_decoded_at": index_payload.get("generated_at") if isinstance(index_payload, dict) else None,
        "index_path": _display_path(index_path),
        "profiles_path": _display_path(BANK_PROFILES_GENERATED_PATH),
        "sales_banks_path": _display_path(SALES_ASSISTANT_BANKS_PATH),
    }


def _reload_sales_assistant_bank_data() -> dict[str, Any]:
    result: dict[str, Any] = {
        "ok": False,
        "banks_status": None,
        "factors_status": None,
        "error": None,
    }
    try:
        banks_status, banks_payload = _sales_backend_request("POST", "/api/banks/reload")
        factors_status, factors_payload = _sales_backend_request("POST", "/api/banks/factors/reload")
        result.update(
            {
                "ok": banks_status < 400 and factors_status < 400,
                "banks_status": banks_status,
                "banks_response": banks_payload,
                "factors_status": factors_status,
                "factors_response": factors_payload,
            }
        )
    except HTTPException as exc:
        result["error"] = exc.detail
    return result


def _run_bank_docs_rebuild(
    *,
    reload_sales_data: bool,
    max_link_depth: int = 1,
    max_links_per_resource: int = 12,
) -> dict[str, Any]:
    command = [
        sys.executable,
        str(ROOT_DIR / "tools" / "rebuild_bank_brain.py"),
        "--bank-root",
        str(BANK_DOCS_ROOT),
        "--decoded-dir",
        str(BANK_DOCS_DECODED_DIR),
        "--index-path",
        str(BANK_DOCS_INDEX_PATH),
        "--generated-index-path",
        str(BANK_DOCS_GENERATED_INDEX_PATH),
        "--profiles-path",
        str(BANK_PROFILES_GENERATED_PATH),
        "--sales-banks-path",
        str(SALES_ASSISTANT_BANKS_PATH),
        "--link-cache-dir",
        str(BANK_DOCS_LINK_CACHE_DIR),
        "--max-link-depth",
        str(max_link_depth),
        "--max-links-per-resource",
        str(max_links_per_resource),
        "--json",
    ]

    try:
        process = subprocess.run(
            command,
            cwd=ROOT_DIR,
            capture_output=True,
            text=True,
            timeout=600,
        )
    except subprocess.TimeoutExpired as exc:
        raise HTTPException(
            status_code=504,
            detail={
                "message": "Bank Brain rebuild timed out while decoding RouteOne documents",
                "timeout_seconds": 600,
                "stdout": exc.stdout,
                "stderr": exc.stderr,
            },
        ) from exc

    stdout = (process.stdout or "").strip()
    stderr = (process.stderr or "").strip()
    rebuild_payload: Any = None
    if stdout:
        try:
            rebuild_payload = json.loads(stdout)
        except ValueError:
            rebuild_payload = {"raw": stdout}

    acceptable_exit = process.returncode in {0, 2}
    sales_reload = _reload_sales_assistant_bank_data() if reload_sales_data else {"ok": None, "skipped": True}
    status = _routeone_docs_status()

    return {
        "ok": acceptable_exit and (not reload_sales_data or sales_reload.get("ok") or sales_reload.get("skipped")),
        "returncode": process.returncode,
        "had_decode_errors": process.returncode == 2,
        "rebuild": rebuild_payload,
        "stderr": stderr,
        "sales_reload": sales_reload,
        "status": status,
    }


def _unique_upload_target(directory: Path, filename: str) -> Path:
    stem = _sanitize_doc_segment(Path(filename).stem, "routeone_document")
    suffix = Path(filename).suffix.lower()
    if not suffix:
        suffix = ".pdf"
    candidate = directory / f"{stem}{suffix}"
    counter = 2
    while candidate.exists():
        candidate = directory / f"{stem}_{counter}{suffix}"
        counter += 1
    return candidate


def _stack_readiness_status() -> dict[str, Any]:
    live_requirements = _live_requirements_status()
    accounts_ready = bool(live_requirements.get("accounts_ready") or live_requirements.get("accounts_with_password", 0))
    live_ready = (
        bool(live_requirements.get("images_dir_exists"))
        and bool(live_requirements.get("chromedriver_found"))
        and accounts_ready
    )

    stack_core_ready = (
        ADMIN_BUNDLE_INDEX.exists()
        and SALES_FRONTEND_INDEX.exists()
        and SALES_BACKEND_ENTRYPOINT.exists()
    )

    return {
        "ok": stack_core_ready,
        "ready_for_live_facebook_posting": live_ready,
        "components": {
            "admin_bundle_exists": ADMIN_BUNDLE_INDEX.exists(),
            "sales_frontend_bundle_exists": SALES_FRONTEND_INDEX.exists(),
            "sales_backend_entrypoint_exists": SALES_BACKEND_ENTRYPOINT.exists(),
            "sales_backend_url": _sales_backend_base_url(),
            "live_requirements": live_requirements,
        },
    }


def _validate_live_requirements(
    *,
    account_id: str | None,
    images: list[str],
) -> list[str]:
    errors: list[str] = []
    session_state = ensure_runtime_session(ROOT_DIR, force_restore=False)

    if not account_id:
        errors.append("account_id is required for live posting")

    accounts = _load_accounts_full()
    if account_id:
        match = next(
            (entry for entry in accounts if str(entry.get("id", "")).strip() == account_id.strip()),
            None,
        )
        if not match:
            errors.append(f"account_id '{account_id}' was not found in accounts.json or FACEBOOK_LOGIN_ACCOUNT_ID")
        elif not match.get("password") and not bool(session_state.get("ok")):
            errors.append(f"account_id '{account_id}' is missing password in accounts.json or FACEBOOK_LOGIN_PASSWORD")

    if not images:
        errors.append("images list is required for live posting")
    FML_IMAGES_DIR.mkdir(parents=True, exist_ok=True)
    for image_name in images:
        clean = str(image_name).strip()
        if not clean:
            continue
        image_path = FML_IMAGES_DIR / clean
        if not image_path.exists():
            errors.append(
                "Missing image file in facebook lister images directory: "
                f"{image_path.name}"
            )

    chromedriver_path = _find_chromedriver()
    if not chromedriver_path:
        errors.append(
            "ChromeDriver was not found. Set CHROMEDRIVER_PATH or install chromedriver."
        )

    return errors


def _marketplace_state_from_publish_result(parsed_result: dict[str, Any] | None) -> dict[str, Any]:
    result = parsed_result if isinstance(parsed_result, dict) else {}
    confirmation = result.get("confirmation") if isinstance(result.get("confirmation"), dict) else {}
    listing_url = str(
        confirmation.get("listing_url")
        or result.get("listing_url")
        or ""
    ).strip()
    confirmation_name = str(confirmation.get("confirmation") or result.get("confirmation") or "").strip()
    explicit_status = str(
        confirmation.get("marketplace_status")
        or result.get("marketplace_status")
        or ""
    ).strip().lower()
    visible = bool(confirmation.get("visible") or result.get("visible"))

    has_item_url = bool(listing_url and "/marketplace/item/" in listing_url.lower())

    if has_item_url:
        status = "live"
        visible = True
    elif explicit_status in {"processing", "needs_review", "failed", "draft"}:
        status = explicit_status
    elif explicit_status == "live":
        status = "processing"
    elif visible and confirmation_name and has_item_url:
        status = "live"
    elif result.get("posted"):
        status = "processing"
    else:
        status = "needs_review"

    return {
        "marketplace_status": status,
        "posted": status == "live",
        "visible": visible,
        "listing_url": listing_url if has_item_url else "",
        "confirmation": confirmation if confirmation else result,
    }


def _publish_live(payload: FacebookPostRequest) -> tuple[bool, str, dict[str, Any]]:
    python_bin = ROOT_DIR / ".venv" / "Scripts" / "python.exe"
    helper_script = ROOT_DIR / "tools" / "facebook_publish.py"
    if not helper_script.exists():
        return False, "Missing tools/facebook_publish.py", {"marketplace_status": "needs_review"}

    account_id = str(payload.account_id or "").strip() or _resolve_default_account_id()
    session_state = ensure_runtime_session(ROOT_DIR, force_restore=False)
    os.environ["FACEBOOK_REQUIRE_SAVED_SESSION"] = "0"
    os.environ["FACEBOOK_SESSION_CHECK_WAIT_SECONDS"] = "0.75"
    os.environ["FACEBOOK_LOGIN_WAIT_SECONDS"] = "16"
    os.environ["FACEBOOK_FORM_SLEEP_SECONDS"] = "0.22"
    os.environ["FACEBOOK_FIELD_WAIT_SECONDS"] = "0.22"
    os.environ["FACEBOOK_ACCOUNT_CHOOSER_WAIT_SECONDS"] = "0.45"
    os.environ["FACEBOOK_POST_NAV_WAIT_SECONDS"] = "1.1"
    os.environ["FACEBOOK_PUBLISH_CONFIRM_SECONDS"] = "60"
    requirement_errors = _validate_live_requirements(
        account_id=account_id,
        images=payload.images,
    )
    if requirement_errors:
        return False, " | ".join(requirement_errors), {"marketplace_status": "needs_review"}

    publish_payload = payload.model_dump()
    publish_payload["account_id"] = account_id
    publish_payload["price"] = payload.suggested_down_payment or _facebook_post_price(payload.price)
    publish_payload["location"] = _facebook_listing_location(payload.location)
    publish_payload["detail_url"] = ""

    request_file = RUNTIME_DIR / "live_publish_request.json"
    status_file = RUNTIME_DIR / "facebook_live_status.json"
    request_file.parent.mkdir(parents=True, exist_ok=True)
    request_file.write_text(json.dumps(publish_payload, indent=2), encoding="utf-8")
    status_file.write_text(
        json.dumps(
            {
                "ok": True,
                "vin": payload.vin,
                "title": payload.title,
                "stage": "Queued live Facebook publish.",
                "type": "main",
                "updated_at": datetime.now(timezone.utc).isoformat(),
            },
            indent=2,
        ),
        encoding="utf-8",
    )

    cmd = [
        str(python_bin if python_bin.exists() else sys.executable),
        str(helper_script),
        "--payload",
        str(request_file),
    ]
    publish_env = os.environ.copy()
    publish_env["FACEBOOK_REQUIRE_SAVED_SESSION"] = "0"
    publish_env["FACEBOOK_SESSION_CHECK_WAIT_SECONDS"] = "0.75"
    publish_env["FACEBOOK_LOGIN_WAIT_SECONDS"] = "16"
    publish_env["FACEBOOK_FORM_SLEEP_SECONDS"] = "0.22"
    publish_env["FACEBOOK_FIELD_WAIT_SECONDS"] = "0.22"
    publish_env["FACEBOOK_ACCOUNT_CHOOSER_WAIT_SECONDS"] = "0.45"
    publish_env["FACEBOOK_POST_NAV_WAIT_SECONDS"] = "1.1"
    publish_env["FACEBOOK_PUBLISH_CONFIRM_SECONDS"] = "60"
    publish_env["FACEBOOK_PUBLISH_STATUS_FILE"] = str(status_file)
    publish_env["FACEBOOK_PUBLISH_VIN"] = str(payload.vin or "")
    publish_env["FACEBOOK_PUBLISH_TITLE"] = str(payload.title or "")
    max_attempts = max(1, min(int(os.getenv("FACEBOOK_PUBLISH_MAX_ATTEMPTS", "2") or "2"), 3))
    state: dict[str, Any] = {"marketplace_status": "needs_review"}
    output = ""
    for attempt in range(1, max_attempts + 1):
        try:
            completed = subprocess.run(
                cmd,
                cwd=str(ROOT_DIR),
                env=publish_env,
                capture_output=True,
                text=True,
                check=False,
                timeout=int(os.getenv("FACEBOOK_PUBLISH_TIMEOUT_SECONDS", "180") or "180"),
            )
        except subprocess.TimeoutExpired as exc:
            output = (exc.stdout or "").strip() if isinstance(exc.stdout, str) else ""
            error_text = (exc.stderr or "").strip() if isinstance(exc.stderr, str) else ""
            detail = _friendly_facebook_publish_detail(
                output=output,
                error_text=error_text,
                fallback="Facebook posting timed out while waiting for login or Marketplace confirmation.",
            )
            if not detail or "timed out" not in detail.lower():
                detail = "Facebook posting timed out while waiting for login or Marketplace confirmation. Xconsole marked the vehicle Needs Review and did not count it Live."
            _write_facebook_live_status(
                status_file,
                {"ok": False, "vin": payload.vin, "title": payload.title, "stage": detail, "type": "failure"},
            )
            return False, detail, {"marketplace_status": "needs_review", "error": detail}

        output = (completed.stdout or "").strip()
        error_text = (completed.stderr or "").strip()
        parsed_result = _last_json_object_from_output(output) if completed.returncode == 0 else None
        state = _marketplace_state_from_publish_result(parsed_result) if isinstance(parsed_result, dict) else {"marketplace_status": "needs_review"}

        if completed.returncode != 0:
            detail = _friendly_facebook_publish_detail(
                output=output,
                error_text=error_text,
                fallback=f"Publish failed with code {completed.returncode}",
            )
            if attempt < max_attempts and _facebook_publish_should_retry(detail, state):
                _write_facebook_live_status(
                    status_file,
                    {
                        "ok": True,
                        "vin": payload.vin,
                        "title": payload.title,
                        "stage": f"Chrome hit a transient Facebook browser crash. Retrying publish ({attempt + 1}/{max_attempts})...",
                        "type": "main",
                    },
                )
                time.sleep(1.0)
                continue
            _write_facebook_live_status(
                status_file,
                {"ok": False, "vin": payload.vin, "title": payload.title, "stage": detail, "type": "failure"},
            )
            return False, detail, {"marketplace_status": "needs_review", "error": detail}

        if not isinstance(parsed_result, dict):
            detail = _friendly_facebook_publish_detail(
                output=output,
                error_text=error_text,
                fallback="Facebook helper returned an empty response payload.",
            )
            if attempt < max_attempts and _facebook_publish_should_retry(detail, state):
                _write_facebook_live_status(
                    status_file,
                    {
                        "ok": True,
                        "vin": payload.vin,
                        "title": payload.title,
                        "stage": f"Facebook helper lost the browser session. Retrying publish ({attempt + 1}/{max_attempts})...",
                        "type": "main",
                    },
                )
                time.sleep(1.0)
                continue
            _write_facebook_live_status(
                status_file,
                {"ok": False, "vin": payload.vin, "title": payload.title, "stage": detail, "type": "failure"},
            )
            return False, detail, {"marketplace_status": "needs_review", "error": detail}

        state = _marketplace_state_from_publish_result(parsed_result)

        if parsed_result.get("ok") is False and state.get("marketplace_status") not in {"processing"}:
            detail = _friendly_facebook_publish_detail(output=output or json.dumps(parsed_result))
            if attempt < max_attempts and _facebook_publish_should_retry(detail, state):
                _write_facebook_live_status(
                    status_file,
                    {
                        "ok": True,
                        "vin": payload.vin,
                        "title": payload.title,
                        "stage": f"Facebook browser crashed during submit. Retrying publish ({attempt + 1}/{max_attempts})...",
                        "type": "main",
                    },
                )
                time.sleep(1.0)
                continue
            _write_facebook_live_status(
                status_file,
                {"ok": False, "vin": payload.vin, "title": payload.title, "stage": detail, "type": "failure"},
            )
            return False, detail, state

        if state.get("marketplace_status") in {"failed", "needs_review", "draft"}:
            detail = _friendly_facebook_publish_detail(output=output or json.dumps(parsed_result))
            if attempt < max_attempts and _facebook_publish_should_retry(detail, state):
                _write_facebook_live_status(
                    status_file,
                    {
                        "ok": True,
                        "vin": payload.vin,
                        "title": payload.title,
                        "stage": f"Facebook browser crashed while finalizing the post. Retrying publish ({attempt + 1}/{max_attempts})...",
                        "type": "main",
                    },
                )
                time.sleep(1.0)
                continue
            _write_facebook_live_status(
                status_file,
                {"ok": False, "vin": payload.vin, "title": payload.title, "stage": detail, "type": "failure"},
            )
            return False, detail, state

        if state.get("marketplace_status") != "live":
            detail = (
                "Facebook submitted the listing, but Marketplace has not exposed a visible listing URL yet. "
                "Xconsole marked it Processing and will keep it out of Live counts until verification succeeds."
            )
            _write_facebook_live_status(
                status_file,
                {"ok": False, "vin": payload.vin, "title": payload.title, "stage": detail, "type": "processing"},
            )
            return False, detail, state
        break

    status_file.write_text(
        json.dumps(
            {
                "ok": True,
                "vin": payload.vin,
                "title": payload.title,
                "stage": "Marketplace listing confirmed.",
                "type": "success",
                "marketplace_status": "live",
                "listing_url": state.get("listing_url"),
                "updated_at": datetime.now(timezone.utc).isoformat(),
            },
            indent=2,
        ),
        encoding="utf-8",
    )
    return True, output or "live publish command completed", state


def _strip_facebook_automation_noise(text: str) -> str:
    clean = re.sub(r"\x1b\[[0-9;]*m", "", str(text or ""))
    clean = clean.replace("\\u001b", "")
    return clean.strip()


def _last_json_object_from_output(text: str) -> dict[str, Any] | None:
    clean = _strip_facebook_automation_noise(text)
    parsed_candidates: list[dict[str, Any]] = []
    for line in reversed(clean.splitlines()):
        line = line.strip()
        if not line.startswith("{"):
            continue
        try:
            parsed = json.loads(line)
        except Exception:
            continue
        if isinstance(parsed, dict):
            parsed_candidates.append(parsed)
    if not parsed_candidates:
        return None

    def score(payload: dict[str, Any]) -> tuple[int, int]:
        rich_keys = {
            "confirmation",
            "marketplace_status",
            "listing_url",
            "posted",
            "submitted",
            "results",
            "failed",
            "live_success",
            "live_detail",
            "error",
            "message",
        }
        richness = sum(1 for key in rich_keys if key in payload)
        session_noise = 1 if {"session_archived", "debug"} & set(payload.keys()) else 0
        return (session_noise, -richness)

    parsed_candidates.sort(key=score)
    return parsed_candidates[0]


def _friendly_facebook_publish_detail(*, output: str = "", error_text: str = "", fallback: str = "") -> str:
    clean = _strip_facebook_automation_noise(error_text or output or fallback)
    parsed = _last_json_object_from_output(clean)
    error_payload: Any = parsed.get("error") if isinstance(parsed, dict) else None
    if isinstance(error_payload, str):
        try:
            nested = json.loads(error_payload)
            if isinstance(nested, dict):
                error_payload = nested
        except Exception:
            pass
    if not error_payload and isinstance(parsed, dict):
        error_payload = parsed

    message = ""
    if isinstance(error_payload, dict):
        message = str(error_payload.get("message") or error_payload.get("error") or "").strip()
    elif isinstance(error_payload, str):
        message = error_payload.strip()
    if not message:
        message = clean.splitlines()[-1].strip() if clean.splitlines() else fallback

    lowered = message.lower()
    if "location" in lowered and ("dropdown" in lowered or "could not be selected" in lowered or "option" in lowered):
        return "Marketplace stopped on the location selector. Retrying now uses ZIP 33317 first and selects Plantation, FL before any photos upload."
    if "login_required" in clean.lower() or "login required" in lowered:
        return "Facebook session is not usable. Open the visible Facebook login once, approve it, then save the session before posting."
    if "account_chooser" in clean.lower() or "saved-account chooser" in lowered or "continue as" in lowered or "log in as" in lowered:
        return "Facebook is showing the saved-account chooser. Click Continue/Log in as the correct Facebook user once in the visible browser, then retry posting."
    if "publish was not confirmed" in lowered:
        return "Facebook clicked Publish but Marketplace did not confirm the listing. Check the visible browser for a required field or review prompt."
    if "vehicle form still has required fields" in lowered:
        return "Facebook still sees a required Marketplace field. The automation captured it and stopped before marking the vehicle posted."
    if "session not created" in lowered and "chrome instance exited" in lowered:
        return "Chrome could not start inside Railway for Facebook posting. Xconsole marked the vehicle Needs Review and did not count it Live."
    if "missing image" in lowered:
        return "Facebook post stopped because one or more selected vehicle photos were not available locally."

    compact = re.sub(r"\s+", " ", message or clean or fallback).strip()
    return compact[:360] + ("..." if len(compact) > 360 else "")


def _facebook_publish_should_retry(detail: str, state: dict[str, Any] | None = None) -> bool:
    text = " ".join(
        [
            str(detail or ""),
            json.dumps(state or {}, ensure_ascii=False) if isinstance(state, dict) else "",
        ]
    ).lower()
    transient_markers = [
        "tab crashed",
        "session not created",
        "chrome instance exited",
        "target frame detached",
        "disconnected: not connected to devtools",
        "invalid session id",
    ]
    return any(marker in text for marker in transient_markers)


def _write_facebook_live_status(status_file: Path, payload: dict[str, Any]) -> None:
    stage = _friendly_facebook_publish_detail(fallback=str(payload.get("stage") or "")) if payload.get("type") == "failure" else str(payload.get("stage") or "")
    if stage.lower().startswith("facebook session archived:"):
        stage = "Facebook session refreshed."
    clean_payload = dict(payload)
    clean_payload["stage"] = stage[:420] + ("..." if len(stage) > 420 else "")
    clean_payload["updated_at"] = datetime.now(timezone.utc).isoformat()
    status_file.write_text(json.dumps(clean_payload, indent=2), encoding="utf-8")


def _queue_live_inventory_request(request: FacebookOneClickPostRequest, *, title: str | None = None) -> dict[str, Any]:
    worker_script = ROOT_DIR / "tools" / "facebook_post_inventory_worker.py"
    if not worker_script.exists():
        raise HTTPException(
            status_code=500,
            detail={"message": "Missing tools/facebook_post_inventory_worker.py"},
        )

    status_file = RUNTIME_DIR / "facebook_live_status.json"
    jobs_dir = RUNTIME_DIR / "facebook_live_jobs"
    jobs_dir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now(timezone.utc).strftime("%Y%m%d-%H%M%S-%f")
    clean_vin = str(request.vin or "").strip().upper()
    listing_title = str(title or clean_vin).strip() or clean_vin
    request_file = jobs_dir / f"{clean_vin}_{stamp}.json"
    request_file.write_text(json.dumps(request.model_dump(), indent=2), encoding="utf-8")
    _write_facebook_live_status(
        status_file,
        {
            "ok": True,
            "vin": clean_vin,
            "title": listing_title,
            "stage": f"Queued Facebook live publish for {listing_title}. Waiting for Marketplace automation to begin...",
            "type": "queued",
        },
    )

    python_bin = ROOT_DIR / ".venv" / "Scripts" / "python.exe"
    cmd = [
        str(python_bin if python_bin.exists() else sys.executable),
        str(worker_script),
    ]
    env = os.environ.copy()
    env["FACEBOOK_PUBLISH_STATUS_FILE"] = str(status_file)
    env["FACEBOOK_PUBLISH_VIN"] = clean_vin
    env["FACEBOOK_PUBLISH_TITLE"] = listing_title
    subprocess.Popen(
        cmd,
        cwd=str(ROOT_DIR),
        env=env,
        stdout=subprocess.DEVNULL,
        stderr=subprocess.DEVNULL,
        stdin=subprocess.DEVNULL,
        close_fds=True,
    )
    return {
        "ok": True,
        "mode": "live",
        "live_success": False,
        "marketplace_status": "queued",
        "live_detail": "Facebook live publish queued. Xconsole is running Marketplace automation in the background.",
        "listing_url": None,
    }


def _publish_live_batch(payloads: list[FacebookPostRequest], account_id: str) -> tuple[bool, str, dict[str, Any]]:
    helper_script = ROOT_DIR / "tools" / "facebook_publish_batch.py"
    if not helper_script.exists():
        return False, "Missing tools/facebook_publish_batch.py", {}

    status_file = RUNTIME_DIR / "facebook_live_status.json"
    request_file = RUNTIME_DIR / "live_publish_batch_request.json"
    session_state = ensure_runtime_session(ROOT_DIR, force_restore=False)
    os.environ["FACEBOOK_REQUIRE_SAVED_SESSION"] = "0"
    os.environ["FACEBOOK_SESSION_CHECK_WAIT_SECONDS"] = "0.75"
    os.environ["FACEBOOK_LOGIN_WAIT_SECONDS"] = "16"
    os.environ["FACEBOOK_FORM_SLEEP_SECONDS"] = "0.22"
    os.environ["FACEBOOK_FIELD_WAIT_SECONDS"] = "0.22"
    os.environ["FACEBOOK_ACCOUNT_CHOOSER_WAIT_SECONDS"] = "0.45"
    os.environ["FACEBOOK_POST_NAV_WAIT_SECONDS"] = "1.1"
    os.environ["FACEBOOK_PUBLISH_CONFIRM_SECONDS"] = "60"
    publish_items: list[dict[str, Any]] = []
    for payload in payloads:
        requirement_errors = _validate_live_requirements(account_id=account_id, images=payload.images)
        if requirement_errors:
            return False, " | ".join(requirement_errors), {}
        item = payload.model_dump()
        item["account_id"] = account_id
        item["price"] = payload.suggested_down_payment or _facebook_post_price(payload.price)
        item["location"] = _facebook_listing_location(payload.location)
        item["detail_url"] = ""
        publish_items.append(item)

    request_file.parent.mkdir(parents=True, exist_ok=True)
    request_file.write_text(json.dumps({"account_id": account_id, "items": publish_items}, indent=2), encoding="utf-8")
    _write_facebook_live_status(
        status_file,
        {
            "ok": True,
            "stage": f"Facebook batch queued: 0/{len(publish_items)} posted.",
            "type": "main",
            "batch_total": len(publish_items),
            "posted": 0,
            "failed": 0,
        },
    )

    python_bin = ROOT_DIR / ".venv" / "Scripts" / "python.exe"
    cmd = [
        str(python_bin if python_bin.exists() else sys.executable),
        str(helper_script),
        "--payload",
        str(request_file),
    ]
    publish_env = os.environ.copy()
    publish_env["FACEBOOK_REQUIRE_SAVED_SESSION"] = "0"
    publish_env["FACEBOOK_SESSION_CHECK_WAIT_SECONDS"] = "0.75"
    publish_env["FACEBOOK_LOGIN_WAIT_SECONDS"] = "16"
    publish_env["FACEBOOK_FORM_SLEEP_SECONDS"] = "0.22"
    publish_env["FACEBOOK_FIELD_WAIT_SECONDS"] = "0.22"
    publish_env["FACEBOOK_ACCOUNT_CHOOSER_WAIT_SECONDS"] = "0.45"
    publish_env["FACEBOOK_POST_NAV_WAIT_SECONDS"] = "1.1"
    publish_env["FACEBOOK_PUBLISH_CONFIRM_SECONDS"] = "60"
    publish_env["FACEBOOK_PUBLISH_STATUS_FILE"] = str(status_file)
    try:
        completed = subprocess.run(
            cmd,
            cwd=str(ROOT_DIR),
            env=publish_env,
            capture_output=True,
            text=True,
            check=False,
            timeout=int(os.getenv("FACEBOOK_BATCH_PUBLISH_TIMEOUT_SECONDS", "900") or "900"),
        )
    except subprocess.TimeoutExpired as exc:
        detail = _friendly_facebook_publish_detail(
            output=(exc.stdout or "") if isinstance(exc.stdout, str) else "",
            error_text=(exc.stderr or "") if isinstance(exc.stderr, str) else "",
            fallback="Facebook batch timed out while waiting for login or Marketplace confirmation.",
        )
        _write_facebook_live_status(
            status_file,
            {
                "ok": False,
                "stage": detail,
                "type": "failure",
                "batch_total": len(publish_items),
                "posted": 0,
                "failed": len(publish_items),
            },
        )
        return False, detail, {"ok": False, "error": detail, "items": []}
    output = (completed.stdout or "").strip()
    error_text = (completed.stderr or "").strip()
    parsed_result = _last_json_object_from_output(output)
    if completed.returncode != 0 or not isinstance(parsed_result, dict) or parsed_result.get("ok") is False:
        detail = _friendly_facebook_publish_detail(output=output, error_text=error_text, fallback=f"Facebook batch failed with code {completed.returncode}")
        _write_facebook_live_status(
            status_file,
            {
                "ok": False,
                "stage": detail,
                "type": "failure",
                "batch_total": len(publish_items),
                "posted": int((parsed_result or {}).get("posted") or 0),
                "failed": int((parsed_result or {}).get("failed") or 0),
            },
        )
        return False, detail, parsed_result or {}

    posted = int(parsed_result.get("posted") or 0)
    failed = int(parsed_result.get("failed") or 0)
    results = parsed_result.get("results") if isinstance(parsed_result.get("results"), list) else []
    processing = sum(
        1
        for item in results
        if isinstance(item, dict) and str(item.get("marketplace_status") or "").strip().lower() == "processing"
    )
    if processing and not posted and not failed:
        detail = f"Facebook batch submitted: {processing} processing, {failed} failed. Marketplace has not exposed a live item URL yet."
    else:
        detail = f"Facebook batch complete: {posted} posted, {failed} failed."
    _write_facebook_live_status(
        status_file,
        {
            "ok": failed == 0 and posted > 0,
            "stage": detail,
            "type": "success" if failed == 0 and posted > 0 else ("processing" if processing and failed == 0 else "failure"),
            "batch_total": len(publish_items),
            "posted": posted,
            "failed": failed,
            "processing": processing,
        },
    )
    return failed == 0 and posted > 0, detail, parsed_result


def _run_live_preflight(
    *,
    account_id: str | None,
    images: list[str],
    vin: str | None = None,
) -> dict[str, Any]:
    normalized_images = [str(item).strip() for item in images if str(item).strip()]
    errors = _validate_live_requirements(
        account_id=account_id,
        images=normalized_images,
    )
    suggestions: list[str] = []
    if vin:
        suggestions = _suggest_images_for_vin(vin=vin, limit=20)

    warnings: list[str] = []
    if vin and not normalized_images and suggestions:
        warnings.append("No image filenames selected yet, but matching VIN image files were found.")

    return {
        "ok": len(errors) == 0,
        "errors": errors,
        "warnings": warnings,
        "account_id": account_id,
        "images_count": len(normalized_images),
        "vin": vin,
        "suggested_images": suggestions,
        "live_requirements": _live_requirements_status(),
    }


def _prepare_live_post(
    *,
    vin: str,
    account_id: str | None,
    import_missing_images: bool,
    image_limit: int,
    overwrite_images: bool,
) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    resolved_account_id = (account_id or "").strip() or _resolve_default_account_id()
    import_result: dict[str, Any] | None = None
    import_error: dict[str, Any] | None = None

    if import_missing_images:
        try:
            import_result = _import_vehicle_images(
                vin=clean_vin,
                limit=image_limit,
                overwrite=overwrite_images,
            )
        except HTTPException as exc:
            import_error = exc.detail if isinstance(exc.detail, dict) else {"message": str(exc.detail)}

    selected_images = _suggest_images_for_vin(vin=clean_vin, limit=image_limit)
    preflight = _run_live_preflight(
        account_id=resolved_account_id,
        images=selected_images,
        vin=clean_vin,
    )

    guidance: list[str] = []
    if not resolved_account_id:
        guidance.append("No usable account_id found. Add an account with id/password in accounts.json.")
    if not selected_images:
        guidance.append(
            "No image filenames selected. Import VIN photos or place image files in "
            "automation/facebook-marketplace-lister/images."
        )
    if not preflight.get("ok"):
        guidance.append("Resolve preflight errors shown below before publishing live.")
    if import_error and not preflight.get("ok"):
        guidance.append("Image import did not complete cleanly; review import_error details.")

    return {
        "ok": bool(preflight.get("ok")),
        "vin": clean_vin,
        "account_id": resolved_account_id,
        "selected_images": selected_images,
        "selected_images_count": len(selected_images),
        "import_result": import_result,
        "import_error": import_error,
        "preflight": preflight,
        "live_requirements": _live_requirements_status(),
        "guidance": guidance,
    }


def _resolve_repair_vin(explicit_vin: str | None) -> str:
    if explicit_vin and str(explicit_vin).strip():
        return str(explicit_vin).strip().upper()

    from_inventory = next(
        (
            str(item.get("vin", "")).strip().upper()
            for item in _load_inventory_candidates()
            if str(item.get("vin", "")).strip()
        ),
        "",
    )
    if from_inventory:
        return from_inventory

    from_posts = next(
        (
            str(item.get("vin", "")).strip().upper()
            for item in _load_runtime_posts()
            if str(item.get("vin", "")).strip()
        ),
        "",
    )
    if from_posts:
        return from_posts

    return "2C4RC1L78NR164218"


def _full_repair_and_relink(
    *,
    vin: str | None,
    ensure_placeholder_images: bool,
    placeholder_count: int,
) -> dict[str, Any]:
    target_vin = _resolve_repair_vin(vin)

    bootstrap = _bootstrap_facebook_lister(create_template_account_if_missing=True)

    placeholder_result: dict[str, Any] | None = None
    if ensure_placeholder_images and not _suggest_images_for_vin(vin=target_vin, limit=1):
        placeholder_result = _seed_placeholder_images_for_vin(
            vin=target_vin,
            count=placeholder_count,
            overwrite=False,
        )

    relink_result: dict[str, Any] | None = None
    if not _suggest_images_for_vin(vin=target_vin, limit=1):
        all_images, _ = _list_facebook_images(limit=2000)
        if all_images:
            relink_result = _relink_images_to_vin(
                vin=target_vin,
                images=all_images[:placeholder_count],
                include_vin_matches=False,
                overwrite=False,
                delete_source=False,
            )

    prepared = _prepare_live_post(
        vin=target_vin,
        account_id=None,
        import_missing_images=True,
        image_limit=max(placeholder_count, 6),
        overwrite_images=False,
    )
    stack = _stack_readiness_status()

    return {
        "ok": bool(prepared.get("ok")) and bool(stack.get("ready_for_live_facebook_posting")),
        "vin": target_vin,
        "bootstrap": bootstrap,
        "placeholder_result": placeholder_result,
        "relink_result": relink_result,
        "prepared": prepared,
        "stack_readiness": stack,
    }


def _wire_everything(
    *,
    vin: str | None,
    ensure_placeholder_images: bool,
    placeholder_count: int,
    reload_sales_data: bool,
) -> dict[str, Any]:
    target_vin = _resolve_repair_vin(vin)
    started_at = datetime.now(timezone.utc).isoformat()

    full_repair = _full_repair_and_relink(
        vin=target_vin,
        ensure_placeholder_images=ensure_placeholder_images,
        placeholder_count=placeholder_count,
    )
    prepared = full_repair.get("prepared") if isinstance(full_repair, dict) else {}
    if not isinstance(prepared, dict):
        prepared = {}

    sales_health_before = _sales_assistant_health_status()
    sales_reload: dict[str, Any] = {
        "requested": reload_sales_data,
        "performed": False,
        "banks_status": None,
        "factors_status": None,
        "banks_response": None,
        "factors_response": None,
        "error": None,
    }

    def _attempt_sales_reload(max_attempts: int = 3) -> dict[str, Any]:
        result: dict[str, Any] = {
            "requested": True,
            "performed": False,
            "banks_status": None,
            "factors_status": None,
            "banks_response": None,
            "factors_response": None,
            "error": None,
            "attempts": 0,
        }
        for attempt in range(1, max_attempts + 1):
            result["attempts"] = attempt
            try:
                banks_status, banks_payload = _sales_backend_request("POST", "/api/banks/reload")
                factors_status, factors_payload = _sales_backend_request(
                    "POST",
                    "/api/banks/factors/reload",
                )
                result.update(
                    {
                        "performed": True,
                        "banks_status": banks_status,
                        "factors_status": factors_status,
                        "banks_response": banks_payload,
                        "factors_response": factors_payload,
                        "error": None,
                    }
                )
                if banks_status < 400 and factors_status < 400:
                    return result
                result["error"] = {
                    "message": "Sales backend reload returned non-success status.",
                    "banks_status": banks_status,
                    "factors_status": factors_status,
                }
            except HTTPException as exc:
                result["error"] = exc.detail

            if attempt < max_attempts:
                time.sleep(0.7)
        return result

    if reload_sales_data:
        sales_reload = _attempt_sales_reload(max_attempts=3)

    sales_health_after = _sales_assistant_health_status()
    stack_after = _stack_readiness_status()
    finished_at = datetime.now(timezone.utc).isoformat()

    selected_images = prepared.get("selected_images") if isinstance(prepared, dict) else []
    if not isinstance(selected_images, list):
        selected_images = []
    account_id = prepared.get("account_id") if isinstance(prepared, dict) else None

    guidance: list[str] = []
    if not stack_after.get("ready_for_live_facebook_posting"):
        guidance.append("Live Facebook posting is still not fully ready. Check live_requirements.")
    if not prepared.get("ok"):
        guidance.append("Prepare live post did not fully pass. Resolve preflight issues before posting.")
    if reload_sales_data and not sales_health_after.get("ok"):
        guidance.append("Sales-assistant service is not healthy after reload. Check /api/sales-assistant/health.")
    if sales_reload.get("error"):
        guidance.append("Sales data reload reported an error.")

    ok = bool(prepared.get("ok")) and bool(stack_after.get("ready_for_live_facebook_posting"))
    if reload_sales_data:
        ok = ok and bool(sales_health_after.get("ok")) and not bool(sales_reload.get("error"))

    return {
        "ok": ok,
        "vin": target_vin,
        "started_at": started_at,
        "finished_at": finished_at,
        "full_repair": full_repair,
        "prepared_account_id": account_id,
        "prepared_images": selected_images,
        "prepared_images_count": len(selected_images),
        "sales_health_before": sales_health_before,
        "sales_reload": sales_reload,
        "sales_health_after": sales_health_after,
        "stack_readiness_after": stack_after,
        "guidance": guidance,
    }


def _append_audit_event(event_type: str, payload: dict[str, Any]) -> None:
    existing = _safe_read_json(BANK_BRAIN_AUDIT_PATH, {"events": []})
    events = existing.get("events", []) if isinstance(existing, dict) else []
    if not isinstance(events, list):
        events = []
    events.append(
        {
            "event_type": event_type,
            "created_at": datetime.now(timezone.utc).isoformat(),
            "payload": payload,
        }
    )
    _safe_write_json(BANK_BRAIN_AUDIT_PATH, {"events": events[-500:]})


def _vehicle_kind_and_buyer(vehicle: dict[str, Any]) -> tuple[str, str, list[str]]:
    title = str(vehicle.get("title") or "").lower()
    body = str(vehicle.get("body_style") or vehicle.get("body") or "").lower()
    text = f"{title} {body}"
    if any(token in text for token in ("pacifica", "minivan", "voyager")):
        return (
            "family van",
            "families, rideshare drivers, and anyone who needs comfortable people-moving space",
            ["family-friendly cabin", "easy daily driving", "room for passengers and cargo"],
        )
    if any(token in text for token in ("ram 3500", "ram 2500", "chassis", "cab", "promaster")):
        return (
            "work truck",
            "business owners, contractors, and buyers who need capability first",
            ["work-ready capability", "commercial-minded utility", "strong hauling/towing presence"],
        )
    if any(token in text for token in ("ram 1500", "pickup", "truck")):
        return (
            "pickup",
            "drivers who want truck capability without giving up everyday comfort",
            ["truck utility", "confident road presence", "daily-driver comfort"],
        )
    if any(token in text for token in ("wrangler", "gladiator")):
        return (
            "adventure vehicle",
            "Jeep shoppers who want open-air capability and weekend flexibility",
            ["Jeep capability", "weekend-ready personality", "strong resale appeal"],
        )
    if any(token in text for token in ("aviator", "black label", "navigator", "grand cherokee", "durango", "wagoneer", "compass", "cherokee", "suv")):
        return (
            "luxury SUV" if any(token in text for token in ("aviator", "black label", "navigator")) else "SUV",
            "families, professionals, and luxury-SUV shoppers who want comfort, space, and confidence",
            ["premium SUV space", "comfortable ride", "daily usability"],
        )
    if any(token in text for token in ("charger", "challenger")):
        return (
            "performance car",
            "drivers who want muscle-car style and a fun daily drive",
            ["bold performance look", "strong engine character", "driver-focused appeal"],
        )
    return (
        "vehicle",
        "buyers who want a clean, inspected vehicle with straightforward numbers",
        ["clean presentation", "daily usability", "strong value"],
    )


def _vehicle_marketing_highlights(vehicle: dict[str, Any]) -> list[str]:
    _, _, defaults = _vehicle_kind_and_buyer(vehicle)
    highlights: list[str] = []
    for key, label in [
        ("drivetrain", "Drivetrain"),
        ("engine", "Engine"),
        ("transmission", "Transmission"),
        ("exterior", "Exterior"),
        ("interior", "Interior"),
    ]:
        value = str(vehicle.get(key) or "").strip()
        if value:
            highlights.append(f"{label}: {value}")
    for item in defaults:
        if item not in highlights:
            highlights.append(item)
    return highlights[:6]


def _carfax_buyer_facing_lines(carfax: dict[str, Any] | None) -> list[str]:
    if not isinstance(carfax, dict):
        return []
    facts = carfax.get("facts") if isinstance(carfax.get("facts"), dict) else carfax
    if not isinstance(facts, dict) or not _has_structured_carfax_facts(facts):
        return []

    lines: list[str] = []
    owner_count = str(facts.get("owner_count") or "").strip()
    accident_damage = str(facts.get("accident_damage") or "").strip()
    accident_confidence = str(facts.get("accident_confidence") or "").strip().lower()
    title_brand = str(facts.get("title_brand") or "").strip()
    service_history = str(facts.get("service_history") or "").strip()
    value_badge = str(facts.get("value_badge") or "").strip()
    market_delta = str(facts.get("market_delta") or "").strip()
    market_position = str(facts.get("market_position") or "").strip()
    carfax_value = str(facts.get("carfax_value") or "").strip()
    last_owned_location = str(facts.get("last_owned_location") or "").strip()
    regular_oil_changes = bool(facts.get("regular_oil_changes"))
    detail_records = facts.get("detail_records_available")
    detail_records_text = f"{detail_records} detailed service records" if isinstance(detail_records, int) and detail_records > 0 else ""

    no_accident_claim = "no accidents reported" in accident_damage.lower() and accident_confidence == "explicit_clear"

    if owner_count and no_accident_claim:
        lines.append(f"{owner_count.capitalize()} and no accidents reported on CARFAX.")
    elif owner_count:
        lines.append(f"{owner_count.capitalize()} on CARFAX.")
    elif no_accident_claim:
        lines.append("No accidents reported on CARFAX.")

    accident_events = facts.get("accident_events") if isinstance(facts.get("accident_events"), list) else []
    if accident_events:
        primary = accident_events[0] if isinstance(accident_events[0], dict) else {}
        event_date = str(primary.get("date") or "").strip()
        event_desc = str(primary.get("description") or "").strip().rstrip(".")
        event_severity = str(primary.get("severity") or "").strip().lower()
        if event_desc:
            friendly_desc = re.sub(r"^Accident reported:\s*", "", event_desc, flags=re.IGNORECASE).strip()
            if event_date and event_date.lower() != "date not parsed":
                lines.append(f"CARFAX shows {friendly_desc[:1].lower() + friendly_desc[1:]} on {event_date}.")
            elif event_severity:
                lines.append(f"CARFAX shows {event_severity} {friendly_desc[:1].lower() + friendly_desc[1:]}.")
        elif accident_damage and "no accidents reported" not in accident_damage.lower():
            lines.append(accident_damage.replace("parsed", "").strip())
    elif accident_confidence in {"conflict", "signal_without_timeline"}:
        lines.append("Accident history needs a quick manual CARFAX check before making accident-free claims.")

    if service_history and "not parsed" not in service_history.lower():
        if detail_records_text and detail_records_text.lower() not in service_history.lower():
            if regular_oil_changes:
                lines.append(f"Service history looks strong with {detail_records_text} and regular oil changes.")
            else:
                lines.append(f"Service history looks solid with {detail_records_text}.")
        else:
            cleaned_service = service_history.replace("good service history", "Service history looks good")
            cleaned_service = cleaned_service.replace("strong service history", "Service history looks strong")
            cleaned_service = cleaned_service.replace("service history available", "Service history is available")
            lines.append(cleaned_service[:1].upper() + cleaned_service[1:] + ("" if cleaned_service.endswith(".") else "."))

    value_bits = [part for part in [value_badge, market_delta, carfax_value] if part]
    if value_badge:
        if market_delta and market_position:
            lines.append(f"CARFAX shows this as {value_badge}, about {market_delta} {market_position}.")
        elif carfax_value:
            lines.append(f"CARFAX values it around {carfax_value} and tags it {value_badge}.")
        else:
            lines.append(f"CARFAX tags it as {value_badge}.")
    elif market_delta and market_position:
        lines.append(f"Priced about {market_delta} {market_position}.")

    if title_brand and "no title brand/issues" in title_brand.lower():
        lines.append("No title issues reported on CARFAX.")
    elif title_brand and "not clearly parsed" not in title_brand.lower():
        lines.append(title_brand.rstrip(".") + ".")

    if last_owned_location:
        lines.append(f"Last owned in {last_owned_location}.")

    deduped: list[str] = []
    seen: set[str] = set()
    for line in lines:
        cleaned = " ".join(str(line or "").split()).strip()
        if not cleaned:
            continue
        key = cleaned.lower()
        if key in seen:
            continue
        seen.add(key)
        deduped.append(cleaned)
    return deduped[:5]


def _caption_asset_lines(vehicle: dict[str, Any]) -> list[str]:
    clean_vin = str(vehicle.get("vin") or vehicle.get("VIN") or vehicle.get("Vin") or "").strip().upper()
    cached_assets = _safe_read_json(_vehicle_assets_cache_path(clean_vin), {}) if clean_vin else {}
    if not isinstance(cached_assets, dict):
        cached_assets = {}
    cached_summary = cached_assets.get("carfax_summary") if isinstance(cached_assets, dict) else {}
    needs_asset_refresh = bool(
        clean_vin
        and (
            not cached_assets
            or not isinstance(cached_summary, dict)
            or _is_generic_carfax_summary(cached_summary)
        )
    )
    if needs_asset_refresh:
        try:
            refreshed_assets = _load_vehicle_assets(clean_vin, refresh=True)
            if isinstance(refreshed_assets, dict):
                cached_assets = refreshed_assets
        except Exception:
            pass
    summary_payload = _vehicle_asset_summary_payload(vehicle=vehicle, assets=cached_assets)
    lines: list[str] = []

    buyer = summary_payload.get("buyer_profile")
    vehicle_kind = str((buyer or {}).get("kind") if isinstance(buyer, dict) else "").lower()
    is_luxury_suv = "luxury" in vehicle_kind or any(
        token in str(vehicle.get("title") or "").lower()
        for token in ("aviator", "black label", "navigator")
    )
    if isinstance(buyer, dict) and buyer.get("buyer") and not is_luxury_suv:
        lines.extend(["", f"Best fit: {buyer.get('buyer')}."])

    sticker_highlights = summary_payload.get("sticker_highlights")
    if isinstance(sticker_highlights, list) and sticker_highlights:
        relevant = []
        for item in sticker_highlights:
            text = str(item or "").strip()
            if not text:
                continue
            lowered = text.lower()
            if is_luxury_suv and any(token in lowered for token in ("engine", "transmission", "drivetrain", "exterior", "interior")):
                relevant.append(text)
            elif not is_luxury_suv:
                relevant.append(text)
        if relevant:
            lines.extend(["", "Key vehicle facts:"])
            lines.extend(f"- {item}" for item in relevant[:5])

    carfax = summary_payload.get("carfax_summary")
    carfax_lines = _carfax_buyer_facing_lines(carfax if isinstance(carfax, dict) else None)
    if carfax_lines:
        lines.extend(["", "CARFAX highlights:"])
        lines.extend(f"- {line}" for line in carfax_lines)
    return lines


def _replace_generic_carfax_caption_line(caption: str, caption_asset_lines: list[str]) -> str:
    latest_line = ""
    for index, line in enumerate(caption_asset_lines):
        if line.strip().lower() in {"carfax summary:", "carfax highlights:"} and index + 1 < len(caption_asset_lines):
            latest_line = str(caption_asset_lines[index + 1] or "").strip()
            break
    if not latest_line:
        return caption

    generic_markers = (
        "official carfax link is available",
        "official carfax report link is present",
        "carfax report link present",
        "confirm accident history",
        "facts not parsed",
    )
    lines = caption.splitlines()
    output: list[str] = []
    index = 0
    replaced = False
    while index < len(lines):
        line = lines[index]
        output.append(line)
        if line.strip().lower() == "carfax summary:":
            index += 1
            while index < len(lines) and not lines[index].strip():
                output.append(lines[index])
                index += 1
            if index < len(lines):
                current = lines[index].strip().lower()
                if any(marker in current for marker in generic_markers):
                    output.append(latest_line)
                    replaced = True
                    index += 1
                    continue
        index += 1

    if not replaced and any(marker in caption.lower() for marker in generic_markers):
        output.extend(["", "CARFAX summary:", latest_line])
    return "\n".join(output)


def _caption_remove_robotic_carfax_language(caption: str) -> str:
    cleaned_lines: list[str] = []
    for line in caption.splitlines():
        text = " ".join(str(line or "").split()).strip()
        lowered = text.lower()
        if not text:
            cleaned_lines.append("")
            continue
        text = text.replace("CARFAX report parsed:", "")
        text = text.replace("parsed CARFAX text", "CARFAX")
        text = text.replace("No accidents reported in CARFAX.", "No accidents reported on CARFAX.")
        text = text.replace("No accidents reported in parsed CARFAX text.", "No accidents reported on CARFAX.")
        text = text.replace("No title brand/issues parsed from CARFAX text.", "No title issues reported on CARFAX.")
        text = re.sub(r"\baccident/damage event\(s\) parsed\b", "accident/damage event(s)", text, flags=re.IGNORECASE)
        text = re.sub(r"\brecord\(s\) parsed\b", "records", text, flags=re.IGNORECASE)
        text = re.sub(r"\bnot parsed from report text\b", "not available", text, flags=re.IGNORECASE)
        cleaned_lines.append(text.strip())
    return "\n".join(cleaned_lines).strip()


def _build_caption_from_vehicle(vehicle: dict[str, Any], caption_override: str | None = None) -> str:
    messenger_cta = (
        "Message my seller page for the full walkaround, extra photos, CARFAX, and the quickest reply: "
        f"{FACEBOOK_SELLER_MESSENGER_LINK}"
    )
    if caption_override and caption_override.strip():
        cleaned = re.sub(r"^https?://\S+\s*$", "", caption_override.strip(), flags=re.IGNORECASE | re.MULTILINE)
        price_line = f"{_facebook_post_price_text(vehicle)} down payment options for qualified buyers."
        if re.search(r"^Price\s*:", cleaned, flags=re.IGNORECASE | re.MULTILINE):
            cleaned = re.sub(r"^Price\s*:.*$", price_line, cleaned, flags=re.IGNORECASE | re.MULTILINE)
        else:
            parts = [line for line in cleaned.splitlines() if line.strip()]
            if parts:
                parts.insert(1, price_line)
                cleaned = "\n".join(parts)
            else:
                cleaned = price_line
        if "rigorous inspection process" not in cleaned.lower():
            cleaned = (
                f"{cleaned.strip()}\n\n"
                "This vehicle has been through a rigorous inspection process, repaired as needed, "
                "and checked so it passes applicable Florida safety requirements."
            )
        lower_cleaned = cleaned.lower()
        caption_asset_lines = _caption_asset_lines(vehicle)
        cleaned = _replace_generic_carfax_caption_line(cleaned, caption_asset_lines)
        lower_cleaned = cleaned.lower()
        if caption_asset_lines and not any(
            marker in lower_cleaned
            for marker in ("ideal buyer:", "best fit:", "sticker highlights:", "key vehicle facts:", "carfax summary:", "carfax highlights:")
        ):
            cleaned = f"{cleaned.strip()}\n" + "\n".join(caption_asset_lines)
        if FACEBOOK_SELLER_MESSENGER_LINK.lower() not in cleaned.lower():
            cleaned = f"{cleaned.strip()}\n\n{messenger_cta}"
        cleaned = re.sub(r"\$[1-9]\d{1,2},\d{3}(?:\.\d+)?(?:\s*(?:plus|\\+)?\s*tax(?:es)?\.?)?", _facebook_post_price_text(vehicle), cleaned, flags=re.IGNORECASE)
        return _caption_remove_robotic_carfax_language(cleaned.strip())

    title = str(vehicle.get("title") or vehicle.get("vin") or "Vehicle").strip()
    price_text = _facebook_post_price_text(vehicle)
    mileage = vehicle.get("mileage")
    vehicle_kind, buyer, _ = _vehicle_kind_and_buyer(vehicle)
    lines = [
        title,
        f"{price_text} down payment options for qualified buyers.",
    ]
    if mileage:
        lines.append(f"Mileage: {_facebook_caption_mileage(mileage)}")
    if vehicle_kind == "luxury SUV":
        lines.extend(
            [
                "",
                "Luxury SUV shoppers: this is for someone who wants a premium cabin, AWD confidence, family space, and a high-end look without stepping into new-car money.",
                "Best fit: families, professionals, rideshare premium drivers, and anyone who wants comfort, space, and confidence every day.",
                "Inspected, serviced as needed, and checked against applicable Florida safety requirements.",
                "Financing options available for qualified buyers.",
                messenger_cta,
            ]
        )
    else:
        lines.extend(
            [
                "",
                f"Built for {buyer}.",
                "Inspected, serviced as needed, and checked against applicable Florida safety requirements.",
                "Financing options available for qualified buyers.",
                messenger_cta,
            ]
        )
    if vehicle_kind in {"family van", "SUV"}:
        lines.extend(["", "Why it stands out:"])
        lines.append("- Practical space for family use, commuting, airport trips, and weekend plans.")
    elif vehicle_kind in {"pickup", "work truck"}:
        lines.extend(["", "Why it stands out:"])
        lines.append("- Great fit for work, towing, hauling, and daily driving.")
    lines.extend(_caption_asset_lines(vehicle))
    caption = "\n".join(lines)
    caption = re.sub(r"\nFinancing options available\.\s*$", "", caption, flags=re.IGNORECASE)
    caption = re.sub(r"\$[1-9]\d{1,2},\d{3}(?:\.\d+)?(?:\s*(?:plus|\\+)?\s*tax(?:es)?\.?)?", price_text, caption, flags=re.IGNORECASE)
    return _caption_remove_robotic_carfax_language(caption.strip())


def _select_vehicle_photo_urls(
    *,
    vehicle: dict[str, Any],
    selected_indexes: list[int],
    skip_indexes: list[int],
    limit: int,
) -> tuple[list[str], list[int], list[str]]:
    photos_raw = vehicle.get("photos") if isinstance(vehicle.get("photos"), list) else []
    urls = [url for url in (_extract_photo_url(entry) for entry in photos_raw) if url]
    skip = {index for index in skip_indexes if index >= 0} if len(urls) > 10 else set()
    if selected_indexes:
        indexes = [index for index in selected_indexes if 0 <= index < len(urls)]
    else:
        indexes = [index for index in range(len(urls)) if index not in skip]
    indexes = indexes[: max(1, int(limit))]
    selected_urls = [urls[index] for index in indexes]
    return selected_urls, indexes, urls


def _import_image_urls_for_vin(
    *,
    vin: str,
    urls: list[str],
    overwrite: bool = False,
) -> dict[str, Any]:
    clean_vin = str(vin or "").strip().upper()
    FML_IMAGES_DIR.mkdir(parents=True, exist_ok=True)

    imported: list[str] = []
    skipped_existing: list[str] = []
    errors: list[str] = []

    with httpx.Client(timeout=20.0, follow_redirects=True) as client:
        for index, url in enumerate(urls, start=1):
            try:
                response = client.get(url)
                if response.status_code >= 400:
                    raise RuntimeError(f"HTTP {response.status_code}")
                extension = _image_extension_from_url_and_content_type(
                    url,
                    response.headers.get("content-type"),
                )
                filename = f"{clean_vin}_FB_{index:02d}{extension}"
                target = FML_IMAGES_DIR / filename
                if target.exists() and not overwrite:
                    skipped_existing.append(filename)
                    continue
                target.write_bytes(response.content)
                imported.append(filename)
            except Exception as exc:
                errors.append(f"{url} -> {exc}")

    return {
        "ok": len(errors) == 0,
        "vin": clean_vin,
        "attempted": len(urls),
        "imported_count": len(imported),
        "imported": imported,
        "skipped_existing_count": len(skipped_existing),
        "skipped_existing": skipped_existing,
        "errors_count": len(errors),
        "errors": errors,
    }


def _run_one_click_post_from_inventory(request: FacebookOneClickPostRequest, *, queue_live: bool = True) -> dict[str, Any]:
    clean_vin = str(request.vin or "").strip().upper()
    vehicle = _find_vehicle_by_vin(clean_vin)
    if not vehicle:
        raise HTTPException(status_code=404, detail={"message": f"Vehicle not found for VIN {clean_vin}"})
    if request.mode == "live" and queue_live:
        return {
            "ok": True,
            "vin": clean_vin,
            "post_result": _queue_live_inventory_request(request, title=str(vehicle.get("title") or clean_vin)),
            "queued": True,
        }

    vehicle_for_post = dict(vehicle)
    vehicle_for_post["vin"] = clean_vin
    valuation = _jd_power_valuation_for_vin(clean_vin)
    jd_trade = _to_float(valuation.get("jd_power_trade_in")) if valuation else None
    if jd_trade and jd_trade > 0:
        pricing = _jd_power_ltv_from_pricing(inventory_price=vehicle_for_post.get("price"), jd_trade_value=jd_trade)
        vehicle_for_post["jd_power_trade_in"] = round(jd_trade, 2)
        vehicle_for_post["jd_power_ltv"] = pricing.get("ltv")
        vehicle_for_post["bank_sale_price"] = pricing.get("bank_sale_price")
        vehicle_for_post["bank_ltv_basis"] = pricing.get("ltv_basis")
    cached_assets = _read_vehicle_assets_cache(clean_vin)
    cached_asset_photos = cached_assets.get("photos") if isinstance(cached_assets.get("photos"), list) else []
    if cached_asset_photos and not vehicle_for_post.get("photos"):
        vehicle_for_post["photos"] = cached_asset_photos
    if request.auto_import_photos and not vehicle_for_post.get("photos"):
        try:
            assets = _load_vehicle_assets(clean_vin, refresh=False)
            if isinstance(assets.get("photos"), list) and assets.get("photos"):
                vehicle_for_post["photos"] = assets.get("photos")
        except Exception:
            pass
    carfax_assets = cached_assets if isinstance(cached_assets, dict) else {}
    quick_specs = carfax_assets.get("quick_specs") if isinstance(carfax_assets.get("quick_specs"), dict) else {}
    if isinstance(quick_specs, dict):
        field_map = {
            "model": "model",
            "body_style": "marketplace_body_style",
            "fuel_type": "marketplace_fuel_type",
            "drivetrain": "drivetrain",
            "engine": "engine",
            "transmission": "transmission",
            "exterior": "exterior",
            "interior": "interior",
            "stock_number": "stock_number",
        }
        for vehicle_key, spec_key in field_map.items():
            spec_value = quick_specs.get(spec_key)
            if spec_value not in (None, "", [], {}) and not vehicle_for_post.get(vehicle_key):
                vehicle_for_post[vehicle_key] = spec_value
    carfax_facts = carfax_assets.get("carfax_facts") if isinstance(carfax_assets.get("carfax_facts"), dict) else {}
    carfax_summary = carfax_assets.get("carfax_summary") if isinstance(carfax_assets.get("carfax_summary"), dict) else {}
    vehicle_for_post.update(
        {
            "carfax_facts": carfax_facts,
            "carfax_summary": carfax_summary,
        }
    )

    selected_urls, selected_indexes, all_urls = _select_vehicle_photo_urls(
        vehicle=vehicle_for_post,
        selected_indexes=request.selected_photo_indexes,
        skip_indexes=request.skip_photo_indexes,
        limit=request.photo_limit,
    )
    selection_fallback_used = False
    selection_fallback_reason: str | None = None
    if not selected_urls and all_urls and not request.selected_photo_indexes:
        # If the default thumbnail skip rule removes every photo (for example only one
        # photo at index 0), relax selection so live posting can still proceed.
        fallback_indexes = [index for index in range(len(all_urls)) if index != 2]
        if not fallback_indexes:
            fallback_indexes = [0]
        fallback_indexes = fallback_indexes[: max(1, int(request.photo_limit))]
        selected_indexes = fallback_indexes
        selected_urls = [all_urls[index] for index in selected_indexes]
        selection_fallback_used = True
        selection_fallback_reason = "skip_indexes_removed_all_photos"

    import_result: dict[str, Any] | None = None
    images_for_post: list[str] = []
    if request.auto_import_photos and selected_urls:
        import_result = _import_image_urls_for_vin(vin=clean_vin, urls=selected_urls, overwrite=False)
        images_for_post = _normalize_image_names(
            list(import_result.get("imported") or []) + list(import_result.get("skipped_existing") or [])
        )

    if not images_for_post:
        images_for_post = _suggest_images_for_vin(vin=clean_vin, limit=request.photo_limit)
    if request.mode == "live" and not images_for_post:
        raise HTTPException(
            status_code=400,
            detail={
                "message": "No postable images resolved for live posting.",
                "vin": clean_vin,
                "all_photo_count": len(all_urls),
                "selected_photo_indexes": selected_indexes,
                "selected_photo_urls_count": len(selected_urls),
                "skip_photo_indexes": request.skip_photo_indexes,
                "guidance": [
                    "Import vehicle photos into automation/facebook-marketplace-lister/images",
                    "Provide explicit selected_photo_indexes in request",
                    "Verify detail page photos are reachable over HTTP",
                ],
            },
        )

    caption = _build_caption_from_vehicle(vehicle_for_post, caption_override=request.caption_override)
    account_id = str(request.account_id or "").strip() or _resolve_default_account_id()
    suggested_down_payment = _facebook_post_price(vehicle_for_post)

    post_request = FacebookPostRequest(
        vin=clean_vin,
        title=str(vehicle_for_post.get("title") or clean_vin),
        price=vehicle_for_post.get("price") or "",
        model=vehicle_for_post.get("model"),
        mileage=vehicle_for_post.get("mileage"),
        body_style=vehicle_for_post.get("body_style"),
        fuel_type=vehicle_for_post.get("fuel_type"),
        condition=vehicle_for_post.get("condition") or "Good",
        drivetrain=vehicle_for_post.get("drivetrain"),
        engine=vehicle_for_post.get("engine"),
        transmission=vehicle_for_post.get("transmission"),
        location=_facebook_listing_location(vehicle_for_post.get("location")),
        exterior=vehicle_for_post.get("exterior"),
        interior=vehicle_for_post.get("interior"),
        detail_url=vehicle_for_post.get("detail_url"),
        description=caption,
        images=images_for_post,
        account_id=account_id,
        suggested_down_payment=suggested_down_payment,
        mode=request.mode,
    )

    result = _facebook_post_impl(post_request)
    if request.mode == "live" and result.get("live_success"):
        _mark_vehicle_posted(
            vin=clean_vin,
            mode="live",
            status_label="Posted",
            detail=str(result.get("live_detail") or ""),
        )

    _append_audit_event(
        "facebook_one_click_post",
        {
            "vin": clean_vin,
            "mode": request.mode,
            "selected_indexes": selected_indexes,
            "selected_urls_count": len(selected_urls),
            "selection_fallback_used": selection_fallback_used,
            "selection_fallback_reason": selection_fallback_reason,
            "images_for_post_count": len(images_for_post),
            "live_success": bool(result.get("live_success")) if request.mode == "live" else None,
        },
    )

    return {
        "ok": True,
        "vin": clean_vin,
        "caption": caption,
        "all_photo_count": len(all_urls),
        "selected_photo_indexes": selected_indexes,
        "selected_photo_urls": selected_urls,
        "selection_fallback_used": selection_fallback_used,
        "selection_fallback_reason": selection_fallback_reason,
        "images_for_post": images_for_post,
        "import_result": import_result,
        "prepared_post_request": post_request.model_dump(),
        "post_result": result,
    }


def _one_click_post_from_inventory(request: FacebookOneClickPostRequest) -> dict[str, Any]:
    return _run_one_click_post_from_inventory(request, queue_live=True)


def _batch_post_from_inventory(request: FacebookBatchPostRequest) -> dict[str, Any]:
    clean_vins: list[str] = []
    seen: set[str] = set()
    for raw_vin in request.vins:
        clean = str(raw_vin or "").strip().upper()
        if clean and clean not in seen:
            clean_vins.append(clean)
            seen.add(clean)
    if not clean_vins:
        raise HTTPException(status_code=400, detail={"message": "No VINs selected for Facebook batch."})

    account_id = str(request.account_id or "").strip() or _resolve_default_account_id()
    prepared: list[FacebookPostRequest] = []
    prep_results: list[dict[str, Any]] = []
    for clean_vin in clean_vins:
        draft = _one_click_post_from_inventory(
            FacebookOneClickPostRequest(
                vin=clean_vin,
                account_id=account_id,
                selected_photo_indexes=[],
                skip_photo_indexes=request.skip_photo_indexes,
                caption_override=None,
                mode="draft",
                auto_import_photos=request.auto_import_photos,
                photo_limit=request.photo_limit,
            )
        )
        prepared_payload = draft.get("prepared_post_request")
        if isinstance(prepared_payload, dict) and prepared_payload.get("vin"):
            prepared_payload = dict(prepared_payload)
            prepared_payload["mode"] = "live"
            prepared_payload["account_id"] = account_id
            prepared_payload["description"] = str(draft.get("caption") or prepared_payload.get("description") or "")
            prepared_payload["images"] = list(draft.get("images_for_post") or prepared_payload.get("images") or [])
            post_request = FacebookPostRequest(**prepared_payload)
        else:
            vehicle = _find_vehicle_by_vin(clean_vin)
            if not vehicle:
                raise HTTPException(status_code=404, detail={"message": f"Vehicle not found for VIN {clean_vin}"})
            vehicle_for_post = dict(vehicle)
            vehicle_for_post["vin"] = clean_vin
            post_request = FacebookPostRequest(
                vin=clean_vin,
                title=str(vehicle_for_post.get("title") or clean_vin),
                price=vehicle_for_post.get("price") or "",
                model=vehicle_for_post.get("model"),
                mileage=vehicle_for_post.get("mileage"),
                body_style=vehicle_for_post.get("body_style"),
                fuel_type=vehicle_for_post.get("fuel_type"),
                condition=vehicle_for_post.get("condition") or "Good",
                drivetrain=vehicle_for_post.get("drivetrain"),
                engine=vehicle_for_post.get("engine"),
                transmission=vehicle_for_post.get("transmission"),
                location=_facebook_listing_location(vehicle_for_post.get("location")),
                exterior=vehicle_for_post.get("exterior"),
                interior=vehicle_for_post.get("interior"),
                detail_url=vehicle_for_post.get("detail_url"),
                description=str(draft.get("caption") or ""),
                images=list(draft.get("images_for_post") or []),
                account_id=account_id,
                suggested_down_payment=_facebook_post_price(vehicle_for_post),
                mode="live",
            )
        prepared.append(post_request)
        prep_results.append(
            {
                "vin": clean_vin,
                "selected_photo_indexes": draft.get("selected_photo_indexes"),
                "images_count": len(post_request.images),
                "down_payment": post_request.suggested_down_payment,
            }
        )

    success, detail, live_result = _publish_live_batch(prepared, account_id)
    result_items = live_result.get("results") if isinstance(live_result.get("results"), list) else []
    posted_vins: set[str] = set()
    for item in result_items:
        if not isinstance(item, dict):
            continue
        clean_vin = str(item.get("vin") or "").strip().upper()
        if not clean_vin:
            continue
        state = _marketplace_state_from_publish_result(item)
        if state.get("marketplace_status") == "live":
            posted_vins.add(clean_vin)
        _set_facebook_vehicle_status(
            vin=clean_vin,
            mode="live",
            marketplace_status=str(state.get("marketplace_status") or "needs_review"),
            detail=detail,
            listing_url=str(state.get("listing_url") or ""),
            confirmation=state.get("confirmation") if isinstance(state.get("confirmation"), dict) else None,
        )

    _append_audit_event(
        "facebook_batch_post",
        {
            "vins": clean_vins,
            "posted_vins": sorted(posted_vins),
            "success": success,
            "detail": detail,
        },
    )
    return {
        "ok": success,
        "mode": "live",
        "requested": len(clean_vins),
        "prepared": prep_results,
        "posted": len(posted_vins),
        "failed": max(0, len(clean_vins) - len(posted_vins)),
        "live_success": bool(success and len(posted_vins) == len(clean_vins)),
        "live_detail": detail,
        "live_result": live_result,
    }


DEFAULT_BANK_PROFILES: list[dict[str, Any]] = [
    {
        "code": "ALLY",
        "name": "Ally",
        "min_score": 620,
        "max_ltv": 130,
        "max_pti": 18,
        "max_dti": 52,
        "max_derogatories": 3,
        "max_utilization": 88,
    },
    {
        "code": "CAP1",
        "name": "Capital One Auto",
        "min_score": 600,
        "max_ltv": 125,
        "max_pti": 17,
        "max_dti": 50,
        "max_derogatories": 4,
        "max_utilization": 90,
    },
    {
        "code": "CHASE",
        "name": "Chase Auto",
        "min_score": 660,
        "max_ltv": 120,
        "max_pti": 16,
        "max_dti": 48,
        "max_derogatories": 2,
        "max_utilization": 75,
    },
    {
        "code": "CU_LOCAL",
        "name": "Local Credit Union",
        "min_score": 640,
        "max_ltv": 115,
        "max_pti": 15,
        "max_dti": 45,
        "max_derogatories": 2,
        "max_utilization": 70,
    },
]


def _coerce_generated_bank_profile(raw: Any) -> dict[str, Any] | None:
    if not isinstance(raw, dict):
        return None
    name = str(raw.get("name") or raw.get("bank") or "").strip()
    code = str(raw.get("code") or "").strip().upper()
    if not name or not code:
        return None

    def as_int(key: str, fallback: int) -> int:
        try:
            return int(float(raw.get(key, fallback)))
        except Exception:
            return fallback

    return {
        "code": code,
        "name": name,
        "min_score": as_int("min_score", 620),
        "max_ltv": as_int("max_ltv", 125),
        "max_pti": as_int("max_pti", 17),
        "max_dti": as_int("max_dti", 48),
        "max_derogatories": as_int("max_derogatories", 3),
        "max_utilization": as_int("max_utilization", 80),
        "max_term_months": as_int("max_term_months", 84),
        "weight": float(raw.get("weight", 1.0) or 1.0),
        "tier": raw.get("tier"),
        "confidence": raw.get("confidence"),
        "source_files": raw.get("source_files") if isinstance(raw.get("source_files"), list) else [],
        "source_links": raw.get("source_links") if isinstance(raw.get("source_links"), list) else [],
        "stips": raw.get("stips") if isinstance(raw.get("stips"), list) else [],
        "restrictions": raw.get("restrictions") if isinstance(raw.get("restrictions"), list) else [],
        "notes": raw.get("notes") if isinstance(raw.get("notes"), list) else [],
        "decoded_doc_count": as_int("decoded_doc_count", 0),
        "decoded_link_count": as_int("decoded_link_count", 0),
        "decoded_evidence_count": as_int("decoded_evidence_count", 0),
    }


def _load_generated_bank_profiles() -> list[dict[str, Any]]:
    candidates = [
        _safe_read_json(BANK_PROFILES_GENERATED_PATH, {}),
        _safe_read_json(SALES_ASSISTANT_BANKS_PATH, {}),
    ]
    for payload in candidates:
        items = []
        if isinstance(payload, dict):
            raw_items = payload.get("profiles")
            if not isinstance(raw_items, list):
                raw_items = payload.get("policies")
            if isinstance(raw_items, list):
                items = raw_items
        elif isinstance(payload, list):
            items = payload
        profiles = [profile for profile in (_coerce_generated_bank_profile(item) for item in items) if profile]
        if profiles:
            return profiles
    return []


def _active_bank_profiles() -> list[dict[str, Any]]:
    generated = _load_generated_bank_profiles()
    return generated or DEFAULT_BANK_PROFILES


def _load_bank_brain_history() -> list[dict[str, Any]]:
    payload = _safe_read_json(BANK_BRAIN_HISTORY_PATH, {"items": []})
    if isinstance(payload, dict):
        items = payload.get("items", [])
    else:
        items = payload
    if not isinstance(items, list):
        return []
    return [entry for entry in items if isinstance(entry, dict)]


def _append_bank_brain_history(entry: dict[str, Any]) -> None:
    items = _load_bank_brain_history()
    items.append(entry)
    _safe_write_json(BANK_BRAIN_HISTORY_PATH, {"items": items[-1000:]})


def _bank_bias_scores() -> dict[str, float]:
    history = _load_bank_brain_history()
    grouped: dict[str, list[str]] = {}
    for entry in history:
        bank = str(entry.get("bank_code", "")).strip().upper()
        outcome = str(entry.get("outcome", "")).strip().lower()
        if not bank or outcome not in {"approved", "declined", "countered"}:
            continue
        grouped.setdefault(bank, []).append(outcome)

    bias: dict[str, float] = {}
    for bank, outcomes in grouped.items():
        approved = outcomes.count("approved")
        ratio = approved / len(outcomes)
        bias[bank] = (ratio - 0.5) * 20
    return bias


def _extract_first_number(text: str, pattern: str) -> float | None:
    match = re.search(pattern, text, flags=re.IGNORECASE)
    if not match:
        return None
    try:
        return float(match.group(1))
    except Exception:
        return None


def _normalize_currency(text: str) -> float | None:
    if text is None:
        return None
    value = str(text).strip().replace(",", "").replace("$", "")
    value = re.sub(r"\s+", "", value)
    if not value:
        return None
    try:
        return float(value)
    except Exception:
        return None


def _split_name_parts(raw_name: str) -> dict[str, str]:
    name = re.sub(r"\s+", " ", str(raw_name or "").strip()).strip(",")
    if not name:
        return {}

    tokens = [part.strip() for part in re.split(r"\s+", name) if part.strip()]
    if not tokens:
        return {}

    suffixes = {"jr", "sr", "ii", "iii", "iv", "v", "md", "esq", "phd"}
    suffix = ""
    if tokens and tokens[-1].lower().replace(".", "") in suffixes:
        suffix = tokens.pop(-1).strip(".")

    if "," in name:
        chunks = [chunk.strip() for chunk in name.split(",", maxsplit=1)]
        last = re.sub(r"\s+", " ", chunks[0]).strip()
        rest = re.sub(r"\s+", " ", chunks[1]).strip() if len(chunks) > 1 else ""
        tokens = [part for part in re.split(r"\s+", f"{rest} {last}") if part]

    if len(tokens) == 1:
        return {"first_name": tokens[0], "suffix": suffix}
    if len(tokens) == 2:
        return {"first_name": tokens[0], "last_name": tokens[1], "suffix": suffix}
    return {
        "first_name": tokens[0],
        "middle_name": " ".join(tokens[1:-1]),
        "last_name": tokens[-1],
        "suffix": suffix,
    }


def _to_routeone_dob_parts(raw_date: str) -> tuple[str, str, str]:
    text = str(raw_date or "").strip()
    if not text:
        return "", "", ""

    for fmt in ("%m/%d/%Y", "%m-%d-%Y", "%m/%d/%y", "%m-%d-%y", "%Y-%m-%d", "%Y/%m/%d"):
        try:
            parsed = datetime.strptime(text, fmt)
            return f"{parsed.month:02d}", f"{parsed.day:02d}", f"{parsed.year:04d}"
        except Exception:
            pass

    compact = re.sub(r"[^0-9/.-]", "", text)
    m = re.match(r"^(\d{1,2})[/.-](\d{1,2})[/.-](\d{2,4})$", compact)
    if m:
        month, day, year = m.groups()
        year = ("20" + year) if len(year) == 2 and int(year) < 40 else ("19" + year) if len(year) == 2 else year
        return month.zfill(2), day.zfill(2), year.zfill(4)
    return "", "", ""


def _extract_labeled_value(text: str, labels: list[str]) -> str:
    if not text:
        return ""
    lower = text.lower()
    for label in labels:
        pattern = rf"(?im)^{re.escape(label)}\s*[:\-]\s*(.+?)$"
        match = re.search(pattern, text, flags=re.MULTILINE)
        if match:
            value = match.group(1).strip()
            if value:
                return value
    return ""


def _extract_phone_like(raw: str) -> str:
    candidates = re.findall(r"\+?\d[\d\-().\s]{7,}\d", str(raw or ""))
    for candidate in candidates:
        digits = re.sub(r"[^0-9]", "", candidate)
        if len(digits) == 10:
            return f"{digits[0:3]}-{digits[3:6]}-{digits[6:]}"
        if len(digits) == 11 and digits.startswith("1"):
            digits = digits[1:]
            return f"{digits[0:3]}-{digits[3:6]}-{digits[6:]}"
    return ""


def _configure_tesseract_binary() -> str | None:
    env_value = str(os.getenv("TESSERACT_CMD") or os.getenv("PYTESSERACT_CMD") or "").strip()
    candidates: list[Path] = []
    if env_value:
        candidates.append(Path(env_value))
    candidates.extend(
        [
            Path("/usr/bin/tesseract"),
            Path("/usr/local/bin/tesseract"),
            Path(r"C:\Program Files\Tesseract-OCR\tesseract.exe"),
            Path(r"C:\Program Files (x86)\Tesseract-OCR\tesseract.exe"),
        ]
    )
    for candidate in candidates:
        try:
            if candidate.exists():
                return str(candidate)
        except Exception:
            continue
    return shutil.which("tesseract")


def _ocr_candidate_score(text: str) -> float:
    clean = str(text or "").strip()
    if not clean:
        return 0.0
    keywords = [
        "credit",
        "score",
        "income",
        "employment",
        "address",
        "residence",
        "tradeline",
        "dti",
        "bankruptcy",
        "auto",
        "payment",
        "ssn",
        "dob",
        "borrower",
    ]
    lowered = clean.lower()
    keyword_hits = sum(1 for keyword in keywords if keyword in lowered)
    digit_hits = len(re.findall(r"\d", clean))
    line_hits = len([line for line in clean.splitlines() if line.strip()])
    alpha_hits = len(re.findall(r"[a-zA-Z]", clean))
    return (
        min(len(clean), 6000) * 0.01
        + keyword_hits * 14
        + min(digit_hits, 400) * 0.08
        + min(line_hits, 120) * 0.8
        + min(alpha_hits, 2000) * 0.01
    )


def _prepare_ocr_images(image: Any) -> list[Any]:
    from PIL import ImageFilter, ImageOps, ImageEnhance  # type: ignore

    grayscale = ImageOps.grayscale(image)
    boosted = ImageEnhance.Contrast(grayscale).enhance(2.2)
    sharpened = boosted.filter(ImageFilter.SHARPEN)
    autocontrast = ImageOps.autocontrast(boosted, cutoff=1)
    binary = autocontrast.point(lambda px: 255 if px > 168 else 0, mode="1").convert("L")
    soft_binary = autocontrast.point(lambda px: 255 if px > 145 else 0, mode="1").convert("L")
    inverted = ImageOps.invert(autocontrast)
    return [grayscale, boosted, sharpened, autocontrast, binary, soft_binary, inverted]


def _ocr_text_from_pil_image(image: Any) -> tuple[str, dict[str, Any]]:
    diagnostics: dict[str, Any] = {"ocr_engine": "tesseract", "ocr_candidates": 0}
    try:
        import pytesseract  # type: ignore
    except Exception as exc:
        diagnostics["ocr_error"] = f"pytesseract import failed: {exc}"
        return "", diagnostics

    tesseract_cmd = _configure_tesseract_binary()
    if not tesseract_cmd:
        diagnostics["ocr_error"] = "tesseract binary not found"
        return "", diagnostics
    pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
    diagnostics["tesseract_cmd"] = tesseract_cmd

    candidates: list[str] = []
    seen: set[str] = set()
    for prepared in _prepare_ocr_images(image):
        for config in (
            "--oem 1 --psm 4",
            "--oem 1 --psm 6",
            "--oem 1 --psm 11",
            "--oem 3 --psm 12",
        ):
            try:
                candidate = (pytesseract.image_to_string(prepared, config=config) or "").strip()
            except Exception:
                continue
            if len(candidate) < 24:
                continue
            normalized = re.sub(r"\s+", " ", candidate).strip().lower()
            if normalized in seen:
                continue
            seen.add(normalized)
            candidates.append(candidate)

    diagnostics["ocr_candidates"] = len(candidates)
    if not candidates:
        diagnostics["ocr_error"] = diagnostics.get("ocr_error") or "tesseract returned no readable text"
        return "", diagnostics
    candidates.sort(key=_ocr_candidate_score, reverse=True)
    best = candidates[0].strip()
    diagnostics["ocr_best_score"] = round(_ocr_candidate_score(best), 2)
    return best, diagnostics


def _ocr_text_from_pdf_region(page: Any, rect: Any, *, scale: float = 4.0) -> tuple[str, dict[str, Any]]:
    try:
        import fitz  # type: ignore
        from PIL import Image  # type: ignore

        pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), clip=rect, alpha=False)
        image = Image.open(BytesIO(pix.tobytes("png")))
        text, diagnostics = _ocr_text_from_pil_image(image)
        diagnostics["region_scale"] = scale
        return text, diagnostics
    except Exception as exc:
        return "", {"ocr_error": str(exc), "region_scale": scale}


def _ncc_clarity_region_map(page_number: int) -> list[tuple[str, tuple[float, float, float, float]]]:
    if page_number == 1:
        return [
            ("identity_header", (0.11, 0.02, 0.46, 0.11)),
            ("personal_info", (0.11, 0.24, 0.41, 0.45)),
            ("credit_score_panel", (0.41, 0.24, 0.90, 0.45)),
            ("alert_and_messages", (0.11, 0.46, 0.91, 0.64)),
            ("credit_summary", (0.11, 0.64, 0.91, 0.76)),
            ("auto_analytics", (0.11, 0.75, 0.91, 0.82)),
            ("account_summary", (0.11, 0.82, 0.91, 0.98)),
        ]
    if page_number == 2:
        return [
            ("score_factors_summary", (0.11, 0.03, 0.91, 0.22)),
            ("employment_summary", (0.11, 0.22, 0.91, 0.34)),
            ("inquiries", (0.11, 0.34, 0.91, 0.97)),
        ]
    if page_number == 3:
        return [
            ("mortgage_accounts", (0.11, 0.00, 0.91, 0.15)),
            ("revolving_accounts", (0.11, 0.15, 0.91, 0.66)),
            ("installment_accounts", (0.11, 0.66, 0.91, 0.98)),
        ]
    if page_number == 4:
        return [
            ("bureau_footer", (0.11, 0.08, 0.45, 0.36)),
            ("generated_timestamp", (0.76, 0.26, 0.92, 0.34)),
        ]
    return []


def _extract_ncc_clarity_pdf_regions(raw: bytes) -> tuple[list[str], dict[str, Any]]:
    diagnostics: dict[str, Any] = {"format": "pdf_region_ocr", "pages": 0, "regions": 0}
    text_parts: list[str] = []
    try:
        import fitz  # type: ignore

        with fitz.open(stream=raw, filetype="pdf") as doc:
            diagnostics["pages"] = len(doc)
            for page_index, page in enumerate(doc, start=1):
                regions = _ncc_clarity_region_map(page_index)
                page_rect = page.rect
                region_payloads: list[str] = []
                page_meta: list[dict[str, Any]] = []
                for label, (x0, y0, x1, y1) in regions:
                    clip = fitz.Rect(
                        page_rect.x0 + page_rect.width * x0,
                        page_rect.y0 + page_rect.height * y0,
                        page_rect.x0 + page_rect.width * x1,
                        page_rect.y0 + page_rect.height * y1,
                    )
                    candidate, meta = _ocr_text_from_pdf_region(page, clip, scale=4.0)
                    page_meta.append({"label": label, **meta})
                    if candidate and len(candidate.strip()) >= 12:
                        region_payloads.append(f"[{label}]\n{candidate.strip()}")
                        diagnostics["regions"] = int(diagnostics.get("regions") or 0) + 1
                if region_payloads:
                    text_parts.append(f"[page {page_index}]\n" + "\n\n".join(region_payloads))
                if page_meta:
                    diagnostics.setdefault("region_details", []).append({"page": page_index, "items": page_meta})
    except Exception as exc:
        diagnostics["region_error"] = str(exc)
    return text_parts, diagnostics


def _extract_route_one_fill_from_text(report_text: str) -> dict[str, Any]:
    text = report_text or ""
    if not text:
        return {}

    parsed: dict[str, Any] = {}
    normalized = text.replace("\r", "")
    text_lower = normalized.lower()

    labeled_name = (
        _extract_labeled_value(normalized, ["applicant", "borrower", "primary applicant", "primary", "name"])
    )
    name_raw = labeled_name or ""
    if not name_raw:
        match = re.search(r"\b([A-Z][A-Za-z'\-]+(?:\s+[A-Z][A-Za-z'\-]+){1,4})\b", normalized)
        if match:
            name_raw = match.group(1)
    parts = _split_name_parts(name_raw)
    if parts.get("first_name"):
        parsed["first_name"] = parts["first_name"]
    if parts.get("middle_name"):
        parsed["middle_name"] = parts["middle_name"]
    if parts.get("last_name"):
        parsed["last_name"] = parts["last_name"]
    if parts.get("suffix"):
        parsed["suffix"] = parts["suffix"]

    dob = (
        _extract_labeled_value(normalized, ["date of birth", "dob", "birth date"])
        or re.search(
            r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})\b",
            normalized,
        ).group(0)
        if re.search(r"\b(\d{1,2}[/-]\d{1,2}[/-]\d{2,4}|\d{4}-\d{2}-\d{2})\b", normalized)
        else ""
    )
    if dob:
        month, day, year = _to_routeone_dob_parts(dob)
        if month:
            parsed["dob_month"] = month
        if day:
            parsed["dob_day"] = day
        if year:
            parsed["dob_year"] = year

    ssn = _extract_labeled_value(normalized, ["ssn", "social security", "social security number"]) or re.search(
        r"\b(\d{3}-\d{2}-\d{4}|\d{9})\b", normalized
    )
    if isinstance(ssn, re.Match):
        value = ssn.group(1).replace(" ", "")
    else:
        value = str(ssn).strip() if ssn else ""
    if value and re.fullmatch(r"\d{9}", value):
        value = f"{value[0:3]}-{value[3:5]}-{value[5:]}"
    if value:
        parsed["ssn"] = value

    email = _extract_labeled_value(normalized, ["email", "e-mail", "email address"]) or re.search(
        r"[a-zA-Z0-9._%+-]+@[a-zA-Z0-9.-]+\.[A-Za-z]{2,}",
        normalized,
    )
    if isinstance(email, re.Match):
        parsed["email"] = email.group(0).strip()
    elif isinstance(email, str) and email:
        parsed["email"] = email.strip()

    address = _extract_labeled_value(normalized, ["address", "mailing address", "current address"])
    if address and len(address) < 6:
        address = ""
    if address:
        parsed["address"] = address
        addr_lines = [part.strip() for part in address.split(",")]
        if len(addr_lines) >= 2:
            city_state_zip = ", ".join(addr_lines[1:]).strip()
            m = re.match(r"^\s*([A-Za-z .'-]+?)\s+([A-Z]{2})\s+(\d{5}(?:-\d{4})?)\s*$", city_state_zip)
            if m:
                parsed["city"] = m.group(1).strip()
                parsed["state"] = m.group(2).strip()
                parsed["zip"] = m.group(3).strip()
        parsed.setdefault("address", address)
    else:
        line_match = re.search(
            r"^\s*([0-9]+ [^,\n]+,\s*[A-Za-z .'-]+\s*,\s*[A-Z]{2}\s+\d{5}(?:-\d{4})?)\s*$",
            normalized,
            flags=re.MULTILINE,
        )
        if line_match:
            address_value = line_match.group(1).strip()
            parsed["address"] = address_value

            chunks = [chunk.strip() for chunk in address_value.split(",")]
            if len(chunks) >= 2:
                parsed["city"] = chunks[-2].strip()
                parsed.setdefault("state", "")
                state_zip = re.search(r"([A-Z]{2})\s+(\d{5}(?:-\d{4})?)", chunks[-1])
                if state_zip:
                    parsed["state"] = state_zip.group(1)
                    parsed["zip"] = state_zip.group(2)

    parsed["home_phone"] = _extract_labeled_value(normalized, ["home phone", "home"]) or _extract_phone_like(
        normalized
    )
    parsed["cellular_phone"] = _extract_labeled_value(normalized, ["cell", "cell phone", "mobile", "phone"]) or (
        _extract_phone_like(normalized) if not parsed.get("home_phone") else ""
    )

    parsed["employment_type"] = _extract_labeled_value(
        normalized, ["employment type", "employment", "empl type", "job type"]
    ) or ""
    parsed["employment_status"] = _extract_labeled_value(normalized, ["employment status", "position type", "job status"]) or ""
    parsed["employment_title"] = _extract_labeled_value(normalized, ["job title", "employment title", "position"]) or ""
    parsed["employer"] = _extract_labeled_value(normalized, ["employer", "employer name", "company"]) or ""

    years_at_address = (
        _extract_first_number(normalized, r"(?:time at (?:address|residence|home))\s*[:\-]?\s*([0-9]{1,2})\s*years?") 
        if re.search(r"time at (?:address|residence|home)", normalized, flags=re.IGNORECASE)
        else None
    )
    if years_at_address is not None:
        parsed["time_at_address_years"] = str(int(years_at_address))

    months_at_address = _extract_first_number(
        normalized,
        r"(?:time at (?:address|residence|home))\s*[:\-]?(?:.*?)([0-9]{1,2})\s*months?",
    )
    if months_at_address is not None:
        parsed["time_at_address_months"] = str(int(months_at_address))

    years_at_job = (
        _extract_first_number(
            normalized,
            r"(?:time at (?:job|work))\s*[:\-]?\s*([0-9]{1,2})\s*years?",
        )
    )
    if years_at_job is not None:
        parsed["time_at_job_years"] = str(int(years_at_job))

    months_at_job = _extract_first_number(
        normalized,
        r"(?:time at (?:job|work))\s*[:\-]?(?:.*?)([0-9]{1,2})\s*months?",
    )
    if months_at_job is not None:
        parsed["time_at_job_months"] = str(int(months_at_job))

    monthly_income = _extract_labeled_value(
        normalized, ["monthly income", "gross monthly income", "base monthly income"]
    ) or re.search(r"\bmonthly\s+income\s*:?\s*\$?\s*([0-9][0-9,]*\.?[0-9]*)", normalized, flags=re.IGNORECASE)
    if isinstance(monthly_income, re.Match):
        value = _normalize_currency(monthly_income.group(1))
        if value is not None:
            parsed["monthly_income"] = round(value, 2)
            parsed["other_income_amount"] = str(int(value))
            parsed["other_income_source"] = parsed.get("other_income_source") or "Monthly Income"
            parsed["income_interval"] = "Monthly"
    elif isinstance(monthly_income, str) and monthly_income:
        normalized_value = _normalize_currency(monthly_income)
        if normalized_value is not None:
            parsed["monthly_income"] = round(normalized_value, 2)
            parsed["other_income_amount"] = str(int(normalized_value))
            parsed["other_income_source"] = parsed.get("other_income_source") or "Monthly Income"
            parsed["income_interval"] = "Monthly"

    annual_income = _extract_labeled_value(normalized, ["annual income", "yearly income"]) or re.search(
        r"\bannual\s+income\s*:?\s*\$?\s*([0-9][0-9,]*\.?[0-9]*)",
        normalized,
        flags=re.IGNORECASE,
    )
    if isinstance(annual_income, re.Match):
        value = _normalize_currency(annual_income.group(1))
        if value is not None and "other_income_amount" not in parsed:
            parsed["monthly_income"] = round(value / 12, 2)
            parsed["other_income_amount"] = str(int(value / 12))
            parsed["other_income_source"] = parsed.get("other_income_source") or "Annual Income"
            parsed["income_interval"] = "Monthly"
    elif isinstance(annual_income, str) and annual_income:
        value = _normalize_currency(annual_income)
        if value is not None and "other_income_amount" not in parsed:
            parsed["monthly_income"] = round(value / 12, 2)
            parsed["other_income_amount"] = str(int(value / 12))
            parsed["other_income_source"] = parsed.get("other_income_source") or "Annual Income"
            parsed["income_interval"] = "Monthly"

    parsed = {key: value for key, value in parsed.items() if value not in (None, "", [], {})}
    parsed["years_at_address"] = _extract_first_number(
        text_lower,
        r"years at address\s*[:\-]?\s*([0-9]{1,2})",
    )
    parsed["years_at_job"] = _extract_first_number(
        text_lower,
        r"years at job\s*[:\-]?\s*([0-9]{1,2})",
    )

    return parsed


def _extract_credit_metrics(report_text: str, structured_data: dict[str, Any]) -> dict[str, Any]:
    text = report_text or ""
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    table_metrics: dict[str, Any] = {}
    for header_line, value_line in zip(lines, lines[1:]):
        if "\t" not in header_line or "\t" not in value_line:
            continue
        headers = [_normalize_header(cell) for cell in header_line.split("\t")]
        values = [cell.strip() for cell in value_line.split("\t")]
        if len(values) < 2:
            continue
        for index, header in enumerate(headers):
            if index >= len(values):
                continue
            value = _to_float(values[index])
            if value is None:
                continue
            if "score" in header or "fico" in header or "beacon" in header:
                table_metrics.setdefault("score", int(value))
            elif "tradeline" in header or "trade line" in header or header in {"trades", "accounts"}:
                table_metrics.setdefault("tradelines", int(value))
            elif "derog" in header or "collection" in header or "charge off" in header or "repo" in header:
                table_metrics.setdefault("derogatories", int(value))
            elif "util" in header:
                table_metrics.setdefault("utilization", float(value))
            elif header == "dti" or "debt to income" in header or "debt ratio" in header:
                table_metrics.setdefault("dti", float(value))

    score = structured_data.get("score")
    if score is None:
        score = table_metrics.get("score")
    if score is None:
        score_match = None
        for pattern in (
            r"(?:fico|credit|beacon|transunion|equifax|experian)\s*(?:score|auto score)?[^0-9\n]{0,20}([3-8][0-9]{2})",
            r"\bscore[^0-9\n]{0,20}([3-8][0-9]{2})",
            r"\b([3-8][0-9]{2})\s*(?:fico|credit score|beacon)\b",
        ):
            score_match = re.search(pattern, text, flags=re.IGNORECASE)
            if score_match:
                break
        if score_match:
            score = int(score_match.group(1))

    tradelines = structured_data.get("tradelines")
    if tradelines is None:
        tradelines = table_metrics.get("tradelines")
    if tradelines is None:
        value = None
        for pattern in (
            r"tradelines?\s*[:\-]?\s*([0-9]{1,3})",
            r"(?:open\s*)?(?:auto\s*)?(?:accounts?|trades?)\s*[:\-]?\s*([0-9]{1,3})",
        ):
            value = _extract_first_number(text, pattern)
            if value is not None:
                break
        tradelines = int(value) if value is not None else None

    derogatories = structured_data.get("derogatories")
    if derogatories is None:
        derogatories = table_metrics.get("derogatories")
    if derogatories is None:
        value = None
        for pattern in (
            r"derogator(?:y|ies)\s*[:\-]?\s*([0-9]{1,2})",
            r"(?:collections?|charge[-\s]?offs?|repos?)\s*[:\-]?\s*([0-9]{1,2})",
        ):
            value = _extract_first_number(text, pattern)
            if value is not None:
                break
        if value is None:
            hits = re.findall(
                r"charge[-\s]?off|collection|repossession|bankruptcy|foreclosure|late\s+90",
                text,
                flags=re.IGNORECASE,
            )
            derogatories = len(hits)
        else:
            derogatories = int(value)

    utilization = structured_data.get("utilization")
    if utilization is None:
        utilization = table_metrics.get("utilization")
    if utilization is None:
        for pattern in (
            r"(?:utilization|revolving\s*utilization|util)\s*[:\-]?\s*([0-9]{1,3}(?:\.[0-9]+)?)\s*%?",
            r"([0-9]{1,3}(?:\.[0-9]+)?)\s*%\s*(?:utilization|revolving)",
        ):
            utilization = _extract_first_number(text, pattern)
            if utilization is not None:
                break

    dti = structured_data.get("dti")
    if dti is None:
        dti = table_metrics.get("dti")
    if dti is None:
        for pattern in (
            r"(?:dti|debt\s*to\s*income|debt\s*ratio|debt\s*income)\s*[:\-]?\s*([0-9]{1,3}(?:\.[0-9]+)?)\s*%?",
            r"([0-9]{1,3}(?:\.[0-9]+)?)\s*%\s*(?:dti|debt\s*to\s*income)",
        ):
            dti = _extract_first_number(text, pattern)
            if dti is not None:
                break

    monthly_income = structured_data.get("monthly_income")
    if monthly_income is None:
        monthly_income = structured_data.get("gross_monthly_income")
    if monthly_income is None:
        for pattern in (
            r"(?:gross\s+)?monthly\s+income\s*[:\-]?\s*\$?\s*([0-9][0-9,]*\.?[0-9]*)",
            r"(?:income|salary|wages)\s*[:\-]?\s*\$?\s*([0-9][0-9,]*\.?[0-9]*)\s*/\s*(?:mo|month|monthly)\b",
        ):
            monthly_income = _extract_first_number(text, pattern)
            if monthly_income is not None:
                break
    if monthly_income is None:
        annual_income = None
        for pattern in (
            r"(?:annual|yearly)\s+income\s*[:\-]?\s*\$?\s*([0-9][0-9,]*\.?[0-9]*)",
            r"\$?\s*([0-9][0-9,]*\.?[0-9]*)\s*/\s*(?:yr|year|annual)\b",
        ):
            annual_income = _extract_first_number(text, pattern)
            if annual_income is not None:
                break
        if annual_income is not None:
            monthly_income = annual_income / 12.0

    years_at_address = structured_data.get("years_at_address")
    if years_at_address is None:
        years_at_address = _extract_first_number(
            text,
            r"(?:years?\s+at\s+(?:address|residence|home)|time at (?:address|residence|home))\s*[:\-]?\s*([0-9]{1,2}(?:\.[0-9]+)?)",
        )

    years_at_job = structured_data.get("years_at_job")
    if years_at_job is None:
        years_at_job = _extract_first_number(
            text,
            r"(?:years?\s+at\s+(?:job|work|employer)|time at (?:job|work|employer))\s*[:\-]?\s*([0-9]{1,2}(?:\.[0-9]+)?)",
        )

    flags: list[str] = []
    for keyword, label in [
        ("thin file", "Thin-file profile"),
        ("bankruptcy", "Bankruptcy noted"),
        ("repossession", "Repossession noted"),
        ("collection", "Collections present"),
        ("late payment", "Late payment history"),
    ]:
        if keyword in text.lower():
            flags.append(label)

    return {
        "score": int(score) if score is not None else None,
        "tradelines": int(tradelines) if tradelines is not None else None,
        "derogatories": int(derogatories) if derogatories is not None else None,
        "utilization": float(utilization) if utilization is not None else None,
        "monthly_income": float(monthly_income) if monthly_income is not None else None,
        "dti": float(dti) if dti is not None else None,
        "current_dti": float(dti) if dti is not None else None,
        "years_at_address": float(years_at_address) if years_at_address is not None else None,
        "years_at_job": float(years_at_job) if years_at_job is not None else None,
        "risk_flags": flags,
    }


def _estimate_monthly_payment(principal: float, apr: float, term_months: int) -> float:
    if principal <= 0 or term_months <= 0:
        return 0.0
    monthly_rate = (apr / 100) / 12
    if monthly_rate <= 0:
        return principal / term_months
    numerator = principal * monthly_rate * ((1 + monthly_rate) ** term_months)
    denominator = ((1 + monthly_rate) ** term_months) - 1
    if denominator == 0:
        return principal / term_months
    return numerator / denominator


def _recommend_banks(metrics: dict[str, Any], structure: dict[str, Any]) -> dict[str, Any]:
    score = metrics.get("score")
    derogatories = metrics.get("derogatories")
    utilization = metrics.get("utilization")
    dti = structure.get("dti")
    ltv = structure.get("ltv")
    pti = structure.get("pti")
    bias = _bank_bias_scores()

    ranked: list[dict[str, Any]] = []
    high_risk_flags: list[str] = list(metrics.get("risk_flags") or [])

    for profile in _active_bank_profiles():
        confidence = 75.0
        reasons: list[str] = []
        if score is not None and score < profile["min_score"]:
            confidence -= (profile["min_score"] - score) * 0.25
            reasons.append(f"Score below {profile['name']} floor ({profile['min_score']})")
        if ltv is not None and ltv > profile["max_ltv"]:
            confidence -= (ltv - profile["max_ltv"]) * 0.9
            reasons.append(f"LTV above {profile['max_ltv']}%")
        if pti is not None and pti > profile["max_pti"]:
            confidence -= (pti - profile["max_pti"]) * 2.0
            reasons.append(f"PTI above {profile['max_pti']}%")
        if dti is not None and dti > profile["max_dti"]:
            confidence -= (dti - profile["max_dti"]) * 1.4
            reasons.append(f"DTI above {profile['max_dti']}%")
        if derogatories is not None and derogatories > profile["max_derogatories"]:
            confidence -= (derogatories - profile["max_derogatories"]) * 6.0
            reasons.append("Derogatory depth above comfort level")
        if utilization is not None and utilization > profile["max_utilization"]:
            confidence -= (utilization - profile["max_utilization"]) * 0.8
            reasons.append("Utilization is elevated")

        confidence += bias.get(profile["code"], 0.0)
        confidence = max(1.0, min(99.0, round(confidence, 1)))

        ranked.append(
            {
                "bank_code": profile["code"],
                "bank_name": profile["name"],
                "confidence": confidence,
                "reasons": reasons,
            }
        )

    ranked.sort(key=lambda item: item["confidence"], reverse=True)
    best = ranked[0] if ranked else None
    backup = ranked[1] if len(ranked) > 1 else None

    suggestions: list[str] = []
    if ltv is not None and ltv > 120:
        suggestions.append("Reduce amount financed or raise down payment to bring LTV closer to 110-120%.")
        high_risk_flags.append("High LTV")
    if pti is not None and pti > 17:
        suggestions.append("Stretch term or reduce payment target to bring PTI under 17%.")
        high_risk_flags.append("High PTI")
    if score is not None and score < 620:
        suggestions.append("Lead with subprime-friendly lenders and add stip-ready docs up front.")
        high_risk_flags.append("Sub-620 score")
    if derogatories is not None and derogatories > 3:
        suggestions.append("Prepare proof-of-income/residence before submission; stip load likely.")
        high_risk_flags.append("Heavy derogatory profile")

    return {
        "ranked_banks": ranked,
        "best_bank": best,
        "backup_bank": backup,
        "high_risk_flags": sorted(set(high_risk_flags)),
        "suggested_changes": suggestions,
    }


def _analyze_bank_brain(report_text: str, structured_data: dict[str, Any]) -> dict[str, Any]:
    metrics = _extract_credit_metrics(report_text=report_text, structured_data=structured_data)
    recommendation = _recommend_banks(metrics=metrics, structure={})
    route_one_fill = _extract_route_one_fill_from_text(report_text)
    if metrics.get("years_at_address") is None:
        route_one_years_address = route_one_fill.get("years_at_address")
        if isinstance(route_one_years_address, (int, float)):
            metrics["years_at_address"] = float(route_one_years_address)
    if metrics.get("years_at_job") is None:
        route_one_years_job = route_one_fill.get("years_at_job")
        if isinstance(route_one_years_job, (int, float)):
            metrics["years_at_job"] = float(route_one_years_job)

    result = {
        "ok": True,
        "metrics": metrics,
        "recommendation": recommendation,
        "route_one_fill": route_one_fill,
        "summary": {
            "best_bank": recommendation.get("best_bank"),
            "backup_bank": recommendation.get("backup_bank"),
            "risk_flags": recommendation.get("high_risk_flags", []),
        },
    }
    _append_audit_event(
        "bank_brain_analyze",
        {
            "metrics": metrics,
            "best_bank": recommendation.get("best_bank"),
        },
    )
    return result


def _extract_text_from_upload(raw: bytes, filename: str, content_type: str | None = None) -> str:
    return str(_extract_upload_document(raw, filename, content_type).get("text") or "")


def _spreadsheet_rows_to_text(rows: list[list[Any]], *, max_rows: int = 5000) -> str:
    output: list[str] = []
    for row in rows[:max_rows]:
        values = [_normalize_table_cell(cell) for cell in row]
        if any(values):
            output.append("\t".join(values).strip())
    return "\n".join(output).strip()


def _extract_docx_upload_text(raw: bytes) -> str:
    try:
        from docx import Document as DocxDocument

        document = DocxDocument(BytesIO(raw))
        rows: list[str] = []
        rows.extend(paragraph.text.strip() for paragraph in document.paragraphs if paragraph.text.strip())
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " ") for cell in row.cells]
                if any(cells):
                    rows.append("\t".join(cells))
        return "\n".join(rows).strip()
    except Exception:
        return ""


def _flatten_json_for_upload(value: Any, *, prefix: str = "", depth: int = 0) -> list[str]:
    if depth > 8:
        return []
    if isinstance(value, dict):
        rows: list[str] = []
        for key, child in value.items():
            child_prefix = f"{prefix}.{key}" if prefix else str(key)
            rows.extend(_flatten_json_for_upload(child, prefix=child_prefix, depth=depth + 1))
        return rows
    if isinstance(value, list):
        rows = []
        for index, child in enumerate(value[:500]):
            rows.extend(_flatten_json_for_upload(child, prefix=f"{prefix}[{index}]", depth=depth + 1))
        return rows
    clean = _normalize_table_cell(value)
    return [f"{prefix}: {clean}" if prefix else clean] if clean else []


def _extract_pdf_upload_text(raw: bytes) -> tuple[str, dict[str, Any]]:
    text_parts: list[str] = []
    diagnostics: dict[str, Any] = {
        "format": "pdf",
        "page_count": 0,
        "image_pages": 0,
        "ocr_attempted": 0,
    }
    try:
        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw))
        diagnostics["pypdf_page_count"] = len(reader.pages)
        pypdf_parts = [f"[page {index}]\n{page.extract_text() or ''}" for index, page in enumerate(reader.pages, start=1)]
        pypdf_text = "\n\n".join(part for part in pypdf_parts if part.strip()).strip()
        if pypdf_text and len(pypdf_text) >= 1800:
            diagnostics["format"] = "pdf_text"
            diagnostics["page_count"] = len(reader.pages)
            diagnostics["fast_text_path"] = "pypdf"
            return pypdf_text, diagnostics
    except Exception as exc:
        diagnostics["pypdf_error"] = str(exc)

    try:
        import fitz  # type: ignore

        with fitz.open(stream=raw, filetype="pdf") as doc:
            diagnostics["page_count"] = len(doc)
            for index, page in enumerate(doc, start=1):
                page_parts: list[str] = []
                page_text = (page.get_text("text", sort=True) or "").strip()
                if not page_text:
                    blocks = page.get_text("blocks", sort=True) or []
                    page_text = "\n".join(str(block[4]).strip() for block in blocks if len(block) >= 5 and str(block[4]).strip())
                if page_text:
                    page_parts.append(page_text)
                try:
                    tables = page.find_tables()
                    for table_index, table in enumerate(tables.tables, start=1):
                        extracted = table.extract()
                        table_rows = _spreadsheet_rows_to_text(extracted)
                        if table_rows:
                            page_parts.append(f"[table {table_index}]\n{table_rows}".strip())
                except Exception:
                    pass
                has_images = bool(page.get_images(full=True))
                if has_images:
                    diagnostics["image_pages"] = int(diagnostics.get("image_pages") or 0) + 1

                current_page_text = "".join(page_parts).strip()
                if has_images or len(current_page_text) < 220:
                    diagnostics["ocr_attempted"] = int(diagnostics.get("ocr_attempted") or 0) + 1
                    try:
                        from PIL import Image  # type: ignore

                        ocr_candidates: list[str] = []
                        ocr_meta: list[dict[str, Any]] = []
                        for scale in (2.0, 3.0):
                            pix = page.get_pixmap(matrix=fitz.Matrix(scale, scale), alpha=False)
                            source = Image.open(BytesIO(pix.tobytes("png")))
                            candidate, candidate_diag = _ocr_text_from_pil_image(source)
                            if candidate:
                                ocr_candidates.append(f"[ocr scale={scale}]\n{candidate}")
                                if len(candidate) >= 900:
                                    ocr_meta.append({"scale": scale, **candidate_diag, "short_circuit": True})
                                    break
                            if candidate_diag:
                                ocr_meta.append({"scale": scale, **candidate_diag})
                        if ocr_candidates:
                            diagnostics["ocr_pages"] = int(diagnostics.get("ocr_pages") or 0) + 1
                            page_parts.extend(ocr_candidates)
                        if ocr_meta:
                            diagnostics.setdefault("ocr_details", []).append({"page": index, "attempts": ocr_meta})
                    except Exception as exc:
                        diagnostics["ocr_error"] = str(exc)

                if page_parts:
                    text_parts.append(f"[page {index}]\n" + "\n".join(part for part in page_parts if part.strip()))
    except Exception as exc:
        diagnostics["fitz_error"] = str(exc)

    current_text = "\n\n".join(part for part in text_parts if part.strip()).strip()
    if (
        int(diagnostics.get("page_count") or 0) in {3, 4, 5}
        and (
            "nccclarity" in current_text.lower()
            or "experian credit report" in current_text.lower()
            or len(current_text) < 3000
        )
    ):
        region_parts, region_diag = _extract_ncc_clarity_pdf_regions(raw)
        if region_parts:
            text_parts.extend(region_parts)
        for key, value in region_diag.items():
            if value not in (None, "", [], {}):
                diagnostics[f"layout_{key}"] = value
    text = "\n\n".join(part for part in text_parts if part.strip()).strip()
    if len(text) < 40 and diagnostics.get("image_pages"):
        diagnostics.setdefault("warnings", []).append("PDF looks scanned or image-based; OCR was attempted but no usable text was found.")
    return text, diagnostics


def _extract_image_upload_text(raw: bytes) -> tuple[str, dict[str, Any]]:
    diagnostics: dict[str, Any] = {"format": "image", "ocr_pages": 0}
    try:
        from PIL import Image  # type: ignore

        image = Image.open(BytesIO(raw))
        text, ocr_diag = _ocr_text_from_pil_image(image)
        diagnostics.update({key: value for key, value in ocr_diag.items() if value not in (None, "", [], {})})
        diagnostics["ocr_pages"] = 1 if text else 0
        diagnostics["extracted_text_chars"] = len(text)
        if not text:
            diagnostics.setdefault("warnings", []).append("Image OCR did not find readable text.")
        return text, diagnostics
    except Exception as exc:
        diagnostics["ocr_error"] = str(exc)
        diagnostics["warnings"] = ["OCR engine is not available or could not read this image."]
        return "", diagnostics


def _extract_upload_document(raw: bytes, filename: str, content_type: str | None = None) -> dict[str, Any]:
    lowered_name = str(filename or "").lower()
    lowered_type = str(content_type or "").lower()
    suffix = Path(lowered_name).suffix
    if lowered_name.endswith(".pdf") or "pdf" in lowered_type or raw[:4] == b"%PDF":
        text, diagnostics = _extract_pdf_upload_text(raw)
        return {"text": text, **diagnostics, "extracted_text_chars": len(text)}

    if suffix in {".png", ".jpg", ".jpeg", ".webp", ".tif", ".tiff"} or lowered_type.startswith("image/"):
        text, diagnostics = _extract_image_upload_text(raw)
        return {"text": text, **diagnostics, "extracted_text_chars": len(text)}

    if lowered_name.endswith(".docx") or "wordprocessingml" in lowered_type:
        text = _extract_docx_upload_text(raw)
        return {"text": text, "format": "docx", "extracted_text_chars": len(text)}

    if suffix in {".xls", ".xlsx", ".xlsm", ".csv", ".tsv"} or "spreadsheet" in lowered_type or "excel" in lowered_type:
        try:
            rows = _rows_from_upload_table(raw, filename, content_type)
            text = _spreadsheet_rows_to_text(rows)
            return {"text": text, "format": suffix.lstrip(".") or "spreadsheet", "rows": len(rows), "extracted_text_chars": len(text)}
        except Exception as exc:
            return {"text": "", "format": suffix.lstrip(".") or "spreadsheet", "error": str(exc), "extracted_text_chars": 0}

    if suffix == ".json" or "json" in lowered_type:
        try:
            payload = json.loads(raw.decode("utf-8-sig", errors="ignore"))
            text = "\n".join(_flatten_json_for_upload(payload))
            return {"text": text, "format": "json", "extracted_text_chars": len(text)}
        except Exception:
            pass

    for encoding in ("utf-8-sig", "utf-8", "latin-1"):
        try:
            text = raw.decode(encoding, errors="ignore").strip()
            if suffix in {".html", ".htm"} or "html" in lowered_type:
                text = _html_to_visible_text(text)
                return {"text": text, "format": "html", "extracted_text_chars": len(text)}
            return {"text": text, "format": suffix.lstrip(".") or "text", "extracted_text_chars": len(text)}
        except Exception:
            continue
    text = raw.decode("utf-8", errors="ignore").strip()
    return {"text": text, "format": "text", "extracted_text_chars": len(text)}


def _simulate_credit_structure(request: CreditStructureRequest) -> dict[str, Any]:
    valuation = _jd_power_valuation_for_vin(request.vin)
    valuation_book = _to_float(valuation.get("jd_power_trade_in")) if valuation else None
    book_value = float(request.book_value or valuation_book or request.vehicle_price or 0.0)
    taxes = float(request.taxes)
    if taxes <= 0 and request.tax_rate > 0:
        taxes = round(float(request.vehicle_price) * float(request.tax_rate), 2)
    financed_amount = max(
        0.0,
        (
            float(request.vehicle_price)
            + taxes
            + float(request.fees)
            + float(request.backend_products)
            - float(request.down_payment)
        ),
    )
    ltv = (financed_amount / book_value * 100.0) if book_value > 0 else None
    payment = _estimate_monthly_payment(financed_amount, request.apr, request.term_months)
    pti = (
        (payment / request.monthly_income * 100.0)
        if request.monthly_income is not None and request.monthly_income > 0
        else None
    )
    dti = None
    if request.current_dti is not None and pti is not None:
        dti = request.current_dti + pti
    elif request.current_dti is not None:
        dti = request.current_dti
    elif pti is not None:
        dti = pti

    metrics = {
        "score": request.credit_score,
        "tradelines": request.tradelines,
        "derogatories": request.derogatories,
        "utilization": request.utilization,
        "risk_flags": [],
    }
    structure = {
        "vin": request.vin,
        "vehicle_price": request.vehicle_price,
        "book_value": round(book_value, 2) if book_value else None,
        "book_value_source": (
            "jd_power_trade_in"
            if valuation_book and (not request.book_value or abs(float(request.book_value) - valuation_book) < 1)
            else ("manual" if request.book_value else "sale_price")
        ),
        "jd_power_trade_in": round(valuation_book, 2) if valuation_book else None,
        "taxes": round(taxes, 2),
        "tax_rate": request.tax_rate,
        "fees": request.fees,
        "backend_products": request.backend_products,
        "down_payment": request.down_payment,
        "term_months": request.term_months,
        "apr": request.apr,
        "financed_amount": round(financed_amount, 2),
        "estimated_payment": round(payment, 2),
        "ltv": round(ltv, 2) if ltv is not None else None,
        "ltv_formula": "amount_financed / jd_power_trade_in" if book_value and book_value != request.vehicle_price else "amount_financed / sale_price",
        "pti": round(pti, 2) if pti is not None else None,
        "dti": round(dti, 2) if dti is not None else None,
    }
    recommendation = _recommend_banks(metrics=metrics, structure=structure)

    result = {
        "ok": True,
        "structure": structure,
        "recommendation": recommendation,
    }
    _append_audit_event(
        "bank_brain_structure",
        {
            "vin": request.vin,
            "structure": structure,
            "best_bank": recommendation.get("best_bank"),
        },
    )
    return result


def _recommend_vehicles_for_customer(request: BankBrainVehicleRecommendationRequest) -> dict[str, Any]:
    valuations = _load_jd_power_valuations()
    candidates: list[dict[str, Any]] = []
    for vehicle in _load_inventory_candidates():
        vin = str(vehicle.get("vin") or "").strip().upper()
        if not vin or vin == "UNKNOWN":
            continue
        inventory_price = _to_float(vehicle.get("price"))
        if inventory_price is None or inventory_price <= 0:
            continue
        sale_price = float(_bank_sale_price_from_inventory_price(inventory_price) or inventory_price)
        valuation = valuations.get(vin)
        book_value = _to_float(valuation.get("jd_power_trade_in")) if valuation else None
        pricing = _jd_power_ltv_from_pricing(inventory_price=inventory_price, jd_trade_value=book_value)
        structure = _simulate_credit_structure(
            CreditStructureRequest(
                vin=vin,
                vehicle_price=sale_price,
                book_value=book_value,
                taxes=pricing["taxes"] or round(sale_price * DEFAULT_BANK_TAX_RATE, 2),
                tax_rate=DEFAULT_BANK_TAX_RATE,
                fees=0,
                backend_products=0,
                down_payment=0,
                term_months=72,
                apr=9.99,
                monthly_income=request.monthly_income,
                current_dti=request.current_dti,
                credit_score=request.score,
            )
        )
        struct = structure.get("structure") or {}
        best = (structure.get("recommendation") or {}).get("best_bank") or {}
        payment = _to_float(struct.get("estimated_payment"))
        ltv = _to_float(pricing.get("ltv"))
        confidence = _to_float(best.get("confidence")) or 0.0
        if request.desired_payment and payment and payment > request.desired_payment * 1.2:
            confidence -= 12
        if ltv and ltv > 140:
            confidence -= 15
        candidates.append(
            {
                "vin": vin,
                "title": vehicle.get("title"),
                "price": vehicle.get("price"),
                "sale_price": pricing["bank_sale_price"] or round(sale_price, 2),
                "taxes": pricing["taxes"],
                "ltv_basis": pricing["ltv_basis"],
                "jd_power_trade_in": book_value,
                "ltv": pricing["ltv"],
                "ltv_formula": "(website_price + 2400 + 6_percent_tax) / jd_power_trade_in",
                "estimated_payment": struct.get("estimated_payment"),
                "best_bank": best,
                "backup_bank": (structure.get("recommendation") or {}).get("backup_bank"),
                "confidence": round(max(1.0, min(99.0, confidence)), 1),
                "reason": (
                    "Strong fit: book value supports the advance and bank confidence is high."
                    if book_value and ltv and ltv <= 120
                    else "Review advance: JD Power value is missing or LTV is elevated."
                ),
            }
        )
    candidates.sort(key=lambda item: (item.get("confidence") or 0, -float(item.get("ltv") or 999)), reverse=True)
    return {"ok": True, "items": candidates[: request.max_results], "count": len(candidates)}


def _eligible_for_posting(vehicle: dict[str, Any]) -> bool:
    status = str(vehicle.get("status_label") or vehicle.get("status") or "").strip().lower()
    photos = vehicle.get("photos") or vehicle.get("images") or []
    return status == "ready" and bool(photos)


def _require_permission(request: Request, permission: str) -> dict[str, Any]:
    user = current_user_from_auth_header(request.headers.get("authorization", ""))
    if not user:
        raise HTTPException(status_code=401, detail={"message": "Authentication required"})
    permissions = set(user.get("permissions") or [])
    if "admin.full" in permissions or permission in permissions:
        return user
    raise HTTPException(
        status_code=403,
        detail={
            "message": f"Missing permission: {permission}",
            "user": user.get("username"),
        },
    )


def _user_has_permission(user: dict[str, Any] | None, permission: str) -> bool:
    permissions = set((user or {}).get("permissions") or [])
    return "admin.full" in permissions or permission in permissions


def _mask_inventory_assets_for_user(items: list[dict[str, Any]], user: dict[str, Any]) -> list[dict[str, Any]]:
    if _user_has_permission(user, "assets.view"):
        return items
    masked: list[dict[str, Any]] = []
    for item in items:
        clean = dict(item)
        clean.pop("photos", None)
        clean.pop("images", None)
        clean.pop("photo_urls", None)
        clean.pop("sticker_url", None)
        clean.pop("carfax_url", None)
        masked.append(clean)
    return masked


@router.get("/me")
def me(request: Request) -> dict[str, Any]:
    user = current_user_from_auth_header(request.headers.get("authorization", ""))
    return {
        "ok": bool(user),
        "user": user,
        "permissions": _public_permission_catalog(),
    }


@router.get("/admin/users")
def admin_users(request: Request) -> dict[str, Any]:
    _require_permission(request, "users.manage")
    return {
        "ok": True,
        "items": list_public_users(),
        "permissions": _public_permission_catalog(),
    }


@router.post("/admin/users")
def admin_users_create(request: Request, payload: XconsoleUserRequest) -> dict[str, Any]:
    _require_permission(request, "users.manage")
    try:
        user = upsert_user(
            username=payload.username,
            password=payload.password,
            display_name=payload.display_name,
            role=payload.role,
            permissions=payload.permissions,
            active=payload.active,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail={"message": str(exc)}) from exc
    return {"ok": True, "user": user, "items": list_public_users()}


@router.put("/admin/users/{username}")
def admin_users_update(username: str, request: Request, payload: XconsoleUserRequest) -> dict[str, Any]:
    _require_permission(request, "users.manage")
    try:
        user = upsert_user(
            username=username,
            password=payload.password,
            display_name=payload.display_name,
            role=payload.role,
            permissions=payload.permissions,
            active=payload.active,
        )
    except ValueError as exc:
        raise HTTPException(status_code=400, detail={"message": str(exc)}) from exc
    return {"ok": True, "user": user, "items": list_public_users()}


@router.delete("/admin/users/{username}")
def admin_users_delete(username: str, request: Request) -> dict[str, Any]:
    _require_permission(request, "users.manage")
    try:
        user = deactivate_user(username)
    except ValueError as exc:
        raise HTTPException(status_code=404, detail={"message": str(exc)}) from exc
    return {"ok": True, "user": user, "items": list_public_users()}


@router.get("/admin/carfax/coverage")
def admin_carfax_coverage(request: Request, sample_size: int = 20) -> dict[str, Any]:
    _require_permission(request, "admin.full")
    return _carfax_coverage_summary(sample_size=max(1, min(sample_size, 100)))


@router.post("/admin/carfax/backfill")
def admin_carfax_backfill(
    request: Request,
    offset: int = 0,
    limit: int = 15,
    only_missing: bool = True,
) -> dict[str, Any]:
    _require_permission(request, "admin.full")
    return _backfill_carfax_inventory(offset=offset, limit=limit, only_missing=only_missing)


@router.get("/leads/inbox")
async def leads_inbox(
    request: Request,
    sync: bool = False,
    source: str = "all",
    force_sync: bool = False,
) -> dict[str, Any]:
    _require_permission(request, "facebook.leads")
    leads = _load_leads()
    connection = _facebook_lead_connection_status()
    sync_source = (source or "all").strip().lower()
    can_sync = sync and (
        connection.get("connected")
        or bool(connection.get("user_token_configured"))
    )
    sync_result = (
        _sync_facebook_leads_if_stale(source=sync_source, force=force_sync) if can_sync else None
    )
    if sync_result is not None and sync_result.get("ok") and sync_result.get("mode") not in {"not_connected", "graph_rejected", "conversation_pull_error", "conversation_permission_missing", "connection_error"}:
        leads = _load_leads()
    return {
        "ok": True,
        "items": leads,
        "count": len(leads),
        "new_count": len([item for item in leads if item.get("status") == "new"]),
        "responded_count": len([item for item in leads if item.get("status") == "responded"]),
        "facebook_connection": connection,
        "sync": sync_result,
    }


@router.post("/leads/manual-add")
def leads_manual_add(request: Request, payload: LeadManualAddRequest) -> dict[str, Any]:
    user = _require_permission(request, "facebook.leads")
    leads = _load_leads()
    entry = {
        "id": _lead_id(f"{payload.customer_name}|{payload.message}|{time.time()}"),
        "customer_name": payload.customer_name.strip() or "Unknown Lead",
        "channel": payload.channel.strip().lower() or "facebook",
        "message": payload.message.strip(),
        "vehicle_vin": str(payload.vehicle_vin or "").strip().upper(),
        "source": payload.source.strip() or "manual",
        "status": "new",
        "created_at": _utc_now(),
        "last_message_at": _utc_now(),
        "created_by": user.get("username"),
    }
    leads.insert(0, entry)
    _save_leads(leads)
    return {"ok": True, "lead": entry, "items": _load_leads()}


def _process_lead_response(
    *,
    user: dict[str, Any],
    payload: LeadRespondRequest,
    upload_name: str | None = None,
    upload_content: bytes | None = None,
    upload_content_type: str | None = None,
) -> Any:
    leads = _load_leads()
    target = next((item for item in leads if str(item.get("id")) == payload.lead_id), None)
    if not target:
        raise HTTPException(status_code=404, detail={"message": "Lead not found"})
    response_text = str(payload.response_text or "").strip()
    attachment_url = str(payload.attachment_url or "").strip()
    attachment_type = str(payload.attachment_type or "").strip().lower() or "image"
    has_upload = bool(upload_content)
    has_attachment = bool(attachment_url) or has_upload
    if not response_text and not has_attachment:
        raise HTTPException(status_code=400, detail={"message": "Response text or an attachment is required."})
    delivery_note = "Response is logged in Xconsole."
    send_result: dict[str, Any] | None = None
    delivery_status = "logged"
    provider_message_id = None
    error_detail = None
    response_attachments: list[dict[str, Any]] = []

    if str(target.get("channel") or "").strip().lower() in {"facebook_marketplace", "facebook_messenger"}:
        recipient_id = str(target.get("profile_id") or "").strip()
        send_steps: list[dict[str, Any]] = []
        if response_text:
            send_steps.append(
                _facebook_send_page_message(
                    recipient_id=recipient_id,
                    message_text=response_text,
                )
            )
        if attachment_url:
            send_steps.append(
                _facebook_send_page_attachment_url(
                    recipient_id=recipient_id,
                    attachment_url=attachment_url,
                    attachment_type=attachment_type,
                )
            )
            response_attachments.append(
                {
                    "type": attachment_type,
                    "url": attachment_url,
                    "title": None,
                }
            )
        if has_upload:
            send_steps.append(
                _facebook_send_page_attachment_upload(
                    recipient_id=recipient_id,
                    filename=upload_name or "attachment",
                    content=upload_content or b"",
                    content_type=upload_content_type,
                )
            )
            response_attachments.append(
                {
                    "type": _guess_messenger_attachment_type(upload_content_type, upload_name),
                    "url": None,
                    "title": upload_name or "attachment",
                }
            )
        send_result = {
            "ok": bool(send_steps) and all(bool(step.get("ok")) for step in send_steps),
            "mode": "sent" if send_steps and all(bool(step.get("ok")) for step in send_steps) else str((next((step for step in send_steps if not step.get("ok")), {}) or {}).get("mode") or ""),
            "steps": send_steps,
        }
        first_success = next((step for step in send_steps if step.get("ok") and step.get("message_id")), None)
        if send_result.get("ok"):
            delivery_status = "sent"
            provider_message_id = str((first_success or {}).get("message_id") or "").strip() or None
            delivery_note = "Reply sent to Messenger."
        else:
            failed_step = next((step for step in send_steps if not step.get("ok")), {}) or {}
            send_mode = str(failed_step.get("mode") or "").strip().lower()
            send_result["message"] = str(failed_step.get("message") or "Messenger send failed.")
            send_mode = str(send_result.get("mode") or "").strip().lower()
            delivery_status = "blocked" if send_mode == "outside_reply_window" else "failed"
            error_detail = str(send_result.get("message") or "Messenger send failed.").strip()
            if send_mode == "outside_reply_window":
                error_detail = (
                    "Messenger blocked this reply because the conversation is outside Facebook's reply window. "
                    "Have the customer send a fresh message, then Xconsole can reply live from here."
                )
            response = _append_lead_response(
                lead_id=payload.lead_id,
                channel=payload.channel,
                response_text=response_text,
                author=str(user.get("username") or "xconsole"),
                delivery_status=delivery_status,
                error_detail=error_detail,
                attachments=response_attachments,
            )
            for item in leads:
                if str(item.get("id")) == payload.lead_id:
                    item["status"] = payload.mark_status or "responded"
                    item["last_response_at"] = response["created_at"]
            _save_leads(leads)
            if send_mode == "outside_reply_window":
                return {
                    "ok": False,
                    "blocked": True,
                    "delivery_note": error_detail,
                    "send_result": send_result,
                    "response": response,
                    "items": _load_leads(),
                }
            return JSONResponse(
                status_code=502,
                content={
                    "ok": False,
                    "message": error_detail,
                    "send_result": send_result,
                    "response": response,
                    "items": _load_leads(),
                },
            )

    response = _append_lead_response(
        lead_id=payload.lead_id,
        channel=payload.channel,
        response_text=response_text,
        author=str(user.get("username") or "xconsole"),
        delivery_status=delivery_status,
        provider_message_id=provider_message_id,
        error_detail=error_detail,
        attachments=response_attachments,
    )
    for item in leads:
        if str(item.get("id")) == payload.lead_id:
            item["status"] = payload.mark_status or "responded"
            item["last_response_at"] = response["created_at"]
    _save_leads(leads)
    return {
        "ok": True,
        "lead_id": payload.lead_id,
        "response": response,
        "delivery_note": delivery_note,
        "send_result": send_result,
        "items": _load_leads(),
    }


@router.post("/leads/respond")
def leads_respond(request: Request, payload: LeadRespondRequest) -> dict[str, Any]:
    user = _require_permission(request, "facebook.leads")
    return _process_lead_response(user=user, payload=payload)


@router.post("/leads/respond-upload")
async def leads_respond_upload(
    request: Request,
    lead_id: str = Form(...),
    response_text: str = Form(default=""),
    channel: str = Form(default="facebook"),
    mark_status: str = Form(default="responded"),
    attachment_url: str = Form(default=""),
    attachment_type: str = Form(default="image"),
    attachment_file: UploadFile | None = File(default=None),
) -> Any:
    user = _require_permission(request, "facebook.leads")
    upload_name = None
    upload_content = None
    upload_content_type = None
    if attachment_file is not None:
        upload_name = attachment_file.filename or "attachment"
        upload_content_type = attachment_file.content_type or "application/octet-stream"
        upload_content = await attachment_file.read()
    payload = LeadRespondRequest(
        lead_id=lead_id,
        response_text=response_text,
        channel=channel,
        mark_status=mark_status,
        attachment_url=attachment_url or None,
        attachment_type=attachment_type or None,
    )
    return _process_lead_response(
        user=user,
        payload=payload,
        upload_name=upload_name,
        upload_content=upload_content,
        upload_content_type=upload_content_type,
    )


@router.post("/leads/sync-facebook")
def leads_sync_facebook(request: Request, source: str = "all") -> dict[str, Any]:
    _require_permission(request, "facebook.leads")
    result = _sync_facebook_leads(source=source)
    result["items"] = _load_leads()
    return result


@router.get("/admin/leads/fb-token-diagnostics")
def leads_facebook_token_diagnostics(request: Request, source: str = "all") -> dict[str, Any]:
    _require_permission(request, "facebook.leads")
    requested_source = (source or "all").strip().lower()
    if requested_source not in {"page", "personal", "all"}:
        requested_source = "all"
    diagnostics = _diagnose_facebook_lead_sync_candidates(requested_source=requested_source)
    diagnostics["ok"] = True
    diagnostics["requested_source"] = requested_source
    diagnostics["connection"] = _facebook_lead_connection_status()
    return diagnostics


@router.get("/offerup/status")
def offerup_status(request: Request) -> dict[str, Any]:
    _require_permission(request, "offerup.post")
    return _load_offerup_status()


@router.post("/offerup/post/from-inventory")
def offerup_post_from_inventory(request: Request, payload: OfferUpPostRequest) -> dict[str, Any]:
    _require_permission(request, "offerup.post")
    return _offerup_post_from_inventory(payload)


@router.get("/vehicles/{vin}/decode")
def vehicles_decode(vin: str, request: Request) -> dict[str, Any]:
    _require_permission(request, "inventory.view")
    return _decode_vin_values(vin)


@router.get("/bank-brain/vehicle/{vin}")
def bank_brain_vehicle(vin: str, request: Request) -> dict[str, Any]:
    _require_permission(request, "bankbrain.view")
    return _vehicle_bank_brain(vin)


@router.get("/health")
async def health() -> dict[str, Any]:
    return {"ok": True, "utc": datetime.now(timezone.utc).isoformat()}


@router.get("/status")
async def status() -> dict[str, Any]:
    vehicles = _load_inventory_candidates()
    posts = _load_runtime_posts()
    accounts = _load_accounts()
    stack_readiness = _stack_readiness_status()
    return {
        "ok": True,
        "vehicles_count": len(vehicles),
        "posts_count": len(posts),
        "accounts_count": len(accounts),
        "deployment": _deployment_fingerprint(),
        "facebook_lister_path": str(FML_DIR),
        "live_requirements": _live_requirements_status(),
        "stack_readiness": stack_readiness,
        "sales_assistant": _sales_assistant_health_status(),
        "inventory_source": _inventory_source_status(),
    }


@router.get("/sales-assistant/health")
def sales_assistant_health() -> dict[str, Any]:
    return _sales_assistant_health_status()


@router.get("/sales-assistant/banks")
def sales_assistant_banks() -> dict[str, Any]:
    return _sales_backend_proxy("GET", "/api/banks")


@router.post("/sales-assistant/banks/reload")
def sales_assistant_banks_reload() -> dict[str, Any]:
    return _sales_backend_proxy("POST", "/api/banks/reload")


@router.get("/sales-assistant/banks/factors")
def sales_assistant_banks_factors() -> dict[str, Any]:
    return _sales_backend_proxy("GET", "/api/banks/factors")


@router.post("/sales-assistant/banks/factors/reload")
def sales_assistant_banks_factors_reload() -> dict[str, Any]:
    return _sales_backend_proxy("POST", "/api/banks/factors/reload")


@router.get("/stack/readiness")
def stack_readiness() -> dict[str, Any]:
    return _stack_readiness_status()


@router.get("/inventory/source-status")
async def inventory_source_status() -> dict[str, Any]:
    return _inventory_source_status()


@router.get("/dealerships")
def dealerships(request: Request) -> dict[str, Any]:
    _require_permission(request, "inventory.view")
    items = _load_dealerships()
    return {
        "ok": True,
        "items": items,
        "count": len(items),
        "active_source_urls": _configured_dealership_source_urls(),
    }


@router.post("/dealerships")
def dealerships_upsert(request: Request, payload: DealershipRequest) -> dict[str, Any]:
    _require_permission(request, "dealerships.manage")
    dealership = _save_dealership(payload)
    return {
        "ok": True,
        "dealership": dealership,
        "items": _load_dealerships(),
        "active_source_urls": _configured_dealership_source_urls(),
    }


@router.post("/inventory/sync-live")
def inventory_sync_live(http_request: Request, request: InventoryLiveSyncRequest) -> dict[str, Any]:
    _require_permission(http_request, "inventory.edit")
    return _sync_live_inventory(
        source_url=request.source_url,
        timeout_seconds=request.timeout_seconds,
        persist=request.persist,
    )


@router.get("/inventory/active")
async def inventory_active(request: Request) -> dict[str, Any]:
    user = _require_permission(request, "inventory.view")
    items = _enrich_inventory_items(_load_inventory_candidates())
    items = _mask_inventory_assets_for_user(items, user)
    counts = _inventory_count_summary(items)
    return {
        "items": items,
        "count": len(items),
        "active_count": counts["active"],
        "in_transit_count": counts["in_transit"],
        "source_status": _inventory_source_status(),
    }


@router.get("/facebook/accounts")
def facebook_accounts() -> dict[str, Any]:
    return {"items": _load_accounts()}


@router.get("/facebook/images")
def facebook_images(limit: int = 200) -> dict[str, Any]:
    normalized_limit = max(1, min(int(limit), 2000))
    items, total = _list_facebook_images(limit=normalized_limit)
    return {
        "items": items,
        "total": total,
        "returned": len(items),
        "limit": normalized_limit,
    }


@router.post("/facebook/bootstrap")
def facebook_bootstrap(request: FacebookBootstrapRequest) -> dict[str, Any]:
    return _bootstrap_facebook_lister(
        create_template_account_if_missing=request.create_template_account_if_missing
    )


@router.post("/facebook/images/import-from-vehicle")
def facebook_images_import_from_vehicle(request: FacebookVehicleImageImportRequest) -> dict[str, Any]:
    return _import_vehicle_images(
        vin=request.vin,
        limit=request.limit,
        overwrite=request.overwrite,
    )


@router.post("/facebook/images/relink")
def facebook_images_relink(request: FacebookRelinkImagesRequest) -> dict[str, Any]:
    return _relink_images_to_vin(
        vin=request.vin,
        images=request.images,
        include_vin_matches=request.include_vin_matches,
        overwrite=request.overwrite,
        delete_source=request.delete_source,
    )


@router.get("/facebook/images/suggest")
def facebook_images_suggest(vin: str, limit: int = 20) -> dict[str, Any]:
    normalized_limit = max(1, min(int(limit), 200))
    items = _suggest_images_for_vin(vin=vin, limit=normalized_limit)
    _, available_total = _list_facebook_images(limit=1)
    return {
        "vin": vin.strip().upper(),
        "items": items,
        "total_matches": len(items),
        "available_total": available_total,
    }


@router.get("/facebook/live-requirements")
def facebook_live_requirements() -> dict[str, Any]:
    return _live_requirements_status()


@router.post("/facebook/live-preflight")
def facebook_live_preflight(request: FacebookPreflightRequest) -> dict[str, Any]:
    return _run_live_preflight(
        account_id=request.account_id,
        images=request.images,
        vin=request.vin,
    )


@router.get("/facebook/live-status")
async def facebook_live_status() -> JSONResponse:
    status_file = RUNTIME_DIR / "facebook_live_status.json"
    headers = {"Cache-Control": "no-store, no-cache, must-revalidate, max-age=0"}
    if not status_file.exists():
        return JSONResponse(content={
            "ok": True,
            "stage": "No live Facebook publish running.",
            "updated_at": None,
        }, headers=headers)
    payload = _safe_read_json(status_file, {})
    if not isinstance(payload, dict):
        return JSONResponse(content={
            "ok": False,
            "stage": "Facebook live status file could not be read.",
            "updated_at": None,
        }, headers=headers)
    if payload.get("stage"):
        payload["stage"] = _friendly_facebook_publish_detail(fallback=str(payload.get("stage"))) if payload.get("type") == "failure" else re.sub(r"\s+", " ", _strip_facebook_automation_noise(str(payload.get("stage")))).strip()[:420]
    return JSONResponse(content={"ok": True, **payload}, headers=headers)


@router.post("/facebook/sync-marketplace")
async def facebook_sync_marketplace(request: Request, payload: FacebookMarketplaceSyncRequest | None = None) -> dict[str, Any]:
    _require_permission(request, "facebook.post")
    options = payload or FacebookMarketplaceSyncRequest()
    return _sync_marketplace_post_statuses(
        verify_live_urls=options.verify_live_urls,
        processing_review_minutes=options.processing_review_minutes,
    )


@router.post("/facebook/prepare-live-post")
def facebook_prepare_live_post(request: FacebookPrepareLivePostRequest) -> dict[str, Any]:
    return _prepare_live_post(
        vin=request.vin,
        account_id=request.account_id,
        import_missing_images=request.import_missing_images,
        image_limit=request.image_limit,
        overwrite_images=request.overwrite_images,
    )


@router.post("/facebook/full-repair")
def facebook_full_repair(request: FacebookFullRepairRequest) -> dict[str, Any]:
    return _full_repair_and_relink(
        vin=request.vin,
        ensure_placeholder_images=request.ensure_placeholder_images,
        placeholder_count=request.placeholder_count,
    )


@router.post("/wire-everything")
def wire_everything(request: WireEverythingRequest) -> dict[str, Any]:
    return _wire_everything(
        vin=request.vin,
        ensure_placeholder_images=request.ensure_placeholder_images,
        placeholder_count=request.placeholder_count,
        reload_sales_data=request.reload_sales_data,
    )


@router.get("/facebook/posts")
def facebook_posts() -> dict[str, Any]:
    return {"items": _load_runtime_posts()}


@router.get("/vehicles")
def vehicles(request: Request) -> dict[str, Any]:
    user = _require_permission(request, "inventory.view")
    items = _enrich_inventory_items(_load_inventory_candidates())
    items = _mask_inventory_assets_for_user(items, user)
    return {
        "items": items,
        "count": len(items),
        "source_status": _inventory_source_status(),
    }


@router.post("/vehicles/manual-add")
def vehicles_manual_add(http_request: Request, request: ManualVehicleAddRequest) -> dict[str, Any]:
    _require_permission(http_request, "inventory.edit")
    payload = _safe_read_json(INVENTORY_MANUAL_PATH, {"items": []})
    items = payload.get("items", []) if isinstance(payload, dict) else []
    if not isinstance(items, list):
        items = []

    clean_vin = str(request.vin).strip().upper()
    items = [
        entry
        for entry in items
        if not (
            isinstance(entry, dict)
            and str(entry.get("vin", "")).strip().upper() == clean_vin
        )
    ]
    items.append(
        {
            "vin": clean_vin,
            "title": request.title.strip(),
            "price": request.price,
            "mileage": request.mileage,
            "drivetrain": request.drivetrain,
            "engine": request.engine,
            "transmission": request.transmission,
            "location": request.location,
            "detail_url": request.detail_url,
            "exterior": request.exterior,
            "interior": request.interior,
            "photos": request.photos,
            "status_label": "Ready",
            "manual_added_at": datetime.now(timezone.utc).isoformat(),
        }
    )
    _safe_write_json(
        INVENTORY_MANUAL_PATH,
        {
            "items": items,
            "updated_at": datetime.now(timezone.utc).isoformat(),
        },
    )
    return {
        "ok": True,
        "vin": clean_vin,
        "items_count": len(items),
        "items": _enrich_inventory_items(_load_inventory_candidates()),
    }


@router.get("/vehicles/{vin}/assets")
def vehicles_assets(vin: str, request: Request, refresh: bool = False) -> dict[str, Any]:
    _require_permission(request, "assets.view")
    return _load_vehicle_assets(vin=vin, refresh=refresh)


@router.post("/vehicles/{vin}/carfax/report-text")
def vehicles_carfax_report_text(vin: str, request: Request, payload: CarfaxReportTextRequest) -> dict[str, Any]:
    _require_permission(request, "carfax.view")
    return _save_carfax_report_text(vin=vin, report_text=payload.report_text, source_url=payload.source_url)


@router.get("/vehicles/{vin}/asset-view/{kind}")
def vehicles_asset_view(vin: str, kind: str, request: Request, refresh: bool = False, visible: bool = False) -> Response:
    clean_kind = str(kind or "").strip().lower()
    if clean_kind not in {"sticker", "carfax"}:
        raise HTTPException(status_code=404, detail={"message": "Unknown vehicle asset view"})

    _require_permission(request, "stickers.view" if clean_kind == "sticker" else "carfax.view")
    clean_vin = str(vin or "").strip().upper()
    vehicle = _find_vehicle_by_vin(clean_vin)
    if not vehicle:
        raise HTTPException(status_code=404, detail={"message": f"Vehicle not found for VIN {clean_vin}"})

    assets = _load_vehicle_assets(vin=clean_vin, refresh=refresh)
    if clean_kind == "carfax":
        if visible:
            assets = _refresh_carfax_report_for_vin(clean_vin, visible_browser=True)
        summary = assets.get("carfax_summary") if isinstance(assets, dict) else {}
        if (
            not assets.get("carfax_facts")
            or (
                not _has_structured_carfax_facts(
                    assets.get("carfax_facts") if isinstance(assets, dict) else None
                )
                and _is_generic_carfax_summary(summary if isinstance(summary, dict) else None)
            )
        ):
            assets = _load_vehicle_assets(vin=clean_vin, refresh=True)
        return HTMLResponse(_carfax_summary_html(vehicle, assets))

    source_url = str(assets.get("sticker_url") or "").strip()
    if not source_url:
        return HTMLResponse(
            _asset_message_html(
                "Sticker not cached",
                "Refresh vehicle assets first. Xconsole has not found a window sticker link for this VIN yet.",
            ),
            status_code=404,
        )

    meta = _cache_remote_vehicle_asset(clean_vin, "sticker", source_url, force=refresh)
    if isinstance(meta, dict) and meta.get("ok") and meta.get("path") and Path(str(meta.get("path"))).exists():
        return FileResponse(
            Path(str(meta["path"])),
            media_type=str(meta.get("content_type") or "application/octet-stream"),
        )

    return HTMLResponse(
        _asset_message_html(
            "Sticker cache failed",
            str(meta.get("error") or "The original sticker service did not return a cacheable file."),
            source_url=source_url,
        ),
        status_code=502,
    )


def _facebook_post_impl(request: FacebookPostRequest) -> dict[str, Any]:
    text = _render_listing_text(request)
    listing_file = _write_listing_text(request.vin.strip().upper(), text)

    response: dict[str, Any] = {
        "ok": True,
        "mode": request.mode,
        "listing_file": str(listing_file.relative_to(ROOT_DIR)).replace("\\", "/"),
        "text": text,
    }

    if request.mode == "live":
        success, detail, marketplace_state = _publish_live(request)
        marketplace_status = str(marketplace_state.get("marketplace_status") or ("live" if success else "needs_review"))
        response["live_success"] = success
        response["live_detail"] = detail
        response["marketplace_status"] = marketplace_status
        response["listing_url"] = marketplace_state.get("listing_url")
        response["marketplace_confirmation"] = marketplace_state.get("confirmation")
        if not success and marketplace_status not in {"processing"}:
            response["ok"] = False
            _set_facebook_vehicle_status(
                vin=request.vin,
                mode="live",
                marketplace_status=marketplace_status,
                detail=str(detail),
                listing_url=str(marketplace_state.get("listing_url") or ""),
                confirmation=marketplace_state.get("confirmation") if isinstance(marketplace_state.get("confirmation"), dict) else None,
            )
            return response
        if marketplace_status == "processing":
            _set_facebook_vehicle_status(
                vin=request.vin,
                mode="live",
                marketplace_status="processing",
                detail=str(detail),
                listing_url=str(marketplace_state.get("listing_url") or ""),
                confirmation=marketplace_state.get("confirmation") if isinstance(marketplace_state.get("confirmation"), dict) else None,
            )
            return response
        _set_facebook_vehicle_status(
            vin=request.vin,
            mode="live",
            marketplace_status="live",
            detail=str(detail),
            listing_url=str(marketplace_state.get("listing_url") or ""),
            confirmation=marketplace_state.get("confirmation") if isinstance(marketplace_state.get("confirmation"), dict) else None,
        )

    return response


def _carfax_coverage_summary(*, sample_size: int = 20) -> dict[str, Any]:
    items = [item for item in _load_inventory_candidates() if isinstance(item, dict)]
    total = len(items)
    with_link = 0
    with_report = 0
    with_structured = 0
    stale: list[str] = []
    for item in items:
        vin = str(item.get("vin") or "").strip().upper()
        if not vin:
            continue
        cached = _safe_read_json(_vehicle_assets_cache_path(vin), {})
        if not isinstance(cached, dict):
            cached = {}
        if cached.get("carfax_url"):
            with_link += 1
        report = cached.get("carfax_report") if isinstance(cached.get("carfax_report"), dict) else {}
        if isinstance(report, dict) and report.get("ok"):
            with_report += 1
        facts = cached.get("carfax_facts") if isinstance(cached.get("carfax_facts"), dict) else {}
        if _has_structured_carfax_facts(facts):
            with_structured += 1
        elif len(stale) < sample_size:
            stale.append(vin)
    return {
        "ok": True,
        "total": total,
        "with_link": with_link,
        "with_report": with_report,
        "with_structured": with_structured,
        "missing_or_generic_sample": stale,
        "updated_at": datetime.now(timezone.utc).isoformat(),
    }


def _backfill_carfax_inventory(
    *,
    offset: int = 0,
    limit: int = 15,
    only_missing: bool = True,
) -> dict[str, Any]:
    items = [item for item in _load_inventory_candidates() if isinstance(item, dict)]
    items = [item for item in items if str(item.get("vin") or "").strip()]
    items.sort(key=lambda item: str(item.get("vin") or ""))
    start = max(0, int(offset or 0))
    end = max(start, min(len(items), start + max(1, min(int(limit or 15), 100))))
    selected = items[start:end]

    results: list[dict[str, Any]] = []
    refreshed = 0
    skipped = 0
    errors = 0
    for item in selected:
        vin = str(item.get("vin") or "").strip().upper()
        try:
            cached = _safe_read_json(_vehicle_assets_cache_path(vin), {})
            cached_facts = cached.get("carfax_facts") if isinstance(cached, dict) and isinstance(cached.get("carfax_facts"), dict) else {}
            cached_report = cached.get("carfax_report") if isinstance(cached, dict) and isinstance(cached.get("carfax_report"), dict) else {}
            if only_missing and (_has_structured_carfax_facts(cached_facts) or cached_report.get("ok")):
                skipped += 1
                results.append({"vin": vin, "status": "skipped"})
                continue

            assets = _load_vehicle_assets(vin=vin, refresh=True)
            facts = assets.get("carfax_facts") if isinstance(assets, dict) and isinstance(assets.get("carfax_facts"), dict) else {}
            report = assets.get("carfax_report") if isinstance(assets, dict) and isinstance(assets.get("carfax_report"), dict) else {}
            refreshed += 1
            results.append(
                {
                    "vin": vin,
                    "status": "refreshed",
                    "carfax_url": bool(assets.get("carfax_url")),
                    "report_ok": bool(report.get("ok")),
                    "structured": _has_structured_carfax_facts(facts),
                    "summary": str((assets.get("carfax_summary") or {}).get("summary") or "")[:220],
                }
            )
        except Exception as exc:
            errors += 1
            results.append({"vin": vin, "status": "error", "error": str(exc)})

    return {
        "ok": True,
        "offset": start,
        "limit": end - start,
        "processed": len(selected),
        "refreshed": refreshed,
        "skipped": skipped,
        "errors": errors,
        "results": results,
        "coverage": _carfax_coverage_summary(sample_size=12),
    }


@router.post("/facebook/post")
def facebook_post(http_request: Request, request: FacebookPostRequest) -> dict[str, Any]:
    _require_permission(http_request, "facebook.post")
    return _facebook_post_impl(request)


@router.post("/facebook/post/from-inventory")
def facebook_post_from_inventory(http_request: Request, request: FacebookOneClickPostRequest) -> dict[str, Any]:
    _require_permission(http_request, "facebook.post")
    return _one_click_post_from_inventory(request)


@router.post("/facebook/post/batch-from-inventory")
def facebook_post_batch_from_inventory(http_request: Request, request: FacebookBatchPostRequest) -> dict[str, Any]:
    _require_permission(http_request, "facebook.post")
    return _batch_post_from_inventory(request)


@router.post("/bank-brain/analyze")
def bank_brain_analyze(request: BankBrainAnalyzeRequest) -> dict[str, Any]:
    return _analyze_bank_brain(
        report_text=request.report_text or "",
        structured_data=request.structured_data or {},
    )


@router.post("/bank-brain/analyze-upload")
async def bank_brain_analyze_upload(file: UploadFile = File(...)) -> dict[str, Any]:
    raw = await file.read()
    extracted = _extract_upload_document(raw, file.filename or "", file.content_type)
    report_text = str(extracted.get("text") or "")
    result = _analyze_bank_brain(report_text=report_text, structured_data={})
    result["source_file"] = file.filename
    result["file_understanding"] = {
        key: value
        for key, value in extracted.items()
        if key not in {"text"} and value not in (None, "", [], {})
    }
    result["extracted_text_chars"] = len(report_text)
    result["extracted_preview"] = report_text[:1200]
    if len(report_text.strip()) < 40:
        warnings = list(extracted.get("warnings") or [])
        result.setdefault("recommendation", {}).setdefault("high_risk_flags", []).append(
            warnings[0] if warnings else "Upload text was not readable"
        )
    return result


@router.post("/bank-brain/structure")
def bank_brain_structure(request: CreditStructureRequest) -> dict[str, Any]:
    return _simulate_credit_structure(request)


@router.post("/bank-brain/recommend-vehicles")
def bank_brain_recommend_vehicles(request: BankBrainVehicleRecommendationRequest) -> dict[str, Any]:
    return _recommend_vehicles_for_customer(request)


@router.post("/bank-brain/decision")
def bank_brain_decision(request: BankBrainDecisionRequest) -> dict[str, Any]:
    entry = {
        "vin": request.vin,
        "bank_code": request.bank_code.strip().upper(),
        "outcome": request.outcome,
        "notes": request.notes,
        "metrics": request.metrics,
        "created_at": datetime.now(timezone.utc).isoformat(),
    }
    _append_bank_brain_history(entry)
    _append_audit_event("bank_brain_decision", entry)
    return {"ok": True, "entry": entry, "history_count": len(_load_bank_brain_history())}


@router.get("/bank-brain/history")
def bank_brain_history(limit: int = 200) -> dict[str, Any]:
    history = _load_bank_brain_history()
    normalized_limit = max(1, min(int(limit), 1000))
    return {"items": history[-normalized_limit:], "total": len(history)}


@router.get("/bank-brain/lenders")
def bank_brain_lenders() -> dict[str, Any]:
    return {"items": _active_bank_profiles()}


@router.get("/bank-brain/valuations/status")
def bank_brain_valuations_status() -> dict[str, Any]:
    payload = _safe_read_json(JD_POWER_VALUATIONS_PATH, {"items": []})
    items = payload.get("items", []) if isinstance(payload, dict) else []
    return {
        "ok": JD_POWER_VALUATIONS_PATH.exists(),
        "count": len(items) if isinstance(items, list) else 0,
        "source_file": payload.get("source_file") if isinstance(payload, dict) else None,
        "updated_at": payload.get("updated_at") if isinstance(payload, dict) else None,
    }


@router.post("/bank-brain/valuations/upload")
async def bank_brain_valuations_upload(file: UploadFile = File(...)) -> dict[str, Any]:
    raw = await file.read()
    parsed = _parse_jd_power_file(raw, file.filename or "jd_power_values.xls", file.content_type)
    items = parsed.get("items", [])
    saved = _save_jd_power_valuations(items, file.filename or "jd_power_values.xls")
    return {
        "ok": True,
        "count": saved.get("count", 0),
        "source_file": saved.get("source_file"),
        "updated_at": saved.get("updated_at"),
        "diagnostics": parsed.get("diagnostics", {}),
    }


@router.get("/bank-brain/docs/status")
def bank_brain_docs_status() -> dict[str, Any]:
    return _routeone_docs_status()


@router.post("/bank-brain/docs/rebuild")
def bank_brain_docs_rebuild(request: BankBrainDocsRebuildRequest) -> dict[str, Any]:
    return _run_bank_docs_rebuild(
        reload_sales_data=request.reload_sales_data,
        max_link_depth=request.max_link_depth,
        max_links_per_resource=request.max_links_per_resource,
    )


@router.post("/bank-brain/docs/upload")
async def bank_brain_docs_upload(
    files: list[UploadFile] = File(...),
    bank: str | None = Form(default=None),
    rebuild: bool = Form(default=True),
    reload_sales_data: bool = Form(default=True),
) -> dict[str, Any]:
    bank_folder = _sanitize_doc_segment(bank or "_Inbox", "_Inbox")
    target_dir = BANK_DOCS_ROOT / bank_folder
    target_dir.mkdir(parents=True, exist_ok=True)

    saved: list[dict[str, Any]] = []
    for upload in files:
        source_name = Path(upload.filename or "routeone_document.pdf").name
        target = _unique_upload_target(target_dir, source_name)
        raw = await upload.read()
        target.write_bytes(raw)
        extracted = _extract_upload_document(raw, source_name, upload.content_type)
        saved.append(
            {
                "filename": source_name,
                "stored_as": _display_path(target),
                "bank": bank_folder,
                "bytes": len(raw),
                "extracted_text_chars": extracted.get("extracted_text_chars", 0),
                "file_understanding": {
                    key: value
                    for key, value in extracted.items()
                    if key not in {"text"} and value not in (None, "", [], {})
                },
            }
        )

    rebuild_result = None
    if rebuild and saved:
        rebuild_result = _run_bank_docs_rebuild(
            reload_sales_data=reload_sales_data,
            max_link_depth=1,
            max_links_per_resource=12,
        )

    return {
        "ok": bool(saved) and (rebuild_result is None or bool(rebuild_result.get("ok"))),
        "saved": saved,
        "rebuild": rebuild_result,
        "status": _routeone_docs_status(),
    }
